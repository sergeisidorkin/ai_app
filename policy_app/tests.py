import csv
import copy
import io
import json
from decimal import Decimal

from docx import Document
from openpyxl import Workbook, load_workbook

from django.contrib.auth import get_user_model
from django.db import connection
from django.db.migrations.executor import MigrationExecutor
from django.core.files.uploadedfile import SimpleUploadedFile
from django.http import QueryDict
from django.test import Client, TestCase, TransactionTestCase
from django.urls import reverse

from classifiers_app.models import OKVCurrency
from experts_app.models import ExpertProfile, ExpertProfileSpecialty, ExpertSpecialty
from group_app.models import GroupMember, OrgUnit
from policy_app.forms import (
    ProductForm,
    ReportStructureForm,
    SectionStructureForm,
    ServiceGoalReportForm,
    TariffForm,
    TypicalServiceCompositionForm,
)
from policy_app.models import (
    ConsultingDirection,
    ConsultingDirectionType,
    ConsultingServiceSubtype,
    ConsultingServiceType,
    ExpertiseDirection,
    Product,
    ReportStructure,
    SectionStructure,
    ServiceGoalReport,
    SpecialtyTariff,
    Tariff,
    TypicalSection,
    TypicalSectionSpecialty,
    TypicalServiceComposition,
    TypicalServiceTerm,
    ensure_system_dsc_section,
)
from users_app.models import Employee


class RemoveTypicalSectionExecutorMigrationTests(TransactionTestCase):
    migrate_from = ("policy_app", "0033_typicalsection_exclude_from_tkp_autofill")
    migrate_to = ("policy_app", "0034_remove_typicalsection_executor")

    def setUp(self):
        super().setUp()
        self.executor = MigrationExecutor(connection)
        self.executor.migrate([self.migrate_from])
        old_apps = self.executor.loader.project_state([self.migrate_from]).apps

        Product = old_apps.get_model("policy_app", "Product")
        TypicalSection = old_apps.get_model("policy_app", "TypicalSection")
        ExpertSpecialty = old_apps.get_model("experts_app", "ExpertSpecialty")
        TypicalSectionSpecialty = old_apps.get_model("policy_app", "TypicalSectionSpecialty")
        if connection.vendor == "postgresql":
            with connection.cursor() as cursor:
                cursor.execute("ALTER TABLE experts_app_expertspecialty ALTER COLUMN specialization_area SET DEFAULT ''")

        product = Product.objects.create(
            short_name="MIG",
            name_en="Migration product",
            display_name="Migration product",
            name_ru="Миграционный продукт",
            service_type="Консалтинг",
            position=1,
        )
        existing_specialty = ExpertSpecialty.objects.create(
            specialty="Юрист",
            specialty_en="",
            position=1,
        )
        self.section = TypicalSection.objects.create(
            product_id=product.pk,
            code="MIG-1",
            short_name="mig-1",
            short_name_ru="mig-1",
            name_en="Migration section",
            name_ru="Миграционный раздел",
            accounting_type="Раздел",
            executor="Партнер; Юрист",
            position=1,
        )
        TypicalSectionSpecialty.objects.create(
            section_id=self.section.pk,
            specialty_id=existing_specialty.pk,
            rank=1,
        )

    def tearDown(self):
        executor = MigrationExecutor(connection)
        executor.migrate(executor.loader.graph.leaf_nodes())
        super().tearDown()

    def test_migration_backfills_executor_into_ranked_specialties(self):
        self.executor.loader.build_graph()
        self.executor.migrate([self.migrate_to])
        new_apps = self.executor.loader.project_state([self.migrate_to]).apps

        ExpertSpecialty = new_apps.get_model("experts_app", "ExpertSpecialty")
        TypicalSectionSpecialty = new_apps.get_model("policy_app", "TypicalSectionSpecialty")

        created_names = list(
            ExpertSpecialty.objects.filter(specialty__in=["Партнер", "Юрист"])
            .order_by("specialty")
            .values_list("specialty", flat=True)
        )
        self.assertEqual(created_names, ["Партнер", "Юрист"])

        links = list(
            TypicalSectionSpecialty.objects.filter(section_id=self.section.pk)
            .select_related("specialty")
            .order_by("rank")
        )
        self.assertEqual([(link.specialty.specialty, link.rank) for link in links], [("Юрист", 1), ("Партнер", 2)])


class SystemDscMigrationTests(TransactionTestCase):
    migrate_from = ("policy_app", "0044_typicalserviceterm_source_data_weeks")
    migrate_to = ("policy_app", "0045_typicalsection_is_system_dsc")

    def setUp(self):
        super().setUp()
        self.executor = MigrationExecutor(connection)
        self.executor.migrate([self.migrate_from])
        old_apps = self.executor.loader.project_state([self.migrate_from]).apps

        Product = old_apps.get_model("policy_app", "Product")
        TypicalSection = old_apps.get_model("policy_app", "TypicalSection")

        product = Product.objects.create(
            short_name="MIG-DSC",
            name_en="Migration DSC",
            display_name="Migration DSC",
            name_ru="Миграция DSC",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        self.lower_dsc_id = TypicalSection.objects.create(
            product_id=product.pk,
            code="dsc",
            short_name="manual-lower",
            short_name_ru="ручной lower",
            name_en="Manual lower",
            name_ru="Ручной lower",
            accounting_type="Услуги",
            position=1,
        ).pk
        self.upper_dsc_id = TypicalSection.objects.create(
            product_id=product.pk,
            code="DSC",
            short_name="manual-upper",
            short_name_ru="ручной upper",
            name_en="Manual upper",
            name_ru="Ручной upper",
            accounting_type="Услуги",
            position=2,
        ).pk
        TypicalSection.objects.create(
            product_id=product.pk,
            code="REG",
            short_name="regular",
            short_name_ru="обычный",
            name_en="Regular",
            name_ru="Обычный",
            accounting_type="Раздел",
            position=3,
        )

    def tearDown(self):
        executor = MigrationExecutor(connection)
        executor.migrate(executor.loader.graph.leaf_nodes())
        super().tearDown()

    def test_migration_deduplicates_case_insensitive_dsc_before_canonical_save(self):
        self.executor.loader.build_graph()
        self.executor.migrate([self.migrate_to])
        new_apps = self.executor.loader.project_state([self.migrate_to]).apps

        TypicalSection = new_apps.get_model("policy_app", "TypicalSection")

        sections = list(TypicalSection.objects.order_by("position", "id"))
        self.assertEqual([section.code for section in sections], ["DSC", "REG"])
        self.assertEqual(sections[0].pk, self.lower_dsc_id)
        self.assertTrue(sections[0].is_system)
        self.assertEqual(sections[0].name_ru, "Описание продукта")
        self.assertFalse(TypicalSection.objects.filter(pk=self.upper_dsc_id).exists())


class ConsultingCatalogBackfillMigrationTests(TransactionTestCase):
    migrate_from = ("policy_app", "0038_product_consulting_service_fields")
    migrate_to = ("policy_app", "0040_backfill_consulting_catalog_refs")

    def setUp(self):
        super().setUp()
        self.executor = MigrationExecutor(connection)
        self.executor.migrate([self.migrate_from])
        old_apps = self.executor.loader.project_state([self.migrate_from]).apps

        Product = old_apps.get_model("policy_app", "Product")
        self.product = Product.objects.create(
            short_name="MIG-CAT",
            name_en="Migration catalog",
            display_name="Migration catalog",
            name_ru="Миграционный каталог",
            consulting_type="Горный",
            service_category="Аудит",
            service_code="A",
            service_subtype="Аудит проектных решений",
            position=1,
        )

    def tearDown(self):
        executor = MigrationExecutor(connection)
        executor.migrate(executor.loader.graph.leaf_nodes())
        super().tearDown()

    def test_migration_seeds_catalog_and_links_products(self):
        self.executor.loader.build_graph()
        self.executor.migrate([self.migrate_to])
        new_apps = self.executor.loader.project_state([self.migrate_to]).apps

        Product = new_apps.get_model("policy_app", "Product")
        ConsultingDirection = new_apps.get_model("policy_app", "ConsultingDirection")
        ConsultingDirectionType = new_apps.get_model("policy_app", "ConsultingDirectionType")
        ConsultingServiceType = new_apps.get_model("policy_app", "ConsultingServiceType")
        ConsultingServiceSubtype = new_apps.get_model("policy_app", "ConsultingServiceSubtype")

        migrated = Product.objects.get(pk=self.product.pk)
        self.assertIsNotNone(migrated.consulting_type_ref_id)
        self.assertIsNotNone(migrated.service_category_ref_id)
        self.assertIsNotNone(migrated.service_subtype_ref_id)
        self.assertEqual(ConsultingDirection.objects.count(), 1)
        self.assertTrue(ConsultingDirectionType.objects.filter(name="Горный").exists())
        self.assertTrue(
            ConsultingServiceType.objects.filter(name="Аудит", code="A").exists()
        )
        self.assertTrue(
            ConsultingServiceSubtype.objects.filter(name="Аудит проектных решений").exists()
        )


class ProductCsvUploadTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-products-admin",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.owner = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )

    def test_product_csv_upload_creates_products_and_derives_code(self):
        csv_file = SimpleUploadedFile(
            "products.csv",
            (
                "Краткое имя;Наименование на английском языке;Наименование на русском языке;"
                "Отображаемое в системе имя;Вид консалтинга;Тип услуг;Код;Подтип услуги;Владелец\n"
                "AUD;Audit;Аудит;Аудит продукта;Горный;Аудит;Z;Аудит проектных решений;IMC Montan\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("product_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(len(response.json()["warnings"]), 1)
        product = Product.objects.get(short_name="AUD")
        self.assertEqual(product.name_en, "Audit")
        self.assertEqual(product.name_ru, "Аудит")
        self.assertEqual(product.display_name, "Аудит продукта")
        self.assertEqual(product.consulting_type, "Горный")
        self.assertEqual(product.service_category, "Аудит")
        self.assertEqual(product.service_code, "A")
        self.assertEqual(product.service_subtype, "Аудит проектных решений")
        self.assertIsNotNone(product.consulting_type_ref_id)
        self.assertIsNotNone(product.service_category_ref_id)
        self.assertIsNotNone(product.service_subtype_ref_id)
        self.assertFalse(product.is_group_owner)
        self.assertEqual(list(product.owners.values_list("short_name", flat=True)), ["IMC Montan"])

    def test_policy_partial_renders_product_csv_download_button(self):
        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Типовые продукты")
        self.assertContains(response, 'id="products-csv-download-btn"', html=False)

    def test_policy_partial_renders_expertise_direction_specialization_area(self):
        ExpertiseDirection.objects.create(
            short_name="ГЭ",
            name="Горная экспертиза",
            pricing_method="vpm",
            specialization_area="Специалисты по горным работам",
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Направления экспертизы")
        self.assertContains(response, "Расчет стоимости услуг")
        self.assertContains(response, "Область специализации")
        self.assertContains(response, "Специалисты по горным работам")

    def test_expertise_direction_form_renders_specialization_area_suffix(self):
        response = self.client.get(reverse("expertise_dir_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Область специализации")
        self.assertContains(response, "Специалисты по")
        self.assertContains(response, 'name="specialization_area_suffix"', html=False)

    def test_expertise_direction_form_saves_specialization_area_with_locked_prefix(self):
        response = self.client.post(
            reverse("expertise_dir_form_create"),
            {
                "short_name": "ГЭ",
                "name": "Горная экспертиза",
                "pricing_method": "vpm",
                "specialization_area_suffix": "горным работам",
                "owner_ids": "__group__",
            },
        )

        self.assertEqual(response.status_code, 200)
        direction = ExpertiseDirection.objects.get(short_name="ГЭ")
        self.assertEqual(direction.specialization_area, "Специалисты по горным работам")
        self.assertContains(response, "Специалисты по горным работам")

    def test_expertise_direction_edit_form_prefills_specialization_area_suffix(self):
        direction = ExpertiseDirection.objects.create(
            short_name="ГЭ",
            name="Горная экспертиза",
            specialization_area="Специалисты по горным работам",
            position=1,
        )

        response = self.client.get(reverse("expertise_dir_form_edit", args=[direction.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Специалисты по")
        self.assertContains(response, 'value="горным работам"', html=False)

    def test_product_csv_download_exports_current_table_columns(self):
        product = Product.objects.create(
            short_name="AUD",
            name_en="Audit",
            name_ru="Аудит",
            display_name="Аудит продукта",
            consulting_type="Горный",
            service_category="Аудит",
            service_code="A",
            service_subtype="Аудит проектных решений",
            position=1,
        )
        product.owners.set([self.owner])

        response = self.client.get(reverse("product_csv_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("typical_products.csv", response["Content-Disposition"])
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(
            rows[0],
            [
                "Краткое имя",
                "Наименование на английском языке",
                "Наименование на русском языке",
                "Отображаемое в системе имя",
                "Вид консалтинга",
                "Тип услуг",
                "Код",
                "Подтип услуги",
                "Владелец",
            ],
        )
        self.assertEqual(
            rows[1],
            [
                "AUD",
                "Audit",
                "Аудит",
                "Аудит продукта",
                "Горный",
                "Аудит",
                "A",
                "Аудит проектных решений",
                "IMC Montan",
            ],
        )

    def test_product_csv_download_respects_product_filter(self):
        tax_product = Product.objects.create(
            short_name="TAX",
            name_en="Tax",
            name_ru="Налоги",
            display_name="Налоговый продукт",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_code="I",
            service_subtype="По международным стандартам",
            position=1,
        )
        Product.objects.create(
            short_name="AUD",
            name_en="Audit",
            name_ru="Аудит",
            display_name="Аудит продукта",
            consulting_type="Горный",
            service_category="Аудит",
            service_code="A",
            service_subtype="Аудит проектных решений",
            position=2,
        )

        response = self.client.get(
            reverse("product_csv_download"),
            {"product": [tax_product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "TAX")


class ProductFormTests(TestCase):
    def setUp(self):
        self.consulting_type = ConsultingDirectionType.objects.create(
            name="Горный ProductForm",
            position=1,
            direction=ConsultingDirection.objects.create(position=1),
        )
        self.service_type = ConsultingServiceType.objects.create(
            direction=self.consulting_type.direction,
            consulting_type=self.consulting_type,
            name="Аудит ProductForm",
            code="A",
            position=1,
        )
        self.service_subtype = ConsultingServiceSubtype.objects.create(
            direction=self.consulting_type.direction,
            service_type=self.service_type,
            name="Аудит проектных решений ProductForm",
            position=1,
        )

    def test_init_ignores_non_numeric_dependent_catalog_ids(self):
        data = QueryDict("", mutable=True)
        data.update(
            {
                "short_name": "AUD",
                "name_en": "Audit",
                "display_name": "Audit",
                "name_ru": "Аудит",
                "consulting_type_ref": "oops",
                "service_category_ref": "nan",
                "service_subtype_ref": str(self.service_subtype.pk),
            }
        )
        form = ProductForm(data=data)

        self.assertEqual(list(form.fields["service_category_ref"].queryset), [])
        self.assertEqual(list(form.fields["service_subtype_ref"].queryset), [])


class PolicyMasterFilterTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-master-filter-admin",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        direction = ConsultingDirection.objects.create(position=1)
        self.consulting_type = ConsultingDirectionType.objects.create(
            direction=direction,
            name="Горный Master",
            position=1,
        )
        self.service_type = ConsultingServiceType.objects.create(
            direction=direction,
            consulting_type=self.consulting_type,
            name="Аудит Master",
            code="AM",
            position=1,
        )
        self.service_subtype = ConsultingServiceSubtype.objects.create(
            direction=direction,
            service_type=self.service_type,
            name="Аудит проектных решений Master",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="MF",
            name_en="Master filter product",
            display_name="Master filter product",
            name_ru="Продукт мастер-фильтра",
            consulting_type_ref=self.consulting_type,
            service_category_ref=self.service_type,
            service_subtype_ref=self.service_subtype,
            position=1,
        )
        self.other_product = Product.objects.create(
            short_name="OTHER-MF",
            name_en="Other product",
            display_name="Other product",
            name_ru="Другой продукт",
            consulting_type_ref=self.consulting_type,
            service_category_ref=self.service_type,
            service_subtype_ref=self.service_subtype,
            position=2,
        )
        self.section = TypicalSection.objects.create(
            product=self.product,
            code="MF-1",
            short_name="mf-1",
            short_name_ru="мф-1",
            name_en="Master section",
            name_ru="Раздел мастер-фильтра",
            accounting_type="Раздел",
            position=1,
        )
        self.other_section = TypicalSection.objects.create(
            product=self.other_product,
            code="OMF-1",
            short_name="omf-1",
            short_name_ru="омф-1",
            name_en="Other section",
            name_ru="Другой раздел",
            accounting_type="Раздел",
            position=1,
        )
        ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Цель",
            service_goal_genitive="Цели",
            report_title="Отчет",
            product_name="Название",
            position=1,
        )
        Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm=Decimal("1.00"),
            service_hours=1,
            service_days_tkp=1,
            created_by=self.user,
            position=1,
        )

    def test_policy_partial_renders_master_filter_row_metadata(self):
        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-target-name="product-select"', html=False)
        self.assertContains(response, 'name="product-select"', html=False)
        self.assertContains(response, 'data-policy-filter-row="1"', html=False)
        self.assertContains(response, f'data-product-id="{self.product.pk}"', html=False)
        self.assertContains(response, 'data-product-label="MF Master filter product"', html=False)
        self.assertContains(response, 'data-consulting-type="Горный Master"', html=False)
        self.assertContains(response, 'data-service-category="Аудит Master"', html=False)
        self.assertContains(response, 'data-service-subtype="Аудит проектных решений Master"', html=False)

    def test_product_create_prefills_catalog_fields_from_selected_product(self):
        response = self.client.get(reverse("product_form_create"), {"product": self.product.pk})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'<option value="{self.consulting_type.pk}" selected', html=False)
        self.assertContains(response, f'<option value="{self.service_type.pk}" selected', html=False)
        self.assertContains(response, f'<option value="{self.service_subtype.pk}" selected', html=False)

    def test_product_create_prefills_catalog_fields_from_direct_refs(self):
        response = self.client.get(
            reverse("product_form_create"),
            {
                "consulting_type_ref": self.consulting_type.pk,
                "service_category_ref": self.service_type.pk,
                "service_subtype_ref": self.service_subtype.pk,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'<option value="{self.consulting_type.pk}" selected', html=False)
        self.assertContains(response, f'<option value="{self.service_type.pk}" selected', html=False)
        self.assertContains(response, f'<option value="{self.service_subtype.pk}" selected', html=False)

    def test_create_forms_prefill_product_from_master_filter_param(self):
        url_names = [
            "section_form_create",
            "structure_form_create",
            "service_goal_report_form_create",
            "typical_service_composition_form_create",
            "typical_service_term_form_create",
            "tariff_form_create",
        ]

        for url_name in url_names:
            with self.subTest(url_name=url_name):
                response = self.client.get(reverse(url_name), {"product": self.product.pk})
                self.assertEqual(response.status_code, 200)
                self.assertContains(response, f'<option value="{self.product.pk}" selected', html=False)

    def test_dependent_section_fields_are_limited_by_selected_product(self):
        structure_form = SectionStructureForm(initial={"product": self.product.pk})
        tariff_form = TariffForm(initial={"product": self.product.pk}, request_user=self.user)

        self.assertEqual(list(structure_form.fields["section"].queryset), [self.section])
        self.assertEqual(list(tariff_form.fields["section"].queryset), [self.section])

    def test_dependent_section_forms_reject_malformed_bound_ids_without_crashing(self):
        form_configs = [
            (
                "structure",
                SectionStructureForm,
                {"subsections": "Подраздел"},
                {},
            ),
            (
                "service_composition",
                TypicalServiceCompositionForm,
                {"service_composition": "Состав услуг", "service_composition_editor_state": ""},
                {},
            ),
            (
                "tariff",
                TariffForm,
                {"base_rate_vpm": "1.00", "service_hours": "1", "service_days_tkp": "1"},
                {"request_user": self.user},
            ),
        ]
        malformed_cases = [
            ({"product": "abc", "section": str(self.section.pk)}, "product"),
            ({"product": str(self.product.pk), "section": "abc"}, "section"),
        ]

        for form_name, form_class, base_data, form_kwargs in form_configs:
            for malformed_data, error_field in malformed_cases:
                with self.subTest(form_name=form_name, error_field=error_field):
                    form = form_class(data={**base_data, **malformed_data}, **form_kwargs)

                    self.assertFalse(form.is_valid())
                    self.assertIn(error_field, form.errors)


class ConsultingDirectionViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-consulting-admin",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        ConsultingDirection.objects.all().delete()

    def test_policy_partial_renders_consulting_direction_table(self):
        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Направления консалтинга")
        self.assertContains(response, 'id="consulting-dir-actions"', html=False)
        self.assertContains(response, 'id="consulting-dir-master"', html=False)

    def test_create_consulting_direction_saves_nested_catalog(self):
        response = self.client.post(
            reverse("consulting_dir_form_create"),
            {
                "consulting_types_payload": json.dumps(
                    [{"id": "", "name": "Финансовый"}], ensure_ascii=False
                ),
                "service_types_payload": json.dumps(
                    [
                        {
                            "id": "",
                            "consulting_type": "Финансовый",
                            "name": "Due diligence",
                            "code": "D",
                        }
                    ],
                    ensure_ascii=False,
                ),
                "service_subtypes_payload": json.dumps(
                    [
                        {
                            "id": "",
                            "consulting_type": "Финансовый",
                            "service_type": "Due diligence",
                            "name": "Экспресс",
                        }
                    ],
                    ensure_ascii=False,
                ),
            },
        )

        self.assertEqual(response.status_code, 200)
        direction = ConsultingDirection.objects.get(
            consulting_types__name="Финансовый"
        )
        self.assertEqual(direction.service_types.get().code, "D")
        self.assertEqual(direction.service_subtypes.get().name, "Экспресс")

    def test_move_up_reorders_consulting_directions(self):
        first = ConsultingDirection.objects.create(position=1)
        second = ConsultingDirection.objects.create(position=2)

        response = self.client.post(reverse("consulting_dir_move_up", args=[second.pk]))

        self.assertEqual(response.status_code, 200)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(first.position, 2)
        self.assertEqual(second.position, 1)


class TypicalSectionViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-sections-admin",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="SEC",
            name_en="Sections",
            display_name="Sections",
            name_ru="Разделы",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )

    def test_policy_partial_renders_tkp_column_for_sections(self):
        TypicalSection.objects.create(
            product=self.product,
            code="SEC-1",
            short_name="section-en",
            short_name_ru="section-ru",
            name_en="Section EN",
            name_ru="Раздел RU",
            accounting_type="Раздел",
            exclude_from_tkp_autofill=True,
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Типовые разделы (услуги)")
        self.assertContains(response, '<th><i class="bi bi-ban me-1"></i>ТКП</th>', html=False)
        self.assertContains(response, 'aria-label="Исключить из автозаполнения в ТКП"', html=False)
        self.assertContains(response, 'id="sections-csv-download-btn"', html=False)

    def test_create_section_saves_tkp_exclusion_flag(self):
        response = self.client.post(
            reverse("section_form_create"),
            {
                "product": self.product.pk,
                "code": "SEC-2",
                "short_name": "audit-en",
                "short_name_ru": "audit-ru",
                "name_en": "Audit EN",
                "name_ru": "Аудит RU",
                "accounting_type": "Раздел",
                "exclude_from_tkp_autofill": "on",
            },
        )

        self.assertEqual(response.status_code, 200)
        section = TypicalSection.objects.get(code="SEC-2")
        self.assertTrue(section.exclude_from_tkp_autofill)

    def test_create_section_auto_creates_system_dsc_first(self):
        response = self.client.post(
            reverse("section_form_create"),
            {
                "product": self.product.pk,
                "code": "SEC-AUTO",
                "short_name": "auto-en",
                "short_name_ru": "auto-ru",
                "name_en": "Auto EN",
                "name_ru": "Авто RU",
                "accounting_type": "Раздел",
            },
        )

        self.assertEqual(response.status_code, 200)
        sections = list(self.product.sections.order_by("position", "id"))
        self.assertEqual([section.code for section in sections], ["DSC", "SEC-AUTO"])
        dsc = sections[0]
        self.assertTrue(dsc.is_system)
        self.assertEqual(dsc.short_name, "Description")
        self.assertEqual(dsc.short_name_ru, "Описание")
        self.assertEqual(dsc.name_en, "Product description")
        self.assertEqual(dsc.name_ru, "Описание продукта")
        self.assertEqual(dsc.accounting_type, "Раздел")
        self.assertFalse(dsc.exclude_from_tkp_autofill)
        self.assertFalse(dsc.ranked_specialties.exists())

    def test_ensure_system_dsc_canonicalizes_existing_row(self):
        existing = TypicalSection.objects.create(
            product=self.product,
            code="dsc",
            short_name="manual",
            short_name_ru="ручной",
            name_en="Manual",
            name_ru="Ручной",
            accounting_type="Услуги",
            exclude_from_tkp_autofill=True,
            position=5,
        )
        TypicalSection.objects.create(
            product=self.product,
            code="SEC-BACKFILL",
            short_name="backfill-en",
            short_name_ru="backfill-ru",
            name_en="Backfill EN",
            name_ru="Backfill RU",
            accounting_type="Раздел",
            position=1,
        )

        dsc = ensure_system_dsc_section(self.product)

        self.assertEqual(dsc.pk, existing.pk)
        self.assertEqual(dsc.code, "DSC")
        self.assertTrue(dsc.is_system)
        self.assertEqual(dsc.short_name, "Description")
        self.assertFalse(dsc.exclude_from_tkp_autofill)
        self.assertEqual(
            list(self.product.sections.order_by("position", "id").values_list("code", flat=True)),
            ["DSC", "SEC-BACKFILL"],
        )

    def test_create_section_rejects_manual_dsc_code(self):
        response = self.client.post(
            reverse("section_form_create"),
            {
                "product": self.product.pk,
                "code": "DSC",
                "short_name": "manual",
                "short_name_ru": "manual",
                "name_en": "Manual",
                "name_ru": "Ручной",
                "accounting_type": "Раздел",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Раздел DSC является системным")
        self.assertFalse(TypicalSection.objects.filter(product=self.product, code="DSC").exists())

    def test_system_dsc_cannot_be_deleted_or_moved(self):
        dsc = ensure_system_dsc_section(self.product)
        regular = TypicalSection.objects.create(
            product=self.product,
            code="SEC-LOCK",
            short_name="lock-en",
            short_name_ru="lock-ru",
            name_en="Lock EN",
            name_ru="Блок RU",
            accounting_type="Раздел",
            position=2,
        )
        ensure_system_dsc_section(self.product)

        delete_response = self.client.post(reverse("section_delete", args=[dsc.pk]))
        move_response = self.client.post(reverse("section_move_down", args=[dsc.pk]))
        regular_move_response = self.client.post(reverse("section_move_up", args=[regular.pk]))

        self.assertEqual(delete_response.status_code, 400)
        self.assertEqual(move_response.status_code, 200)
        self.assertEqual(regular_move_response.status_code, 200)
        self.assertEqual(
            list(self.product.sections.order_by("position", "id").values_list("code", flat=True)),
            ["DSC", "SEC-LOCK"],
        )

    def test_section_form_renders_product_options_with_display_name(self):
        response = self.client.get(reverse("section_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<select name="product"', html=False)
        self.assertContains(response, "policy-product-select")
        self.assertContains(response, 'data-short-label="SEC"', html=False)
        self.assertContains(response, "SEC Sections")

    def test_section_csv_upload_accepts_rows_without_legacy_executor_column(self):
        csv_file = SimpleUploadedFile(
            "sections.csv",
            (
                "Продукт;Код;Краткое имя EN;Краткое имя RU;Наименование EN;Наименование RU;Тип учета;Направление экспертизы\n"
                "SEC;SEC-3;section-3;section-3-ru;Section 3;Раздел 3;Раздел;\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("section_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        self.assertTrue(TypicalSection.objects.filter(code="SEC-3").exists())

    def test_section_csv_upload_skips_manual_dsc_and_updates_system_row(self):
        csv_file = SimpleUploadedFile(
            "sections.csv",
            (
                "Продукт;Код;Краткое имя EN;Краткое имя RU;Наименование EN;Наименование RU;Тип учета;Направление экспертизы\n"
                "SEC;DSC;manual;ручной;Manual;Ручной;Услуги;\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("section_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertIn("раздел DSC является системным", response.json()["warnings"][0])
        dsc = TypicalSection.objects.get(product=self.product, code="DSC")
        self.assertTrue(dsc.is_system)
        self.assertEqual(dsc.short_name, "Description")
        self.assertEqual(dsc.position, 1)

    def test_section_csv_download_exports_current_table_columns(self):
        owner = GroupMember.objects.create(
            short_name="IMC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        department = OrgUnit.objects.create(
            company=owner,
            level=1,
            department_name="Налоговый департамент",
            short_name="TAX-DEPT",
            unit_type="expertise",
            position=1,
        )
        expertise = ExpertiseDirection.objects.create(
            name="Налоги",
            short_name="TAX",
            position=1,
        )
        specialty = ExpertSpecialty.objects.create(
            expertise_direction=department,
            expertise_dir=expertise,
            specialty="Налоговый due diligence",
            position=1,
        )
        section = TypicalSection.objects.create(
            product=self.product,
            code="SEC-4",
            short_name="tax-dd",
            short_name_ru="нал-dd",
            name_en="Tax DD",
            name_ru="Налоговый ДД",
            accounting_type="Услуги",
            expertise_dir=expertise,
            expertise_direction=department,
            exclude_from_tkp_autofill=True,
            position=1,
        )
        TypicalSectionSpecialty.objects.create(section=section, specialty=specialty, rank=1)

        response = self.client.get(reverse("section_csv_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("typical_sections.csv", response["Content-Disposition"])
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(
            rows[0],
            [
                "Продукт",
                "Код",
                "Краткое имя EN",
                "Краткое имя RU",
                "Наименование раздела (услуги) EN",
                "Наименование раздела (услуги) RU",
                "Тип учета",
                "Исполнитель",
                "Экспертиза",
                "Подразделение",
                "ТКП",
            ],
        )
        self.assertEqual(
            next(row for row in rows[1:] if row[1] == "SEC-4"),
            [
                "SEC",
                "SEC-4",
                "tax-dd",
                "нал-dd",
                "Tax DD",
                "Налоговый ДД",
                "Услуги",
                "Налоговый due diligence",
                "TAX",
                "Налоговый департамент",
                "Да",
            ],
        )

    def test_section_csv_download_respects_product_filter(self):
        other_product = Product.objects.create(
            short_name="SEC2",
            name_en="Sections 2",
            display_name="Sections 2",
            name_ru="Разделы 2",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        TypicalSection.objects.create(
            product=self.product,
            code="SEC-A",
            short_name="section-a",
            short_name_ru="section-a-ru",
            name_en="Section A",
            name_ru="Раздел A",
            accounting_type="Раздел",
            position=1,
        )
        TypicalSection.objects.create(
            product=other_product,
            code="SEC-B",
            short_name="section-b",
            short_name_ru="section-b-ru",
            name_en="Section B",
            name_ru="Раздел B",
            accounting_type="Раздел",
            position=1,
        )

        response = self.client.get(
            reverse("section_csv_download"),
            {"product": [self.product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "SEC")
        self.assertEqual(rows[1][1], "SEC-A")

    def test_section_csv_upload_accepts_current_table_export_columns(self):
        owner = GroupMember.objects.create(
            short_name="IMC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        department = OrgUnit.objects.create(
            company=owner,
            level=1,
            department_name="Налоговый департамент",
            short_name="TAX-DEPT",
            unit_type="expertise",
            position=1,
        )
        expertise = ExpertiseDirection.objects.create(
            name="Налоги",
            short_name="TAX",
            position=1,
        )
        first_specialty = ExpertSpecialty.objects.create(
            expertise_direction=department,
            expertise_dir=expertise,
            specialty="Налоговый due diligence",
            position=1,
        )
        second_specialty = ExpertSpecialty.objects.create(
            expertise_direction=department,
            expertise_dir=expertise,
            specialty="Трансфертное ценообразование",
            position=2,
        )
        csv_file = SimpleUploadedFile(
            "sections.csv",
            (
                "Продукт;Код;Краткое имя EN;Краткое имя RU;"
                "Наименование раздела (услуги) EN;Наименование раздела (услуги) RU;"
                "Тип учета;Исполнитель;Экспертиза;Подразделение;ТКП\n"
                "SEC;SEC-5;tax-dd;нал-dd;Tax DD;Налоговый ДД;Услуги;"
                "Налоговый due diligence, Трансфертное ценообразование;TAX;Налоговый департамент;Да\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("section_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        section = TypicalSection.objects.get(code="SEC-5")
        self.assertEqual(section.expertise_dir, expertise)
        self.assertEqual(section.expertise_direction, department)
        self.assertTrue(section.exclude_from_tkp_autofill)
        self.assertEqual(
            list(section.ranked_specialties.values_list("specialty", flat=True)),
            [first_specialty.pk, second_specialty.pk],
        )

    def test_section_csv_upload_updates_existing_section_by_product_and_code(self):
        section = TypicalSection.objects.create(
            product=self.product,
            code="SEC-UPD",
            short_name="old-en",
            short_name_ru="old-ru",
            name_en="Old EN",
            name_ru="Старый раздел",
            accounting_type="Раздел",
            exclude_from_tkp_autofill=False,
            position=1,
        )
        csv_file = SimpleUploadedFile(
            "sections.csv",
            (
                "Продукт;Код;Краткое имя EN;Краткое имя RU;"
                "Наименование раздела (услуги) EN;Наименование раздела (услуги) RU;"
                "Тип учета;Исполнитель;Экспертиза;Подразделение;ТКП\n"
                "SEC;SEC-UPD;new-en;new-ru;New EN;Новый раздел;Услуги;;;;Да\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("section_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(response.json()["updated"], 1)
        section.refresh_from_db()
        self.assertEqual(section.short_name, "new-en")
        self.assertEqual(section.short_name_ru, "new-ru")
        self.assertEqual(section.name_en, "New EN")
        self.assertEqual(section.name_ru, "Новый раздел")
        self.assertEqual(section.accounting_type, "Услуги")
        self.assertTrue(section.exclude_from_tkp_autofill)

    def test_section_csv_upload_rolls_back_row_when_specialty_insert_fails(self):
        owner = GroupMember.objects.create(
            short_name="IMC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        department = OrgUnit.objects.create(
            company=owner,
            level=1,
            department_name="Налоговый департамент",
            short_name="TAX-DEPT",
            unit_type="expertise",
            position=1,
        )
        expertise = ExpertiseDirection.objects.create(
            name="Налоги",
            short_name="TAX",
            position=1,
        )
        specialty = ExpertSpecialty.objects.create(
            expertise_direction=department,
            expertise_dir=expertise,
            specialty="Налоговый due diligence",
            position=1,
        )
        csv_file = SimpleUploadedFile(
            "sections.csv",
            (
                "Продукт;Код;Краткое имя EN;Краткое имя RU;"
                "Наименование раздела (услуги) EN;Наименование раздела (услуги) RU;"
                "Тип учета;Исполнитель;Экспертиза;Подразделение;ТКП\n"
                "SEC;SEC-6;tax-dd;нал-dd;Tax DD;Налоговый ДД;Услуги;"
                "Налоговый due diligence, Налоговый due diligence;TAX;Налоговый департамент;Да\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("section_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(len(response.json()["warnings"]), 1)
        self.assertIn("ошибка сохранения", response.json()["warnings"][0])
        self.assertFalse(TypicalSection.objects.filter(code="SEC-6").exists())
        self.assertFalse(TypicalSectionSpecialty.objects.filter(specialty=specialty).exists())


class SectionStructureViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-structures-admin",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="STR",
            name_en="Structure",
            display_name="Structure System",
            name_ru="Структура",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )
        self.section = TypicalSection.objects.create(
            product=self.product,
            code="STR-1",
            short_name="section-en",
            short_name_ru="section-ru",
            name_en="Section EN",
            name_ru="Раздел RU",
            accounting_type="Раздел",
            position=1,
        )

    def test_policy_partial_renders_structure_csv_buttons(self):
        SectionStructure.objects.create(
            product=self.product,
            section=self.section,
            subsections="Подраздел 1\nПодраздел 2",
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Типовая структура раздела (состава услуг)")
        self.assertContains(response, ">Код<", html=False)
        self.assertContains(response, "STR-1")
        self.assertContains(response, "Подраздел 1")
        self.assertContains(response, 'id="structures-csv-download-btn"', html=False)
        self.assertContains(response, 'id="structures-csv-upload-btn"', html=False)

    def test_structure_form_renders_product_options_with_display_name(self):
        response = self.client.get(reverse("structure_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<select name="product"', html=False)
        self.assertContains(response, "policy-product-select")
        self.assertContains(response, 'data-short-label="STR"', html=False)
        self.assertContains(response, "STR Structure System")
        self.assertContains(response, 'name="section_code"', html=False)
        self.assertContains(response, "readonly-field", html=False)
        self.assertContains(response, 'tabindex="-1"', html=False)
        self.assertContains(response, "policy-section-select")
        self.assertContains(response, '"label": "STR-1 Раздел RU"', html=False)
        self.assertContains(response, '"displayLabel": "Раздел RU"', html=False)

    def test_structure_csv_download_exports_current_table_columns(self):
        SectionStructure.objects.create(
            product=self.product,
            section=self.section,
            subsections="Подраздел 1\nПодраздел 2",
            position=1,
        )

        response = self.client.get(reverse("structure_csv_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("section_structures.csv", response["Content-Disposition"])
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(rows[0], ["Продукт", "Код", "Раздел (услуга)", "Подразделы"])
        self.assertEqual(rows[1], ["STR", "STR-1", "Раздел RU", "Подраздел 1\nПодраздел 2"])

    def test_structure_csv_download_respects_product_filter(self):
        other_product = Product.objects.create(
            short_name="STR2",
            name_en="Structure 2",
            display_name="Structure 2",
            name_ru="Структура 2",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        other_section = TypicalSection.objects.create(
            product=other_product,
            code="STR-2",
            short_name="other-en",
            short_name_ru="other-ru",
            name_en="Other EN",
            name_ru="Другой раздел",
            accounting_type="Раздел",
            position=1,
        )
        SectionStructure.objects.create(
            product=self.product,
            section=self.section,
            subsections="Подраздел 1",
            position=1,
        )
        SectionStructure.objects.create(
            product=other_product,
            section=other_section,
            subsections="Подраздел A",
            position=2,
        )

        response = self.client.get(
            reverse("structure_csv_download"),
            {"product": [self.product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "STR")

    def test_structure_csv_upload_creates_rows(self):
        csv_file = SimpleUploadedFile(
            "section_structures.csv",
            (
                "Продукт;Код;Раздел (услуга);Подразделы\n"
                "STR;STR-1;Раздел RU;Подраздел 1\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("structure_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        item = SectionStructure.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.section, self.section)
        self.assertEqual(item.subsections, "Подраздел 1")
        self.assertEqual(item.position, 1)

    def test_structure_csv_upload_accepts_legacy_rows_without_code(self):
        csv_file = SimpleUploadedFile(
            "section_structures.csv",
            (
                "Продукт;Раздел (услуга);Подразделы\n"
                "STR;Раздел RU;Подраздел 1\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("structure_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        self.assertEqual(SectionStructure.objects.get().section, self.section)


class ReportStructureViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-report-structures-admin",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="REP",
            name_en="Report",
            display_name="Report Product",
            name_ru="Отчет",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )

    def _create_report_structure(self, level, code, name, position):
        return ReportStructure.objects.create(
            product=self.product,
            level=level,
            code=code,
            name=name,
            position=position,
        )

    def test_policy_partial_renders_report_structure_table_and_actions(self):
        self._create_report_structure(0, "RPT", "Итоговый отчет", 1)
        self._create_report_structure(1, "SEC", "Раздел", 2)
        self._create_report_structure(2, "SUB", "Подраздел", 3)

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Типовая структура отчета")
        self.assertContains(response, 'id="report-structures-master"', html=False)
        self.assertContains(response, 'name="report-structure-select"', html=False)
        self.assertContains(response, 'id="report-structures-actions"', html=False)
        self.assertContains(response, 'id="report-structures-csv-download-btn"', html=False)
        self.assertContains(response, 'id="report-structures-csv-upload-btn"', html=False)
        self.assertContains(response, "Итоговый отчет")
        self.assertContains(response, "1.1")

    def test_report_structure_form_and_level_validation(self):
        response = self.client.get(reverse("report_structure_form_create"), {"product": self.product.pk})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<select name="product"', html=False)
        self.assertContains(response, "policy-product-select")
        self.assertContains(response, 'name="level"', html=False)
        self.assertContains(response, 'name="number"', html=False)
        self.assertContains(response, "readonly-field", html=False)
        self.assertContains(response, "var itemsByProduct", html=False)

        form = ReportStructureForm(data={
            "product": self.product.pk,
            "level": 10,
            "code": "BAD",
            "name": "Недопустимый уровень",
        })
        self.assertFalse(form.is_valid())
        self.assertIn("level", form.errors)

    def test_create_and_edit_report_structure_rows(self):
        response = self.client.post(
            reverse("report_structure_form_create"),
            {
                "product": self.product.pk,
                "level": 0,
                "code": "RPT",
                "name": "Итоговый отчет",
            },
        )

        self.assertEqual(response.status_code, 200)
        item = ReportStructure.objects.get()
        self.assertEqual(item.position, 1)
        self.assertEqual(item.level, 0)

        response = self.client.post(
            reverse("report_structure_form_edit", args=[item.pk]),
            {
                "product": self.product.pk,
                "level": 1,
                "code": "SEC",
                "name": "Раздел отчета",
            },
        )

        self.assertEqual(response.status_code, 200)
        item.refresh_from_db()
        self.assertEqual(item.level, 1)
        self.assertEqual(item.code, "SEC")
        self.assertEqual(item.name, "Раздел отчета")

    def test_report_structure_numbering_resets_after_level_zero(self):
        self._create_report_structure(0, "RPT1", "Первый отчет", 1)
        self._create_report_structure(1, "SEC1", "Первый раздел", 2)
        self._create_report_structure(2, "SUB1", "Первый подраздел", 3)
        self._create_report_structure(1, "SEC2", "Второй раздел", 4)
        self._create_report_structure(0, "RPT2", "Второй отчет", 5)
        self._create_report_structure(1, "SEC3", "Раздел второго отчета", 6)

        response = self.client.get(reverse("report_structure_csv_download"))

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(rows[0], [
            "Продукт",
            "Уровень",
            "Номер",
            "Код",
            "Наименование отчета, раздела (подраздела)",
        ])
        self.assertEqual([row[2] for row in rows[1:]], ["0", "1", "1.1", "2", "0", "1"])

    def test_report_structure_csv_upload_and_download_filter(self):
        other_product = Product.objects.create(
            short_name="REP2",
            name_en="Report 2",
            display_name="Report Product 2",
            name_ru="Отчет 2",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        ReportStructure.objects.create(
            product=other_product,
            level=0,
            code="OTHER",
            name="Другой отчет",
            position=1,
        )
        csv_file = SimpleUploadedFile(
            "report_structures.csv",
            (
                "Продукт;Уровень;Номер;Код;Наименование отчета, раздела (подраздела)\n"
                "REP;0;0;RPT;Итоговый отчет\n"
                "REP;1;1;SEC;Раздел отчета\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("report_structure_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 2)
        self.assertEqual(response.json()["warnings"], [])
        self.assertEqual(ReportStructure.objects.filter(product=self.product).count(), 2)

        response = self.client.get(
            reverse("report_structure_csv_download"),
            {"product": [self.product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 3)
        self.assertEqual({row[0] for row in rows[1:]}, {"REP"})

    def test_report_structure_move_down_reorders_within_product(self):
        first = self._create_report_structure(0, "RPT", "Отчет", 1)
        second = self._create_report_structure(1, "SEC", "Раздел", 2)
        other_product = Product.objects.create(
            short_name="REP2",
            name_en="Report 2",
            display_name="Report Product 2",
            name_ru="Отчет 2",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        other = ReportStructure.objects.create(
            product=other_product,
            level=0,
            code="OTHER",
            name="Другой отчет",
            position=1,
        )

        response = self.client.post(reverse("report_structure_move_down", args=[first.pk]))

        self.assertEqual(response.status_code, 200)
        first.refresh_from_db()
        second.refresh_from_db()
        other.refresh_from_db()
        self.assertEqual(first.position, 2)
        self.assertEqual(second.position, 1)
        self.assertEqual(other.position, 1)


class ServiceGoalReportViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-admin",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="TAX",
            name_en="Tax",
            display_name="Tax",
            name_ru="Налоги",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )

    def test_policy_partial_renders_service_goal_reports_table(self):
        ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Подготовка заключения",
            service_goal_genitive="Подготовки заключения",
            report_title="Итоговый отчет",
            product_name="Налоговый обзор",
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Цели услуг и названия отчетов")
        self.assertContains(response, "Титул отчета/ТКП")
        self.assertContains(response, "Название продукта")
        self.assertContains(response, "Подготовка заключения")
        self.assertContains(response, "Подготовки заключения")
        self.assertContains(response, "Налоговый обзор")
        self.assertContains(response, 'id="service-goal-reports-actions"', html=False)
        self.assertContains(response, 'id="service-goal-reports-csv-download-btn"', html=False)
        self.assertContains(response, 'id="service-goal-reports-csv-upload-btn"', html=False)

    def test_create_service_goal_report_saves_row(self):
        response = self.client.post(
            reverse("service_goal_report_form_create"),
            {
                "product": self.product.pk,
                "service_goal": "Подготовка документов",
                "service_goal_genitive": "Подготовки документов",
                "report_title": "Отчет по документам",
                "product_name": "Документарная проверка",
            },
        )

        self.assertEqual(response.status_code, 200)
        item = ServiceGoalReport.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.service_goal, "Подготовка документов")
        self.assertEqual(item.service_goal_genitive, "Подготовки документов")
        self.assertEqual(item.report_title, "Отчет по документам")
        self.assertEqual(item.product_name, "Документарная проверка")
        self.assertEqual(item.position, 1)

    def test_service_goal_report_form_renders_product_picker_with_display_name(self):
        response = self.client.get(reverse("service_goal_report_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<select name="product"', html=False)
        self.assertContains(response, "policy-product-select")
        self.assertContains(response, 'data-short-label="TAX"', html=False)
        self.assertContains(response, "TAX Tax")
        self.assertContains(response, "Титул отчета/ТКП")
        self.assertContains(response, "Название продукта")

        form = ServiceGoalReportForm()
        labels = [label for _, label in form.fields["product"].choices]
        self.assertIn("TAX Tax", labels)

    def test_service_goal_report_csv_download_exports_current_table_columns(self):
        ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Подготовка заключения",
            service_goal_genitive="Подготовки заключения",
            report_title="Итоговый отчет",
            product_name="Налоговый обзор",
            position=1,
        )

        response = self.client.get(reverse("service_goal_report_csv_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("service_goal_reports.csv", response["Content-Disposition"])
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(
            rows[0],
            [
                "Продукт",
                "Цели оказания услуг",
                "Цели оказания услуг в родительном падеже",
                "Титул отчета/ТКП",
                "Название продукта",
            ],
        )
        self.assertEqual(
            rows[1],
            [
                "TAX",
                "Подготовка заключения",
                "Подготовки заключения",
                "Итоговый отчет",
                "Налоговый обзор",
            ],
        )

    def test_service_goal_report_csv_download_respects_product_filter(self):
        other_product = Product.objects.create(
            short_name="AUD",
            name_en="Audit",
            display_name="Audit",
            name_ru="Аудит",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Цель TAX",
            service_goal_genitive="Цели TAX",
            report_title="Отчет TAX",
            product_name="Продукт TAX",
            position=1,
        )
        ServiceGoalReport.objects.create(
            product=other_product,
            service_goal="Цель AUD",
            service_goal_genitive="Цели AUD",
            report_title="Отчет AUD",
            product_name="Продукт AUD",
            position=2,
        )

        response = self.client.get(
            reverse("service_goal_report_csv_download"),
            {"product": [self.product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "TAX")

    def test_service_goal_report_csv_upload_creates_rows(self):
        csv_file = SimpleUploadedFile(
            "service_goal_reports.csv",
            (
                "Продукт;Цели оказания услуг;Цели оказания услуг в родительном падеже;Титул отчета/ТКП;Название продукта\n"
                "TAX;Подготовка документов;Подготовки документов;Отчет по документам;Документарная проверка\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("service_goal_report_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        item = ServiceGoalReport.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.service_goal, "Подготовка документов")
        self.assertEqual(item.service_goal_genitive, "Подготовки документов")
        self.assertEqual(item.report_title, "Отчет по документам")
        self.assertEqual(item.product_name, "Документарная проверка")
        self.assertEqual(item.position, 1)

    def test_service_goal_report_csv_upload_updates_existing_product_row(self):
        item = ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Старая цель",
            service_goal_genitive="Старой цели",
            report_title="Старый титул",
            product_name="Старый продукт",
            position=1,
        )
        csv_file = SimpleUploadedFile(
            "service_goal_reports.csv",
            (
                "Продукт;Цели оказания услуг;Цели оказания услуг в родительном падеже;Титул отчета/ТКП;Название продукта\n"
                "TAX;Новая цель;Новой цели;Новый титул;Новый продукт\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("service_goal_report_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(response.json()["updated"], 1)
        self.assertEqual(ServiceGoalReport.objects.count(), 1)
        item.refresh_from_db()
        self.assertEqual(item.service_goal, "Новая цель")
        self.assertEqual(item.service_goal_genitive, "Новой цели")
        self.assertEqual(item.report_title, "Новый титул")
        self.assertEqual(item.product_name, "Новый продукт")

    def test_move_up_reorders_globally_across_table(self):
        other_product = Product.objects.create(
            short_name="AUD",
            name_en="Audit",
            display_name="Audit",
            name_ru="Аудит",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        first = ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Первая цель",
            service_goal_genitive="Первой цели",
            report_title="Первый отчет",
            position=1,
        )
        second = ServiceGoalReport.objects.create(
            product=other_product,
            service_goal="Вторая цель",
            service_goal_genitive="Второй цели",
            report_title="Второй отчет",
            position=2,
        )

        response = self.client.post(reverse("service_goal_report_move_up", args=[second.pk]))

        self.assertEqual(response.status_code, 200)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(first.position, 2)
        self.assertEqual(second.position, 1)

    def test_non_staff_user_cannot_reorder_service_goal_reports(self):
        first = ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Первая цель",
            service_goal_genitive="Первой цели",
            report_title="Первый отчет",
            position=1,
        )
        second = ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Вторая цель",
            service_goal_genitive="Второй цели",
            report_title="Второй отчет",
            position=2,
        )
        non_staff = get_user_model().objects.create_user(
            username="policy-user",
            password="secret123",
            is_staff=False,
        )
        client = self.client_class()
        client.force_login(non_staff)

        response = client.post(reverse("service_goal_report_move_up", args=[second.pk]))

        self.assertEqual(response.status_code, 302)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(first.position, 1)
        self.assertEqual(second.position, 2)


class TypicalServiceCompositionViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-admin-2",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="TAX2",
            name_en="Tax 2",
            display_name="Tax 2",
            name_ru="Налоги 2",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )
        self.other_product = Product.objects.create(
            short_name="AUD2",
            name_en="Audit 2",
            display_name="Audit 2",
            name_ru="Аудит 2",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        self.section = TypicalSection.objects.create(
            product=self.product,
            code="S1",
            short_name="sec-en",
            short_name_ru="sec-ru",
            name_en="Section EN",
            name_ru="Раздел RU",
            accounting_type="Раздел",
            position=1,
        )
        self.other_section = TypicalSection.objects.create(
            product=self.other_product,
            code="S2",
            short_name="other-en",
            short_name_ru="other-ru",
            name_en="Other EN",
            name_ru="Другой раздел RU",
            accounting_type="Раздел",
            position=1,
        )

    def _docx_upload_file(self, sections, include_code=True):
        document = Document()
        for section in sections:
            document.add_heading(section["product"], level=1)
            headers = ["ID", "Продукт", "Код", "Раздел (услуга)", "Состав услуг"] if include_code else [
                "ID",
                "Продукт",
                "Раздел (услуга)",
                "Состав услуг",
            ]
            table = document.add_table(rows=1, cols=len(headers))
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header
            for row in section["rows"]:
                cells = table.add_row().cells
                cells[0].text = str(row.get("id", ""))
                cells[1].text = row.get("product", section["product"])
                if include_code:
                    cells[2].text = row.get("section_code", "")
                    cells[3].text = row.get("section", "")
                    cells[4].text = row.get("service_composition", "")
                else:
                    cells[2].text = row.get("section", "")
                    cells[3].text = row.get("service_composition", "")
        buffer = io.BytesIO()
        document.save(buffer)
        return SimpleUploadedFile(
            "typical_service_compositions.docx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_policy_partial_renders_typical_service_compositions_table(self):
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="Подготовка,\nанализ,\nвыпуск отчета",
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Типовой состав услуг")
        self.assertContains(response, ">Код<", html=False)
        self.assertContains(response, "S1")
        self.assertContains(response, "Раздел RU")
        self.assertContains(response, "<p>Подготовка,<br>анализ,<br>выпуск отчета</p>", html=False)
        self.assertContains(response, 'id="typical-service-compositions-wrap-toggle"', html=False)
        self.assertContains(response, 'id="typical-service-compositions-table"', html=False)
        self.assertContains(response, 'class="policy-service-composition-cell"', html=False)
        self.assertContains(
            response,
            'class="policy-service-composition-content policy-service-composition-content--rich ql-editor"',
            html=False,
        )
        self.assertContains(response, 'id="typical-service-compositions-csv-download-btn"', html=False)
        self.assertContains(response, 'id="typical-service-compositions-csv-upload-btn"', html=False)
        self.assertContains(response, 'id="typical-service-compositions-docx-download-btn"', html=False)
        self.assertContains(response, 'id="typical-service-compositions-docx-upload-btn"', html=False)
        self.assertNotContains(response, 'id="typical-service-compositions-xlsx-download-btn"', html=False)
        self.assertNotContains(response, 'id="typical-service-compositions-xlsx-upload-btn"', html=False)

    def test_policy_partial_renders_typical_service_composition_rich_html(self):
        editor_state = {
            "html": (
                '<p class="ql-align-justify"><strong>Подготовка</strong></p>'
                '<ol><li data-list="ordered">Этап 1</li>'
                '<li class="ql-indent-1" data-list="ordered">Подэтап 1.1</li></ol>'
                '<ul><li data-list="bullet">Маркер</li>'
                '<li data-list="dash"><span class="ql-ui" contenteditable="false"></span>Дефис</li></ul>'
            ),
            "plain_text": "Подготовка\nЭтап 1\nПодэтап 1.1\nМаркер\nДефис",
        }
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition=editor_state["plain_text"],
            service_composition_editor_state=editor_state,
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            '<p class="ql-align-justify"><strong>Подготовка</strong></p>',
            html=False,
        )
        self.assertContains(
            response,
            '<li data-list="ordered"><span class="ql-ui"></span>Этап 1</li>',
            html=False,
        )
        self.assertContains(
            response,
            '<li class="ql-indent-1" data-list="ordered"><span class="ql-ui"></span>Подэтап 1.1</li>',
            html=False,
        )
        self.assertContains(
            response,
            '<li data-list="bullet"><span class="ql-ui"></span>Маркер</li>',
            html=False,
        )
        self.assertContains(
            response,
            '<li data-list="dash"><span class="ql-ui"></span>Дефис</li>',
            html=False,
        )
        self.assertNotContains(response, 'contenteditable="false"', html=False)

    def test_policy_partial_renders_typical_service_composition_plain_text_fallback(self):
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="Первый блок\nстрока\n\nВторой блок",
            service_composition_editor_state={},
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "<p>Первый блок<br>строка</p><p>Второй блок</p>",
            html=False,
        )

    def test_policy_partial_sanitizes_typical_service_composition_html(self):
        editor_state = {
            "html": (
                '<p onclick="alert(1)">'
                '<script>alert(1)</script>'
                '<span class="ql-font-calibri bad-class" '
                'style="color: #ff0000; background-image: url(javascript:alert(1)); background-color: rgb(1, 2, 3)">'
                'Безопасно'
                '</span>'
                '</p>'
            ),
            "plain_text": "Безопасно",
        }
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition=editor_state["plain_text"],
            service_composition_editor_state=editor_state,
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            '<span class="ql-font-calibri" style="color: #ff0000; background-color: rgb(1, 2, 3)">Безопасно</span>',
            html=False,
        )
        self.assertNotContains(response, "<script", html=False)
        self.assertNotContains(response, "alert(1)", html=False)
        self.assertNotContains(response, "onclick", html=False)
        self.assertNotContains(response, "bad-class", html=False)
        self.assertNotContains(response, "background-image", html=False)
        self.assertNotContains(response, "javascript:", html=False)

    def test_create_typical_service_composition_saves_row(self):
        editor_state = {
            "html": "<p><strong>Этап 1</strong></p><p>Этап 2</p>",
            "plain_text": "Этап 1\nЭтап 2",
        }
        response = self.client.post(
            reverse("typical_service_composition_form_create"),
            {
                "product": self.product.pk,
                "section": self.section.pk,
                "service_composition": "",
                "service_composition_editor_state": json.dumps(editor_state, ensure_ascii=False),
            },
        )

        self.assertEqual(response.status_code, 200)
        item = TypicalServiceComposition.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.section, self.section)
        self.assertEqual(item.service_composition, "Этап 1\nЭтап 2")
        self.assertEqual(item.service_composition_editor_state, editor_state)
        self.assertEqual(item.position, 1)

    def test_create_typical_service_composition_preserves_custom_list_markers(self):
        editor_state = {
            "html": (
                '<ol><li data-list="dash">Этап с дефисом</li>'
                '<li data-list="check">Этап с галочкой</li></ol>'
            ),
            "plain_text": "Этап с дефисом\nЭтап с галочкой",
        }
        response = self.client.post(
            reverse("typical_service_composition_form_create"),
            {
                "product": self.product.pk,
                "section": self.section.pk,
                "service_composition": "",
                "service_composition_editor_state": json.dumps(editor_state, ensure_ascii=False),
            },
        )

        self.assertEqual(response.status_code, 200)
        item = TypicalServiceComposition.objects.get()
        self.assertEqual(item.service_composition, "Этап с дефисом\nЭтап с галочкой")
        self.assertEqual(item.service_composition_editor_state, editor_state)

    def test_typical_service_composition_form_renders_product_options_with_display_name(self):
        response = self.client.get(reverse("typical_service_composition_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<select name="product"', html=False)
        self.assertContains(response, "policy-product-select")
        self.assertContains(response, 'data-short-label="TAX2"', html=False)
        self.assertContains(response, "TAX2 Tax 2")
        self.assertContains(response, 'name="section_code"', html=False)
        self.assertContains(response, "readonly-field", html=False)
        self.assertContains(response, 'tabindex="-1"', html=False)
        self.assertContains(response, "policy-section-select")
        self.assertContains(response, '"label": "S1 Раздел RU"', html=False)
        self.assertContains(response, '"displayLabel": "Раздел RU"', html=False)

    def test_typical_service_composition_csv_download_exports_current_table_columns(self):
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="Подготовка\nАнализ\nВыпуск отчета",
            position=1,
        )

        response = self.client.get(reverse("typical_service_composition_csv_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("typical_service_compositions.csv", response["Content-Disposition"])
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(rows[0], ["Продукт", "Код", "Раздел (услуга)", "Состав услуг"])
        self.assertEqual(rows[1], ["TAX2", "S1", "Раздел RU", "Подготовка\nАнализ\nВыпуск отчета"])

    def test_typical_service_composition_csv_upload_creates_rows(self):
        csv_file = SimpleUploadedFile(
            "typical_service_compositions.csv",
            (
                "Продукт;Код;Раздел (услуга);Состав услуг\n"
                "TAX2;S1;Раздел RU;Подготовка и выпуск отчета\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("typical_service_composition_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        item = TypicalServiceComposition.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.section, self.section)
        self.assertEqual(item.service_composition, "Подготовка и выпуск отчета")
        self.assertEqual(
            item.service_composition_editor_state,
            {"html": "", "plain_text": "Подготовка и выпуск отчета"},
        )
        self.assertEqual(item.position, 1)

    def test_typical_service_composition_csv_upload_accepts_legacy_rows_without_code(self):
        csv_file = SimpleUploadedFile(
            "typical_service_compositions.csv",
            (
                "Продукт;Раздел (услуга);Состав услуг\n"
                "TAX2;Раздел RU;Подготовка и выпуск отчета\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("typical_service_composition_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        self.assertEqual(TypicalServiceComposition.objects.get().section, self.section)

    def test_typical_service_composition_docx_download_exports_editable_table(self):
        editor_state = {
            "html": (
                '<p><strong>Подготовка</strong></p>'
                '<ol><li data-list="ordered">Анализ</li>'
                '<li class="ql-indent-1" data-list="ordered">Выпуск отчета</li></ol>'
            ),
            "plain_text": "Подготовка\nАнализ\nВыпуск отчета",
        }
        item = TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition=editor_state["plain_text"],
            service_composition_editor_state=editor_state,
            position=1,
        )

        response = self.client.get(reverse("typical_service_composition_docx_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("typical_service_compositions.docx", response["Content-Disposition"])
        document = Document(io.BytesIO(response.content))
        self.assertEqual(document.paragraphs[0].text, "TAX2")
        self.assertIn(document.paragraphs[0].style.name.lower(), {"heading 1", "заголовок 1"})
        self.assertEqual(len(document.tables), 1)
        table = document.tables[0]
        self.assertEqual([cell.text for cell in table.rows[0].cells], ["ID", "Продукт", "Код", "Раздел (услуга)", "Состав услуг"])
        self.assertEqual(table.rows[1].cells[0].text, str(item.pk))
        self.assertEqual(table.rows[1].cells[1].text, "TAX2")
        self.assertEqual(table.rows[1].cells[2].text, "S1")
        self.assertEqual(table.rows[1].cells[3].text, "Раздел RU")
        self.assertEqual(table.rows[1].cells[4].text, editor_state["plain_text"])
        self.assertIn('w:tblW w:type="pct" w:w="5000"', table._tbl.xml)
        for width in ("300", "650", "500", "1100", "2450"):
            self.assertIn(f'w:tcW w:type="pct" w:w="{width}"', table._tbl.xml)
        list_paragraphs = table.rows[1].cells[4].paragraphs[1:]
        self.assertTrue(all("w:numPr" in paragraph._element.xml for paragraph in list_paragraphs))

    def test_typical_service_composition_docx_download_restarts_numbering_per_cell(self):
        editor_state = {
            "html": (
                '<ol><li data-list="ordered">Первый пункт</li>'
                '<li class="ql-indent-1" data-list="ordered">Подпункт</li></ol>'
            ),
            "plain_text": "Первый пункт\nПодпункт",
        }
        for position in (1, 2):
            TypicalServiceComposition.objects.create(
                product=self.product,
                section=self.section,
                service_composition=editor_state["plain_text"],
                service_composition_editor_state=editor_state,
                position=position,
            )

        response = self.client.get(reverse("typical_service_composition_docx_download"))

        self.assertEqual(response.status_code, 200)
        document = Document(io.BytesIO(response.content))
        table = document.tables[0]
        first_cell_num_id = table.rows[1].cells[4].paragraphs[0]._p.pPr.numPr.numId.val
        second_cell_num_id = table.rows[2].cells[4].paragraphs[0]._p.pPr.numPr.numId.val
        self.assertNotEqual(first_cell_num_id, second_cell_num_id)

    def test_typical_service_composition_docx_download_groups_rows_by_product(self):
        editor_state = {
            "html": "<p>Состав TAX2</p>",
            "plain_text": "Состав TAX2",
        }
        other_editor_state = {
            "html": "<p>Состав AUD2</p>",
            "plain_text": "Состав AUD2",
        }
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition=editor_state["plain_text"],
            service_composition_editor_state=editor_state,
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=self.other_product,
            section=self.other_section,
            service_composition=other_editor_state["plain_text"],
            service_composition_editor_state=other_editor_state,
            position=1,
        )

        response = self.client.get(reverse("typical_service_composition_docx_download"))

        self.assertEqual(response.status_code, 200)
        document = Document(io.BytesIO(response.content))
        self.assertEqual(len(document.tables), 2)
        self.assertEqual(document.paragraphs[0].text, "TAX2")
        self.assertEqual(document.paragraphs[1].text, "AUD2")
        self.assertEqual(document.tables[0].rows[1].cells[1].text, "TAX2")
        self.assertEqual(document.tables[1].rows[1].cells[1].text, "AUD2")

    def test_typical_service_composition_csv_download_respects_product_filter(self):
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="TAX2 состав",
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=self.other_product,
            section=self.other_section,
            service_composition="AUD2 состав",
            position=1,
        )

        response = self.client.get(
            reverse("typical_service_composition_csv_download"),
            {"product": [self.product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "TAX2")

    def test_typical_service_composition_docx_download_respects_product_filter(self):
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="TAX2 состав",
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=self.other_product,
            section=self.other_section,
            service_composition="AUD2 состав",
            position=1,
        )

        response = self.client.get(
            reverse("typical_service_composition_docx_download"),
            {"product": [self.product.pk, self.other_product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        document = Document(io.BytesIO(response.content))
        self.assertEqual(len(document.tables), 2)
        self.assertEqual(document.paragraphs[0].text, "TAX2")
        self.assertEqual(document.paragraphs[1].text, "AUD2")

    def test_typical_service_composition_csv_download_respects_consulting_filter(self):
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="TAX2 состав",
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=self.other_product,
            section=self.other_section,
            service_composition="AUD2 состав",
            position=1,
        )

        response = self.client.get(
            reverse("typical_service_composition_csv_download"),
            {"category": ["Аудит"]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "AUD2")

    def test_typical_service_composition_docx_upload_reads_multiple_product_tables(self):
        tax_item = TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="TAX2 старый",
            service_composition_editor_state={"html": "<p>TAX2 старый</p>", "plain_text": "TAX2 старый"},
            position=1,
        )
        aud_item = TypicalServiceComposition.objects.create(
            product=self.other_product,
            section=self.other_section,
            service_composition="AUD2 старый",
            service_composition_editor_state={"html": "<p>AUD2 старый</p>", "plain_text": "AUD2 старый"},
            position=1,
        )
        docx_file = self._docx_upload_file([
            {
                "product": "TAX2",
                "rows": [{
                    "id": tax_item.pk,
                    "section": "Раздел RU",
                    "service_composition": "TAX2 новый",
                }],
            },
            {
                "product": "AUD2",
                "rows": [{
                    "id": aud_item.pk,
                    "section": "Другой раздел RU",
                    "service_composition": "AUD2 новый",
                }],
            },
        ])

        response = self.client.post(reverse("typical_service_composition_docx_upload"), {"csv_file": docx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(response.json()["updated"], 2)
        tax_item.refresh_from_db()
        aud_item.refresh_from_db()
        self.assertEqual(tax_item.service_composition, "TAX2 новый")
        self.assertEqual(aud_item.service_composition, "AUD2 новый")

    def test_typical_service_composition_docx_upload_updates_existing_row_by_id(self):
        item = TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="Старый текст",
            service_composition_editor_state={"html": "<p>Старый текст</p>", "plain_text": "Старый текст"},
            position=1,
        )
        docx_file = self._docx_upload_file([
            {
                "product": "TAX2",
                "rows": [{
                    "id": item.pk,
                    "section": "Раздел RU",
                    "service_composition": "Новый текст из Word",
                }],
            },
        ])

        response = self.client.post(reverse("typical_service_composition_docx_upload"), {"csv_file": docx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(response.json()["updated"], 1)
        self.assertEqual(response.json()["warnings"], [])
        item.refresh_from_db()
        self.assertEqual(item.service_composition, "Новый текст из Word")
        self.assertEqual(
            item.service_composition_editor_state,
            {"html": "<p>Новый текст из Word</p>", "plain_text": "Новый текст из Word"},
        )

    def test_typical_service_composition_docx_upload_creates_row_when_id_is_blank(self):
        docx_file = self._docx_upload_file([
            {
                "product": "TAX2",
                "rows": [{
                    "id": "",
                    "section_code": "S1",
                    "section": "Раздел RU",
                    "service_composition": "Новая строка из Word",
                }],
            },
        ])

        response = self.client.post(reverse("typical_service_composition_docx_upload"), {"csv_file": docx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["updated"], 0)
        item = TypicalServiceComposition.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.section, self.section)
        self.assertEqual(item.service_composition, "Новая строка из Word")
        self.assertEqual(item.position, 1)

    def test_typical_service_composition_docx_upload_accepts_legacy_table_without_code(self):
        docx_file = self._docx_upload_file(
            [
                {
                    "product": "TAX2",
                    "rows": [{
                        "id": "",
                        "section": "Раздел RU",
                        "service_composition": "Старая структура Word",
                    }],
                },
            ],
            include_code=False,
        )

        response = self.client.post(reverse("typical_service_composition_docx_upload"), {"csv_file": docx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        self.assertEqual(TypicalServiceComposition.objects.get().section, self.section)

    def test_typical_service_composition_docx_upload_preserves_multilevel_lists(self):
        editor_state = {
            "html": (
                '<ol><li data-list="ordered">Основной пункт</li>'
                '<li class="ql-indent-1" data-list="ordered">Подпункт</li></ol>'
                '<ul><li data-list="dash">Пункт с дефисом</li></ul>'
            ),
            "plain_text": "Основной пункт\nПодпункт\nПункт с дефисом",
        }
        item = TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition=editor_state["plain_text"],
            service_composition_editor_state=editor_state,
            position=1,
        )
        download_response = self.client.get(reverse("typical_service_composition_docx_download"))
        docx_file = SimpleUploadedFile(
            "typical_service_compositions.docx",
            download_response.content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response = self.client.post(reverse("typical_service_composition_docx_upload"), {"csv_file": docx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["updated"], 1)
        item.refresh_from_db()
        self.assertEqual(item.service_composition, editor_state["plain_text"])
        self.assertIn('data-list="ordered"', item.service_composition_editor_state["html"])
        self.assertIn('class="ql-indent-1" data-list="ordered"', item.service_composition_editor_state["html"])
        self.assertIn('data-list="dash"', item.service_composition_editor_state["html"])

    def test_typical_service_composition_docx_upload_warns_on_invalid_references(self):
        docx_file = self._docx_upload_file([
            {
                "product": "TAX2",
                "rows": [
                    {
                        "id": "not-id",
                        "section": "Раздел RU",
                        "service_composition": "Некорректный ID",
                    },
                    {
                        "id": "",
                        "product": "UNKNOWN",
                        "section": "Раздел RU",
                        "service_composition": "Некорректный продукт",
                    },
                ],
            },
        ])

        response = self.client.post(reverse("typical_service_composition_docx_upload"), {"csv_file": docx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(response.json()["updated"], 0)
        self.assertEqual(len(response.json()["warnings"]), 2)
        self.assertFalse(TypicalServiceComposition.objects.exists())

    def test_typical_service_composition_xlsx_download_preserves_editor_state(self):
        editor_state = {
            "html": (
                '<p><strong>Подготовка</strong></p>'
                '<ol><li data-list="dash">Анализ</li>'
                '<li class="ql-indent-1" data-list="check">Выпуск отчета</li></ol>'
            ),
            "plain_text": "Подготовка\nАнализ\nВыпуск отчета",
        }
        TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition=editor_state["plain_text"],
            service_composition_editor_state=editor_state,
            position=1,
        )

        response = self.client.get(reverse("typical_service_composition_xlsx_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("typical_service_compositions.xlsx", response["Content-Disposition"])
        workbook = load_workbook(io.BytesIO(response.content))
        sheet = workbook.active
        self.assertEqual(sheet["A1"].value, "Продукт")
        self.assertEqual(sheet["B1"].value, "Код")
        self.assertEqual(sheet["C1"].value, "Раздел (услуга)")
        self.assertEqual(sheet["D1"].value, "Состав услуг")
        self.assertEqual(sheet["E1"].value, "Состояние редактора (JSON)")
        self.assertTrue(sheet.column_dimensions["E"].hidden)
        self.assertEqual(sheet["A2"].value, "TAX2")
        self.assertEqual(sheet["B2"].value, "S1")
        self.assertEqual(sheet["C2"].value, "Раздел RU")
        self.assertEqual(sheet["D2"].value, editor_state["plain_text"])
        self.assertEqual(json.loads(sheet["E2"].value), editor_state)

    def test_typical_service_composition_xlsx_upload_preserves_editor_state(self):
        editor_state = {
            "html": (
                '<p><span style="color:#ff0000"><strong>Подготовка</strong></span></p>'
                '<ol><li data-list="dash">Анализ</li>'
                '<li class="ql-indent-1" data-list="check">Выпуск отчета</li></ol>'
            ),
            "plain_text": "Подготовка\nАнализ\nВыпуск отчета",
        }
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Продукт", "Код", "Раздел (услуга)", "Состав услуг", "Состояние редактора (JSON)"])
        sheet.column_dimensions["E"].hidden = True
        sheet.append([
            "TAX2",
            "S1",
            "Раздел RU",
            editor_state["plain_text"],
            json.dumps(editor_state, ensure_ascii=False),
        ])
        buffer = io.BytesIO()
        workbook.save(buffer)
        xlsx_file = SimpleUploadedFile(
            "typical_service_compositions.xlsx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        response = self.client.post(reverse("typical_service_composition_xlsx_upload"), {"xlsx_file": xlsx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        item = TypicalServiceComposition.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.section, self.section)
        self.assertEqual(item.service_composition, editor_state["plain_text"])
        self.assertEqual(item.service_composition_editor_state, editor_state)

    def test_typical_service_composition_xlsx_upload_falls_back_to_plain_text(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Продукт", "Раздел (услуга)", "Состав услуг"])
        sheet.append(["TAX2", "Раздел RU", "Подготовка без скрытого состояния"])
        buffer = io.BytesIO()
        workbook.save(buffer)
        xlsx_file = SimpleUploadedFile(
            "typical_service_compositions.xlsx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        response = self.client.post(reverse("typical_service_composition_xlsx_upload"), {"xlsx_file": xlsx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        item = TypicalServiceComposition.objects.get()
        self.assertEqual(item.service_composition, "Подготовка без скрытого состояния")
        self.assertEqual(
            item.service_composition_editor_state,
            {"html": "", "plain_text": "Подготовка без скрытого состояния"},
        )

    def test_typical_service_composition_xlsx_upload_uses_visible_text_when_state_is_stale(self):
        editor_state = {
            "html": "<p><strong>Старый текст</strong></p>",
            "plain_text": "Старый текст",
        }
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Продукт", "Код", "Раздел (услуга)", "Состав услуг", "Состояние редактора (JSON)"])
        sheet.column_dimensions["E"].hidden = True
        sheet.append([
            "TAX2",
            "S1",
            "Раздел RU",
            "Новый текст из Excel",
            json.dumps(editor_state, ensure_ascii=False),
        ])
        buffer = io.BytesIO()
        workbook.save(buffer)
        xlsx_file = SimpleUploadedFile(
            "typical_service_compositions.xlsx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        response = self.client.post(reverse("typical_service_composition_xlsx_upload"), {"xlsx_file": xlsx_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(len(response.json()["warnings"]), 1)
        item = TypicalServiceComposition.objects.get()
        self.assertEqual(item.service_composition, "Новый текст из Excel")
        self.assertEqual(
            item.service_composition_editor_state,
            {"html": "", "plain_text": "Новый текст из Excel"},
        )

    def test_create_typical_service_composition_rejects_section_from_other_product(self):
        response = self.client.post(
            reverse("typical_service_composition_form_create"),
            {
                "product": self.product.pk,
                "section": self.other_section.pk,
                "service_composition": "Некорректная связь",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Раздел должен относиться к выбранному продукту.")
        self.assertFalse(TypicalServiceComposition.objects.exists())

    def test_non_staff_user_cannot_reorder_typical_service_compositions(self):
        first = TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="Этап 1",
            position=1,
        )
        second = TypicalServiceComposition.objects.create(
            product=self.product,
            section=self.section,
            service_composition="Этап 2",
            position=2,
        )
        non_staff = get_user_model().objects.create_user(
            username="policy-user-2",
            password="secret123",
            is_staff=False,
        )
        client = self.client_class()
        client.force_login(non_staff)

        response = client.post(reverse("typical_service_composition_move_down", args=[first.pk]))

        self.assertEqual(response.status_code, 302)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(first.position, 1)
        self.assertEqual(second.position, 2)


class TypicalServiceTermViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-admin-terms",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="TERM",
            name_en="Terms",
            display_name="Terms",
            name_ru="Сроки",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )
        self.other_product = Product.objects.create(
            short_name="TERM2",
            name_en="Terms 2",
            display_name="Terms 2",
            name_ru="Сроки 2",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )

    def test_policy_partial_renders_typical_service_terms_table(self):
        TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.5"),
            final_report_weeks=3,
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Типовые сроки оказания услуг")
        self.assertContains(response, "Сроки предоставления исходных данных")
        self.assertContains(response, "Срок подготовки Предварительного отчёта")
        self.assertContains(response, "Срок подготовки Итогового отчёта")
        self.assertContains(response, ">0,0 нед.<", html=False)
        self.assertContains(response, ">1,5 мес.<", html=False)
        self.assertContains(response, ">3,0 нед.<", html=False)
        self.assertContains(response, 'id="typical-service-terms-actions"', html=False)
        self.assertContains(response, 'id="typical-service-terms-gantt-edit-btn"', html=False)
        self.assertContains(response, 'id="typical-service-term-gantt-editor"', html=False)
        self.assertContains(response, 'id="typical-service-term-gantt-cancel-btn"', html=False)
        self.assertContains(response, 'id="typical-service-term-gantt-resources-btn"', html=False)
        self.assertContains(response, 'id="typical-service-term-gantt-resources"', html=False)
        self.assertContains(response, reverse("typical_service_term_gantt", args=[TypicalServiceTerm.objects.get().pk]))
        self.assertContains(response, 'id="typical-service-terms-csv-download-btn"', html=False)
        self.assertContains(response, 'id="typical-service-terms-csv-upload-btn"', html=False)

    def test_create_typical_service_term_saves_row(self):
        response = self.client.post(
            reverse("typical_service_term_form_create"),
            {
                "product": self.product.pk,
                "preliminary_report_months": "2.5",
                "final_report_weeks": "4",
            },
        )

        self.assertEqual(response.status_code, 200)
        item = TypicalServiceTerm.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.preliminary_report_months, Decimal("2.5"))
        self.assertEqual(item.final_report_weeks, 4)
        self.assertEqual(item.position, 1)

    def test_typical_service_term_form_renders_product_options_with_display_name(self):
        response = self.client.get(reverse("typical_service_term_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<select name="product"', html=False)
        self.assertContains(response, "policy-product-select")
        self.assertContains(response, 'data-short-label="TERM"', html=False)
        self.assertContains(response, "TERM Terms")

    def test_create_typical_service_term_accepts_comma_decimal(self):
        response = self.client.post(
            reverse("typical_service_term_form_create"),
            {
                "product": self.product.pk,
                "preliminary_report_months": "1,5",
                "final_report_weeks": "2",
            },
        )

        self.assertEqual(response.status_code, 200)
        item = TypicalServiceTerm.objects.get()
        self.assertEqual(item.preliminary_report_months, Decimal("1.5"))

    def test_edit_form_renders_comma_decimal_value(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.5"),
            final_report_weeks=3,
            position=1,
        )

        response = self.client.get(reverse("typical_service_term_form_edit", args=[item.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'value="1,5"', html=False)

    def test_typical_service_term_gantt_get_returns_default_diagram(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.5"),
            final_report_weeks=3,
            position=1,
        )
        TypicalSection.objects.create(
            product=self.product,
            code="TERM-S1",
            short_name="term-s1",
            short_name_ru="term-s1",
            name_en="Term section",
            name_ru="Раздел продукта",
            accounting_type="Раздел",
            position=1,
        )
        TypicalSection.objects.create(
            product=self.other_product,
            code="TERM2-S1",
            short_name="term2-s1",
            short_name_ru="term2-s1",
            name_en="Other term section",
            name_ru="Раздел другого продукта",
            accounting_type="Раздел",
            position=1,
        )

        response = self.client.get(reverse("typical_service_term_gantt", args=[item.pk]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["term"]["product"], "TERM")
        tasks = payload["gantt"]["data"]
        self.assertEqual(
            {task["system_key"] for task in tasks},
            {
                "source_data",
                "source_data_asset",
                "preliminary_report",
                "preliminary_report_asset",
                "preliminary_report_submission",
                "final_report",
            },
        )
        source_data = next(task for task in tasks if task["system_key"] == "source_data")
        source_asset = next(task for task in tasks if task["system_key"] == "source_data_asset")
        preliminary = next(task for task in tasks if task["system_key"] == "preliminary_report")
        asset = next(task for task in tasks if task["system_key"] == "preliminary_report_asset")
        submission = next(task for task in tasks if task["system_key"] == "preliminary_report_submission")
        self.assertEqual(source_asset["text"], "Актив")
        self.assertEqual(source_asset["parent"], source_data["id"])
        self.assertEqual(asset["text"], "Актив")
        self.assertEqual(asset["parent"], preliminary["id"])
        self.assertEqual(asset["start_date"], preliminary["start_date"])
        self.assertEqual(asset["end_date"], preliminary["end_date"])
        self.assertEqual(submission["text"], "Отправка Предварительного отчёта")
        self.assertEqual(submission["type"], "milestone")
        self.assertEqual(submission["start_date"], preliminary["end_date"])
        self.assertEqual(submission["end_date"], preliminary["end_date"])
        self.assertEqual(len(payload["gantt"]["links"]), 3)
        self.assertEqual(payload["gantt"]["meta"]["project_start"], payload["gantt"]["meta"]["base_date"])
        self.assertIn("project_end", payload["gantt"]["meta"])
        self.assertEqual(payload["gantt"]["meta"]["calendar_kind"], "abstract")
        self.assertEqual(payload["gantt"]["meta"]["executor_display"], "resource_name")
        self.assertEqual(
            payload["section_options"],
            [{"id": TypicalSection.objects.get(product=self.product).pk, "label": "Раздел продукта", "specialties": []}],
        )

    def test_typical_service_term_gantt_get_adds_asset_task_to_existing_diagram_without_saving(self):
        source_gantt = {
            "data": [
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01", "calendar_kind": "abstract"},
        }
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
            gantt_data=copy.deepcopy(source_gantt),
        )

        response = self.client.get(reverse("typical_service_term_gantt", args=[item.pk]))

        self.assertEqual(response.status_code, 200)
        tasks = response.json()["gantt"]["data"]
        source_data = next(task for task in tasks if task.get("system_key") == "source_data")
        source_asset = next(task for task in tasks if task.get("system_key") == "source_data_asset")
        self.assertEqual(source_data["text"], "Исходные данные")
        self.assertEqual(source_asset["text"], "Актив")
        self.assertEqual(source_asset["parent"], source_data["id"])
        submission = next(task for task in tasks if task.get("system_key") == "preliminary_report_submission")
        self.assertEqual(submission["text"], "Отправка Предварительного отчёта")
        self.assertEqual(submission["type"], "milestone")
        asset = next(task for task in tasks if task.get("system_key") == "preliminary_report_asset")
        self.assertEqual(asset["text"], "Актив")
        self.assertEqual(asset["parent"], "preliminary")
        item.refresh_from_db()
        self.assertEqual(item.gantt_data, source_gantt)

    def test_typical_service_term_gantt_post_rejects_production_calendar_payload(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        payload = {
            "data": [
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01", "calendar_kind": "production"},
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("условном календаре", response.json()["error"])

    def test_typical_service_term_gantt_returns_assignment_options_and_autofills_section_specialty(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        geology = ExpertSpecialty.objects.create(specialty="Геология", position=1)
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=2)
        section = TypicalSection.objects.create(
            product=self.product,
            code="TERM-S1",
            short_name="term-s1",
            short_name_ru="term-s1",
            name_en="Term section",
            name_ru="Раздел продукта",
            accounting_type="Раздел",
            position=1,
        )
        TypicalSectionSpecialty.objects.create(section=section, specialty=geology, rank=2)
        TypicalSectionSpecialty.objects.create(section=section, specialty=mining, rank=1)
        executor_user = get_user_model().objects.create_user(
            username="executor-terms",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=executor_user, patronymic="Петрович")
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=mining, rank=1)
        other_executor_user = get_user_model().objects.create_user(
            username="executor-terms-geology",
            first_name="Петр",
            last_name="Петров",
        )
        other_employee = Employee.objects.create(user=other_executor_user, patronymic="Сергеевич")
        other_profile = ExpertProfile.objects.create(employee=other_employee, position=2)
        ExpertProfileSpecialty.objects.create(profile=other_profile, specialty=geology, rank=1)

        response = self.client.get(reverse("typical_service_term_gantt", args=[item.pk]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["specialty_options"], ["Геология", "Горное дело"])
        self.assertEqual(
            payload["executor_options"],
            [
                {
                    "id": profile.pk,
                    "value": f"expert-profile:{profile.pk}",
                    "label": "Иванов И.П.",
                    "specialties": ["Горное дело"],
                },
                {
                    "id": other_profile.pk,
                    "value": f"expert-profile:{other_profile.pk}",
                    "label": "Петров П.С.",
                    "specialties": ["Геология"],
                },
            ],
        )
        self.assertEqual(
            payload["section_options"][0]["specialties"],
            [{"label": "Горное дело", "rank": 1}, {"label": "Геология", "rank": 2}],
        )

        save_payload = {
            "data": [
                {
                    "id": "section-task",
                    "text": "Раздел продукта",
                    "type": "service_section",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                    "specialty": "",
                    "executor": f"expert-profile:{profile.pk}",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01"},
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(save_payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        item.refresh_from_db()
        self.assertEqual(item.gantt_data["data"][0]["specialty"], "Горное дело")
        self.assertEqual(item.gantt_data["data"][0]["executor"], f"expert-profile:{profile.pk}")

    def test_typical_service_term_gantt_distinguishes_duplicate_executor_labels(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=1)
        geology = ExpertSpecialty.objects.create(specialty="Геология", position=2)
        mining_user = get_user_model().objects.create_user(
            username="duplicate-executor-mining",
            first_name="Иван",
            last_name="Иванов",
        )
        mining_employee = Employee.objects.create(user=mining_user, patronymic="Петрович")
        mining_profile = ExpertProfile.objects.create(employee=mining_employee, position=1)
        ExpertProfileSpecialty.objects.create(profile=mining_profile, specialty=mining, rank=1)
        geology_user = get_user_model().objects.create_user(
            username="duplicate-executor-geology",
            first_name="Иван",
            last_name="Иванов",
        )
        geology_employee = Employee.objects.create(user=geology_user, patronymic="Петрович")
        geology_profile = ExpertProfile.objects.create(employee=geology_employee, position=2)
        ExpertProfileSpecialty.objects.create(profile=geology_profile, specialty=geology, rank=1)

        response = self.client.get(reverse("typical_service_term_gantt", args=[item.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json()["executor_options"],
            [
                {
                    "id": mining_profile.pk,
                    "value": f"expert-profile:{mining_profile.pk}",
                    "label": "Иванов И.П.",
                    "specialties": ["Горное дело"],
                },
                {
                    "id": geology_profile.pk,
                    "value": f"expert-profile:{geology_profile.pk}",
                    "label": "Иванов И.П.",
                    "specialties": ["Геология"],
                },
            ],
        )

        payload = {
            "data": [
                {
                    "id": "task",
                    "text": "Проектная задача",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                    "type": "task",
                    "specialty": mining.specialty,
                    "executor": f"expert-profile:{geology_profile.pk}",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-08",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01"},
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("Выберите исполнителя, связанного с выбранной специальностью.", response.json()["error"])

    def test_typical_service_term_gantt_post_saves_project_resources_meta(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=1)
        executor_user = get_user_model().objects.create_user(
            username="resource-executor-terms",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=executor_user, patronymic="Петрович")
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=mining, rank=1)
        payload = {
            "data": [
                {
                    "id": "resource-task",
                    "text": "Проектная задача",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                    "type": "task",
                    "specialty": "Горное дело",
                    "executor": "Иванов И.П.",
                    "resource_id": "resource-1",
                    "resource_name": "Сотрудник 1",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {
                "base_date": "2026-01-01",
                "resources": [
                    {
                        "id": "resource-1",
                        "specialty": "Горное дело",
                        "executor": "Иванов И.П.",
                        "resource_name": "Горный эксперт",
                        "task_ids": ["resource-task"],
                    }
                ],
            },
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        item.refresh_from_db()
        self.assertEqual(
            item.gantt_data["meta"]["resources"],
            [
                {
                    "id": "resource-1",
                    "specialty": "Горное дело",
                    "executor": "Иванов И.П.",
                    "resource_name": "Сотрудник 1",
                    "task_ids": ["resource-task"],
                    "position": 1,
                }
            ],
        )
        self.assertEqual(item.gantt_data["data"][0]["resource_name"], "Сотрудник 1")

    def test_typical_service_term_gantt_post_renames_duplicate_resource_ids(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=1)
        geology = ExpertSpecialty.objects.create(specialty="Геология", position=2)
        first_user = get_user_model().objects.create_user(
            username="resource-dup-first",
            first_name="Иван",
            last_name="Иванов",
        )
        second_user = get_user_model().objects.create_user(
            username="resource-dup-second",
            first_name="Петр",
            last_name="Петров",
        )
        first_employee = Employee.objects.create(user=first_user, patronymic="Петрович")
        second_employee = Employee.objects.create(user=second_user, patronymic="Петрович")
        first_profile = ExpertProfile.objects.create(employee=first_employee, position=1)
        second_profile = ExpertProfile.objects.create(employee=second_employee, position=2)
        ExpertProfileSpecialty.objects.create(profile=first_profile, specialty=mining, rank=1)
        ExpertProfileSpecialty.objects.create(profile=second_profile, specialty=geology, rank=1)
        payload = {
            "data": [
                {
                    "id": "resource-task-1",
                    "text": "Первая задача",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                    "type": "task",
                    "specialty": "Горное дело",
                    "executor": "Иванов И.П.",
                    "resource_id": "resource-1",
                },
                {
                    "id": "resource-task-2",
                    "text": "Вторая задача",
                    "start_date": "2026-01-08",
                    "end_date": "2026-01-15",
                    "type": "task",
                    "specialty": "Геология",
                    "executor": "Петров П.П.",
                    "resource_id": "resource-1",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {
                "base_date": "2026-01-01",
                "resources": [
                    {
                        "id": "resource-1",
                        "specialty": "Горное дело",
                        "executor": "Иванов И.П.",
                        "task_ids": ["resource-task-1"],
                    },
                    {
                        "id": "resource-1",
                        "specialty": "Геология",
                        "executor": "Петров П.П.",
                        "task_ids": ["resource-task-2"],
                    },
                ],
            },
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200, response.content.decode("utf-8"))
        item.refresh_from_db()
        resource_ids = [resource["id"] for resource in item.gantt_data["meta"]["resources"]]
        self.assertEqual(resource_ids, ["resource-1", "resource-1-2"])
        second_task = next(task for task in item.gantt_data["data"] if task["id"] == "resource-task-2")
        self.assertEqual(second_task["resource_id"], "resource-1-2")

    def test_typical_service_term_gantt_post_rejects_resource_executor_from_other_specialty(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=1)
        geology = ExpertSpecialty.objects.create(specialty="Геология", position=2)
        executor_user = get_user_model().objects.create_user(
            username="resource-executor-invalid",
            first_name="Петр",
            last_name="Петров",
        )
        employee = Employee.objects.create(user=executor_user, patronymic="Сергеевич")
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=geology, rank=1)
        payload = {
            "data": [
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {
                "base_date": "2026-01-01",
                "resources": [
                    {
                        "id": "resource-1",
                        "specialty": mining.specialty,
                        "executor": "Петров П.С.",
                        "resource_name": "Некорректный ресурс",
                        "task_ids": [],
                    }
                ],
            },
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("Выберите исполнителя, связанного с выбранной специальностью.", response.json()["error"])
        item.refresh_from_db()
        self.assertEqual(item.gantt_data, {})

    def test_typical_service_term_gantt_post_rejects_duplicate_project_resource_pair(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=1)
        executor_user = get_user_model().objects.create_user(
            username="resource-duplicate-pair",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=executor_user, patronymic="Петрович")
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=mining, rank=1)
        payload = {
            "data": [
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {
                "base_date": "2026-01-01",
                "resources": [
                    {"id": "resource-1", "specialty": mining.specialty, "executor": "Иванов И.П.", "task_ids": []},
                    {"id": "resource-2", "specialty": mining.specialty, "executor": "Иванов И.П.", "task_ids": []},
                ],
            },
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("Ресурс с такой специальностью и ФИО уже есть в таблице.", response.json()["error"])

    def test_typical_service_term_gantt_post_rejects_resource_task_from_unavailable_section_specialty(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=1)
        geology = ExpertSpecialty.objects.create(specialty="Геология", position=2)
        section = TypicalSection.objects.create(
            product=self.product,
            code="TERM-S2",
            short_name="term-s2",
            short_name_ru="term-s2",
            name_en="Term section 2",
            name_ru="Раздел геологов",
            accounting_type="Раздел",
            position=1,
        )
        TypicalSectionSpecialty.objects.create(section=section, specialty=geology, rank=1)
        executor_user = get_user_model().objects.create_user(
            username="resource-unavailable-section",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=executor_user, patronymic="Петрович")
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=mining, rank=1)
        payload = {
            "data": [
                {
                    "id": "section-task",
                    "text": "Раздел геологов",
                    "type": "service_section",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                },
                {
                    "id": "child-task",
                    "text": "Операция раздела",
                    "parent": "section-task",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                    "type": "task",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-08",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {
                "base_date": "2026-01-01",
                "resources": [
                    {
                        "id": "resource-1",
                        "specialty": mining.specialty,
                        "executor": "Иванов И.П.",
                        "task_ids": ["child-task"],
                    }
                ],
            },
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn(
            "Выберите задачу из раздела (услуги), доступного выбранной специальности ресурса.",
            response.json()["error"],
        )

    def test_typical_service_term_gantt_post_rejects_resource_parent_task_assignment(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        mining = ExpertSpecialty.objects.create(specialty="Горное дело", position=1)
        executor_user = get_user_model().objects.create_user(
            username="resource-parent-assignment",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=executor_user, patronymic="Петрович")
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=mining, rank=1)
        payload = {
            "data": [
                {
                    "id": "parent-task",
                    "text": "Родительская задача",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                    "type": "project",
                    "specialty": mining.specialty,
                    "executor": "Иванов И.П.",
                    "resource_id": "resource-1",
                },
                {
                    "id": "child-task",
                    "text": "Дочерняя задача",
                    "parent": "parent-task",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                    "type": "task",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-08",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {
                "base_date": "2026-01-01",
                "resources": [
                    {
                        "id": "resource-1",
                        "specialty": mining.specialty,
                        "executor": "Иванов И.П.",
                        "resource_name": "Горный эксперт",
                        "task_ids": ["parent-task"],
                    }
                ],
            },
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("Родительскую задачу нельзя назначить ресурсу проекта.", response.json()["error"])
        item.refresh_from_db()
        self.assertEqual(item.gantt_data, {})

    def test_typical_service_term_gantt_post_saves_diagram_and_updates_terms(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        payload = {
            "data": [
                {
                    "id": "preliminary",
                    "text": "Переименованный предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-04-01",
                    "progress": 0,
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Переименованный итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-04-01",
                    "end_date": "2026-05-13",
                    "progress": 0,
                    "type": "task",
                },
            ],
            "links": [{"id": "link-1", "source": "preliminary", "target": "final", "type": "0"}],
            "meta": {
                "base_date": "2026-01-01",
                "project_start": "2026-01-01",
                "project_end": "2026-06-01",
            },
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        item.refresh_from_db()
        self.assertEqual(item.source_data_weeks, 0)
        self.assertEqual(item.preliminary_report_months, Decimal("2.0"))
        self.assertEqual(item.final_report_weeks, 6)
        source_data = next(
            task for task in item.gantt_data["data"]
            if task.get("system_key") == "source_data"
        )
        self.assertEqual(source_data["text"], "Исходные данные")
        preliminary = next(
            task for task in item.gantt_data["data"]
            if task.get("system_key") == "preliminary_report"
        )
        self.assertEqual(preliminary["text"], "Предварительный отчёт")
        asset = next(
            task for task in item.gantt_data["data"]
            if task.get("system_key") == "preliminary_report_asset"
        )
        self.assertEqual(asset["text"], "Актив")
        self.assertEqual(asset["parent"], "preliminary")
        self.assertEqual(asset["start_date"], preliminary["start_date"])
        self.assertEqual(asset["end_date"], preliminary["end_date"])
        submission = next(
            task for task in item.gantt_data["data"]
            if task.get("system_key") == "preliminary_report_submission"
        )
        self.assertEqual(submission["text"], "Отправка Предварительного отчёта")
        self.assertEqual(submission["type"], "milestone")
        final = next(
            task for task in item.gantt_data["data"]
            if task.get("system_key") == "final_report"
        )
        self.assertEqual(final["text"], "Итоговый отчёт")
        self.assertEqual(item.gantt_data["links"][0]["source"], "preliminary")
        self.assertEqual(item.gantt_data["meta"]["project_start"], "2026-01-01")
        self.assertEqual(item.gantt_data["meta"]["project_end"], "2026-06-01")

    def test_typical_service_term_gantt_post_accepts_product_section_task_type(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        TypicalSection.objects.create(
            product=self.product,
            code="TERM-S1",
            short_name="term-s1",
            short_name_ru="term-s1",
            name_en="Term section",
            name_ru="Раздел продукта",
            accounting_type="Раздел",
            position=1,
        )
        payload = {
            "data": [
                {
                    "id": "section-task",
                    "text": "Раздел продукта",
                    "type": "service_section",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01"},
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        item.refresh_from_db()
        self.assertEqual(item.gantt_data["data"][0]["type"], "service_section")
        self.assertEqual(item.gantt_data["data"][0]["text"], "Раздел продукта")

    def test_typical_service_term_gantt_post_rejects_section_from_other_product(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        TypicalSection.objects.create(
            product=self.other_product,
            code="TERM2-S1",
            short_name="term2-s1",
            short_name_ru="term2-s1",
            name_en="Other term section",
            name_ru="Чужой раздел",
            accounting_type="Раздел",
            position=1,
        )
        payload = {
            "data": [
                {
                    "id": "section-task",
                    "text": "Чужой раздел",
                    "type": "service_section",
                    "start_date": "2026-01-01",
                    "end_date": "2026-01-08",
                },
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01"},
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("Выберите раздел", response.json()["error"])
        item.refresh_from_db()
        self.assertEqual(item.gantt_data, {})

    def test_typical_service_term_gantt_post_rolls_up_parent_dates_from_children(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        payload = {
            "data": [
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "progress": 0,
                    "type": "task",
                },
                {
                    "id": "preliminary-child",
                    "text": "Подзадача предварительного отчёта",
                    "parent": "preliminary",
                    "start_date": "2026-01-15",
                    "end_date": "2026-04-01",
                    "progress": 0,
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "start_date": "2026-04-01",
                    "end_date": "2026-04-15",
                    "progress": 0,
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01"},
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        item.refresh_from_db()
        parent_task = next(
            task for task in item.gantt_data["data"]
            if task.get("system_key") == "preliminary_report"
        )
        self.assertEqual(parent_task["end_date"], "2026-04-01")
        self.assertEqual(parent_task["type"], "project")
        asset_task = next(
            task for task in item.gantt_data["data"]
            if task.get("system_key") == "preliminary_report_asset"
        )
        self.assertEqual(asset_task["start_date"], parent_task["start_date"])
        self.assertEqual(asset_task["end_date"], parent_task["end_date"])
        self.assertEqual(item.preliminary_report_months, Decimal("3.0"))

    def test_typical_service_term_gantt_post_rejects_cyclic_parent_chain(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        payload = {
            "data": [
                {
                    "id": "preliminary",
                    "text": "Предварительный отчёт",
                    "system_key": "preliminary_report",
                    "parent": "final",
                    "start_date": "2026-01-01",
                    "end_date": "2026-02-01",
                    "type": "task",
                },
                {
                    "id": "final",
                    "text": "Итоговый отчёт",
                    "system_key": "final_report",
                    "parent": "preliminary",
                    "start_date": "2026-02-01",
                    "end_date": "2026-02-15",
                    "type": "task",
                },
            ],
            "links": [],
            "meta": {"base_date": "2026-01-01"},
        }

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps(payload, ensure_ascii=False),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("циклическая связь", response.json()["error"])
        item.refresh_from_db()
        self.assertEqual(item.gantt_data, {})

    def test_typical_service_term_gantt_post_requires_system_tasks(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )

        response = self.client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps({"data": [{"id": "task", "start_date": "2026-01-01", "end_date": "2026-01-08"}]}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        item.refresh_from_db()
        self.assertEqual(item.preliminary_report_months, Decimal("1.0"))
        self.assertEqual(item.final_report_weeks, 2)

    def test_non_staff_user_cannot_edit_typical_service_term_gantt(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        non_staff = get_user_model().objects.create_user(
            username="policy-user-terms-gantt",
            password="secret123",
            is_staff=False,
        )
        client = self.client_class()
        client.force_login(non_staff)

        response = client.post(
            reverse("typical_service_term_gantt", args=[item.pk]),
            data=json.dumps({"data": []}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 302)
        item.refresh_from_db()
        self.assertEqual(item.gantt_data, {})

    def test_typical_service_term_csv_download_exports_current_table_columns(self):
        TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.5"),
            final_report_weeks=3,
            position=1,
        )

        response = self.client.get(reverse("typical_service_term_csv_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("typical_service_terms.csv", response["Content-Disposition"])
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(
            rows[0],
            [
                "Продукт",
                "Сроки предоставления исходных данных",
                "Единица срока предоставления исходных данных",
                "Срок подготовки Предварительного отчёта",
                "Единица срока подготовки Предварительного отчёта",
                "Срок подготовки Итогового отчёта",
                "Единица срока подготовки Итогового отчёта",
            ],
        )
        self.assertEqual(rows[1], ["TERM", "0,0", "нед.", "1,5", "мес.", "3,0", "нед."])

    def test_typical_service_term_csv_download_respects_product_filter(self):
        TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.5"),
            final_report_weeks=3,
            position=1,
        )
        TypicalServiceTerm.objects.create(
            product=self.other_product,
            preliminary_report_months=Decimal("2.0"),
            final_report_weeks=5,
            position=2,
        )

        response = self.client.get(
            reverse("typical_service_term_csv_download"),
            {"product": [self.product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "TERM")

    def test_typical_service_term_csv_upload_creates_rows(self):
        csv_file = SimpleUploadedFile(
            "typical_service_terms.csv",
            (
                "Продукт;Срок подготовки Предварительного отчёта, мес.;Срок подготовки Итогового отчёта, нед.\n"
                "TERM;2,5;4\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("typical_service_term_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        item = TypicalServiceTerm.objects.get()
        self.assertEqual(item.product, self.product)
        self.assertEqual(item.source_data_weeks, 0)
        self.assertEqual(item.preliminary_report_months, Decimal("2.5"))
        self.assertEqual(item.final_report_weeks, 4)
        self.assertEqual(item.position, 1)

    def test_typical_service_term_csv_upload_accepts_source_data_weeks(self):
        csv_file = SimpleUploadedFile(
            "typical_service_terms.csv",
            (
                "Продукт;Сроки предоставления исходных данных;Единица срока предоставления исходных данных;"
                "Срок подготовки Предварительного отчёта;Единица срока подготовки Предварительного отчёта;"
                "Срок подготовки Итогового отчёта;Единица срока подготовки Итогового отчёта\n"
                "TERM;3;дн.;2,5;нед.;4;мес.\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("typical_service_term_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        item = TypicalServiceTerm.objects.get()
        self.assertEqual(item.source_data_weeks, 3)
        self.assertEqual(item.source_data_term_unit, "days")
        self.assertEqual(item.preliminary_report_months, Decimal("2.5"))
        self.assertEqual(item.preliminary_report_term_unit, "weeks")
        self.assertEqual(item.final_report_weeks, 4)
        self.assertEqual(item.final_report_term_unit, "months")

    def test_typical_service_term_csv_upload_updates_existing_product_row(self):
        item = TypicalServiceTerm.objects.create(
            product=self.product,
            source_data_weeks=1,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        csv_file = SimpleUploadedFile(
            "typical_service_terms.csv",
            (
                "Продукт;Сроки предоставления исходных данных, нед.;Срок подготовки Предварительного отчёта, мес.;Срок подготовки Итогового отчёта, нед.\n"
                "TERM;4;3,5;7\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("typical_service_term_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(response.json()["updated"], 1)
        self.assertEqual(TypicalServiceTerm.objects.count(), 1)
        item.refresh_from_db()
        self.assertEqual(item.source_data_weeks, 4)
        self.assertEqual(item.preliminary_report_months, Decimal("3.5"))
        self.assertEqual(item.final_report_weeks, 7)

    def test_non_staff_user_cannot_reorder_typical_service_terms(self):
        first = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months=Decimal("1.0"),
            final_report_weeks=2,
            position=1,
        )
        second = TypicalServiceTerm.objects.create(
            product=self.other_product,
            preliminary_report_months=Decimal("2.0"),
            final_report_weeks=4,
            position=2,
        )
        non_staff = get_user_model().objects.create_user(
            username="policy-user-terms",
            password="secret123",
            is_staff=False,
        )
        client = self.client_class()
        client.force_login(non_staff)

        response = client.post(reverse("typical_service_term_move_down", args=[first.pk]))

        self.assertEqual(response.status_code, 302)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(first.position, 1)
        self.assertEqual(second.position, 2)


class SpecialtyTariffViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-admin-3",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.company = GroupMember.objects.create(
            short_name="ACME",
            country_name="Россия",
            position=1,
        )
        self.direction = OrgUnit.objects.create(
            company=self.company,
            department_name="Налоговое консультирование",
            unit_type="expertise",
            position=1,
        )
        self.direction_2 = OrgUnit.objects.create(
            company=self.company,
            department_name="Сопровождение сделок",
            unit_type="expertise",
            position=2,
        )
        self.policy_direction_1 = ExpertiseDirection.objects.create(
            name="Налоги",
            short_name="TAX",
            position=1,
        )
        self.policy_direction_2 = ExpertiseDirection.objects.create(
            name="Сделки",
            short_name="M&A",
            position=2,
        )
        self.specialty_1 = ExpertSpecialty.objects.create(
            specialty="Налоговый due diligence",
            expertise_direction=self.direction,
            expertise_dir=self.policy_direction_1,
            position=1,
        )
        self.specialty_2 = ExpertSpecialty.objects.create(
            specialty="Трансфертное ценообразование",
            expertise_direction=self.direction,
            expertise_dir=self.policy_direction_1,
            position=2,
        )
        self.specialty_3 = ExpertSpecialty.objects.create(
            specialty="M&A",
            expertise_direction=self.direction_2,
            expertise_dir=self.policy_direction_2,
            position=3,
        )
        self.currency = OKVCurrency.objects.create(
            code_numeric="978",
            code_alpha="EUR",
            name="Евро",
            position=1,
        )

    def test_policy_partial_renders_specialty_tariffs_table(self):
        item = SpecialtyTariff.objects.create(
            specialty_group="Налоговые специалисты",
            daily_rate_tkp_eur="1500.00",
            daily_rate_ss="1250.00",
            currency=self.currency,
            position=1,
        )
        item.specialties.set([self.specialty_1, self.specialty_2, self.specialty_3])

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Тарифы специальностей")
        self.assertContains(response, "Налоговые специалисты")
        self.assertContains(response, "Налоговый due diligence")
        self.assertContains(response, "Трансфертное ценообразование")
        self.assertContains(response, "TAX, M&amp;A")
        self.assertContains(response, "1\xa0250,00 EUR")
        self.assertNotContains(response, "<th style=\"vertical-align: top;\">Валюта</th>", html=False)
        self.assertContains(response, 'id="specialty-tariffs-actions"', html=False)
        self.assertContains(response, 'id="specialty-tariffs-specialties-toggle"', html=False)

    def test_create_specialty_tariff_saves_row(self):
        response = self.client.post(
            reverse("specialty_tariff_form_create"),
            {
                "specialty_group": "Сделки",
                "specialties": [str(self.specialty_1.pk), str(self.specialty_2.pk), str(self.specialty_3.pk)],
                "daily_rate_tkp_eur": "1 500,50",
                "daily_rate_ss": "1200,25",
                "currency": self.currency.pk,
            },
        )

        self.assertEqual(response.status_code, 200)
        item = SpecialtyTariff.objects.get()
        self.assertEqual(item.specialty_group, "Сделки")
        self.assertEqual(item.expertise_direction_display, "TAX, M&A")
        self.assertEqual(str(item.daily_rate_tkp_eur), "1500.50")
        self.assertEqual(str(item.daily_rate_ss), "1200.25")
        self.assertEqual(item.currency, self.currency)
        self.assertEqual(item.created_by, self.user)
        self.assertEqual(item.position, 1)
        self.assertCountEqual(
            item.specialties.values_list("pk", flat=True),
            [self.specialty_1.pk, self.specialty_2.pk, self.specialty_3.pk],
        )

    def test_create_specialty_tariff_uses_unique_expertise_values(self):
        response = self.client.post(
            reverse("specialty_tariff_form_create"),
            {
                "specialty_group": "Налоги",
                "specialties": [str(self.specialty_1.pk), str(self.specialty_2.pk)],
                "daily_rate_tkp_eur": "900,00",
                "daily_rate_ss": "800,00",
                "currency": self.currency.pk,
            },
        )

        self.assertEqual(response.status_code, 200)
        item = SpecialtyTariff.objects.get(specialty_group="Налоги")
        self.assertEqual(item.expertise_direction_display, "TAX")


class TariffViewsTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="policy-admin-4",
            password="secret123",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="TAR",
            name_en="Tariff product",
            display_name="Tariff product",
            name_ru="Тарифный продукт",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )
        self.section = TypicalSection.objects.create(
            product=self.product,
            code="TS1",
            short_name="tariff-section",
            short_name_ru="tariff-section-ru",
            name_en="Tariff section EN",
            name_ru="Тарифный раздел",
            accounting_type="Раздел",
            position=1,
        )

    def test_policy_partial_renders_tariff_days_column(self):
        Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm="10.00",
            service_hours=8,
            service_days_tkp=5,
            created_by=self.user,
            position=1,
        )

        response = self.client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Объем услуг в днях для ТКП")
        self.assertContains(response, ">Код<", html=False)
        self.assertContains(response, "TS1")
        self.assertContains(response, ">5<", html=False)
        self.assertContains(response, 'id="tariffs-csv-download-btn"', html=False)
        self.assertContains(response, 'id="tariffs-csv-upload-btn"', html=False)

    def test_create_tariff_saves_service_days_tkp(self):
        response = self.client.post(
            reverse("tariff_form_create"),
            {
                "product": self.product.pk,
                "section": self.section.pk,
                "base_rate_vpm": "15.00",
                "service_hours": "12",
                "service_days_tkp": "7",
            },
        )

        self.assertEqual(response.status_code, 200)
        tariff = Tariff.objects.get()
        self.assertEqual(tariff.service_hours, 12)
        self.assertEqual(tariff.service_days_tkp, 7)

    def test_tariff_form_renders_product_options_with_display_name(self):
        response = self.client.get(reverse("tariff_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<select name="product"', html=False)
        self.assertContains(response, "policy-product-select")
        self.assertContains(response, 'data-short-label="TAR"', html=False)
        self.assertContains(response, "TAR Tariff product")
        self.assertContains(response, 'name="section_code"', html=False)
        self.assertContains(response, "readonly-field", html=False)
        self.assertContains(response, 'tabindex="-1"', html=False)
        self.assertContains(response, "policy-section-select")
        self.assertContains(response, '"label": "TS1 Тарифный раздел"', html=False)
        self.assertContains(response, '"displayLabel": "Тарифный раздел"', html=False)

    def test_tariff_csv_download_exports_current_table_columns(self):
        Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm=Decimal("10.00"),
            service_hours=8,
            service_days_tkp=5,
            created_by=self.user,
            position=1,
        )

        response = self.client.get(reverse("tariff_csv_download"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("section_tariffs.csv", response["Content-Disposition"])
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(
            rows[0],
            [
                "Продукт",
                "Код",
                "Раздел (услуга)",
                "Базовая ставка в ВПМ",
                "Объем услуг в часах",
                "Объем услуг в днях для ТКП",
                "Руководитель направления",
            ],
        )
        self.assertEqual(rows[1], ["TAR", "TS1", "Тарифный раздел", "10,00", "8", "5", "policy-admin-4"])

    def test_tariff_csv_download_respects_product_filter(self):
        other_product = Product.objects.create(
            short_name="AUD",
            name_en="Audit product",
            display_name="Audit product",
            name_ru="Аудитный продукт",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        other_section = TypicalSection.objects.create(
            product=other_product,
            code="AS1",
            short_name="audit-section",
            short_name_ru="audit-section-ru",
            name_en="Audit section EN",
            name_ru="Аудитный раздел",
            accounting_type="Раздел",
            position=1,
        )
        Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm=Decimal("10.00"),
            service_hours=8,
            service_days_tkp=5,
            created_by=self.user,
            position=1,
        )
        Tariff.objects.create(
            product=other_product,
            section=other_section,
            base_rate_vpm=Decimal("20.00"),
            service_hours=16,
            service_days_tkp=10,
            created_by=self.user,
            position=2,
        )

        response = self.client.get(
            reverse("tariff_csv_download"),
            {"product": [self.product.pk]},
        )

        self.assertEqual(response.status_code, 200)
        rows = list(csv.reader(io.StringIO(response.content.decode("utf-8-sig")), delimiter=";"))
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[1][0], "TAR")

    def test_tariff_csv_upload_creates_rows_for_current_user(self):
        csv_file = SimpleUploadedFile(
            "section_tariffs.csv",
            (
                "Продукт;Код;Раздел (услуга);Базовая ставка в ВПМ;Объем услуг в часах;Объем услуг в днях для ТКП\n"
                "TAR;TS1;Тарифный раздел;12,50;16;6\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("tariff_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        tariff = Tariff.objects.get()
        self.assertEqual(tariff.product, self.product)
        self.assertEqual(tariff.section, self.section)
        self.assertEqual(tariff.base_rate_vpm, Decimal("12.50"))
        self.assertEqual(tariff.service_hours, 16)
        self.assertEqual(tariff.service_days_tkp, 6)
        self.assertEqual(tariff.created_by, self.user)
        self.assertEqual(tariff.position, 1)

    def test_tariff_csv_upload_updates_existing_product_section_owner_row(self):
        tariff = Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm=Decimal("10.00"),
            service_hours=8,
            service_days_tkp=5,
            created_by=self.user,
            position=1,
        )
        csv_file = SimpleUploadedFile(
            "section_tariffs.csv",
            (
                "Продукт;Код;Раздел (услуга);Базовая ставка в ВПМ;Объем услуг в часах;Объем услуг в днях для ТКП\n"
                "TAR;TS1;Тарифный раздел;22,50;18;9\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("tariff_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 0)
        self.assertEqual(response.json()["updated"], 1)
        self.assertEqual(Tariff.objects.count(), 1)
        tariff.refresh_from_db()
        self.assertEqual(tariff.base_rate_vpm, Decimal("22.50"))
        self.assertEqual(tariff.service_hours, 18)
        self.assertEqual(tariff.service_days_tkp, 9)
        self.assertEqual(tariff.position, 1)

    def test_tariff_csv_upload_accepts_legacy_rows_without_code(self):
        csv_file = SimpleUploadedFile(
            "section_tariffs.csv",
            (
                "Продукт;Раздел (услуга);Базовая ставка в ВПМ;Объем услуг в часах;Объем услуг в днях для ТКП\n"
                "TAR;Тарифный раздел;12,50;16;6\n"
            ).encode("utf-8"),
            content_type="text/csv",
        )

        response = self.client.post(reverse("tariff_csv_upload"), {"csv_file": csv_file})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["created"], 1)
        self.assertEqual(response.json()["warnings"], [])
        self.assertEqual(Tariff.objects.get().section, self.section)

    def test_move_up_normalizes_positions_before_reorder(self):
        first = Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm="10.00",
            service_hours=8,
            service_days_tkp=2,
            created_by=self.user,
            position=1,
        )
        second = Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm="12.00",
            service_hours=10,
            service_days_tkp=3,
            created_by=self.user,
            position=3,
        )

        response = self.client.post(reverse("tariff_move_up", args=[second.pk]))

        self.assertEqual(response.status_code, 200)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(first.position, 2)
        self.assertEqual(second.position, 1)

    def test_policy_partial_orders_tariff_groups_by_employee_position(self):
        other_user = get_user_model().objects.create_user(
            username="policy-admin-5",
            password="secret123",
            is_staff=True,
        )
        Employee.objects.create(user=self.user, job_title="Руководитель 1", position=2)
        Employee.objects.create(user=other_user, job_title="Руководитель 2", position=1)

        first = Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm="10.00",
            service_hours=8,
            service_days_tkp=2,
            created_by=self.user,
            position=1,
        )
        second = Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm="12.00",
            service_hours=10,
            service_days_tkp=3,
            created_by=other_user,
            position=1,
        )

        admin_user = get_user_model().objects.create_superuser(
            username="policy-superuser",
            email="superuser@example.com",
            password="secret123",
        )
        client = self.client_class()
        client.force_login(admin_user)

        response = client.get(reverse("policy_partial"))

        self.assertEqual(response.status_code, 200)
        tariffs = list(response.context["tariffs"])
        self.assertEqual([item.pk for item in tariffs], [second.pk, first.pk])

    def test_tariff_verbose_name_plural_matches_application(self):
        self.assertEqual(Tariff._meta.verbose_name_plural, "Тарифы разделов (услуг)")


class TariffAdminTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.superuser = user_model.objects.create_superuser(
            username="admin-root",
            email="root@example.com",
            password="secret123",
        )
        self.owner_first = user_model.objects.create_user(
            username="dept-head-first",
            password="secret123",
            is_staff=True,
        )
        self.owner_second = user_model.objects.create_user(
            username="dept-head-second",
            password="secret123",
            is_staff=True,
        )
        self.first_profile = Employee.objects.create(
            user=self.owner_first,
            job_title="Первый руководитель",
            position=1,
        )
        self.second_profile = Employee.objects.create(
            user=self.owner_second,
            job_title="Второй руководитель",
            position=2,
        )
        self.product = Product.objects.create(
            short_name="TAR-ADMIN",
            name_en="Tariff product admin",
            display_name="Tariff product admin",
            name_ru="Тарифный продукт админ",
            consulting_type="Горный",
            service_category="Инжиниринг",
            service_subtype="По международным стандартам",
            position=1,
        )
        self.section = TypicalSection.objects.create(
            product=self.product,
            code="TSA",
            short_name="tariff-section-admin",
            short_name_ru="tariff-section-admin-ru",
            name_en="Tariff section admin EN",
            name_ru="Тарифный раздел админ",
            accounting_type="Раздел",
            position=1,
        )
        self.first_tariff = Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm="10.00",
            service_hours=8,
            service_days_tkp=2,
            created_by=self.owner_first,
            position=1,
        )
        self.second_tariff = Tariff.objects.create(
            product=self.product,
            section=self.section,
            base_rate_vpm="12.00",
            service_hours=10,
            service_days_tkp=3,
            created_by=self.owner_second,
            position=1,
        )
        self.client.force_login(self.superuser)

    def test_admin_move_owner_down_swaps_group_positions(self):
        response = self.client.post(
            reverse("admin:policy_app_tariff_move_owner_down", args=[self.first_tariff.pk])
        )

        self.assertEqual(response.status_code, 302)
        self.first_profile.refresh_from_db()
        self.second_profile.refresh_from_db()
        self.assertEqual(self.first_profile.position, 2)
        self.assertEqual(self.second_profile.position, 1)

    def test_admin_move_owner_down_rejects_get(self):
        response = self.client.get(
            reverse("admin:policy_app_tariff_move_owner_down", args=[self.first_tariff.pk])
        )

        self.assertEqual(response.status_code, 405)
        self.first_profile.refresh_from_db()
        self.second_profile.refresh_from_db()
        self.assertEqual(self.first_profile.position, 1)
        self.assertEqual(self.second_profile.position, 2)

    def test_admin_move_owner_down_requires_csrf_token(self):
        csrf_client = Client(enforce_csrf_checks=True)
        csrf_client.force_login(self.superuser)

        response = csrf_client.post(
            reverse("admin:policy_app_tariff_move_owner_down", args=[self.first_tariff.pk])
        )

        self.assertEqual(response.status_code, 403)
        self.first_profile.refresh_from_db()
        self.second_profile.refresh_from_db()
        self.assertEqual(self.first_profile.position, 1)
        self.assertEqual(self.second_profile.position, 2)
