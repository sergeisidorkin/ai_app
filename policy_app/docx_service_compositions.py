from __future__ import annotations

import html
import io
from collections import OrderedDict

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from contracts_app.docx_processor import (
    _ensure_marker_numberings,
    _ensure_multilevel_numbering,
    _ensure_quill_ordered_numbering,
    _find_builtin_list_paragraph_style_id,
    _find_bullet_style_id,
    _find_list_paragraph_style_id,
    _insert_rich_paragraphs,
    _rich_paragraphs_from_html,
)


DOCX_HEADERS = ["ID", "Продукт", "Раздел (услуга)", "Состав услуг"]
_LIST_TYPES = {"ordered", "bullet", "circle", "square", "dash", "ndash", "check"}
_BULLET_MARKER_MAP = {
    "\uf0b7": "bullet",
    "\u2022": "bullet",
    "•": "bullet",
    "-": "dash",
    "\u2013": "ndash",
    "–": "ndash",
    "\uf0fc": "check",
    "✓": "check",
    "o": "circle",
    "\uf0a7": "square",
    "▪": "square",
}


def build_typical_service_compositions_docx(rows: list[dict[str, object]]) -> bytes:
    document = Document()
    marker_num_ids = _ensure_marker_numberings(document)
    try:
        multilevel_num_id = _ensure_multilevel_numbering(document)
    except Exception:
        multilevel_num_id = ""
    try:
        rich_ordered_num_id = _ensure_quill_ordered_numbering(document)
    except Exception:
        rich_ordered_num_id = ""
    bullet_style_id = _find_bullet_style_id(document)
    list_paragraph_style_id = _find_list_paragraph_style_id(document)
    builtin_list_paragraph_style_id = _find_builtin_list_paragraph_style_id(document)

    product_groups: OrderedDict[str, list[dict[str, object]]] = OrderedDict()
    for item in rows:
        product_name = str(item.get("product") or "").strip()
        product_groups.setdefault(product_name, []).append(item)

    for product_name, product_rows in product_groups.items():
        document.add_heading(product_name or "—", level=1)
        table = document.add_table(rows=1, cols=len(DOCX_HEADERS))
        table.style = "Table Grid"
        _write_table_header_row(table)

        for item in product_rows:
            cells = table.add_row().cells
            cells[0].text = str(item.get("id") or "")
            cells[1].text = product_name
            cells[2].text = str(item.get("section") or "")
            _write_rich_service_cell(
                cells[3],
                html_value=str(item.get("html") or ""),
                plain_text=str(item.get("plain_text") or ""),
                marker_num_ids=marker_num_ids,
                multilevel_num_id=multilevel_num_id,
                rich_ordered_num_id=rich_ordered_num_id,
                bullet_style_id=bullet_style_id,
                list_paragraph_style_id=list_paragraph_style_id,
                builtin_list_paragraph_style_id=builtin_list_paragraph_style_id,
            )

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def parse_typical_service_compositions_docx(file_obj) -> list[dict[str, object]]:
    document = Document(file_obj)
    sections = _iter_composition_sections(document)
    if not sections:
        raise ValueError("В DOCX не найдена таблица с заголовками: ID, Продукт, Раздел (услуга), Состав услуг.")

    numbering = _build_numbering_lookup(document)
    rows: list[dict[str, object]] = []

    for section_index, (product_heading, table) in enumerate(sections, start=1):
        header_indexes = _table_header_indexes(table)
        for data_row_index, row in enumerate(table.rows[1:], start=1):
            cells = row.cells
            if not any(_cell_text(cell) for cell in cells):
                continue
            try:
                id_cell = cells[header_indexes[_normalize_header("ID")]]
                product_cell = cells[header_indexes[_normalize_header("Продукт")]]
                section_cell = cells[header_indexes[_normalize_header("Раздел (услуга)")]]
                service_cell = cells[header_indexes[_normalize_header("Состав услуг")]]
            except (IndexError, KeyError):
                raise ValueError("В DOCX нарушена структура таблицы типового состава услуг.")

            product_name = _cell_text(product_cell) or product_heading
            editor_state = _cell_editor_state(service_cell, numbering)
            rows.append(
                {
                    "row_number": f"{section_index}.{data_row_index}",
                    "id": _cell_text(id_cell),
                    "product": product_name,
                    "section": _cell_text(section_cell),
                    "editor_state": editor_state,
                }
            )

    return rows


def _write_rich_service_cell(
    cell,
    *,
    html_value: str,
    plain_text: str,
    marker_num_ids: dict[str, str],
    multilevel_num_id: str,
    rich_ordered_num_id: str,
    bullet_style_id: str,
    list_paragraph_style_id: str,
    builtin_list_paragraph_style_id: str,
) -> None:
    rich_items = _rich_paragraphs_from_html(html_value) if html_value.strip() else _plain_text_rich_items(plain_text)
    if not rich_items:
        cell.text = ""
        return
    anchor = cell.paragraphs[0]._p
    parent = anchor.getparent()
    _insert_rich_paragraphs(
        anchor,
        parent,
        rich_items,
        marker_num_ids=marker_num_ids,
        multilevel_num_id=multilevel_num_id,
        rich_ordered_num_id=rich_ordered_num_id,
        bullet_style_id=bullet_style_id,
        list_paragraph_style_id=list_paragraph_style_id,
        builtin_list_paragraph_style_id=builtin_list_paragraph_style_id,
        source_rPr=None,
        language_code="ru-RU",
    )


def _plain_text_rich_items(value: str) -> list[dict[str, object]]:
    text = str(value or "")
    if not text:
        return []
    lines = text.splitlines() or [text]
    return [
        {"runs": [{"text": line}], "alignment": None, "list_type": "", "list_level": 0}
        for line in lines
    ]


def _write_table_header_row(table) -> None:
    for index, header in enumerate(DOCX_HEADERS):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _table_header_indexes(table) -> dict[str, int]:
    header_cells = table.rows[0].cells
    return {
        _normalize_header(_cell_text(cell)): index
        for index, cell in enumerate(header_cells)
    }


def _is_compositions_table(table) -> bool:
    if not table.rows:
        return False
    headers = [_normalize_header(_cell_text(cell)) for cell in table.rows[0].cells]
    expected = [_normalize_header(header) for header in DOCX_HEADERS]
    return all(header in headers for header in expected)


def _iter_composition_sections(document) -> list[tuple[str, DocxTable]]:
    sections: list[tuple[str, DocxTable]] = []
    current_heading = ""

    for child in document.element.body:
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            if _is_heading_level_1(paragraph):
                current_heading = paragraph.text.strip()
            continue
        if child.tag != qn("w:tbl"):
            continue
        table = DocxTable(child, document)
        if not _is_compositions_table(table):
            continue
        sections.append((current_heading, table))

    return sections


def _is_heading_level_1(paragraph: Paragraph) -> bool:
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
    if style_name in {"heading 1", "заголовок 1"}:
        return True
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is not None and str(outline.get(qn("w:val")) or "") == "0":
        return True
    return False


def _cell_text(cell) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def _normalize_header(value: str) -> str:
    return " ".join(str(value or "").strip().lower().split())


def _cell_editor_state(cell, numbering: dict[str, dict[int, str]]) -> dict[str, str]:
    blocks: list[str] = []
    plain_lines: list[str] = []
    current_list_tag = ""

    def close_list() -> None:
        nonlocal current_list_tag
        if current_list_tag:
            blocks.append(f"</{current_list_tag}>")
            current_list_tag = ""

    for paragraph in cell.paragraphs:
        paragraph_text = paragraph.text.strip()
        if not paragraph_text and not any(run.text for run in paragraph.runs):
            continue

        plain_lines.append(paragraph_text)
        list_type, list_level = _paragraph_list_info(paragraph, numbering)
        content = _paragraph_runs_html(paragraph)
        if list_type:
            list_tag = "ol" if list_type == "ordered" else "ul"
            if current_list_tag != list_tag:
                close_list()
                blocks.append(f"<{list_tag}>")
                current_list_tag = list_tag
            attrs = []
            if list_level:
                attrs.append(f'class="ql-indent-{list_level}"')
            attrs.append(f'data-list="{list_type}"')
            blocks.append(f"<li {' '.join(attrs)}>{content}</li>")
            continue

        close_list()
        attrs = _paragraph_attrs(paragraph, include_indent=False)
        blocks.append(f"<p{attrs}>{content}</p>")

    close_list()
    plain_text = "\n".join(line for line in plain_lines if line).strip()
    return {"html": "".join(blocks), "plain_text": plain_text}


def _paragraph_runs_html(paragraph) -> str:
    chunks = []
    for run in paragraph.runs:
        if not run.text:
            continue
        chunks.append(_run_html(run))
    if chunks:
        return "".join(chunks)
    return html.escape(paragraph.text or "")


def _run_html(run) -> str:
    content = html.escape(run.text or "").replace("\n", "<br>")
    style_parts = []
    color = getattr(getattr(run.font, "color", None), "rgb", None)
    if color:
        style_parts.append(f"color:#{color}")
    background = _run_background(run)
    if background:
        style_parts.append(f"background-color:#{background}")
    if run.font.name:
        style_parts.append(f"font-family:{html.escape(run.font.name, quote=True)}")
    if style_parts:
        content = f'<span style="{";".join(style_parts)}">{content}</span>'
    if run.font.strike:
        content = f"<s>{content}</s>"
    if run.underline:
        content = f"<u>{content}</u>"
    if run.italic:
        content = f"<em>{content}</em>"
    if run.bold:
        content = f"<strong>{content}</strong>"
    return content


def _run_background(run) -> str:
    r_pr = run._r.rPr
    if r_pr is None:
        return ""
    shading = r_pr.find(qn("w:shd"))
    if shading is None:
        return ""
    fill = str(shading.get(qn("w:fill")) or "").strip()
    if not fill or fill.lower() in {"auto", "ffffff"}:
        return ""
    return fill.upper()


def _paragraph_attrs(paragraph, *, include_indent: bool) -> str:
    classes = []
    alignment = _paragraph_alignment(paragraph)
    if alignment and alignment != "left":
        classes.append(f"ql-align-{alignment}")
    if include_indent:
        _, list_level = _paragraph_list_info(paragraph, {})
        if list_level:
            classes.append(f"ql-indent-{list_level}")
    if not classes:
        return ""
    return f' class="{" ".join(classes)}"'


def _paragraph_alignment(paragraph) -> str:
    alignment = paragraph.alignment
    if alignment is None:
        p_pr = paragraph._p.pPr
        jc = p_pr.find(qn("w:jc")) if p_pr is not None else None
        raw = str(jc.get(qn("w:val")) or "") if jc is not None else ""
    else:
        raw = str(alignment).split(".")[-1].lower()
    return {
        "center": "center",
        "right": "right",
        "both": "justify",
        "justify": "justify",
        "left": "left",
    }.get(raw, "")


def _paragraph_list_info(paragraph, numbering: dict[str, dict[int, str]]) -> tuple[str, int]:
    p_pr = paragraph._p.pPr
    num_pr = p_pr.numPr if p_pr is not None else None
    if num_pr is None or num_pr.numId is None:
        return "", 0
    num_id = str(num_pr.numId.val)
    try:
        level = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
    except (TypeError, ValueError):
        level = 0
    level_map = numbering.get(num_id) or {}
    list_type = level_map.get(level) or level_map.get(0) or ""
    if list_type not in _LIST_TYPES:
        list_type = "ordered"
    return list_type, max(0, level)


def _build_numbering_lookup(document) -> dict[str, dict[int, str]]:
    numbering_part = getattr(document.part, "numbering_part", None)
    if numbering_part is None:
        return {}
    numbering_elm = numbering_part._element
    abstract_levels: dict[str, dict[int, tuple[str, str]]] = {}
    for abstract_num in numbering_elm.findall(qn("w:abstractNum")):
        abstract_id = abstract_num.get(qn("w:abstractNumId"))
        if not abstract_id:
            continue
        levels: dict[int, tuple[str, str]] = {}
        for level in abstract_num.findall(qn("w:lvl")):
            try:
                ilvl = int(level.get(qn("w:ilvl")) or 0)
            except (TypeError, ValueError):
                ilvl = 0
            num_fmt = level.find(qn("w:numFmt"))
            lvl_text = level.find(qn("w:lvlText"))
            levels[ilvl] = (
                str(num_fmt.get(qn("w:val")) or "") if num_fmt is not None else "",
                str(lvl_text.get(qn("w:val")) or "") if lvl_text is not None else "",
            )
        abstract_levels[abstract_id] = levels

    lookup: dict[str, dict[int, str]] = {}
    for num in numbering_elm.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abstract_ref = num.find(qn("w:abstractNumId"))
        abstract_id = abstract_ref.get(qn("w:val")) if abstract_ref is not None else ""
        if not num_id or not abstract_id:
            continue
        levels = abstract_levels.get(abstract_id) or {}
        default_marker = _classify_bullet_marker((levels.get(0) or ("", ""))[1])
        lookup[num_id] = {
            level: (
                "ordered"
                if fmt and fmt != "bullet"
                else default_marker or _classify_bullet_marker(marker) or "bullet"
            )
            for level, (fmt, marker) in levels.items()
        }
    return lookup


def _classify_bullet_marker(marker: str) -> str:
    return _BULLET_MARKER_MAP.get(str(marker or ""))
