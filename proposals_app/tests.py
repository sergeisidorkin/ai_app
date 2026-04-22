from __future__ import annotations

import base64
import json
import tempfile
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock, patch
from urllib.parse import parse_qs, quote, urlparse

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from classifiers_app.models import BusinessEntityRecord, BusinessEntityIdentifierRecord, LegalEntityRecord, OKSMCountry, OKVCurrency
from contacts_app.models import CitizenshipRecord, PersonRecord
from core.models import CloudStorageSettings
from experts_app.models import ExpertContractDetails, ExpertProfile, ExpertProfileSpecialty, ExpertSpecialty
from group_app.models import GroupMember, OrgUnit
from letters_app.models import LetterTemplate
from nextcloud_app.api import NextcloudShare
from nextcloud_app.models import NextcloudUserLink
from notifications_app.email_delivery import EmailDeliveryError
from policy_app.models import (
    ExpertiseDirection,
    Grade,
    Product,
    ServiceGoalReport,
    SpecialtyTariff,
    Tariff,
    TypicalSection,
    TypicalSectionSpecialty,
    TypicalServiceComposition,
    TypicalServiceTerm,
)
from projects_app.models import ProjectRegistration, ProjectRegistrationProduct
from users_app.models import Employee
from yandexdisk_app.models import YandexDiskAccount

from contracts_app.docx_processor import insert_floating_image_at_placeholder, process_template
from .document_generation import (
    build_proposal_docx_source_token,
    generate_and_store_proposal_pdf,
    get_proposal_docx_source_token_payload,
    store_generated_documents,
)
from .forms import ProposalDispatchForm, ProposalRegistrationForm, _proposal_variable_column_choices
from .forms import ProposalVariableForm
from .models import (
    ProposalAsset,
    ProposalCommercialOffer,
    ProposalLegalEntity,
    ProposalObject,
    ProposalRegistration,
    ProposalTemplate,
    ProposalVariable,
)
from .services import _load_proposal_pdf_attachment
from .views import PROPOSAL_NEXTCLOUD_TARGETS_SESSION_KEY
from .variable_resolver import resolve_variables

TEST_FACSIMILE_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMBAAZ/qh8AAAAASUVORK5CYII="
)


@override_settings(MEDIA_URL="/media/")
class ProposalDocumentGenerationTests(TestCase):
    def setUp(self):
        self.temp_media = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_media.cleanup)
        self.override = override_settings(MEDIA_ROOT=self.temp_media.name)
        self.override.enable()
        self.addCleanup(self.override.disable)

        self.user = get_user_model().objects.create_user(
            username="proposal-staff",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)

        self.country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        self.group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        self.country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        self.proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=self.group_member,
            type=self.product,
            name="Приморское",
            year=2026,
            status=ProposalRegistration.ProposalStatus.PRELIMINARY,
            customer='ООО "Приморское"',
            asset_owner='ООО "Приморское"',
            asset_owner_matches_customer=True,
            country=self.country,
            proposal_project_name='Due Diligence ООО "Приморское"',
            service_composition_mode="sections",
            service_sections_editor_state=[
                {
                    "code": "S-101",
                    "service_name": "Раздел 1",
                    "html": "<p><strong>Этап 1</strong> и описание</p><p><span style=\"color:#ff0000\">Красный текст</span></p>",
                    "plain_text": "Этап 1 и описание\n\nКрасный текст",
                },
                {
                    "code": "S-102",
                    "service_name": "Раздел 2",
                    "html": "<p class=\"ql-align-center\"><u>Этап 2</u></p>",
                    "plain_text": "Этап 2",
                },
            ],
            service_term_months="4.5",
            advance_percent="50",
            preliminary_report_percent="20",
            identifier="ОГРН",
            registration_number="1174910001683",
            proposal_workspace_disk_path="/Corporate Root/ТКП/2026/33330RU DD Приморское",
        )
        ServiceGoalReport.objects.create(
            product=self.product,
            service_goal="Проведение due diligence",
            service_goal_genitive="Проведения due diligence",
            report_title='Due Diligence ООО "Приморское"',
            position=1,
        )

        template_doc = Document()
        template_doc.add_paragraph("Заказчик: {{name}}")
        template_doc.add_paragraph("Страна: {{client_country_full_name}}")
        template_doc.add_paragraph("Страна legacy: {{country_full_name}}")
        template_doc.add_paragraph("Титул: {{client_owner_name}}")
        template_doc.add_paragraph("Краткое название: {{service_type_short}}")
        template_doc.add_paragraph("Цель в родительном: {{service_goal_genitive}}")
        template_doc.add_paragraph("Титул ТКП: {{tkp_preliminary}}")
        template_doc.add_paragraph("Предварительная оплата всего: {{preliminary_payment_percentage_full}}")
        template_doc.add_paragraph("Срок до Предварительного отчёта: {{preliminary_report_term_month}}")
        template_doc.add_paragraph("Состав услуг:")
        template_doc.add_paragraph("[[scope_of_work]]")
        template_doc.add_paragraph("Расчёт вознаграждения:")
        template_doc.add_paragraph("[[budget_table]]")
        template_doc.add_paragraph("Активы:")
        template_doc.add_paragraph("[[actives_name]]")
        buffer = BytesIO()
        template_doc.save(buffer)
        buffer.seek(0)

        self.template = ProposalTemplate.objects.create(
            group_member=self.group_member,
            product=self.product,
            sample_name="RU Шаблон ТКП_IMC_DD",
            version="2",
            file=SimpleUploadedFile(
                "proposal-template.docx",
                buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            position=1,
        )

        ProposalVariable.objects.create(
            key="{{name}}",
            description="Наименование Заказчика",
            source_section="proposals",
            source_table="registry",
            source_column="customer",
            position=1,
        )
        ProposalVariable.objects.create(
            key="{{client_country_full_name}}",
            description="Наименование страны Заказчика (полное)",
            is_computed=True,
            position=2,
        )
        ProposalVariable.objects.create(
            key="{{client_owner_name}}",
            description="Наименование заказчика и владельца активов на титуле",
            is_computed=True,
            position=3,
        )
        ProposalVariable.objects.create(
            key="{{service_type_short}}",
            description="Наименование ТКП (проекта) (краткое)",
            is_computed=True,
            position=4,
        )
        ProposalVariable.objects.create(
            key="{{service_goal_genitive}}",
            description="Цель оказания услуг в родительном падеже",
            is_computed=True,
            position=5,
        )
        ProposalVariable.objects.create(
            key="{{tkp_preliminary}}",
            description="Предварительное ТКП на титуле",
            is_computed=True,
            position=6,
        )
        ProposalVariable.objects.create(
            key="{{preliminary_payment_percentage_full}}",
            description="Размер оплаты Предварительного отчёта в процентах (с учетом предоплаты)",
            is_computed=True,
            position=7,
        )
        ProposalVariable.objects.create(
            key="{{preliminary_report_term_month}}",
            description="Срок оказания услуг от получения исходных данных до сдачи Предварительного отчёта в месяцах",
            is_computed=True,
            position=8,
        )
        ProposalVariable.objects.create(
            key="[[actives_name]]",
            description="Список наименований активов",
            is_computed=True,
            position=9,
        )
        ProposalVariable.objects.create(
            key="[[scope_of_work]]",
            description="Состав оказываемых услуг",
            is_computed=True,
            position=10,
        )
        ProposalVariable.objects.create(
            key="[[budget_table]]",
            description="Расчёт вознаграждения за оказание услуг",
            is_computed=True,
            position=11,
        )
        ProposalAsset.objects.create(
            proposal=self.proposal,
            short_name="Месторождение Приморское",
            position=1,
        )
        ProposalAsset.objects.create(
            proposal=self.proposal,
            short_name="Фабрика Приморская",
            position=2,
        )
        ProposalAsset.objects.create(
            proposal=self.proposal,
            short_name="Месторождение Приморское",
            position=3,
        )
        ProposalCommercialOffer.objects.create(
            proposal=self.proposal,
            specialist="Иванов И.И.",
            job_title="Геолог",
            professional_status="Партнер",
            service_name="Раздел 1",
            rate_eur_per_day="1200",
            asset_day_counts=[2, 3, 1],
            total_eur_without_vat="7200",
            position=1,
        )
        ProposalCommercialOffer.objects.create(
            proposal=self.proposal,
            specialist="Петров П.П.",
            job_title="Инженер",
            professional_status="Эксперт",
            service_name="Раздел 2",
            rate_eur_per_day="800",
            asset_day_counts=[1, 1, 2],
            total_eur_without_vat="3200",
            position=2,
        )
        ProposalCommercialOffer.objects.create(
            proposal=self.proposal,
            specialist="",
            job_title="",
            professional_status="",
            service_name="Командировочные расходы",
            rate_eur_per_day=None,
            asset_day_counts=["", "", ""],
            total_eur_without_vat="500",
            position=3,
        )
        self.proposal.commercial_totals_json = {
            "exchange_rate": "96.5",
            "discount_percent": "7.5",
            "contract_total": "900000",
            "contract_total_auto": "900000",
            "rub_total_service_text": "Курс евро Банка России на текущую дату:",
            "discounted_total_service_text": "Размер скидки:",
        }
        self.proposal.save(update_fields=["commercial_totals_json"])

    @patch("ai_app.proposals_app.document_generation._get_cloud_upload_user")
    @patch("ai_app.proposals_app.document_generation.cloud_upload_file", return_value=True)
    def test_create_documents_generates_docx_and_uploads_it_to_workspace_folder(
        self,
        mocked_cloud_upload,
        mocked_get_cloud_upload_user,
    ):
        mocked_get_cloud_upload_user.return_value = self.user
        response = self.client.post(
            reverse("proposal_dispatch_create_documents"),
            {"proposal_ids[]": [self.proposal.pk]},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["generated"], 1)
        self.assertEqual(payload["updates"][0]["id"], self.proposal.pk)
        self.assertTrue(payload["updates"][0]["docx_file_name"].endswith(".docx"))

        self.proposal.refresh_from_db()
        self.assertTrue(self.proposal.docx_file_name.endswith(".docx"))
        expected_docx_path = (
            f"{self.proposal.proposal_workspace_disk_path}/{self.proposal.docx_file_name}"
        )
        self.assertEqual(self.proposal.docx_file_link, expected_docx_path)
        self.assertEqual(self.proposal.pdf_file_name, "")
        self.assertEqual(self.proposal.pdf_file_link, "")

        mocked_cloud_upload.assert_called_once()
        self.assertEqual(mocked_cloud_upload.call_args.args[0], self.user)
        self.assertEqual(mocked_cloud_upload.call_args.args[1], expected_docx_path)
        generated_doc = Document(BytesIO(mocked_cloud_upload.call_args.args[2]))
        full_text = "\n".join(paragraph.text for paragraph in generated_doc.paragraphs)
        self.assertIn('Заказчик: ООО "Приморское"', full_text)
        self.assertIn("Страна: Российская Федерация", full_text)
        self.assertIn("Страна legacy: Российская Федерация", full_text)
        self.assertIn('Титул: ООО "Приморское"', full_text)
        self.assertIn("Краткое название: Due Diligence", full_text)
        self.assertIn("Цель в родительном: Проведения due diligence", full_text)
        self.assertIn("Титул ТКП: (предварительное)", full_text)
        self.assertIn("Предварительная оплата всего: 70%", full_text)
        self.assertIn("Срок до Предварительного отчёта: 4,5 месяца", full_text)
        self.assertIn("Состав услуг:", full_text)
        self.assertIn("Этап 1 и описание", full_text)
        self.assertIn("Красный текст", full_text)
        self.assertIn("Этап 2", full_text)
        self.assertIn("Месторождение Приморское", full_text)
        self.assertIn("Фабрика Приморская", full_text)
        scope_paragraph = next(paragraph for paragraph in generated_doc.paragraphs if paragraph.text == "Этап 1 и описание")
        self.assertTrue(any(run.bold for run in scope_paragraph.runs if "Этап 1" in run.text))
        red_paragraph = next(paragraph for paragraph in generated_doc.paragraphs if paragraph.text == "Красный текст")
        self.assertTrue(any((run.font.color.rgb and str(run.font.color.rgb) == "FF0000") for run in red_paragraph.runs))
        centered_paragraph = next(paragraph for paragraph in generated_doc.paragraphs if paragraph.text == "Этап 2")
        self.assertTrue(any(run.underline for run in centered_paragraph.runs))
        asset_paragraphs = [
            paragraph
            for paragraph in generated_doc.paragraphs
            if paragraph.text in {"Месторождение Приморское", "Фабрика Приморская"}
        ]
        self.assertEqual(len(asset_paragraphs), 2)
        for paragraph in asset_paragraphs:
            self.assertTrue(
                "w:numPr" in paragraph._element.xml or "w:pStyle" in paragraph._element.xml
            )
        budget_table = next(
            table
            for table in generated_doc.tables
            if any("Специалист" in cell.text for cell in table.rows[0].cells)
        )
        budget_rows = [[cell.text.strip() for cell in row.cells] for row in budget_table.rows]
        self.assertIn("Месторождение Приморское", budget_rows[0])
        self.assertIn("Фабрика Приморская", budget_rows[0])
        self.assertIn("Ставка,\n€/дн", budget_rows[0])
        self.assertIn("Кол-во\nдней", budget_rows[0])
        self.assertIn("Итого,\n€ без НДС", budget_rows[0])
        self.assertIn("Иванов И.И.", budget_rows[1])
        self.assertIn("Геолог, Партнер", budget_rows[1])
        self.assertIn("6", budget_rows[1])
        self.assertIn("7\u00a0200,00", budget_rows[1])
        self.assertIn("ИТОГО, по расчёту", budget_rows[3][0])
        self.assertIn("10\u00a0400,00", budget_rows[3][-1])
        self.assertIn("Командировочные расходы, евро", budget_rows[4][0])
        self.assertIn("ИТОГО, евро с командировочными по расчёту", budget_rows[5][0])
        self.assertIn("10\u00a0900,00", budget_rows[5][-1])
        self.assertIn("ИТОГО, рубли без НДС", budget_rows[6][0])
        self.assertIn("96,5", budget_rows[6][2])
        self.assertIn("1\u00a0051\u00a0850,00", budget_rows[6][-1])
        self.assertIn("ИТОГО, рубли без НДС с учетом скидки", budget_rows[7][0])
        self.assertIn("7,5%", budget_rows[7][2])
        self.assertIn("972\u00a0961,25", budget_rows[7][-1])
        self.assertIn("ИТОГО в договор, рубли без НДС с учётом доп. скидки", budget_rows[8][0])
        self.assertIn("900\u00a0000,00", budget_rows[8][-1])
        self.assertIn('w:tblLayout w:type="fixed"', budget_table._tbl.xml)
        self.assertIn('w:tblW w:type="pct" w:w="5000"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="775"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="2000"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="350"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="358"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="300"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="500"', budget_table._tbl.xml)
        empty_fixed_cell = budget_table.rows[6].cells[3]
        self.assertTrue(empty_fixed_cell.paragraphs)
        self.assertTrue(empty_fixed_cell.paragraphs[0].runs)
        self.assertEqual(empty_fixed_cell.paragraphs[0].runs[0].font.size.pt, 7)
        self.assertIn('w:sz w:val="14"', empty_fixed_cell.paragraphs[0]._element.xml)
        budget_run_sizes = [
            run.font.size.pt
            for row in budget_table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        self.assertTrue(budget_run_sizes)
        self.assertTrue(all(size == 7 for size in budget_run_sizes))

    @patch("ai_app.proposals_app.document_generation._get_cloud_upload_user")
    @patch("ai_app.proposals_app.document_generation.cloud_upload_file", return_value=True)
    def test_store_generated_documents_uses_connected_cloud_user(
        self,
        mocked_cloud_upload,
        mocked_get_cloud_upload_user,
    ):
        connected_user = get_user_model().objects.create_user(
            username="connected-cloud-user",
            password="secret",
            is_staff=True,
        )
        mocked_get_cloud_upload_user.return_value = connected_user

        store_generated_documents(self.user, self.proposal, b"docx-bytes")

        mocked_cloud_upload.assert_called_once()
        self.assertEqual(mocked_cloud_upload.call_args.args[0], connected_user)

    def test_scope_of_work_preserves_rich_list_markers(self):
        template_doc = Document()
        template_doc.add_paragraph("[[scope_of_work]]")
        buffer = BytesIO()
        template_doc.save(buffer)

        generated_bytes = process_template(
            buffer.getvalue(),
            {},
            list_replacements={
                "[[scope_of_work]]": [
                    {"html": "<ul><li>Первый пункт</li><li>Второй пункт</li></ul>"}
                ]
            },
            default_language_code="ru-RU",
        )

        generated_doc = Document(BytesIO(generated_bytes))
        scope_paragraphs = [paragraph for paragraph in generated_doc.paragraphs if paragraph.text in {"Первый пункт", "Второй пункт"}]

        self.assertEqual(len(scope_paragraphs), 2)
        for paragraph in scope_paragraphs:
            self.assertTrue(
                "w:numPr" in paragraph._element.xml or "w:pStyle" in paragraph._element.xml
            )
            self.assertIn('w:lang w:val="ru-RU"', paragraph._element.xml)

    def test_scope_of_work_plain_text_fallback_remains_plain_paragraphs(self):
        proposal = ProposalRegistration(
            service_composition_mode="sections",
            service_composition="Первый абзац\n\nВторой абзац",
        )
        template_doc = Document()
        template_doc.add_paragraph("[[scope_of_work]]")
        buffer = BytesIO()
        template_doc.save(buffer)

        replacements, lists, _ = resolve_variables(
            proposal,
            [ProposalVariable(key="[[scope_of_work]]", is_computed=True)],
        )

        generated_bytes = process_template(
            buffer.getvalue(),
            replacements,
            list_replacements=lists,
            default_language_code="ru-RU",
        )

        generated_doc = Document(BytesIO(generated_bytes))
        scope_paragraphs = [
            paragraph
            for paragraph in generated_doc.paragraphs
            if paragraph.text in {"Первый абзац", "Второй абзац"}
        ]

        self.assertEqual(len(scope_paragraphs), 2)
        for paragraph in scope_paragraphs:
            self.assertNotIn("w:numPr", paragraph._element.xml)
            self.assertNotIn("w:pStyle", paragraph._element.xml)
            self.assertIn('w:lang w:val="ru-RU"', paragraph._element.xml)

    def test_process_template_applies_default_language_to_scalar_replacements(self):
        template_doc = Document()
        template_doc.add_paragraph("Заказчик: {{name}}")
        buffer = BytesIO()
        template_doc.save(buffer)

        generated_bytes = process_template(
            buffer.getvalue(),
            {"{{name}}": 'ООО "Приморское"'},
            default_language_code="ru-RU",
        )

        generated_doc = Document(BytesIO(generated_bytes))
        paragraph = next(paragraph for paragraph in generated_doc.paragraphs if 'ООО "Приморское"' in paragraph.text)

        self.assertIn('w:lang w:val="ru-RU"', paragraph._element.xml)

    def test_process_template_preserves_existing_language_when_default_not_provided(self):
        template_doc = Document()
        paragraph = template_doc.add_paragraph()
        run = paragraph.add_run("Заказчик: {{name}}")
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run._element.insert(0, r_pr)
        lang = OxmlElement("w:lang")
        lang.set(qn("w:val"), "en-US")
        lang.set(qn("w:bidi"), "en-US")
        lang.set(qn("w:eastAsia"), "en-US")
        r_pr.append(lang)

        buffer = BytesIO()
        template_doc.save(buffer)

        generated_bytes = process_template(
            buffer.getvalue(),
            {"{{name}}": 'ООО "Приморское"'},
        )

        generated_doc = Document(BytesIO(generated_bytes))
        generated_paragraph = next(
            item for item in generated_doc.paragraphs if 'ООО "Приморское"' in item.text
        )

        self.assertIn('w:lang w:val="en-US"', generated_paragraph._element.xml)

    def test_resolve_budget_table_returns_table_spec(self):
        _, lists, tables = resolve_variables(
            self.proposal,
            [ProposalVariable(key="[[budget_table]]", is_computed=True)],
        )

        self.assertEqual(lists, {})
        self.assertIn("[[budget_table]]", tables)
        table_spec = tables["[[budget_table]]"]
        self.assertEqual(table_spec["font_size_pt"], 7)
        self.assertEqual(table_spec["style"], "Table Grid")
        self.assertEqual(len(table_spec["column_widths_pct"]), 8)
        self.assertEqual(table_spec["column_widths_pct"][:3], [15.5, 40.0, 7.0])
        self.assertEqual(table_spec["column_widths_pct"][-2:], [6.0, 10.0])
        self.assertTrue(all(width == table_spec["column_widths_pct"][3] for width in table_spec["column_widths_pct"][3:6]))
        self.assertEqual(table_spec["rows"][0][0]["text"], "Специалист")
        self.assertEqual(table_spec["rows"][0][2]["text"], "Ставка,\n€/дн")
        self.assertEqual(table_spec["rows"][0][-2]["text"], "Кол-во\nдней")
        self.assertEqual(table_spec["rows"][0][-1]["text"], "Итого,\n€ без НДС")
        self.assertEqual(table_spec["rows"][0][3]["text"], "Месторождение Приморское")
        self.assertTrue(table_spec["rows"][0][1]["no_wrap"])
        self.assertEqual(table_spec["rows"][1][0]["text"], "Иванов И.И.")
        self.assertEqual(table_spec["rows"][1][1]["text"], "Геолог, Партнер")
        self.assertTrue(table_spec["rows"][1][1]["no_wrap"])
        self.assertEqual(table_spec["rows"][1][-2]["text"], "6")
        self.assertEqual(table_spec["rows"][8][0]["text"], "ИТОГО в договор, рубли без НДС с учётом доп. скидки")

    def test_resolve_budget_table_omits_asset_column_for_single_asset(self):
        proposal = ProposalRegistration.objects.create(
            number=4444,
            group_member=self.group_member,
            type=self.product,
            name="Один актив",
            year=2026,
            status=ProposalRegistration.ProposalStatus.PRELIMINARY,
            customer='ООО "Приморское"',
        )
        ProposalAsset.objects.create(
            proposal=proposal,
            short_name="Единственный актив",
            position=1,
        )
        ProposalCommercialOffer.objects.create(
            proposal=proposal,
            specialist="Иванов И.И.",
            job_title="Геолог",
            professional_status="Партнер",
            service_name="Раздел 1",
            rate_eur_per_day="1200",
            asset_day_counts=[2],
            total_eur_without_vat="2400",
            position=1,
        )

        _, _, tables = resolve_variables(
            proposal,
            [ProposalVariable(key="[[budget_table]]", is_computed=True)],
        )

        header_row = tables["[[budget_table]]"]["rows"][0]
        header_texts = [cell["text"] for cell in header_row]
        self.assertEqual(
            header_texts,
            [
                "Специалист",
                "Должность/направление",
                "Ставка,\n€/дн",
                "Кол-во\nдней",
                "Итого,\n€ без НДС",
            ],
        )
        self.assertEqual(tables["[[budget_table]]"]["font_size_pt"], 8)
        self.assertEqual(tables["[[budget_table]]"]["column_widths_pct"], [18, 46, 12, 12, 12])
        data_row = tables["[[budget_table]]"]["rows"][1]
        self.assertEqual(len(data_row), 5)
        self.assertEqual(data_row[3]["text"], "2")

    def test_process_template_inserts_budget_table(self):
        template_doc = Document()
        template_doc.add_paragraph("[[budget_table]]")
        buffer = BytesIO()
        template_doc.save(buffer)

        replacements, lists, tables = resolve_variables(
            self.proposal,
            [ProposalVariable(key="[[budget_table]]", is_computed=True)],
        )

        generated_bytes = process_template(
            buffer.getvalue(),
            replacements,
            table_replacements=tables,
            list_replacements=lists,
            default_language_code="ru-RU",
        )

        generated_doc = Document(BytesIO(generated_bytes))
        self.assertEqual(len(generated_doc.tables), 1)
        budget_table = generated_doc.tables[0]
        budget_rows = [[cell.text.strip() for cell in row.cells] for row in budget_table.rows]

        self.assertIn("Месторождение Приморское", budget_rows[0])
        self.assertIn("Фабрика Приморская", budget_rows[0])
        self.assertIn("Ставка,\n€/дн", budget_rows[0])
        self.assertIn("Кол-во\nдней", budget_rows[0])
        self.assertIn("Итого,\n€ без НДС", budget_rows[0])
        self.assertIn("Иванов И.И.", budget_rows[1])
        self.assertIn("Геолог, Партнер", budget_rows[1])
        self.assertIn("6", budget_rows[1])
        self.assertIn("7\u00a0200,00", budget_rows[1])
        self.assertIn("ИТОГО, по расчёту", budget_rows[3][0])
        self.assertIn("10\u00a0400,00", budget_rows[3][-1])
        self.assertIn("Командировочные расходы, евро", budget_rows[4][0])
        self.assertIn("ИТОГО, евро с командировочными по расчёту", budget_rows[5][0])
        self.assertIn("10\u00a0900,00", budget_rows[5][-1])
        self.assertIn("ИТОГО, рубли без НДС", budget_rows[6][0])
        self.assertIn("96,5", budget_rows[6][2])
        self.assertIn("1\u00a0051\u00a0850,00", budget_rows[6][-1])
        self.assertIn("ИТОГО, рубли без НДС с учетом скидки", budget_rows[7][0])
        self.assertIn("7,5%", budget_rows[7][2])
        self.assertIn("972\u00a0961,25", budget_rows[7][-1])
        self.assertIn("ИТОГО в договор, рубли без НДС с учётом доп. скидки", budget_rows[8][0])
        self.assertIn("900\u00a0000,00", budget_rows[8][-1])
        self.assertIn('w:tblLayout w:type="fixed"', budget_table._tbl.xml)
        self.assertIn('w:tblW w:type="pct" w:w="5000"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="775"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="2000"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="350"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="358"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="300"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="500"', budget_table._tbl.xml)
        budget_run_sizes = [
            run.font.size.pt
            for row in budget_table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        self.assertTrue(budget_run_sizes)
        self.assertTrue(all(size == 7 for size in budget_run_sizes))

    def test_process_template_inserts_single_asset_budget_table_with_pct_widths(self):
        proposal = ProposalRegistration.objects.create(
            number=5555,
            group_member=self.group_member,
            type=self.product,
            name="Один актив",
            year=2026,
            status=ProposalRegistration.ProposalStatus.PRELIMINARY,
            customer='ООО "Приморское"',
        )
        ProposalAsset.objects.create(
            proposal=proposal,
            short_name="Единственный актив",
            position=1,
        )
        ProposalCommercialOffer.objects.create(
            proposal=proposal,
            specialist="Иванов И.И.",
            job_title="Геолог",
            professional_status="Партнер",
            service_name="Раздел 1",
            rate_eur_per_day="1200",
            asset_day_counts=[2],
            total_eur_without_vat="2400",
            position=1,
        )

        template_doc = Document()
        template_doc.add_paragraph("[[budget_table]]")
        buffer = BytesIO()
        template_doc.save(buffer)

        replacements, lists, tables = resolve_variables(
            proposal,
            [ProposalVariable(key="[[budget_table]]", is_computed=True)],
        )

        generated_bytes = process_template(
            buffer.getvalue(),
            replacements,
            table_replacements=tables,
            list_replacements=lists,
            default_language_code="ru-RU",
        )

        generated_doc = Document(BytesIO(generated_bytes))
        budget_table = generated_doc.tables[0]
        self.assertIn('w:tblLayout w:type="fixed"', budget_table._tbl.xml)
        self.assertIn('w:tblW w:type="pct" w:w="5000"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="900"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="2300"', budget_table._tbl.xml)
        self.assertIn('w:tcW w:type="pct" w:w="600"', budget_table._tbl.xml)
        budget_run_sizes = [
            run.font.size.pt
            for row in budget_table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        self.assertTrue(budget_run_sizes)
        self.assertTrue(all(size == 8 for size in budget_run_sizes))

    def test_insert_floating_image_at_placeholder_replaces_marker_with_anchor(self):
        template_doc = Document()
        template_doc.add_paragraph("Подпись [[facsimile]]")
        buffer = BytesIO()
        template_doc.save(buffer)

        generated_bytes = insert_floating_image_at_placeholder(
            buffer.getvalue(),
            TEST_FACSIMILE_PNG_BYTES,
        )

        generated_doc = Document(BytesIO(generated_bytes))
        paragraph = generated_doc.paragraphs[0]
        self.assertEqual(paragraph.text, "Подпись ")
        self.assertIn("wp:anchor", paragraph._p.xml)
        self.assertIn('behindDoc="1"', paragraph._p.xml)
        self.assertIn('wp:positionH relativeFrom="page"', paragraph._p.xml)
        self.assertIn("<wp:align>center</wp:align>", paragraph._p.xml)
        self.assertNotIn("[[facsimile]]", paragraph._p.xml)

    def test_insert_floating_image_at_placeholder_rejects_truncated_image(self):
        template_doc = Document()
        template_doc.add_paragraph("[[facsimile]]")
        buffer = BytesIO()
        template_doc.save(buffer)

        with self.assertRaisesMessage(
            RuntimeError,
            "Факсимиле должно быть изображением в поддерживаемом формате Word",
        ):
            insert_floating_image_at_placeholder(
                buffer.getvalue(),
                TEST_FACSIMILE_PNG_BYTES[:16],
            )

    @patch("ai_app.proposals_app.document_generation.get_any_connected_service_user")
    @patch("ai_app.proposals_app.document_generation.is_nextcloud_primary", return_value=False)
    @patch("ai_app.proposals_app.document_generation.cloud_upload_file", return_value=True)
    def test_store_generated_documents_prefers_request_user_yandex_account(
        self,
        mocked_cloud_upload,
        _mocked_is_nextcloud_primary,
        mocked_get_any_connected_service_user,
    ):
        YandexDiskAccount.objects.create(user=self.user, access_token="request-user-token")
        fallback_user = get_user_model().objects.create_user(
            username="fallback-cloud-user",
            password="secret",
            is_staff=True,
        )
        mocked_get_any_connected_service_user.return_value = fallback_user

        store_generated_documents(self.user, self.proposal, b"docx-bytes")

        mocked_get_any_connected_service_user.assert_not_called()
        mocked_cloud_upload.assert_called_once()
        self.assertEqual(mocked_cloud_upload.call_args.args[0], self.user)

    def test_create_documents_returns_error_when_workspace_folder_is_missing(self):
        self.proposal.proposal_workspace_disk_path = ""
        self.proposal.save(update_fields=["proposal_workspace_disk_path"])

        response = self.client.post(
            reverse("proposal_dispatch_create_documents"),
            {"proposal_ids[]": [self.proposal.pk]},
        )

        self.assertEqual(response.status_code, 400)
        payload = response.json()
        self.assertFalse(payload["ok"])
        self.assertIn("рабочая папка", payload["error"])

    @override_settings(
        ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com",
        ONLYOFFICE_VERIFY_SSL=True,
    )
    @patch("proposals_app.document_generation.cloud_upload_file", return_value=True)
    @patch("proposals_app.document_generation.requests.get")
    @patch("proposals_app.document_generation.requests.post")
    @patch("proposals_app.document_generation._get_cloud_upload_user")
    def test_generate_and_store_pdf_uses_onlyoffice_and_uploads_result_to_workspace(
        self,
        mocked_get_cloud_upload_user,
        mocked_requests_post,
        mocked_requests_get,
        mocked_cloud_upload_file,
    ):
        mocked_get_cloud_upload_user.return_value = self.user
        mocked_requests_post.return_value = Mock(
            json=Mock(
                return_value={
                    "endConvert": True,
                    "fileType": "pdf",
                    "fileUrl": "https://docs.example.com/cache/generated.pdf",
                    "percent": 100,
                }
            ),
            raise_for_status=Mock(),
        )
        mocked_requests_get.return_value = Mock(
            content=b"pdf-bytes",
            raise_for_status=Mock(),
        )
        self.proposal.docx_file_name = "existing-offer.docx"
        self.proposal.docx_file_link = (
            "/Corporate Root/ТКП/2026/33330RU DD Приморское/existing-offer.docx"
        )
        self.proposal.save(update_fields=["docx_file_name", "docx_file_link"])

        result = generate_and_store_proposal_pdf(
            self.user,
            self.proposal,
            source_url="https://app.example.com/proposals/docx-source/1/?token=abc",
        )

        self.assertEqual(result["pdf_name"], "existing-offer.pdf")
        self.assertEqual(
            result["pdf_path"],
            "/Corporate Root/ТКП/2026/33330RU DD Приморское/existing-offer.pdf",
        )
        mocked_requests_post.assert_called_once()
        self.assertTrue(
            mocked_requests_post.call_args.args[0].startswith("https://docs.example.com/converter?shardkey=")
        )
        self.assertEqual(
            mocked_requests_post.call_args.kwargs["json"]["url"],
            "https://app.example.com/proposals/docx-source/1/?token=abc",
        )
        self.assertEqual(mocked_requests_post.call_args.kwargs["json"]["outputtype"], "pdf")
        mocked_requests_get.assert_called_once_with(
            "https://docs.example.com/cache/generated.pdf",
            timeout=120,
            verify=True,
        )
        mocked_cloud_upload_file.assert_called_once_with(
            self.user,
            "/Corporate Root/ТКП/2026/33330RU DD Приморское/existing-offer.pdf",
            b"pdf-bytes",
        )


@override_settings(ROOT_URLCONF="proposals_app.urls")
class ProposalDispatchSendTests(TestCase):
    def setUp(self):
        self.temp_media = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_media.cleanup)
        self.media_override = override_settings(MEDIA_ROOT=self.temp_media.name, MEDIA_URL="/media/")
        self.media_override.enable()
        self.addCleanup(self.media_override.disable)

        self.user = get_user_model().objects.create_user(
            username="proposal-dispatch-staff",
            email="staff@example.com",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)

        self.group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        self.country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
        )
        LetterTemplate.objects.update_or_create(
            template_type="proposal_sending",
            user=None,
            defaults={
                "subject_template": "Технико-коммерческое предложение IMC Montan {{tkp_id}}",
                "body_html": "<p>Направляем проект ТКП {{tkp_id}}</p>",
                "is_default": True,
            },
        )
        self.successful_pdf_name = "ТКП_333300RU_DD_Приморское.pdf"
        self.successful_pdf_link = f"/media/proposal_documents/2026/333300RU DD Приморское/{self.successful_pdf_name}"
        self.successful_pdf_bytes = b"%PDF-1.4 successful proposal pdf"
        successful_pdf_path = Path(self.temp_media.name) / self.successful_pdf_link.removeprefix("/media/")
        successful_pdf_path.parent.mkdir(parents=True, exist_ok=True)
        successful_pdf_path.write_bytes(self.successful_pdf_bytes)

        self.failed_pdf_name = "ТКП_333400RU_DD_Балтика.pdf"
        self.failed_pdf_link = f"/media/proposal_documents/2026/333400RU DD Балтика/{self.failed_pdf_name}"
        self.failed_pdf_bytes = b"%PDF-1.4 failed proposal pdf"
        failed_pdf_path = Path(self.temp_media.name) / self.failed_pdf_link.removeprefix("/media/")
        failed_pdf_path.parent.mkdir(parents=True, exist_ok=True)
        failed_pdf_path.write_bytes(self.failed_pdf_bytes)

        self.successful_proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=self.group_member,
            type=self.product,
            name="Приморское",
            year=2026,
            contact_email="recipient@example.com",
            docx_file_name="ТКП_333300RU_DD_Приморское.docx",
            docx_file_link="/Corporate Root/ТКП/2026/333300RU DD Приморское/ТКП_333300RU_DD_Приморское.docx",
            pdf_file_name=self.successful_pdf_name,
            pdf_file_link=self.successful_pdf_link,
            proposal_workspace_disk_path="/Corporate Root/ТКП/2026/333300RU DD Приморское",
        )
        self.failed_proposal = ProposalRegistration.objects.create(
            number=3334,
            group_member=self.group_member,
            type=self.product,
            name="Балтика",
            year=2026,
            contact_email="",
            docx_file_name="ТКП_333400RU_DD_Балтика.docx",
            docx_file_link="/Corporate Root/ТКП/2026/333400RU DD Балтика/ТКП_333400RU_DD_Балтика.docx",
            pdf_file_name=self.failed_pdf_name,
            pdf_file_link=self.failed_pdf_link,
            proposal_workspace_disk_path="/Corporate Root/ТКП/2026/333400RU DD Балтика",
        )

    def test_dispatch_transfer_to_contract_creates_project_registration_and_completes_proposal(self):
        self.successful_proposal.customer = 'ООО "Приморское"'
        self.successful_proposal.country = self.country
        self.successful_proposal.identifier = "ОГРН"
        self.successful_proposal.registration_number = "1174910001683"
        self.successful_proposal.registration_date = date(2017, 11, 1)
        self.successful_proposal.save(
            update_fields=[
                "customer",
                "country",
                "identifier",
                "registration_number",
                "registration_date",
            ]
        )

        response = self.client.post(
            reverse("proposal_dispatch_transfer_to_contract"),
            {"proposal_ids[]": [self.successful_proposal.pk]},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["created"], 1)

        project = ProjectRegistration.objects.get(
            number=self.successful_proposal.number,
            group_member=self.successful_proposal.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            agreement_number=f"IMCM/{self.successful_proposal.number}",
        )
        self.assertEqual(project.type, self.successful_proposal.type)
        self.assertEqual(project.type_short_display, self.successful_proposal.type.short_name)
        self.assertEqual(
            list(
                ProjectRegistrationProduct.objects
                .filter(registration=project)
                .order_by("rank")
                .values_list("product_id", flat=True)
            ),
            [self.successful_proposal.type_id],
        )
        self.assertEqual(project.name, self.successful_proposal.name)
        self.assertEqual(project.year, self.successful_proposal.year)
        self.assertEqual(project.customer, self.successful_proposal.customer)
        self.assertEqual(project.country, self.successful_proposal.country)
        self.assertEqual(project.identifier, self.successful_proposal.identifier)
        self.assertEqual(project.registration_number, self.successful_proposal.registration_number)
        self.assertEqual(project.registration_date, self.successful_proposal.registration_date)
        self.assertEqual(BusinessEntityRecord.objects.count(), 0)

        self.successful_proposal.refresh_from_db()
        self.assertEqual(
            self.successful_proposal.status,
            ProposalRegistration.ProposalStatus.COMPLETED,
        )
        self.assertTrue(self.successful_proposal.transfer_to_contract_date)

    def test_dispatch_transfer_to_contract_reuses_existing_main_contract_for_ru_group(self):
        existing_project = ProjectRegistration.objects.create(
            number=self.successful_proposal.number,
            group_member=self.successful_proposal.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Уже существует",
            year=2026,
            position=1,
        )

        response = self.client.post(
            reverse("proposal_dispatch_transfer_to_contract"),
            {"proposal_ids[]": [self.successful_proposal.pk]},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["created"], 0)
        self.assertEqual(payload["existing"], 1)
        self.assertEqual(
            ProjectRegistration.objects.filter(
                number=self.successful_proposal.number,
                group_member=self.successful_proposal.group_member,
                agreement_type=ProjectRegistration.AgreementType.MAIN,
            ).count(),
            1,
        )
        self.assertEqual(
            ProjectRegistration.objects.get(pk=existing_project.pk).agreement_number,
            f"IMCM/{self.successful_proposal.number}",
        )

    @override_settings(PROPOSAL_SYSTEM_FROM_EMAIL="ai@imcmontanai.ru")
    @patch("proposals_app.services.send_notification_email")
    def test_dispatch_send_updates_only_successfully_sent_rows(self, mocked_send_notification_email):
        mocked_send_notification_email.return_value = {
            "recipient_email": "recipient@example.com",
            "subject": "ignored",
            "is_html": True,
        }

        response = self.client.post(
            reverse("proposal_dispatch_send"),
            {
                "proposal_ids[]": [self.successful_proposal.pk, self.failed_proposal.pk],
                "delivery_channels[]": ["system_email"],
                "sent_at": "2026-04-04T12:30",
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["updated"], 1)
        self.assertEqual(payload["email_delivery"]["sent"], 1)
        self.assertEqual(payload["email_delivery"]["failed"], 1)
        self.assertEqual(payload["proposal_ids"], [self.successful_proposal.pk])
        self.assertEqual(payload["status"], ProposalRegistration.ProposalStatus.SENT)
        self.assertEqual(payload["updates"][0]["id"], self.successful_proposal.pk)
        self.assertEqual(payload["updates"][0]["status"], ProposalRegistration.ProposalStatus.SENT)
        self.assertEqual(mocked_send_notification_email.call_count, 1)

        call_kwargs = mocked_send_notification_email.call_args.kwargs
        expected_tkp_id = " ".join(
            [
                self.successful_proposal.short_uid,
                self.product.short_name,
                self.successful_proposal.name,
            ]
        )
        self.assertEqual(
            call_kwargs["subject"],
            f"Технико-коммерческое предложение IMC Montan {expected_tkp_id}",
        )
        self.assertEqual(
            call_kwargs["content"],
            f"<p>Направляем проект ТКП {expected_tkp_id}</p>",
        )
        self.assertEqual(call_kwargs["from_email"], "ai@imcmontanai.ru")
        self.assertEqual(
            call_kwargs["attachments"],
            [(self.successful_pdf_name, self.successful_pdf_bytes, "application/pdf")],
        )

        self.successful_proposal.refresh_from_db()
        self.failed_proposal.refresh_from_db()
        self.assertEqual(self.successful_proposal.sent_date, "04.04.2026 12:30")
        self.assertEqual(self.successful_proposal.status, ProposalRegistration.ProposalStatus.SENT)
        self.assertEqual(self.failed_proposal.sent_date, "")
        self.assertNotEqual(self.failed_proposal.status, ProposalRegistration.ProposalStatus.SENT)
        self.assertEqual(self.successful_proposal.pdf_file_name, self.successful_pdf_name)
        self.assertEqual(self.successful_proposal.pdf_file_link, self.successful_pdf_link)

    @override_settings(PROPOSAL_SYSTEM_FROM_EMAIL="ai@imcmontanai.ru")
    @patch("proposals_app.services.send_notification_email")
    def test_dispatch_send_preserves_completed_status(self, mocked_send_notification_email):
        mocked_send_notification_email.return_value = {
            "recipient_email": "recipient@example.com",
            "subject": "ignored",
            "is_html": True,
        }
        self.successful_proposal.status = ProposalRegistration.ProposalStatus.COMPLETED
        self.successful_proposal.transfer_to_contract_date = "08.04.2026 11:00"
        self.successful_proposal.save(update_fields=["status", "transfer_to_contract_date"])

        response = self.client.post(
            reverse("proposal_dispatch_send"),
            {
                "proposal_ids[]": [self.successful_proposal.pk],
                "delivery_channels[]": ["system_email"],
                "sent_at": "2026-04-04T12:30",
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["updates"][0]["status"], ProposalRegistration.ProposalStatus.COMPLETED)
        self.successful_proposal.refresh_from_db()
        self.assertEqual(self.successful_proposal.sent_date, "04.04.2026 12:30")
        self.assertEqual(self.successful_proposal.status, ProposalRegistration.ProposalStatus.COMPLETED)

    @patch("proposals_app.views.generate_and_store_proposal_pdf")
    @patch("proposals_app.services.send_notification_email")
    def test_dispatch_send_does_not_generate_pdf(self, mocked_send_notification_email, mocked_generate_pdf):
        mocked_send_notification_email.return_value = {
            "recipient_email": "recipient@example.com",
            "subject": "ignored",
            "is_html": True,
        }

        response = self.client.post(
            reverse("proposal_dispatch_send"),
            {
                "proposal_ids[]": [self.successful_proposal.pk],
                "delivery_channels[]": ["system_email"],
                "sent_at": "2026-04-04T12:30",
            },
        )

        self.assertEqual(response.status_code, 200)
        mocked_generate_pdf.assert_not_called()
        self.successful_proposal.refresh_from_db()
        self.assertEqual(self.successful_proposal.pdf_file_name, self.successful_pdf_name)
        self.assertEqual(self.successful_proposal.pdf_file_link, self.successful_pdf_link)


    @patch("proposals_app.services.send_notification_email")
    @patch(
        "proposals_app.services.get_user_notification_email_options",
        return_value={},
    )
    def test_dispatch_send_marks_row_sent_when_at_least_one_channel_succeeds(
        self,
        _mocked_email_options,
        mocked_send_notification_email,
    ):
        mocked_send_notification_email.return_value = {
            "recipient_email": "recipient@example.com",
            "subject": "ignored",
            "is_html": True,
        }

        response = self.client.post(
            reverse("proposal_dispatch_send"),
            {
                "proposal_ids[]": [self.successful_proposal.pk],
                "delivery_channels[]": ["system_email", "connected_email"],
                "sent_at": "2026-04-04T12:30",
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["updated"], 1)
        self.assertEqual(payload["email_delivery"]["sent"], 1)
        self.assertEqual(payload["email_delivery"]["failed"], 1)
        self.assertEqual(
            payload["email_delivery"]["channels"]["connected_email"]["errors"][0]["error"],
            "У отправителя не настроен активный внешний SMTP-аккаунт.",
        )
        self.assertEqual(mocked_send_notification_email.call_count, 1)
        self.successful_proposal.refresh_from_db()
        self.assertEqual(self.successful_proposal.sent_date, "04.04.2026 12:30")

    @override_settings(PROPOSAL_SYSTEM_FROM_EMAIL="ai@imcmontanai.ru")
    @patch("proposals_app.services.send_notification_email")
    def test_dispatch_send_uses_rendered_fallback_subject_when_template_subject_empty(self, mocked_send_notification_email):
        mocked_send_notification_email.return_value = {
            "recipient_email": "recipient@example.com",
            "subject": "ignored",
            "is_html": True,
        }
        LetterTemplate.objects.filter(
            template_type="proposal_sending",
            user__isnull=True,
        ).update(
            subject_template="",
            body_html="<p>Направляем проект ТКП {{tkp_id}}</p>",
        )

        response = self.client.post(
            reverse("proposal_dispatch_send"),
            {
                "proposal_ids[]": [self.successful_proposal.pk],
                "delivery_channels[]": ["system_email"],
                "sent_at": "2026-04-04T12:30",
            },
        )

        self.assertEqual(response.status_code, 200)
        call_kwargs = mocked_send_notification_email.call_args.kwargs
        expected_tkp_id = " ".join(
            [
                self.successful_proposal.short_uid,
                self.product.short_name,
                self.successful_proposal.name,
            ]
        )
        self.assertEqual(
            call_kwargs["subject"],
            f"Технико-коммерческое предложение IMC Montan {expected_tkp_id}",
        )
        self.assertEqual(call_kwargs["from_email"], "ai@imcmontanai.ru")

    @patch("proposals_app.services.send_notification_email")
    @patch("proposals_app.services.get_user_notification_email_options", return_value={})
    def test_dispatch_send_returns_error_when_connected_email_unavailable(
        self,
        _mocked_email_options,
        mocked_send_notification_email,
    ):
        response = self.client.post(
            reverse("proposal_dispatch_send"),
            {
                "proposal_ids[]": [self.successful_proposal.pk],
                "delivery_channels[]": ["connected_email"],
                "sent_at": "2026-04-04T12:30",
            },
        )

        self.assertEqual(response.status_code, 400)
        payload = response.json()
        self.assertFalse(payload["ok"])
        self.assertEqual(payload["error"], "Не удалось отправить ни одного письма.")
        self.assertEqual(payload["email_delivery"]["failed"], 1)
        self.assertEqual(
            payload["email_delivery"]["channels"]["connected_email"]["errors"][0]["error"],
            "У отправителя не настроен активный внешний SMTP-аккаунт.",
        )
        mocked_send_notification_email.assert_not_called()
        self.successful_proposal.refresh_from_db()
        self.assertEqual(self.successful_proposal.sent_date, "")

    @patch("proposals_app.services.send_notification_email")
    def test_dispatch_send_returns_error_when_pdf_attachment_missing(self, mocked_send_notification_email):
        self.successful_proposal.pdf_file_name = ""
        self.successful_proposal.pdf_file_link = ""
        self.successful_proposal.save(update_fields=["pdf_file_name", "pdf_file_link"])

        response = self.client.post(
            reverse("proposal_dispatch_send"),
            {
                "proposal_ids[]": [self.successful_proposal.pk],
                "delivery_channels[]": ["system_email"],
                "sent_at": "2026-04-04T12:30",
            },
        )

        self.assertEqual(response.status_code, 400)
        payload = response.json()
        self.assertFalse(payload["ok"])
        self.assertEqual(payload["error"], "Не удалось отправить ни одного письма.")
        self.assertEqual(payload["email_delivery"]["failed"], 1)
        self.assertEqual(
            payload["email_delivery"]["channels"]["system_email"]["errors"][0]["error"],
            "Для ТКП не указан PDF-файл для вложения.",
        )
        mocked_send_notification_email.assert_not_called()

    def test_load_proposal_pdf_attachment_rejects_media_path_traversal(self):
        escaped_file = Path(self.temp_media.name).parent / "outside-secret.pdf"
        escaped_file.write_bytes(b"secret")
        self.successful_proposal.pdf_file_link = "/media/../outside-secret.pdf"
        self.successful_proposal.save(update_fields=["pdf_file_link"])

        with self.assertRaisesMessage(
            EmailDeliveryError,
            "Некорректный путь к локальному PDF-файлу ТКП для вложения.",
        ):
            _load_proposal_pdf_attachment(self.successful_proposal, sender=self.user)

    @patch("proposals_app.services.requests.get")
    def test_load_proposal_pdf_attachment_rejects_external_http_link(self, mocked_requests_get):
        self.successful_proposal.pdf_file_link = "https://internal.example.local/secret.pdf"
        self.successful_proposal.save(update_fields=["pdf_file_link"])

        with self.assertRaisesMessage(
            EmailDeliveryError,
            "Прямые внешние ссылки на PDF-файл ТКП для вложения не поддерживаются.",
        ):
            _load_proposal_pdf_attachment(self.successful_proposal, sender=self.user)

        mocked_requests_get.assert_not_called()


@override_settings(ROOT_URLCONF="proposals_app.urls")
class ProposalDispatchSignTests(TestCase):
    def setUp(self):
        self.temp_media = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_media.cleanup)
        self.media_override = override_settings(MEDIA_ROOT=self.temp_media.name)
        self.media_override.enable()
        self.addCleanup(self.media_override.disable)

        self.user = get_user_model().objects.create_user(
            username="proposal-dispatch-pdf-staff",
            email="staff@example.com",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.employee = Employee.objects.create(user=self.user, patronymic="Иванович")
        self.person = PersonRecord.objects.create(
            last_name="Иванов",
            first_name="Иван",
            middle_name="Иванович",
            position=1,
        )
        self.employee.person_record = self.person
        self.employee.save(update_fields=["person_record"])
        self.expert_profile = ExpertProfile.objects.create(employee=self.employee, position=1)
        self.contract_citizenship = CitizenshipRecord.objects.create(
            person=self.person,
            identifier="Паспорт",
            number="123456",
            position=1,
        )
        self.contract_detail = ExpertContractDetails.objects.create(
            expert_profile=self.expert_profile,
            citizenship_record=self.contract_citizenship,
        )
        self.contract_detail.facsimile_file.save(
            "facsimile.png",
            ContentFile(TEST_FACSIMILE_PNG_BYTES),
            save=True,
        )

        self.group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        LetterTemplate.objects.update_or_create(
            template_type="proposal_sending",
            user=None,
            defaults={
                "subject_template": "Технико-коммерческое предложение IMC Montan {{tkp_id}}",
                "body_html": "<p>Направляем проект ТКП {{tkp_id}}</p>",
                "is_default": True,
            },
        )
        self.proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=self.group_member,
            type=self.product,
            name="Приморское",
            year=2026,
            contact_email="recipient@example.com",
            docx_file_name="existing-offer.docx",
            docx_file_link="/Corporate Root/ТКП/2026/333300RU DD Приморское/existing-offer.docx",
            proposal_workspace_disk_path="/Corporate Root/ТКП/2026/333300RU DD Приморское",
        )

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    @patch("proposals_app.views.store_existing_proposal_docx_bytes")
    @patch("proposals_app.views.load_existing_proposal_docx_bytes")
    @patch("proposals_app.views.generate_and_store_proposal_pdf")
    def test_dispatch_sign_generates_pdf_via_dedicated_endpoint(
        self,
        mocked_generate_and_store_proposal_pdf,
        mocked_load_existing_proposal_docx_bytes,
        mocked_store_existing_proposal_docx_bytes,
    ):
        template_doc = Document()
        template_doc.add_paragraph("[[facsimile]]")
        buffer = BytesIO()
        template_doc.save(buffer)
        mocked_load_existing_proposal_docx_bytes.return_value = buffer.getvalue()
        mocked_store_existing_proposal_docx_bytes.return_value = {
            "docx_name": "existing-offer.docx",
            "docx_path": "/Corporate Root/ТКП/2026/333300RU DD Приморское/existing-offer.docx",
            "output_dir": "/Corporate Root/ТКП/2026/333300RU DD Приморское",
        }
        mocked_generate_and_store_proposal_pdf.return_value = {
            "pdf_name": "existing-offer.pdf",
            "pdf_path": "/Corporate Root/ТКП/2026/333300RU DD Приморское/existing-offer.pdf",
        }

        response = self.client.post(
            reverse("proposal_dispatch_sign_documents"),
            {
                "proposal_ids[]": [self.proposal.pk],
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["generated"], 1)
        self.assertEqual(payload["updates"][0]["id"], self.proposal.pk)
        self.assertEqual(payload["updates"][0]["pdf_file_name"], "existing-offer.pdf")

        self.proposal.refresh_from_db()
        self.assertEqual(self.proposal.pdf_file_name, "existing-offer.pdf")
        self.assertEqual(
            self.proposal.pdf_file_link,
            "/Corporate Root/ТКП/2026/333300RU DD Приморское/existing-offer.pdf",
        )
        mocked_load_existing_proposal_docx_bytes.assert_called_once_with(self.user, self.proposal)
        mocked_store_existing_proposal_docx_bytes.assert_called_once()
        mocked_generate_and_store_proposal_pdf.assert_called_once()
        self.assertEqual(mocked_generate_and_store_proposal_pdf.call_args.args[0], self.user)
        self.assertEqual(mocked_generate_and_store_proposal_pdf.call_args.args[1], self.proposal)
        source_url = mocked_generate_and_store_proposal_pdf.call_args.kwargs["source_url"]
        self.assertTrue(
            source_url.startswith(
                f"http://testserver{reverse('proposal_onlyoffice_docx_source', args=[self.proposal.pk])}?token="
            )
        )
        token = parse_qs(urlparse(source_url).query)["token"][0]
        token_payload = get_proposal_docx_source_token_payload(self.proposal, token)
        self.assertIsNotNone(token_payload)
        self.assertEqual(token_payload["signer_user_id"], self.user.pk)

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    @patch("proposals_app.views.store_existing_proposal_docx_bytes")
    @patch("proposals_app.views.load_existing_proposal_docx_bytes")
    @patch("proposals_app.views.generate_and_store_proposal_pdf")
    def test_dispatch_sign_inserts_facsimile_into_saved_docx(
        self,
        mocked_generate_and_store_proposal_pdf,
        mocked_load_existing_proposal_docx_bytes,
        mocked_store_existing_proposal_docx_bytes,
    ):
        template_doc = Document()
        template_doc.add_paragraph("Подпись [[facsimile]]")
        buffer = BytesIO()
        template_doc.save(buffer)
        mocked_load_existing_proposal_docx_bytes.return_value = buffer.getvalue()
        mocked_store_existing_proposal_docx_bytes.return_value = {
            "docx_name": "existing-offer.docx",
            "docx_path": "/Corporate Root/ТКП/2026/333300RU DD Приморское/existing-offer.docx",
            "output_dir": "/Corporate Root/ТКП/2026/333300RU DD Приморское",
        }
        mocked_generate_and_store_proposal_pdf.return_value = {
            "pdf_name": "existing-offer.pdf",
            "pdf_path": "/Corporate Root/ТКП/2026/333300RU DD Приморское/existing-offer.pdf",
        }

        response = self.client.post(
            reverse("proposal_dispatch_sign_documents"),
            {
                "proposal_ids[]": [self.proposal.pk],
            },
        )

        self.assertEqual(response.status_code, 200)
        saved_docx_bytes = mocked_store_existing_proposal_docx_bytes.call_args.args[2]
        saved_doc = Document(BytesIO(saved_docx_bytes))
        paragraph = saved_doc.paragraphs[0]
        self.assertEqual(paragraph.text, "Подпись ")
        self.assertIn("wp:anchor", paragraph._p.xml)
        self.assertIn('behindDoc="1"', paragraph._p.xml)
        self.assertNotIn("[[facsimile]]", paragraph._p.xml)

    @patch("proposals_app.views.generate_and_store_proposal_pdf")
    def test_dispatch_sign_returns_error_without_current_user_facsimile(
        self,
        mocked_generate_and_store_proposal_pdf,
    ):
        self.contract_detail.facsimile_file.delete(save=True)

        response = self.client.post(
            reverse("proposal_dispatch_sign_documents"),
            {
                "proposal_ids[]": [self.proposal.pk],
            },
        )

        self.assertEqual(response.status_code, 400)
        payload = response.json()
        self.assertFalse(payload["ok"])
        self.assertIn("Факсимиле", payload["error"])
        mocked_generate_and_store_proposal_pdf.assert_not_called()

    @patch("proposals_app.views.load_existing_proposal_docx_bytes")
    def test_onlyoffice_docx_source_inserts_facsimile_behind_text(
        self,
        mocked_load_existing_proposal_docx_bytes,
    ):
        template_doc = Document()
        template_doc.add_paragraph("[[facsimile]]")
        buffer = BytesIO()
        template_doc.save(buffer)
        mocked_load_existing_proposal_docx_bytes.return_value = buffer.getvalue()

        response = self.client.get(
            reverse("proposal_onlyoffice_docx_source", args=[self.proposal.pk]),
            {
                "token": build_proposal_docx_source_token(self.proposal, signer_user_id=self.user.pk),
            },
        )

        self.assertEqual(response.status_code, 200)
        generated_doc = Document(BytesIO(response.content))
        paragraph = generated_doc.paragraphs[0]
        self.assertEqual(paragraph.text, "")
        self.assertIn("wp:anchor", paragraph._p.xml)
        self.assertIn('behindDoc="1"', paragraph._p.xml)
        self.assertIn('wp:positionH relativeFrom="page"', paragraph._p.xml)
        self.assertIn("<wp:align>center</wp:align>", paragraph._p.xml)
        self.assertNotIn("[[facsimile]]", paragraph._p.xml)


@override_settings(ROOT_URLCONF="proposals_app.urls")
class ProposalDispatchFormTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="proposal-dispatch-form-staff",
            email="staff@example.com",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        self.country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        self.proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=self.group_member,
            type=self.product,
            name="Приморское",
            year=2026,
            contact_full_name="Иванов Иван Иванович",
            contact_email="contact@example.com",
            position=1,
        )

    def test_dispatch_form_edit_saves_contact_name_parts_to_person_registry(self):
        response = self.client.post(
            reverse("proposal_dispatch_form_edit", args=[self.proposal.pk]),
            {
                "docx_file_name": "",
                "docx_file_link": "",
                "pdf_file_name": "",
                "pdf_file_link": "",
                "sent_date": "08.04.2026 12:00",
                "recipient": 'ООО "Альфа"',
                "recipient_country": str(self.country.pk),
                "recipient_identifier": "ОГРН",
                "recipient_registration_number": "1234567890123",
                "recipient_registration_date": "10.04.2026",
                "recipient_job_title": "Генеральный директор",
                "contact_last_name": "Петров",
                "contact_first_name": "Петр",
                "contact_middle_name": "Петрович",
                "contact_email": "contact@example.com",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.proposal.refresh_from_db()
        self.assertEqual(self.proposal.recipient, 'ООО "Альфа"')
        self.assertEqual(self.proposal.recipient_country, self.country)
        self.assertEqual(self.proposal.recipient_identifier, "ОГРН")
        self.assertEqual(self.proposal.recipient_registration_number, "1234567890123")
        self.assertEqual(self.proposal.recipient_registration_date, date(2026, 4, 10))
        self.assertEqual(self.proposal.recipient_job_title, "Генеральный директор")
        self.assertEqual(self.proposal.contact_full_name, "Петров Петр Петрович")
        self.assertEqual(self.proposal.contact_short_name, "Петров П.П.")

        person = PersonRecord.objects.get(last_name="Петров")
        self.assertEqual(person.first_name, "Петр")
        self.assertEqual(person.middle_name, "Петрович")

    def test_dispatch_form_defaults_recipient_country_to_russia(self):
        form = ProposalDispatchForm()
        self.assertEqual(form.fields["recipient_country"].initial, self.country.pk)

    def test_dispatch_form_edit_does_not_overwrite_existing_person_with_same_last_name(self):
        existing_person = PersonRecord.objects.create(
            last_name="Иванов",
            first_name="Иван",
            middle_name="Иванович",
            position=1,
        )

        response = self.client.post(
            reverse("proposal_dispatch_form_edit", args=[self.proposal.pk]),
            {
                "docx_file_name": "",
                "docx_file_link": "",
                "pdf_file_name": "",
                "pdf_file_link": "",
                "sent_date": "08.04.2026 12:00",
                "recipient": 'ООО "Альфа"',
                "recipient_country": str(self.country.pk),
                "recipient_identifier": "ОГРН",
                "recipient_registration_number": "1234567890123",
                "recipient_registration_date": "10.04.2026",
                "recipient_job_title": "Генеральный директор",
                "contact_last_name": "Иванов",
                "contact_first_name": "Петр",
                "contact_middle_name": "Петрович",
                "contact_email": "contact@example.com",
            },
        )

        self.assertEqual(response.status_code, 200)
        existing_person.refresh_from_db()
        self.assertEqual(existing_person.first_name, "Иван")
        self.assertEqual(existing_person.middle_name, "Иванович")

        matching_people = list(PersonRecord.objects.filter(last_name="Иванов").order_by("position", "id"))
        self.assertEqual(len(matching_people), 2)
        self.assertTrue(
            any(
                person.first_name == "Петр" and person.middle_name == "Петрович"
                for person in matching_people
            )
        )

    def test_dispatch_form_edit_does_not_allow_overwriting_transfer_date(self):
        self.proposal.transfer_to_contract_date = "08.04.2026 11:00"
        self.proposal.save(update_fields=["transfer_to_contract_date"])

        response = self.client.post(
            reverse("proposal_dispatch_form_edit", args=[self.proposal.pk]),
            {
                "docx_file_name": "",
                "docx_file_link": "",
                "pdf_file_name": "",
                "pdf_file_link": "",
                "sent_date": "08.04.2026 12:00",
                "transfer_to_contract_date": "99.99.9999 99:99",
                "recipient": 'ООО "Альфа"',
                "recipient_country": str(self.country.pk),
                "recipient_identifier": "ОГРН",
                "recipient_registration_number": "1234567890123",
                "recipient_registration_date": "10.04.2026",
                "recipient_job_title": "Генеральный директор",
                "contact_last_name": "Петров",
                "contact_first_name": "Петр",
                "contact_middle_name": "Петрович",
                "contact_email": "contact@example.com",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.proposal.refresh_from_db()
        self.assertEqual(self.proposal.transfer_to_contract_date, "08.04.2026 11:00")


class ProposalRegistrationFormTests(TestCase):
    def _base_form_payload(self, **overrides):
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        payload = {
            "number": 3333,
            "group_member": group_member.pk,
            "type": product.pk,
            "name": "Тестовое ТКП",
            "kind": ProposalRegistration.ProposalKind.REGULAR,
            "status": ProposalRegistration.ProposalStatus.FINAL,
            "year": 2026,
        }
        payload.update(overrides)
        return payload

    def test_new_form_uses_rub_as_default_currency(self):
        OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
            position=1,
        )

        form = ProposalRegistrationForm()

        currency_id = form.fields["currency"].initial
        self.assertIsNotNone(currency_id)
        self.assertEqual(form.fields["currency"].queryset.get(pk=currency_id).code_alpha, "RUB")

    def test_bound_form_ignores_invalid_country_ids_when_building_region_choices(self):
        form = ProposalRegistrationForm(
            data=self._base_form_payload(
                country="not-a-number",
                registration_region="Приморский край",
                asset_owner_country="still-not-a-number",
                asset_owner_region="Хабаровский край",
            )
        )

        self.assertEqual(
            form.fields["registration_region"].choices,
            [("", "---------"), ("Приморский край", "Приморский край")],
        )
        self.assertEqual(
            form.fields["asset_owner_region"].choices,
            [("", "---------"), ("Хабаровский край", "Хабаровский край")],
        )

    def test_new_form_initializes_single_empty_asset_row(self):
        form = ProposalRegistrationForm()

        self.assertEqual(
            form.fields["assets_payload"].initial,
            '[{"short_name": "", "country_id": "", "country_name": "", "identifier": "", "registration_number": "", "registration_date": ""}]',
        )

    @patch("proposals_app.forms.timezone.now")
    def test_new_form_sets_evaluation_date_to_january_first_before_july(self, mocked_now):
        mocked_now.return_value = timezone.make_aware(datetime(2026, 4, 9, 10, 0, 0))

        form = ProposalRegistrationForm()

        self.assertEqual(form.fields["evaluation_date"].initial, date(2026, 1, 1))
        self.assertNotIn("readonly", form.fields["evaluation_date"].widget.attrs)

    @patch("proposals_app.forms.timezone.now")
    def test_new_form_sets_evaluation_date_to_june_first_from_july_onward(self, mocked_now):
        mocked_now.return_value = timezone.make_aware(datetime(2026, 7, 2, 10, 0, 0))

        form = ProposalRegistrationForm()

        self.assertEqual(form.fields["evaluation_date"].initial, date(2026, 6, 1))
        self.assertNotIn("readonly", form.fields["evaluation_date"].widget.attrs)

    def test_form_preserves_explicit_autocomplete_flags_in_related_payload(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        form = ProposalRegistrationForm(
            data=self._base_form_payload(
                assets_payload=json.dumps(
                    [
                        {
                            "short_name": 'ООО "Актив"',
                            "country_id": str(country.pk),
                            "country_name": "Россия",
                            "identifier": "ОГРН",
                            "registration_number": "1234567890",
                            "registration_date": "01.04.2026",
                            "selected_identifier_record_id": "55",
                            "selected_from_autocomplete": True,
                        }
                    ],
                    ensure_ascii=False,
                )
            )
        )

        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_assets[0]["selected_identifier_record_id"], "55")
        self.assertTrue(form.cleaned_assets[0]["selected_from_autocomplete"])

    @patch("proposals_app.forms.get_cbr_eur_rate_for_today")
    def test_new_form_prefills_cbr_exchange_rate_in_commercial_totals(self, mocked_rate):
        mocked_rate.return_value = Decimal("96.5432")

        form = ProposalRegistrationForm()

        self.assertEqual(
            json.loads(form.fields["commercial_totals_payload"].initial),
            {
                "discount_percent": "5",
                "rub_total_service_text": f"Курс евро Банка России на {timezone.localdate().strftime('%d.%m.%Y')}:",
                "discounted_total_service_text": "Размер скидки:",
                "exchange_rate": "96.5432",
                "travel_expenses_mode": "actual",
            },
        )

    @patch("proposals_app.forms.get_cbr_eur_rate_for_today")
    def test_existing_form_keeps_stored_exchange_rate_in_commercial_totals(self, mocked_rate):
        mocked_rate.return_value = Decimal("101.1111")
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="FIX",
            name_en="Fixed rate",
            name_ru="Фиксированный курс",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=group_member,
            type=product,
            name="Историческое ТКП",
            kind=ProposalRegistration.ProposalKind.REGULAR,
            status=ProposalRegistration.ProposalStatus.FINAL,
            year=2026,
            customer='ООО "История"',
            commercial_totals_json={
                "exchange_rate": "95.4321",
                "discount_percent": "0",
            },
        )

        form = ProposalRegistrationForm(instance=proposal)

        self.assertEqual(
            json.loads(form.fields["commercial_totals_payload"].initial)["exchange_rate"],
            "95.4321",
        )

    @patch("proposals_app.forms.get_cbr_eur_rate_for_today")
    def test_existing_form_keeps_stored_exchange_rate_label_text(self, mocked_rate):
        mocked_rate.return_value = Decimal("101.1111")
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="FIXTXT",
            name_en="Fixed text",
            name_ru="Фиксированный текст",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        proposal = ProposalRegistration.objects.create(
            number=3334,
            group_member=group_member,
            type=product,
            name="Исторический текст курса",
            kind=ProposalRegistration.ProposalKind.REGULAR,
            status=ProposalRegistration.ProposalStatus.FINAL,
            year=2026,
            customer='ООО "История"',
            commercial_totals_json={
                "exchange_rate": "95.4321",
                "discount_percent": "0",
                "rub_total_service_text": "Курс евро Банка России на 15.01.2026:",
            },
        )

        form = ProposalRegistrationForm(instance=proposal)

        self.assertEqual(
            json.loads(form.fields["commercial_totals_payload"].initial)["rub_total_service_text"],
            "Курс евро Банка России на 15.01.2026:",
        )

    @patch("proposals_app.forms.get_cbr_eur_rate_for_today")
    def test_existing_form_does_not_autofill_missing_exchange_rate(self, mocked_rate):
        mocked_rate.return_value = Decimal("101.1111")
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="FIXBLANK",
            name_en="Blank rate",
            name_ru="Пустой курс",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        proposal = ProposalRegistration.objects.create(
            number=3335,
            group_member=group_member,
            type=product,
            name="Пустой курс",
            kind=ProposalRegistration.ProposalKind.REGULAR,
            status=ProposalRegistration.ProposalStatus.FINAL,
            year=2026,
            customer='ООО "История"',
            commercial_totals_json={
                "exchange_rate": "",
                "discount_percent": "0",
                "rub_total_service_text": "Курс евро Банка России на 15.01.2026:",
            },
        )

        form = ProposalRegistrationForm(instance=proposal)

        self.assertEqual(
            json.loads(form.fields["commercial_totals_payload"].initial)["exchange_rate"],
            "",
        )

    def test_new_form_defaults_report_languages_to_russian(self):
        form = ProposalRegistrationForm()

        self.assertEqual(form.initial["report_languages"], "русский")

    def test_new_form_defaults_status_to_final(self):
        form = ProposalRegistrationForm()

        self.assertEqual(form.fields["status"].initial, ProposalRegistration.ProposalStatus.FINAL)

    def test_new_form_requires_year(self):
        form = ProposalRegistrationForm()

        self.assertTrue(form.fields["year"].required)

    def test_form_requires_name(self):
        form = ProposalRegistrationForm()

        self.assertTrue(form.fields["name"].required)

    def test_form_uses_russian_required_error_for_required_fields(self):
        form = ProposalRegistrationForm(data={})

        self.assertFalse(form.is_valid())
        self.assertEqual(form.errors["type"][0], "Обязательное поле.")

    def test_form_uses_russian_integer_error_message(self):
        form = ProposalRegistrationForm(data=self._base_form_payload(number="abc"))

        self.assertFalse(form.is_valid())
        self.assertEqual(form.errors["number"][0], "Введите целое число.")

    def test_form_uses_russian_date_error_message(self):
        form = ProposalRegistrationForm(data=self._base_form_payload(registration_date="not-a-date"))

        self.assertFalse(form.is_valid())
        self.assertEqual(form.errors["registration_date"][0], "Введите дату в формате ДД.ММ.ГГГГ.")

    def test_form_uses_russian_choice_error_message(self):
        form = ProposalRegistrationForm(data=self._base_form_payload(kind="unknown"))

        self.assertFalse(form.is_valid())
        self.assertEqual(form.errors["kind"][0], "Выберите корректное значение.")

    def test_form_renders_non_editable_statuses_as_disabled_options(self):
        form = ProposalRegistrationForm()
        choices = form.fields["status"].widget.optgroups("status", [])
        options = [option for _group_name, group_options, _index in choices for option in group_options]
        option_by_value = {str(option["value"]): option for option in options if option.get("value") is not None}

        self.assertEqual(option_by_value["sent"]["label"], "Отправленное")
        self.assertEqual(option_by_value["completed"]["label"], "Завершённое")
        self.assertTrue(option_by_value["sent"]["attrs"].get("disabled"))
        self.assertTrue(option_by_value["completed"]["attrs"].get("disabled"))

    def test_form_rejects_manual_selection_of_non_editable_status(self):
        form = ProposalRegistrationForm(
            data=self._base_form_payload(status=ProposalRegistration.ProposalStatus.SENT)
        )

        self.assertFalse(form.is_valid())
        self.assertEqual(form.errors["status"][0], "Выберите корректное значение.")

    def test_existing_form_allows_empty_year(self):
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=group_member,
            type=product,
            name="Тестовое ТКП",
            year=None,
        )

        form = ProposalRegistrationForm(instance=proposal)

        self.assertFalse(form.fields["year"].required)

    def test_new_form_defaults_percent_fields_to_40_with_integer_step(self):
        form = ProposalRegistrationForm()

        self.assertEqual(form.fields["advance_percent"].initial, 40)
        self.assertEqual(form.fields["preliminary_report_percent"].initial, 40)
        self.assertEqual(form.fields["advance_term_days"].initial, 10)
        self.assertEqual(form.fields["preliminary_report_term_days"].initial, 7)
        self.assertEqual(form.fields["final_report_term_days"].initial, 15)
        self.assertEqual(form.fields["advance_percent"].widget.attrs["step"], "1")
        self.assertEqual(form.fields["preliminary_report_percent"].widget.attrs["step"], "1")

    def test_form_includes_final_report_term_weeks_decimal_field(self):
        form = ProposalRegistrationForm()

        self.assertEqual(form.fields["service_term_months"].label, "Срок подготовки Предварительного отчёта, мес.")
        self.assertEqual(form.fields["preliminary_report_date"].label, "Дата Предварительного отчёта")
        self.assertEqual(form.fields["final_report_date"].label, "Дата Итогового отчёта")
        self.assertEqual(form.fields["final_report_term_weeks"].label, "Срок подготовки Итогового отчёта, нед.")
        self.assertTrue(form.fields["service_term_months"].widget.attrs["readonly"])
        self.assertEqual(form.fields["service_term_months"].widget.attrs["tabindex"], "-1")
        self.assertIn("readonly-field", form.fields["service_term_months"].widget.attrs["class"])
        self.assertEqual(form.fields["final_report_term_weeks"].widget.attrs["step"], "0.1")
        self.assertTrue(form.fields["final_report_term_weeks"].widget.attrs["readonly"])
        self.assertEqual(form.fields["final_report_term_weeks"].widget.attrs["tabindex"], "-1")
        self.assertIn("readonly-field", form.fields["final_report_term_weeks"].widget.attrs["class"])

    def test_form_accepts_comma_in_final_report_term_weeks(self):
        form = ProposalRegistrationForm(
            data=self._base_form_payload(final_report_term_weeks="2,5")
        )

        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data["final_report_term_weeks"], Decimal("2.5"))

    def test_form_accepts_formatted_service_cost_without_js_raw_submit(self):
        form = ProposalRegistrationForm(
            data=self._base_form_payload(service_cost="1 250 000,50")
        )

        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data["service_cost"], Decimal("1250000.50"))

    def test_invalid_bound_percent_does_not_raise_during_init(self):
        form = ProposalRegistrationForm(
            data={
                "advance_percent": "abc",
                "preliminary_report_percent": "20",
            }
        )

        self.assertIsNone(form.initial.get("final_report_percent"))
        self.assertFalse(form.is_valid())
        self.assertIn("advance_percent", form.errors)

    def test_report_languages_normalizes_legacy_codes(self):
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "report_languages": "RU, EN, zh",
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data["report_languages"], "русский, английский, китайский")

    def test_proposal_variable_form_accepts_registry_binding(self):
        form = ProposalVariableForm(
            data={
                "key": "{{country}}",
                "description": "Страна",
                "source_section": "proposals",
                "source_table": "registry",
                "source_column": "country",
            }
        )

        self.assertTrue(form.is_valid(), form.errors)

    def test_proposal_variable_column_choices_use_current_registry_headers(self):
        self.assertIn(
            ("customer", "Заказчик: наименование"),
            _proposal_variable_column_choices("proposals", "registry"),
        )
        self.assertIn(
            ("registration_region", "Заказчик: регион"),
            _proposal_variable_column_choices("proposals", "registry"),
        )
        self.assertIn(
            ("asset_owner_region", "Владелец: регион"),
            _proposal_variable_column_choices("proposals", "registry"),
        )
        self.assertIn(
            ("country_full_name", "Наименование страны (полное)"),
            _proposal_variable_column_choices("proposals", "registry"),
        )

    def test_proposal_variable_form_locks_computed_variable_fields(self):
        variable = ProposalVariable.objects.create(
            key="{{client_country_full_name}}",
            description="Наименование страны Заказчика (полное)",
            is_computed=True,
            position=1,
        )

        form = ProposalVariableForm(
            data={
                "key": "{{tampered}}",
                "description": "Обновлённое описание",
                "source_section": "proposals",
                "source_table": "registry",
                "source_column": "country_full_name",
            },
            instance=variable,
        )

        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(form.cleaned_data["key"], "{{client_country_full_name}}")
        self.assertEqual(form.cleaned_data["source_section"], "")
        self.assertEqual(form.cleaned_data["source_table"], "")
        self.assertEqual(form.cleaned_data["source_column"], "")
        self.assertEqual(form.fields["source_section"].choices, [("", "---")])
        self.assertEqual(form.fields["source_table"].choices, [("", "---")])
        self.assertEqual(form.fields["source_column"].choices, [("", "---")])

    def test_asset_owner_copies_customer_when_checkbox_enabled(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "registration_region": "Приморский край",
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "asset_owner": 'ООО "Другое"',
                "asset_owner_country": "",
                "asset_owner_region": "",
                "asset_owner_identifier": "",
                "asset_owner_registration_number": "",
                "asset_owner_registration_date": "",
                "asset_owner_matches_customer": "on",
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()

        self.assertEqual(proposal.asset_owner, proposal.customer)
        self.assertEqual(proposal.asset_owner_country, proposal.country)
        self.assertEqual(proposal.asset_owner_region, proposal.registration_region)
        self.assertEqual(proposal.asset_owner_identifier, proposal.identifier)
        self.assertEqual(proposal.asset_owner_registration_number, proposal.registration_number)
        self.assertEqual(proposal.asset_owner_registration_date, proposal.registration_date)

    def test_customer_tz_mode_keeps_service_composition_in_sync(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "service_composition": "Старый текст разделов",
                "service_composition_customer_tz": "Новый текст ТЗ заказчика",
                "service_customer_tz_editor_state": '{"html":"<p><strong>Новый</strong> текст ТЗ заказчика</p>","plain_text":"Новый текст ТЗ заказчика"}',
                "service_composition_mode": "customer_tz",
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()

        self.assertEqual(proposal.service_composition_mode, "customer_tz")
        self.assertEqual(proposal.service_composition_customer_tz, "Новый текст ТЗ заказчика")
        self.assertEqual(proposal.service_composition, "Новый текст ТЗ заказчика")
        self.assertEqual(
            proposal.service_customer_tz_editor_state,
            {
                "html": "<p><strong>Новый</strong> текст ТЗ заказчика</p>",
                "plain_text": "Новый текст ТЗ заказчика",
            },
        )

    def test_service_sections_editor_state_is_saved_to_model(self):
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "service_sections_editor_state": (
                    '[{"code":"S-101","service_name":"Раздел 1","html":"<p><u>Этап 1</u></p>",'
                    '"plain_text":"Этап 1"}]'
                ),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()

        self.assertEqual(
            proposal.service_sections_editor_state,
            [
                {
                    "code": "S-101",
                    "service_name": "Раздел 1",
                    "html": "<p><u>Этап 1</u></p>",
                    "plain_text": "Этап 1",
                }
            ],
        )

    def test_service_sections_payload_saves_code_by_selected_service(self):
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        TypicalSection.objects.create(
            product=product,
            code="S-101",
            short_name="service-1",
            short_name_ru="Услуга 1",
            name_en="Service 1",
            name_ru="Раздел 1",
            accounting_type="Услуги",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "service_sections_payload": '[{"service_name":"Раздел 1","code":""}]',
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()

        self.assertEqual(
            proposal.service_sections_json,
            [{"service_name": "Раздел 1", "code": "S-101"}],
        )

    def test_resolve_variables_for_new_registry_columns(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        currency = OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=group_member,
            type=product,
            name="Приморское",
            status=ProposalRegistration.ProposalStatus.PRELIMINARY,
            proposal_project_name="Due Diligence АО «Полиметалл УК»",
            customer='АО «Полиметалл УК»',
            asset_owner='АО «Полиметалл УК»',
            asset_owner_matches_customer=True,
            country=country,
            asset_owner_country=country,
            registration_region="Тюменская область",
            asset_owner_region="Тюменская область",
            purpose="Проверка актива",
            service_composition_mode="sections",
            service_composition="Этап 1\nЭтап 2",
            service_sections_editor_state=[
                {
                    "code": "S-101",
                    "service_name": "Раздел 1",
                    "html": "<p><strong>Этап 1</strong></p>",
                    "plain_text": "Этап 1",
                },
                {
                    "code": "S-102",
                    "service_name": "Раздел 2",
                    "html": "<p><u>Этап 2</u></p>",
                    "plain_text": "Этап 2",
                },
            ],
            evaluation_date="2026-04-01",
            service_term_months="4.5",
            preliminary_report_date="2026-04-15",
            final_report_term_weeks="2.0",
            final_report_date="2026-05-20",
            report_languages="RU, EN",
            service_cost="1250000.50",
            currency=currency,
            advance_percent="50",
            advance_term_days=5,
            preliminary_report_percent="20",
            preliminary_report_term_days=10,
            final_report_percent="30",
            final_report_term_days=15,
        )
        ServiceGoalReport.objects.create(
            product=product,
            service_goal="Проверка актива",
            service_goal_genitive="Подготовки коммерческого предложения",
            report_title="ТКП по проекту DD",
            position=1,
        )
        variables = [
            ProposalVariable(
                key="{{proposal_project_name}}",
                source_section="proposals",
                source_table="registry",
                source_column="proposal_project_name",
            ),
            ProposalVariable(key="{{purpose}}", source_section="proposals", source_table="registry", source_column="purpose"),
            ProposalVariable(
                key="{{service_cost}}",
                source_section="proposals",
                source_table="registry",
                source_column="service_cost",
            ),
            ProposalVariable(
                key="{{evaluation_date}}",
                source_section="proposals",
                source_table="registry",
                source_column="evaluation_date",
            ),
            ProposalVariable(
                key="{{final_report_term_weeks}}",
                source_section="proposals",
                source_table="registry",
                source_column="final_report_term_weeks",
            ),
            ProposalVariable(
                key="{{status}}",
                source_section="proposals",
                source_table="registry",
                source_column="status",
            ),
            ProposalVariable(
                key="{{advance_percent}}",
                source_section="proposals",
                source_table="registry",
                source_column="advance_percent",
            ),
            ProposalVariable(
                key="{{preliminary_report_percent}}",
                source_section="proposals",
                source_table="registry",
                source_column="preliminary_report_percent",
            ),
            ProposalVariable(
                key="{{final_report_percent}}",
                source_section="proposals",
                source_table="registry",
                source_column="final_report_percent",
            ),
            ProposalVariable(
                key="{{currency}}",
                source_section="proposals",
                source_table="registry",
                source_column="currency",
            ),
            ProposalVariable(
                key="{{country}}",
                source_section="proposals",
                source_table="registry",
                source_column="country",
            ),
            ProposalVariable(
                key="{{registration_region}}",
                source_section="proposals",
                source_table="registry",
                source_column="registration_region",
            ),
            ProposalVariable(
                key="{{asset_owner_region}}",
                source_section="proposals",
                source_table="registry",
                source_column="asset_owner_region",
            ),
            ProposalVariable(
                key="{{client_country_full_name}}",
                is_computed=True,
            ),
            ProposalVariable(
                key="{{client_owner_name}}",
                is_computed=True,
            ),
            ProposalVariable(
                key="{{service_type_short}}",
                is_computed=True,
            ),
            ProposalVariable(
                key="{{service_goal_genitive}}",
                is_computed=True,
            ),
            ProposalVariable(
                key="{{tkp_preliminary}}",
                is_computed=True,
            ),
            ProposalVariable(
                key="{{preliminary_payment_percentage_full}}",
                is_computed=True,
            ),
            ProposalVariable(
                key="{{preliminary_report_term_month}}",
                is_computed=True,
            ),
            ProposalVariable(
                key="[[scope_of_work]]",
                is_computed=True,
            ),
            ProposalVariable(
                key="[[actives_name]]",
                is_computed=True,
            ),
            ProposalVariable(
                key="{{owner_country_full_name}}",
                is_computed=True,
            ),
            ProposalVariable(key="{{year}}", is_computed=True),
            ProposalVariable(key="{{day}}", is_computed=True),
            ProposalVariable(key="{{month}}", is_computed=True),
        ]

        with patch("proposals_app.variable_resolver._today", return_value=date(2026, 4, 9)):
            replacements, lists, tables = resolve_variables(proposal, variables)

        self.assertEqual(replacements["{{proposal_project_name}}"], "Due Diligence АО «Полиметалл УК»")
        self.assertEqual(replacements["{{purpose}}"], "Проверка актива")
        self.assertEqual(replacements["{{service_cost}}"], "1\u00a0250\u00a0000,50")
        self.assertEqual(replacements["{{evaluation_date}}"], "01.04.2026")
        self.assertEqual(replacements["{{final_report_term_weeks}}"], "2,0")
        self.assertEqual(replacements["{{status}}"], "Предварительное")
        self.assertEqual(replacements["{{advance_percent}}"], "50%")
        self.assertEqual(replacements["{{preliminary_report_percent}}"], "20%")
        self.assertEqual(replacements["{{final_report_percent}}"], "30%")
        self.assertEqual(replacements["{{currency}}"], "RUB")
        self.assertEqual(replacements["{{country}}"], "Россия")
        self.assertEqual(replacements["{{registration_region}}"], "Тюменская область")
        self.assertEqual(replacements["{{asset_owner_region}}"], "Тюменская область")
        self.assertEqual(replacements["{{client_country_full_name}}"], "Российская Федерация")
        self.assertEqual(replacements["{{country_full_name}}"], "Российская Федерация")
        self.assertEqual(replacements["{{client_owner_name}}"], 'АО «Полиметалл УК»')
        self.assertEqual(replacements["{{service_type_short}}"], "Due Diligence")
        self.assertEqual(replacements["{{service_goal_genitive}}"], "Подготовки коммерческого предложения")
        self.assertEqual(replacements["{{tkp_preliminary}}"], "(предварительное)")
        self.assertEqual(replacements["{{preliminary_payment_percentage_full}}"], "70%")
        self.assertEqual(replacements["{{preliminary_report_term_month}}"], "4,5 месяца")
        self.assertEqual(replacements["{{owner_country_full_name}}"], "Российская Федерация")
        self.assertEqual(replacements["{{year}}"], "2026")
        self.assertEqual(replacements["{{day}}"], "09")
        self.assertEqual(replacements["{{month}}"], "апреля")
        self.assertEqual(lists["[[scope_of_work]]"], [{"html": "<p><strong>Этап 1</strong></p>"}, {"html": "<p><u>Этап 2</u></p>"}])
        self.assertEqual(lists["[[actives_name]]"], [])
        self.assertEqual(tables, {})

        self.assertEqual(
            ProposalVariable(key="{{year}}", is_computed=True).binding_display,
            "Расчётное поле",
        )

    def test_resolve_actives_name_list_uses_proposal_assets_short_names(self):
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        proposal = ProposalRegistration.objects.create(
            number=1111,
            group_member=group_member,
            type=product,
            name="Активы",
            year=2026,
            customer='ООО "Заказчик"',
        )
        ProposalAsset.objects.create(proposal=proposal, short_name="Актив 1", position=1)
        ProposalAsset.objects.create(proposal=proposal, short_name="Актив 2", position=2)
        ProposalAsset.objects.create(proposal=proposal, short_name="Актив 1", position=3)

        replacements, lists, tables = resolve_variables(
            proposal,
            [ProposalVariable(key="[[actives_name]]", is_computed=True)],
        )

        self.assertEqual(replacements, {})
        self.assertEqual(lists["[[actives_name]]"], ["Актив 1", "Актив 2"])
        self.assertEqual(tables, {})

    def test_resolve_client_owner_name_depends_on_asset_owner_checkbox(self):
        proposal = ProposalRegistration(customer='АО «Полиметалл УК»')
        variables = [ProposalVariable(key="{{client_owner_name}}", is_computed=True)]

        replacements, _, _ = resolve_variables(proposal, variables)
        self.assertEqual(replacements["{{client_owner_name}}"], 'АО «Полиметалл УК»')

        proposal.asset_owner_matches_customer = False
        proposal.asset_owner = 'ООО «КАП»'
        replacements, _, _ = resolve_variables(proposal, variables)
        self.assertEqual(replacements["{{client_owner_name}}"], 'АО «Полиметалл УК» / ООО «КАП»')

        proposal.proposal_project_name = 'Due Diligence ООО «КАП»'
        replacements, _, _ = resolve_variables(
            proposal,
            [ProposalVariable(key="{{service_type_short}}", is_computed=True)],
        )
        self.assertEqual(replacements["{{service_type_short}}"], "Due Diligence")

    def test_resolve_tkp_preliminary_depends_on_status(self):
        proposal = ProposalRegistration(status=ProposalRegistration.ProposalStatus.PRELIMINARY)

        replacements, _, _ = resolve_variables(
            proposal,
            [ProposalVariable(key="{{tkp_preliminary}}", is_computed=True)],
        )
        self.assertEqual(replacements["{{tkp_preliminary}}"], "(предварительное)")

        proposal.status = ProposalRegistration.ProposalStatus.FINAL
        replacements, _, _ = resolve_variables(
            proposal,
            [ProposalVariable(key="{{tkp_preliminary}}", is_computed=True)],
        )
        self.assertEqual(replacements["{{tkp_preliminary}}"], "")

    def test_resolve_preliminary_payment_percentage_full_sums_existing_percent_fields(self):
        proposal = ProposalRegistration(advance_percent="50", preliminary_report_percent="20")

        replacements, _, _ = resolve_variables(
            proposal,
            [ProposalVariable(key="{{preliminary_payment_percentage_full}}", is_computed=True)],
        )
        self.assertEqual(replacements["{{preliminary_payment_percentage_full}}"], "70%")

    def test_resolve_preliminary_report_term_month_uses_month_declension(self):
        cases = [
            ("1", "1 месяц"),
            ("4.5", "4,5 месяца"),
            ("5", "5 месяцев"),
            ("6", "6 месяцев"),
        ]

        for source_value, expected in cases:
            with self.subTest(service_term_months=source_value):
                proposal = ProposalRegistration(service_term_months=source_value)
                replacements, _, _ = resolve_variables(
                    proposal,
                    [ProposalVariable(key="{{preliminary_report_term_month}}", is_computed=True)],
                )
                self.assertEqual(replacements["{{preliminary_report_term_month}}"], expected)

    def test_resolve_scope_of_work_prefers_customer_tz_editor_state(self):
        proposal = ProposalRegistration(
            service_composition_mode="customer_tz",
            service_composition="Текст из разделов",
            service_composition_customer_tz="Текст ТЗ заказчика",
            service_customer_tz_editor_state={
                "html": "<p><strong>Текст ТЗ заказчика</strong></p>",
                "plain_text": "Текст ТЗ заказчика",
            },
        )

        _, lists, _ = resolve_variables(
            proposal,
            [ProposalVariable(key="[[scope_of_work]]", is_computed=True)],
        )

        self.assertEqual(lists["[[scope_of_work]]"], [{"html": "<p><strong>Текст ТЗ заказчика</strong></p>"}])

    def test_form_saves_assets_from_payload(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "assets_payload": (
                    '[{"short_name":"ООО \\"Актив\\"","country_id":"%s","identifier":"ОГРН",'
                    '"registration_number":"123","registration_date":"02.04.2026"}]'
                ) % country.pk,
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()
        form.save_assets(proposal)

        asset = ProposalAsset.objects.get(proposal=proposal)
        self.assertEqual(asset.short_name, 'ООО "Актив"')
        self.assertEqual(asset.country, country)
        self.assertEqual(asset.identifier, "ОГРН")
        self.assertEqual(asset.registration_number, "123")
        self.assertEqual(asset.registration_date.isoformat(), "2026-04-02")

    def test_form_saves_legal_entities_from_payload(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "assets_payload": '[{"short_name":"ООО \\"Актив\\""}]',
                "legal_entities_payload": (
                    '[{"asset_short_name":"ООО \\"Актив\\"","short_name":"ООО \\"Юрлицо\\"","country_id":"%s",'
                    '"identifier":"ОГРН","registration_number":"456","registration_date":"03.04.2026"}]'
                ) % country.pk,
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()
        form.save_assets(proposal)
        form.save_legal_entities(proposal)

        legal_entity = ProposalLegalEntity.objects.get(proposal=proposal)
        self.assertEqual(legal_entity.asset_short_name, 'ООО "Актив"')
        self.assertEqual(legal_entity.short_name, 'ООО "Юрлицо"')
        self.assertEqual(legal_entity.country, country)
        self.assertEqual(legal_entity.identifier, "ОГРН")
        self.assertEqual(legal_entity.registration_number, "456")
        self.assertEqual(legal_entity.registration_date.isoformat(), "2026-04-03")

    def test_form_allows_legal_entity_without_short_name(self):
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "assets_payload": '[{"short_name":"ООО \\"Актив\\""}]',
                "legal_entities_payload": '[{"asset_short_name":"ООО \\"Актив\\"","short_name":"","identifier":"ОГРН"}]',
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()
        form.save_assets(proposal)
        form.save_legal_entities(proposal)

        legal_entity = ProposalLegalEntity.objects.get(proposal=proposal)
        self.assertEqual(legal_entity.asset_short_name, 'ООО "Актив"')
        self.assertEqual(legal_entity.short_name, "")
        self.assertEqual(legal_entity.identifier, "ОГРН")

    def test_form_saves_objects_from_payload(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "assets_payload": '[{"short_name":"ООО \\"Актив\\""}]',
                "legal_entities_payload": '[{"asset_short_name":"ООО \\"Актив\\"","short_name":"ООО \\"Юрлицо\\""}]',
                "objects_payload": (
                    '[{"legal_entity_short_name":"ООО \\"Юрлицо\\"","short_name":"Объект 1","region":"Россия",'
                    '"object_type":"ОГРН","license":"Лицензия 123","registration_date":"04.04.2026"}]'
                ),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()
        form.save_assets(proposal)
        form.save_legal_entities(proposal)
        form.save_objects(proposal)

        proposal_object = ProposalObject.objects.get(proposal=proposal)
        self.assertEqual(proposal_object.legal_entity_short_name, 'ООО "Юрлицо"')
        self.assertEqual(proposal_object.short_name, "Объект 1")
        self.assertEqual(proposal_object.region, "Россия")
        self.assertEqual(proposal_object.object_type, "ОГРН")
        self.assertEqual(proposal_object.license, "Лицензия 123")
        self.assertEqual(proposal_object.registration_date.isoformat(), "2026-04-04")

    def test_form_ignores_hidden_placeholder_object_rows_without_short_name(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "assets_payload": '[{"short_name":"ООО \\"Актив 1\\""},{"short_name":"ООО \\"Актив 2\\""}]',
                "legal_entities_payload": (
                    '[{"asset_short_name":"ООО \\"Актив 1\\"","short_name":"ООО \\"Юрлицо 1\\""},'
                    '{"asset_short_name":"ООО \\"Актив 2\\"","short_name":"ООО \\"Юрлицо 2\\""}]'
                ),
                "objects_payload": (
                    '[{"legal_entity_short_name":"ООО \\"Юрлицо 1\\"","short_name":"Объект 1"},'
                    '{"legal_entity_short_name":"ООО \\"Юрлицо 2\\"","short_name":""}]'
                ),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()
        form.save_assets(proposal)
        form.save_legal_entities(proposal)
        form.save_objects(proposal)

        objects = list(ProposalObject.objects.filter(proposal=proposal).order_by("position"))
        self.assertEqual(len(objects), 1)
        self.assertEqual(objects[0].legal_entity_short_name, 'ООО "Юрлицо 1"')
        self.assertEqual(objects[0].short_name, "Объект 1")

    def test_form_saves_commercial_offer_from_payload(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "commercial_offer_payload": (
                    '[{"specialist":"Иванов","job_title":"Партнер","professional_status":"ACCA","service_name":"Финансы",'
                    '"rate_eur_per_day":"1200.50","asset_day_counts":["2","3"],'
                    '"total_eur_without_vat":"6002.50"}]'
                ),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()
        form.save_commercial_offers(proposal)

        offer = ProposalCommercialOffer.objects.get(proposal=proposal)
        self.assertEqual(offer.specialist, "Иванов")
        self.assertEqual(offer.job_title, "Партнер")
        self.assertEqual(offer.professional_status, "ACCA")
        self.assertEqual(offer.service_name, "Финансы")
        self.assertEqual(str(offer.rate_eur_per_day), "1200.50")
        self.assertEqual(offer.asset_day_counts, [2, 3])
        self.assertEqual(str(offer.total_eur_without_vat), "6002.50")

    def test_form_saves_commercial_totals_payload(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "commercial_totals_payload": (
                    '{"exchange_rate":"96.50","discount_percent":"7.50",'
                    '"contract_total":"1200000","contract_total_auto":"1300000",'
                    '"rub_total_service_text":"Курс ЦБ","discounted_total_service_text":"Скидка проекта",'
                    '"travel_expenses_mode":"calculation"}'
                ),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()

        self.assertEqual(
            proposal.commercial_totals_json,
            {
                "exchange_rate": "96.50",
                "discount_percent": "7.50",
                "contract_total": "1200000",
                "contract_total_auto": "1300000",
                "rub_total_service_text": "Курс ЦБ",
                "discounted_total_service_text": "Скидка проекта",
                "travel_expenses_mode": "calculation",
            },
        )

    def test_form_preserves_legacy_travel_total_in_actual_mode(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "Тестовое ТКП",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Приморское"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "commercial_offer_payload": (
                    '[{"specialist":"","job_title":"","professional_status":"",'
                    '"service_name":"Командировочные расходы, евро","rate_eur_per_day":"",'
                    '"asset_day_counts":["",""],"total_eur_without_vat":"500.00"}]'
                ),
                "commercial_totals_payload": (
                    '{"discount_percent":"5","travel_expenses_mode":"actual"}'
                ),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()
        form.save_commercial_offers(proposal)

        offer = ProposalCommercialOffer.objects.get(proposal=proposal)
        self.assertEqual(offer.service_name, "Командировочные расходы, евро")
        self.assertEqual(offer.asset_day_counts, ["", ""])
        self.assertEqual(str(offer.total_eur_without_vat), "500.00")

    def test_form_preserves_zero_values_in_commercial_totals_payload(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        product = Product.objects.create(
            short_name="ZERO",
            name_en="Zero values",
            name_ru="Нулевые значения",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

        form = ProposalRegistrationForm(
            data={
                "number": 3333,
                "group_member": group_member.pk,
                "type": product.pk,
                "name": "ТКП с нулями",
                "kind": ProposalRegistration.ProposalKind.REGULAR,
                "status": ProposalRegistration.ProposalStatus.FINAL,
                "year": 2026,
                "customer": 'ООО "Ноль"',
                "country": country.pk,
                "identifier": "ОГРН",
                "registration_number": "1174910001683",
                "registration_date": "01.04.2026",
                "commercial_totals_payload": (
                    '{"exchange_rate":"0","discount_percent":"0",'
                    '"contract_total":"0","contract_total_auto":"0",'
                    '"rub_total_service_text":"Курс ЦБ","discounted_total_service_text":"Скидка проекта",'
                    '"travel_expenses_mode":"actual"}'
                ),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        proposal = form.save()

        self.assertEqual(
            proposal.commercial_totals_json,
            {
                "exchange_rate": "0",
                "discount_percent": "0",
                "contract_total": "0",
                "contract_total_auto": "0",
                "rub_total_service_text": "Курс ЦБ",
                "discounted_total_service_text": "Скидка проекта",
                "travel_expenses_mode": "actual",
            },
        )


class ProposalFormContextTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="proposal-context-staff",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )

    def _create_staff_employee(self, *, username, first_name, last_name, patronymic="", department=None, role=""):
        user = get_user_model().objects.create_user(
            username=username,
            password="secret",
            first_name=first_name,
            last_name=last_name,
            is_staff=True,
        )
        employee = Employee.objects.create(
            user=user,
            patronymic=patronymic,
            department=department,
            role=role,
        )
        return user, employee

    def test_typical_sections_json_prefers_highest_grade_within_best_rank(self):
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        section = TypicalSection.objects.create(
            product=product,
            code="FIN",
            short_name="FIN",
            name_en="Finance",
            name_ru="Финансы",
            position=1,
        )
        specialty = ExpertSpecialty.objects.create(specialty="Партнер", position=1)
        TypicalSectionSpecialty.objects.create(section=section, specialty=specialty, rank=1)
        grade_low = Grade.objects.create(
            grade_en="G1",
            grade_ru="G1",
            qualification=1,
            qualification_levels=5,
            created_by=self.user,
            position=1,
        )
        grade_high = Grade.objects.create(
            grade_en="G3",
            grade_ru="G3",
            qualification=3,
            qualification_levels=5,
            created_by=self.user,
            position=2,
        )
        grade_top_other_rank = Grade.objects.create(
            grade_en="G5",
            grade_ru="G5",
            qualification=5,
            qualification_levels=5,
            created_by=self.user,
            position=3,
        )

        _, employee_low = self._create_staff_employee(
            username="candidate-low",
            first_name="Иван",
            last_name="Петров",
        )
        profile_low = ExpertProfile.objects.create(
            employee=employee_low,
            professional_status="Association of Chartered Certified Accountants",
            professional_status_short="ACCA",
            grade=grade_low,
            position=1,
        )
        ExpertProfileSpecialty.objects.create(profile=profile_low, specialty=specialty, rank=1)

        _, employee_high = self._create_staff_employee(
            username="candidate-high",
            first_name="Анна",
            last_name="Сидорова",
        )
        profile_high = ExpertProfile.objects.create(
            employee=employee_high,
            professional_status="Chartered Financial Analyst",
            professional_status_short="CFA",
            grade=grade_high,
            position=2,
        )
        ExpertProfileSpecialty.objects.create(profile=profile_high, specialty=specialty, rank=1)

        _, employee_other_rank = self._create_staff_employee(
            username="candidate-rank-two",
            first_name="Петр",
            last_name="Иванов",
        )
        profile_other_rank = ExpertProfile.objects.create(
            employee=employee_other_rank,
            professional_status="Financial Risk Manager",
            professional_status_short="FRM",
            grade=grade_top_other_rank,
            position=3,
        )
        ExpertProfileSpecialty.objects.create(profile=profile_other_rank, specialty=specialty, rank=2)

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        entries = response.context["typical_sections_json"][str(product.pk)]
        finance_entry = next(item for item in entries if item["name"] == section.name_ru)
        self.assertEqual(finance_entry["default_specialist"], "Анна Сидорова")
        self.assertEqual(finance_entry["default_professional_status"], "CFA")
        self.assertEqual(
            [item["name"] for item in finance_entry["specialist_options"]],
            ["Анна Сидорова", "Иван Петров", "Петр Иванов"],
        )

    def test_typical_sections_json_uses_ranked_section_specialties_for_executor_payload(self):
        product = Product.objects.create(
            short_name="CUR",
            name_en="Current specialties",
            name_ru="Актуальные специальности",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        section = TypicalSection.objects.create(
            product=product,
            code="CUR-1",
            short_name="CUR-1",
            name_en="Current section",
            name_ru="Актуальный раздел",
            position=1,
        )
        specialty = ExpertSpecialty.objects.create(specialty="Новая специальность", position=1)
        TypicalSectionSpecialty.objects.create(section=section, specialty=specialty, rank=1)
        grade = Grade.objects.create(
            grade_en="G2",
            grade_ru="G2",
            qualification=2,
            qualification_levels=5,
            created_by=self.user,
            position=1,
        )
        _, employee = self._create_staff_employee(
            username="current-specialty-candidate",
            first_name="Елена",
            last_name="Экспертова",
        )
        profile = ExpertProfile.objects.create(
            employee=employee,
            professional_status="ACCA",
            grade=grade,
            position=1,
        )
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=specialty, rank=1)

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        entries = response.context["typical_sections_json"][str(product.pk)]
        section_entry = next(item for item in entries if item["name"] == section.name_ru)
        self.assertEqual(section_entry["executor"], "Новая специальность")
        self.assertEqual(section_entry["default_specialist"], "Елена Экспертова")

    def test_typical_sections_json_marks_sections_excluded_from_tkp_autofill(self):
        product = Product.objects.create(
            short_name="TKP",
            name_en="TKP product",
            name_ru="ТКП продукт",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        TypicalSection.objects.create(
            product=product,
            code="IN",
            short_name="IN",
            name_en="Included",
            name_ru="Включенный раздел",
            position=1,
        )
        TypicalSection.objects.create(
            product=product,
            code="OUT",
            short_name="OUT",
            name_en="Excluded",
            name_ru="Исключенный раздел",
            exclude_from_tkp_autofill=True,
            position=2,
        )

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        entries = response.context["typical_sections_json"][str(product.pk)]
        self.assertEqual(
            [item["name"] for item in entries],
            ["Включенный раздел", "Исключенный раздел"],
        )
        excluded_entry = next(item for item in entries if item["name"] == "Исключенный раздел")
        self.assertTrue(excluded_entry["exclude_from_tkp_autofill"])
        included_entry = next(item for item in entries if item["name"] == "Включенный раздел")
        self.assertFalse(included_entry["exclude_from_tkp_autofill"])

    def test_typical_sections_json_exposes_accounting_type_for_tz_editor(self):
        product = Product.objects.create(
            short_name="SEC",
            name_en="Section product",
            name_ru="Продукт разделов",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        TypicalSection.objects.create(
            product=product,
            code="SEC-1",
            short_name="section-1",
            short_name_ru="Раздел 1",
            name_en="Section 1",
            name_ru="Раздел 1",
            accounting_type="Раздел",
            position=1,
        )
        TypicalSection.objects.create(
            product=product,
            code="SRV-1",
            short_name="service-1",
            short_name_ru="Услуга 1",
            name_en="Service 1",
            name_ru="Услуга 1",
            accounting_type="Услуги",
            position=2,
        )

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        entries = response.context["typical_sections_json"][str(product.pk)]
        self.assertEqual(
            {item["name"]: item["accounting_type"] for item in entries},
            {
                "Раздел 1": "Раздел",
                "Услуга 1": "Услуги",
            },
        )

    def test_typical_sections_json_uses_direction_head_for_special_expertise_sections(self):
        product = Product.objects.create(
            short_name="VAL",
            name_en="Valuation",
            name_ru="Оценка",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=2,
        )
        expertise = ExpertiseDirection.objects.create(
            name="Финансовая экспертиза",
            short_name="FIN",
            position=1,
        )
        direction = OrgUnit.objects.create(
            company=self.group_member,
            department_name="Финансовое направление",
            short_name="Финансы",
            expertise=expertise,
            unit_type="expertise",
            position=1,
        )
        section = TypicalSection.objects.create(
            product=product,
            code="VAL",
            short_name="VAL",
            name_en="Valuation",
            name_ru="Оценка бизнеса",
            expertise_dir=expertise,
            expertise_direction=direction,
            position=1,
        )
        specialty = ExpertSpecialty.objects.create(specialty="Оценщик", position=2)
        TypicalSectionSpecialty.objects.create(section=section, specialty=specialty, rank=1)
        _, employee_candidate = self._create_staff_employee(
            username="valuation-candidate",
            first_name="Мария",
            last_name="Экспертова",
        )
        profile_candidate = ExpertProfile.objects.create(
            employee=employee_candidate,
            professional_status="Accredited Senior Appraiser",
            professional_status_short="ASA",
            position=4,
        )
        ExpertProfileSpecialty.objects.create(profile=profile_candidate, specialty=specialty, rank=1)

        _, employee_head = self._create_staff_employee(
            username="direction-head",
            first_name="Олег",
            last_name="Руководитель",
            department=direction,
            role="Руководитель направления",
        )
        ExpertProfile.objects.create(
            employee=employee_head,
            professional_status="Certified Public Accountant",
            professional_status_short="CPA",
            position=5,
        )

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        entries = response.context["typical_sections_json"][str(product.pk)]
        section_entry = next(item for item in entries if item["name"] == section.name_ru)
        self.assertEqual(section_entry["default_specialist"], "Олег Руководитель")
        self.assertEqual(section_entry["default_professional_status"], "CPA")
        self.assertEqual(section_entry["specialist_options"][0]["name"], "Олег Руководитель")

    def test_typical_sections_json_includes_commercial_rate_and_days_autofill_data(self):
        product = Product.objects.create(
            short_name="TAX",
            name_en="Tax",
            name_ru="Налоги",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=5,
        )
        section = TypicalSection.objects.create(
            product=product,
            code="TX-1",
            short_name="tax-1",
            short_name_ru="Налог 1",
            name_en="Tax Service",
            name_ru="Налоговый анализ",
            position=1,
        )
        specialty = ExpertSpecialty.objects.create(
            specialty="Партнер",
            position=1,
        )
        TypicalSectionSpecialty.objects.create(section=section, specialty=specialty, rank=1)
        grade = Grade.objects.create(
            grade_en="G2",
            grade_ru="G2",
            qualification=2,
            qualification_levels=5,
            base_rate_share=25,
            created_by=self.user,
            position=10,
        )
        _, employee = self._create_staff_employee(
            username="commercial-rate-candidate",
            first_name="Ирина",
            last_name="Смирнова",
        )
        profile = ExpertProfile.objects.create(
            employee=employee,
            professional_status="ACCA",
            grade=grade,
            position=10,
        )
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=specialty, rank=1)
        SpecialtyTariff.objects.create(
            specialty_group="Партнеры",
            daily_rate_tkp_eur="1000.00",
            position=1,
        ).specialties.set([specialty])
        Tariff.objects.create(
            product=product,
            section=section,
            base_rate_vpm="1.00",
            service_hours=8,
            service_days_tkp=7,
            created_by=self.user,
            position=1,
        )

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        entry = response.context["typical_sections_json"][str(product.pk)][0]
        self.assertEqual(entry["specialty_tariff_rate_eur"], "1000.00")
        self.assertEqual(entry["service_days_tkp"], 7)
        self.assertTrue(entry["specialty_is_director"])
        self.assertEqual(entry["default_base_rate_share"], 25)
        self.assertEqual(entry["specialist_options"][0]["base_rate_share"], 25)

    def test_service_goal_reports_json_exposes_first_product_defaults(self):
        product = Product.objects.create(
            short_name="COM",
            name_en="Commercial Offer",
            name_ru="Коммерческое предложение",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=3,
        )
        ServiceGoalReport.objects.create(
            product=product,
            service_goal="Подготовка коммерческого предложения",
            service_goal_genitive="Подготовки коммерческого предложения",
            report_title="ТКП по проекту COM",
            position=2,
        )
        ServiceGoalReport.objects.create(
            product=product,
            service_goal="Не должно попасть в автозаполнение",
            service_goal_genitive="Не должно попасть в автозаполнение в родительном падеже",
            report_title="Не использовать",
            position=3,
        )

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["service_goal_reports_json"][str(product.pk)],
            {
                "report_title": "ТКП по проекту COM",
                "service_goal": "Подготовка коммерческого предложения",
                "service_goal_genitive": "Подготовки коммерческого предложения",
            },
        )

    def test_typical_service_compositions_json_exposes_text_by_product_and_section(self):
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=4,
        )
        section = TypicalSection.objects.create(
            product=product,
            code="S-101",
            short_name="service-1",
            short_name_ru="Услуга 1",
            name_en="Service 1",
            name_ru="Раздел 1",
            accounting_type="Услуги",
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=product,
            section=section,
            service_composition="Этап 1\nЭтап 2",
            service_composition_editor_state={
                "html": "<p><strong>Этап 1</strong></p><p>Этап 2</p>",
                "plain_text": "Этап 1\nЭтап 2",
            },
            position=1,
        )

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["typical_service_compositions_json"][str(product.pk)],
            [
                {
                    "code": "S-101",
                    "service_name": "Раздел 1",
                    "service_composition": "Этап 1\nЭтап 2",
                    "service_composition_editor_state": {
                        "html": "<p><strong>Этап 1</strong></p><p>Этап 2</p>",
                        "plain_text": "Этап 1\nЭтап 2",
                    },
                }
            ],
        )

    def test_typical_service_terms_json_exposes_first_product_default(self):
        product = Product.objects.create(
            short_name="TERM",
            name_en="Terms",
            name_ru="Сроки",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=5,
        )
        TypicalServiceTerm.objects.create(
            product=product,
            preliminary_report_months=Decimal("1.5"),
            final_report_weeks=2,
            position=2,
        )
        TypicalServiceTerm.objects.create(
            product=product,
            preliminary_report_months=Decimal("3.0"),
            final_report_weeks=4,
            position=3,
        )

        response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["typical_service_terms_json"][str(product.pk)],
            {
                "preliminary_report_months": "1.5",
                "final_report_weeks": "2.0",
            },
        )
        self.assertContains(response, 'id="proposal-typical-service-terms-data"', html=False)

    def test_proposal_form_renders_composite_project_name_inputs(self):
        with patch("proposals_app.forms.get_cbr_eur_rate_for_today", return_value=Decimal("95.1111")):
            response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="proposal-project-name-prefix"', html=False)
        self.assertContains(response, 'id="proposal-project-name-suffix"', html=False)
        self.assertContains(response, 'type="hidden" name="proposal_project_name"', html=False)
        self.assertContains(response, 'id="proposal-purpose-prefix"', html=False)
        self.assertContains(response, 'id="proposal-purpose-suffix"', html=False)
        self.assertContains(response, 'type="hidden" name="purpose"', html=False)
        self.assertContains(response, 'id="proposal-commercial-totals-payload"', html=False)
        self.assertContains(response, 'name="status"', html=False)
        self.assertContains(response, 'value="final"', html=False)
        self.assertContains(response, 'id="proposal-report-languages-dropdown"', html=False)
        self.assertContains(response, 'id="proposal-report-languages-toggle"', html=False)
        self.assertContains(response, 'id="proposal-report-language-en"', html=False)
        self.assertContains(response, 'value="русский"', html=False)
        self.assertNotContains(response, 'id="proposal-kind-filter-toggle"', html=False)

    def test_proposal_form_renders_report_term_and_date_inputs_for_type_autofill(self):
        with patch("proposals_app.forms.get_cbr_eur_rate_for_today", return_value=Decimal("95.1111")):
            response = self.client.get(reverse("proposal_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'name="service_term_months"', html=False)
        self.assertContains(response, 'name="preliminary_report_date"', html=False)
        self.assertContains(response, 'name="final_report_term_weeks"', html=False)
        self.assertContains(response, 'name="final_report_date"', html=False)
        self.assertContains(response, 'id="proposal-typical-service-terms-data"', html=False)
        self.assertContains(response, 'js-proposal-report-terms-lock', count=2, html=False)


class ProposalNextcloudWorkspaceHookTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="proposal-nextcloud-staff",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )

    def _payload(self, **overrides):
        payload = {
            "number": "3333",
            "group_member": str(self.group_member.pk),
            "type": str(self.product.pk),
            "name": "Тестовое ТКП",
            "kind": ProposalRegistration.ProposalKind.REGULAR,
            "status": ProposalRegistration.ProposalStatus.FINAL,
            "year": "2026",
            "report_languages": "русский",
        }
        payload.update(overrides)
        return payload

    @patch("ai_app.proposals_app.views.create_proposal_workspace")
    @patch("ai_app.proposals_app.views.is_nextcloud_primary", return_value=True)
    def test_create_view_triggers_nextcloud_workspace_for_new_proposal(self, _mocked_is_nextcloud, mocked_workspace):
        mocked_workspace.return_value = "/Corporate Root/ТКП/2026/333300RU DD Тестовое ТКП"
        response = self.client.post(reverse("proposal_form_create"), self._payload())

        self.assertEqual(response.status_code, 200)
        self.assertEqual(ProposalRegistration.objects.count(), 1)
        mocked_workspace.assert_called_once()
        self.assertEqual(mocked_workspace.call_args.args[0], self.user)
        proposal = ProposalRegistration.objects.first()
        self.assertEqual(mocked_workspace.call_args.args[1].pk, proposal.pk)
        self.assertEqual(proposal.proposal_workspace_disk_path, mocked_workspace.return_value)
        self.assertEqual(proposal.proposal_workspace_public_url, "")

    def test_create_view_rejects_new_proposal_without_year(self):
        response = self.client.post(reverse("proposal_form_create"), self._payload(year=""))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(ProposalRegistration.objects.count(), 0)
        self.assertContains(response, "year: Укажите год.", html=False)

    def test_create_view_does_not_duplicate_bsn_for_synced_customer_owner_and_default_asset(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )

        response = self.client.post(
            reverse("proposal_form_create"),
            self._payload(
                customer='ООО "Заказчик"',
                country=str(country.pk),
                identifier="ОГРН",
                registration_number="1234567890",
                registration_date="01.04.2026",
                asset_owner='ООО "Заказчик"',
                asset_owner_country=str(country.pk),
                asset_owner_identifier="ОГРН",
                asset_owner_registration_number="1234567890",
                asset_owner_registration_date="01.04.2026",
                asset_owner_matches_customer="on",
                assets_payload=json.dumps(
                    [
                        {
                            "short_name": 'ООО "Заказчик"',
                            "country_id": str(country.pk),
                            "country_name": "Россия",
                            "identifier": "ОГРН",
                            "registration_number": "1234567890",
                            "registration_date": "01.04.2026",
                            "selected_identifier_record_id": "",
                            "selected_from_autocomplete": False,
                            "user_edited": False,
                        }
                    ],
                    ensure_ascii=False,
                ),
            ),
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(BusinessEntityRecord.objects.count(), 1)
        entity = BusinessEntityRecord.objects.get()
        self.assertEqual(entity.name, 'ООО "Заказчик"')
        self.assertEqual(entity.source, "[ТКП / Заказчик]")

    def test_create_view_syncs_customer_region_into_registry_chain(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )

        response = self.client.post(
            reverse("proposal_form_create"),
            self._payload(
                customer='ООО "Регион"',
                country=str(country.pk),
                registration_region="Тюменская область",
                identifier="ОГРН",
                registration_number="1234567890",
                registration_date="01.04.2026",
            ),
        )

        self.assertEqual(response.status_code, 200)
        proposal = ProposalRegistration.objects.get()
        self.assertEqual(proposal.registration_region, "Тюменская область")

        identifier_record = BusinessEntityIdentifierRecord.objects.get(
            business_entity__name='ООО "Регион"',
        )
        self.assertEqual(identifier_record.registration_region, "Тюменская область")

        name_record = LegalEntityRecord.objects.get(
            attribute=LegalEntityRecord.ATTRIBUTE_NAME,
            identifier_record=identifier_record,
        )
        self.assertEqual(name_record.registration_region, "Тюменская область")

        address_record = LegalEntityRecord.objects.get(
            attribute=LegalEntityRecord.ATTRIBUTE_LEGAL_ADDRESS,
            identifier_record=identifier_record,
        )
        self.assertEqual(address_record.registration_region, "Тюменская область")

    @patch("ai_app.proposals_app.views.create_proposal_workspace")
    @patch("ai_app.proposals_app.views.is_nextcloud_primary", return_value=True)
    def test_edit_view_does_not_trigger_nextcloud_workspace(self, _mocked_is_nextcloud, mocked_workspace):
        proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=self.group_member,
            type=self.product,
            name="Черновик",
            year=2026,
        )

        response = self.client.post(
            reverse("proposal_form_edit", args=[proposal.pk]),
            self._payload(name="Обновленное ТКП", year=""),
        )

        self.assertEqual(response.status_code, 200)
        mocked_workspace.assert_not_called()
        proposal.refresh_from_db()
        self.assertIsNone(proposal.year)


@override_settings(
    NEXTCLOUD_PROVISIONING_BASE_URL="https://cloud.example.com",
    NEXTCLOUD_PROVISIONING_USERNAME="cloud-admin",
    NEXTCLOUD_PROVISIONING_TOKEN="token",
    NEXTCLOUD_OIDC_PROVIDER_ID=1,
    ROOT_URLCONF="proposals_app.urls",
)
class ProposalDispatchDiskColumnTests(TestCase):
    def setUp(self):
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.nextcloud_root_path = "/Corporate Root"
        settings_obj.save()

        self.user = get_user_model().objects.create_user(
            username="proposal-disk-user",
            email="proposal-disk-user@example.com",
            password="secret",
            is_staff=True,
            is_active=True,
        )
        self.client.force_login(self.user)
        self.user_link = NextcloudUserLink.objects.create(
            user=self.user,
            nextcloud_user_id=f"ncstaff-{self.user.pk}",
            nextcloud_username=f"ncstaff-{self.user.pk}",
            nextcloud_email=self.user.email,
        )
        self.group_member = GroupMember.objects.create(
            short_name="IMC Montan",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        self.proposal = ProposalRegistration.objects.create(
            number=3333,
            group_member=self.group_member,
            type=self.product,
            name="Тестовое ТКП",
            year=2026,
            proposal_workspace_disk_path="/Corporate Root/ТКП/2026/333300RU DD Тестовое ТКП",
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_renders_disk_icon_with_nextcloud_share_target(self, mocked_list_user_shares):
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="77",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Shared/333300RU DD Тестовое ТКП",
            )
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, ">Облако<", html=False)
        self.assertContains(response, 'title="Открыть папку на Nextcloud"', html=False)
        self.assertContains(
            response,
            "/apps/files/files?dir=/Shared/333300RU%20DD%20%D0%A2%D0%B5%D1%81%D1%82%D0%BE%D0%B2%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F",
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_builds_docx_link_from_viewer_share_target(
        self,
        mocked_list_user_shares,
        mocked_list_resources,
    ):
        self.proposal.docx_file_name = "ТКП_333300RU_DD_Тестовое_ТКП.docx"
        self.proposal.docx_file_link = (
            f"{self.proposal.proposal_workspace_disk_path}/{self.proposal.docx_file_name}"
        )
        self.proposal.save(update_fields=["docx_file_name", "docx_file_link"])
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="77-docx",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Shared/333300RU DD Тестовое ТКП",
            )
        }
        mocked_list_resources.return_value = [
            {
                "name": self.proposal.docx_file_name,
                "path": self.proposal.docx_file_link,
                "type": "file",
                "size": 42,
                "modified": "Mon, 30 Mar 2026 10:00:00 GMT",
                "file_id": "2068",
            }
        ]

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_url = "https://cloud.example.com/apps/files/files/2068?dir=/Shared/333300RU%20DD%20%D0%A2%D0%B5%D1%81%D1%82%D0%BE%D0%B2%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F&amp;openfile=true"
        owner_url = f"https://cloud.example.com/apps/files/files?dir={quote(self.proposal.docx_file_link, safe='/')}"
        self.assertContains(response, f'href="{expected_url}"', html=False)
        self.assertNotContains(response, f'href="{owner_url}"', html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_builds_pdf_link_from_viewer_share_target(
        self,
        mocked_list_user_shares,
        mocked_list_resources,
    ):
        self.proposal.pdf_file_name = "ТКП_333300RU_DD_Тестовое_ТКП.pdf"
        self.proposal.pdf_file_link = (
            f"{self.proposal.proposal_workspace_disk_path}/{self.proposal.pdf_file_name}"
        )
        self.proposal.save(update_fields=["pdf_file_name", "pdf_file_link"])
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="77-pdf",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Shared/333300RU DD Тестовое ТКП",
            )
        }
        mocked_list_resources.return_value = [
            {
                "name": self.proposal.pdf_file_name,
                "path": self.proposal.pdf_file_link,
                "type": "file",
                "size": 42,
                "modified": "Mon, 30 Mar 2026 10:00:00 GMT",
                "file_id": "3099",
            }
        ]

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_url = "https://cloud.example.com/apps/files/files/3099?dir=/Shared/333300RU%20DD%20%D0%A2%D0%B5%D1%81%D1%82%D0%BE%D0%B2%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F&amp;openfile=true"
        owner_url = f"https://cloud.example.com/apps/files/files?dir={quote(self.proposal.pdf_file_link, safe='/')}"
        self.assertContains(response, f'href="{expected_url}"', html=False)
        self.assertNotContains(response, f'href="{owner_url}"', html=False)
        self.assertContains(response, "bi-file-pdf-fill", html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_preserves_legacy_media_docx_link(self, _mocked_list_user_shares):
        self.proposal.docx_file_name = "legacy.docx"
        self.proposal.docx_file_link = "/media/proposal_documents/2026/33330RU/legacy.docx"
        self.proposal.save(update_fields=["docx_file_name", "docx_file_link"])

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            'href="/media/proposal_documents/2026/33330RU/legacy.docx"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_keeps_docx_link_when_filename_is_missing(self, _mocked_list_user_shares):
        self.proposal.docx_file_name = ""
        self.proposal.docx_file_link = "/media/proposal_documents/2026/33330RU/legacy.docx"
        self.proposal.save(update_fields=["docx_file_name", "docx_file_link"])

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            'href="/media/proposal_documents/2026/33330RU/legacy.docx"',
            html=False,
        )
        self.assertContains(response, "DOCX файл", html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_keeps_pdf_link_when_filename_is_missing(self, _mocked_list_user_shares):
        self.proposal.pdf_file_name = ""
        self.proposal.pdf_file_link = "/media/proposal_documents/2026/33330RU/legacy.pdf"
        self.proposal.save(update_fields=["pdf_file_name", "pdf_file_link"])

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            'href="/media/proposal_documents/2026/33330RU/legacy.pdf"',
            html=False,
        )
        self.assertContains(response, "PDF файл", html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_generated_docx_download_redirects_to_cloud_file_when_local_copy_missing(
        self,
        mocked_list_user_shares,
        mocked_list_resources,
    ):
        self.proposal.docx_file_name = "ТКП_333300RU_DD_Тестовое_ТКП.docx"
        self.proposal.docx_file_link = (
            f"{self.proposal.proposal_workspace_disk_path}/{self.proposal.docx_file_name}"
        )
        self.proposal.save(update_fields=["docx_file_name", "docx_file_link"])
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="77-download",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Shared/333300RU DD Тестовое ТКП",
            )
        }
        mocked_list_resources.return_value = [
            {
                "name": self.proposal.docx_file_name,
                "path": self.proposal.docx_file_link,
                "type": "file",
                "size": 42,
                "modified": "Mon, 30 Mar 2026 10:00:00 GMT",
                "file_id": "2068",
            }
        ]

        response = self.client.get(reverse("proposal_generated_docx_download", args=[self.proposal.pk]))

        expected_url = "https://cloud.example.com/apps/files/files/2068?dir=/Shared/333300RU%20DD%20%D0%A2%D0%B5%D1%81%D1%82%D0%BE%D0%B2%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F&openfile=true"
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], expected_url)

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_uses_editor_files_url_instead_of_saved_public_link(
        self,
        mocked_list_user_shares,
    ):
        self.proposal.proposal_workspace_public_url = "https://cloud.example.com/s/readonly"
        self.proposal.save(update_fields=["proposal_workspace_public_url"])
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="82",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Shared/333300RU DD Тестовое ТКП",
            )
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'href="https://cloud.example.com/s/readonly"', html=False)
        self.assertContains(
            response,
            "/apps/files/files?dir=/Shared/333300RU%20DD%20%D0%A2%D0%B5%D1%81%D1%82%D0%BE%D0%B2%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F",
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_falls_back_to_direct_share_lookup_for_new_workspace(
        self,
        _mocked_list_user_shares,
        mocked_get_user_share,
    ):
        mocked_get_user_share.return_value = NextcloudShare(
            share_id="78",
            path=self.proposal.proposal_workspace_disk_path,
            share_with=self.user_link.nextcloud_user_id,
            permissions=15,
            target_path="/Shared/333300RU DD Тестовое ТКП",
        )

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "/apps/files/files?dir=/Shared/333300RU%20DD%20%D0%A2%D0%B5%D1%81%D1%82%D0%BE%D0%B2%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F",
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_resolves_new_workspace_from_parent_share_lookup_when_direct_share_is_missing(
        self,
        _mocked_list_user_shares,
        mocked_get_user_share,
    ):
        def _share_lookup(_owner_user_id, lookup_path, _share_with_user_id):
            normalized_lookup_path = str(lookup_path or "").strip()
            if normalized_lookup_path == self.proposal.proposal_workspace_disk_path:
                return None
            if normalized_lookup_path == "/Corporate Root/ТКП":
                return NextcloudShare(
                    share_id="78-parent",
                    path="/Corporate Root/ТКП",
                    share_with=self.user_link.nextcloud_user_id,
                    permissions=15,
                    target_path="/ТКП",
                )
            return None

        mocked_get_user_share.side_effect = _share_lookup

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_suffix = self.proposal.proposal_workspace_disk_path.split("/ТКП/", 1)[1]
        expected_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(f'/ТКП/{expected_suffix}', safe='/')}"
        )
        self.assertContains(
            response,
            f'href="{expected_url}"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_resolves_editor_cloud_link_from_parent_shared_folder_for_director(
        self,
        mocked_list_user_shares,
    ):
        Employee.objects.create(user=self.user, role="Директор")
        mocked_list_user_shares.return_value = {
            "/Corporate Root/ТКП": NextcloudShare(
                share_id="79",
                path="/Corporate Root/ТКП",
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Shared/ТКП",
            )
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        director_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/Shared/ТКП/2026/333300RU DD Тестовое ТКП', safe='/')}"
        )
        self.assertContains(
            response,
            f'href="{director_url}"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_strips_owner_root_from_direct_share_target_path_for_viewer(
        self,
        mocked_list_user_shares,
    ):
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="79-direct-owner-path",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Corporate Root/ТКП/2026/333300RU DD Тестовое ТКП",
            )
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        stripped_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/ТКП/2026/333300RU DD Тестовое ТКП', safe='/')}"
        )
        owner_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/Corporate Root/ТКП/2026/333300RU DD Тестовое ТКП', safe='/')}"
        )
        self.assertContains(
            response,
            f'href="{stripped_url}"',
            html=False,
        )
        self.assertNotContains(
            response,
            f'href="{owner_url}"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_resolves_editor_cloud_link_from_parent_share_without_owner_root_prefix(
        self,
        mocked_list_user_shares,
    ):
        mocked_list_user_shares.return_value = {
            "/ТКП": NextcloudShare(
                share_id="83",
                path="/ТКП",
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/ТКП",
            )
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_suffix = self.proposal.proposal_workspace_disk_path.split("/ТКП/", 1)[1]
        expected_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(f'/ТКП/{expected_suffix}', safe='/')}"
        )
        self.assertContains(
            response,
            f'href="{expected_url}"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_resolves_parent_shared_folder_when_direct_share_has_no_target_path(
        self,
        mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="80",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="",
            ),
            "/Corporate Root/ТКП": NextcloudShare(
                share_id="81",
                path="/Corporate Root/ТКП",
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/Shared/ТКП",
            ),
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_suffix = self.proposal.proposal_workspace_disk_path.split("/Corporate Root/ТКП/", 1)[1]
        expected_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(f'/Shared/ТКП/{expected_suffix}', safe='/')}"
        )
        self.assertContains(
            response,
            expected_url,
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_prefers_earlier_parent_share_when_same_depth_candidates_exist(
        self,
        mocked_list_user_shares,
    ):
        mocked_list_user_shares.return_value = {
            "/2026": NextcloudShare(
                share_id="84",
                path="/2026",
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/2026",
            ),
            "/ТКП": NextcloudShare(
                share_id="85",
                path="/ТКП",
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/ТКП",
            ),
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_suffix = self.proposal.proposal_workspace_disk_path.split("/ТКП/", 1)[1]
        competing_suffix = self.proposal.proposal_workspace_disk_path.rsplit("/2026/", 1)[1]
        expected_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(f'/ТКП/{expected_suffix}', safe='/')}"
        )
        competing_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(f'/2026/{competing_suffix}', safe='/')}"
        )
        self.assertContains(
            response,
            f'href="{expected_url}"',
            html=False,
        )
        self.assertNotContains(
            response,
            f'href="{competing_url}"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share")
    def test_proposals_partial_falls_back_to_direct_cloud_link_when_viewer_has_no_nextcloud_link(
        self,
        mocked_ensure_public_link_share,
    ):
        self.user_link.delete()

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        owner_url = f"https://cloud.example.com/apps/files/files?dir={quote(self.proposal.proposal_workspace_disk_path, safe='/')}"
        self.assertContains(
            response,
            f'href="{owner_url}"',
            html=False,
        )
        mocked_ensure_public_link_share.assert_not_called()

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_uses_stored_target_path_when_share_api_is_silent(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.proposal.proposal_workspace_target_path = "/Shared/333300RU DD Тестовое ТКП"
        self.proposal.save(update_fields=["proposal_workspace_target_path"])

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/Shared/333300RU DD Тестовое ТКП', safe='/')}"
        )
        guessed_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/ТКП/2026/333300RU DD Тестовое ТКП', safe='/')}"
        )
        owner_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(self.proposal.proposal_workspace_disk_path, safe='/')}"
        )
        self.assertContains(response, f'href="{expected_url}"', html=False)
        self.assertNotContains(response, f'href="{guessed_url}"', html=False)
        self.assertNotContains(response, f'href="{owner_url}"', html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_uses_cached_session_target_path_when_share_api_lags(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        session = self.client.session
        session[PROPOSAL_NEXTCLOUD_TARGETS_SESSION_KEY] = {
            self.proposal.proposal_workspace_disk_path: "/333300RU DD Тестовое ТКП",
        }
        session.save()

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        cached_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/333300RU DD Тестовое ТКП', safe='/')}"
        )
        stripped_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/ТКП/2026/333300RU DD Тестовое ТКП', safe='/')}"
        )
        self.assertContains(
            response,
            f'href="{cached_url}"',
            html=False,
        )
        self.assertNotContains(
            response,
            f'href="{stripped_url}"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_keeps_gray_icon_when_viewer_target_cannot_be_resolved(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        owner_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(self.proposal.proposal_workspace_disk_path, safe='/')}"
        )
        guessed_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote('/ТКП/2026/333300RU DD Тестовое ТКП', safe='/')}"
        )
        self.assertNotContains(response, f'href="{owner_url}"', html=False)
        self.assertNotContains(response, f'href="{guessed_url}"', html=False)
        self.assertContains(response, 'style="color: #ccc;"', html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_falls_back_to_saved_public_workspace_url_for_viewer_with_nextcloud(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.proposal.proposal_workspace_public_url = "https://cloud.example.com/s/saved-proposal-folder"
        self.proposal.save(update_fields=["proposal_workspace_public_url"])
        owner_url = f"https://cloud.example.com/apps/files/files?dir={quote(self.proposal.proposal_workspace_disk_path, safe='/')}"

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            'href="https://cloud.example.com/s/saved-proposal-folder"',
            html=False,
        )
        self.assertNotContains(
            response,
            f'href="{owner_url}"',
            html=False,
        )

    def test_proposals_partial_preserves_saved_public_workspace_url_in_mixed_rows_without_viewer_link(self):
        self.proposal.proposal_workspace_public_url = "https://cloud.example.com/s/saved-proposal-folder"
        self.proposal.save(update_fields=["proposal_workspace_public_url"])
        ProposalRegistration.objects.create(
            number=3334,
            group_member=self.group_member,
            type=self.product,
            name="Второе ТКП",
            year=2026,
            proposal_workspace_disk_path="/Corporate Root/ТКП/2026/333400RU DD Второе ТКП",
        )
        self.user_link.delete()

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            'href="https://cloud.example.com/s/saved-proposal-folder"',
            html=False,
        )
        self.assertContains(
            response,
            'href="https://cloud.example.com/apps/files/files?dir=/Corporate%20Root/%D0%A2%D0%9A%D0%9F/2026/333400RU%20DD%20%D0%92%D1%82%D0%BE%D1%80%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_builds_direct_cloud_link_from_expected_workspace_path_for_legacy_rows(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.user_link.delete()
        self.proposal.proposal_workspace_disk_path = ""
        self.proposal.save(update_fields=["proposal_workspace_disk_path"])

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        expected_workspace_path = f"/Corporate Root/ТКП/2026/{self.proposal.short_uid} DD Тестовое ТКП"
        owner_url = (
            "https://cloud.example.com/apps/files/files?dir="
            f"{quote(expected_workspace_path, safe='/')}"
        )
        self.assertContains(
            response,
            f'href="{owner_url}"',
            html=False,
        )

    def test_proposals_partial_preserves_saved_public_workspace_url_when_path_is_missing(self):
        self.user_link.delete()
        self.proposal.year = None
        self.proposal.proposal_workspace_disk_path = ""
        self.proposal.proposal_workspace_public_url = "https://cloud.example.com/s/legacy-proposal-folder"
        self.proposal.save(update_fields=["year", "proposal_workspace_disk_path", "proposal_workspace_public_url"])

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            'href="https://cloud.example.com/s/legacy-proposal-folder"',
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_uses_stored_file_id_editor_url_when_share_api_silent(
        self,
        _mocked_list_user_shares,
        mocked_list_resources,
        _mocked_get_user_share,
    ):
        self.proposal.proposal_workspace_target_path = "/Shared/333300RU DD Тестовое ТКП"
        self.proposal.docx_file_name = "ТКП_333300RU_DD_Тестовое_ТКП.docx"
        self.proposal.docx_file_link = (
            f"{self.proposal.proposal_workspace_disk_path}/{self.proposal.docx_file_name}"
        )
        self.proposal.docx_file_id = "2068"
        self.proposal.save(
            update_fields=[
                "proposal_workspace_target_path",
                "docx_file_name",
                "docx_file_link",
                "docx_file_id",
            ]
        )

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        # No extra PROPFIND when the file id is already persisted in the DB.
        mocked_list_resources.assert_not_called()
        expected_url = (
            "https://cloud.example.com/apps/files/files/2068?dir="
            f"{quote('/Shared/333300RU DD Тестовое ТКП', safe='/')}"
            "&amp;openfile=true"
        )
        self.assertContains(response, f'href="{expected_url}"', html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_proposals_partial_falls_back_to_file_redirect_url_when_target_unknown(
        self,
        _mocked_list_user_shares,
        mocked_list_resources,
        _mocked_get_user_share,
    ):
        # Target path is unknown (API silent, no DB value) but we still have a
        # stable file id; Nextcloud's /f/<id> redirect works regardless of the
        # viewer's mountpoint.
        self.proposal.docx_file_name = "ТКП_333300RU_DD_Тестовое_ТКП.docx"
        self.proposal.docx_file_link = (
            f"{self.proposal.proposal_workspace_disk_path}/{self.proposal.docx_file_name}"
        )
        self.proposal.docx_file_id = "2068"
        self.proposal.save(
            update_fields=["docx_file_name", "docx_file_link", "docx_file_id"]
        )

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        mocked_list_resources.assert_not_called()
        self.assertContains(response, 'href="https://cloud.example.com/f/2068"', html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("proposals_app.views.create_proposal_workspace")
    @patch("proposals_app.views.is_nextcloud_primary", return_value=True)
    def test_maybe_create_workspace_persists_viewer_target_path(
        self,
        _mocked_is_nextcloud,
        mocked_workspace,
        mocked_list_resources,
    ):
        from types import SimpleNamespace

        from proposals_app.views import _maybe_create_nextcloud_proposal_workspace

        class _FakeSession(dict):
            """dict with attribute support so `session.modified = True` works."""

            modified = False

        mocked_workspace.return_value = (
            "/Corporate Root/ТКП/2026/333500RU DD Новое ТКП",
            "/Shared/333500RU DD Новое ТКП",
        )
        mocked_list_resources.return_value = [
            {
                "name": "333500RU DD Новое ТКП",
                "path": "/Corporate Root/ТКП/2026/333500RU DD Новое ТКП",
                "type": "folder",
                "file_id": "9100",
            }
        ]
        request = SimpleNamespace(user=self.user, session=_FakeSession())

        self.proposal.proposal_workspace_disk_path = ""
        self.proposal.proposal_workspace_target_path = ""
        self.proposal.proposal_workspace_public_url = "https://stale.example.com/s/readonly"
        self.proposal.save(
            update_fields=[
                "proposal_workspace_disk_path",
                "proposal_workspace_target_path",
                "proposal_workspace_public_url",
            ]
        )

        _maybe_create_nextcloud_proposal_workspace(request, self.proposal)

        self.proposal.refresh_from_db()
        self.assertEqual(
            self.proposal.proposal_workspace_disk_path,
            "/Corporate Root/ТКП/2026/333500RU DD Новое ТКП",
        )
        self.assertEqual(
            self.proposal.proposal_workspace_target_path,
            "/Shared/333500RU DD Новое ТКП",
        )
        self.assertEqual(self.proposal.proposal_workspace_public_url, "")
        # Persisted folder id is used by the table to build a robust
        # ``/f/<file_id>`` URL regardless of how the viewer has the folder
        # mounted in their file tree.
        self.assertEqual(self.proposal.proposal_workspace_file_id, "9100")
        self.assertTrue(request.session.modified)
        self.assertEqual(
            request.session[PROPOSAL_NEXTCLOUD_TARGETS_SESSION_KEY],
            {"/Corporate Root/ТКП/2026/333500RU DD Новое ТКП": "/Shared/333500RU DD Новое ТКП"},
        )

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_proposals_partial_prefers_file_redirect_url_when_workspace_file_id_persisted(
        self,
        mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        # Reproduces the real-world scenario: the direct share returns a
        # ``target_path`` that collides with a broader parent share already
        # mounted in the viewer's tree, leaving the ``?dir=`` URL pointing at
        # a non-existent top-level mount. The file-id redirect URL follows
        # the viewer's actual mount and always resolves.
        self.proposal.proposal_workspace_file_id = "9100"
        self.proposal.save(update_fields=["proposal_workspace_file_id"])
        mocked_list_user_shares.return_value = {
            self.proposal.proposal_workspace_disk_path: NextcloudShare(
                share_id="90-direct",
                path=self.proposal.proposal_workspace_disk_path,
                share_with=self.user_link.nextcloud_user_id,
                permissions=15,
                target_path="/333300RU DD Тестовое ТКП",
            ),
        }

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'href="https://cloud.example.com/f/9100"', html=False)
        self.assertNotContains(
            response,
            "/apps/files/files?dir=/333300RU%20DD%20%D0%A2%D0%B5%D1%81%D1%82%D0%BE%D0%B2%D0%BE%D0%B5%20%D0%A2%D0%9A%D0%9F",
            html=False,
        )


class ProposalAccessTests(TestCase):
    def test_proposals_partial_requires_staff(self):
        user = get_user_model().objects.create_user(
            username="proposal-nonstaff",
            password="secret",
            is_staff=False,
        )
        self.client.force_login(user)

        response = self.client.get(reverse("proposals_partial"))

        self.assertEqual(response.status_code, 302)
