import json
import os
import shutil
import tempfile
import uuid
from datetime import date, datetime, timedelta
from decimal import Decimal
from io import BytesIO
from urllib.parse import quote
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Group
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from classifiers_app.models import OKSMCountry
from contacts_app.models import CitizenshipRecord, PersonRecord
from core.models import CloudStorageSettings
from contracts_app.docx_processor import process_template
from contracts_app.forms import ContractProjectRegistrationForm, ContractTemplateForm
from contracts_app.models import (
    ContractProjectRegistration,
    ContractProjectRegistrationProduct,
    ContractReturnComment,
    ContractSubject,
    ContractTemplate,
    ContractVariable,
)
from contracts_app.variable_resolver import resolve_variables
from experts_app.models import ExpertContractDetails, ExpertProfile, ExpertProfileSpecialty, ExpertSpecialty
from group_app.models import GroupMember, OrgUnit
from nextcloud_app.api import NextcloudApiError, NextcloudShare
from nextcloud_app.models import NextcloudUserLink
from notifications_app.models import Notification, NotificationPerformerLink
from policy_app.models import (
    ADMIN_GROUP,
    DEPARTMENT_HEAD_GROUP,
    EXPERT_GROUP,
    LAWYER_GROUP,
    Product,
    SectionStructure,
    ServiceGoalReport,
    TypicalSection,
    TypicalSectionSpecialty,
    TypicalServiceComposition,
)
from projects_app.models import Performer, ProjectRegistration, ProjectRegistrationProduct
from proposals_app.models import ProposalRegistration
from users_app.forms import FREELANCER_LABEL
from users_app.models import Employee


@override_settings(
    NEXTCLOUD_PROVISIONING_BASE_URL="https://cloud.example.com",
    NEXTCLOUD_PROVISIONING_USERNAME="cloud-admin",
    NEXTCLOUD_PROVISIONING_TOKEN="token",
    NEXTCLOUD_OIDC_PROVIDER_ID=1,
)
class ContractsCloudLabelTests(TestCase):
    def setUp(self):
        self.media_root = tempfile.mkdtemp(prefix="contracts-tests-")
        self.addCleanup(lambda: shutil.rmtree(self.media_root, ignore_errors=True))

        self.user = get_user_model().objects.create_user(
            username="contracts-admin",
            password="secret",
            is_staff=True,
        )
        admin_group, _ = Group.objects.get_or_create(name=ADMIN_GROUP)
        self.user.groups.add(admin_group)
        self.client.force_login(self.user)

        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.nextcloud_root_path = "/Corporate Root"
        settings_obj.save()

        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.group_member = GroupMember.objects.create(
            short_name="Россия",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.project = ProjectRegistration.objects.create(
            number=7001,
            group_member=self.group_member,
            type=self.product,
            name="Договорный проект",
            year=2026,
        )
        self.proposal = ProposalRegistration.objects.create(
            number=7001,
            group_member=self.group_member,
            type=self.product,
            name="ТКП проект",
            year=2026,
        )
        self.project.proposal_registration = self.proposal
        self.project.save(update_fields=["proposal_registration"])
        self.employee_user = get_user_model().objects.create_user(
            username="expert@example.com",
            email="expert@example.com",
            password="secret",
            is_staff=True,
        )
        self.employee = Employee.objects.create(
            user=self.employee_user,
            patronymic="Иванович",
            employment=FREELANCER_LABEL,
        )
        expert_group, _ = Group.objects.get_or_create(name=EXPERT_GROUP)
        self.employee_user.groups.add(expert_group)
        self.employee_link = NextcloudUserLink.objects.create(
            user=self.employee_user,
            nextcloud_user_id="nc-expert",
            nextcloud_username="nc-expert",
            nextcloud_email=self.employee_user.email,
        )
        self.performer = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor="Иванов Иван Иванович",
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            contract_batch_id=uuid.uuid4(),
            contract_file="Договор 7001_Иванов ИИ.docx",
            contract_project_link="https://cloud.example.com/s/contract-docx",
            contract_project_disk_folder="/Corporate Root/2026/Project/09 Договоры/000 Иванов ИИ",
        )

    def test_contracts_partial_renders_contract_projects_before_contract_conclusion(self):
        self.performer.contract_signing_note = "Скрытая заметка отправки договора"
        self.performer.save(update_fields=["contract_signing_note"])

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Составление проекта договора")
        self.assertContains(response, "Отправка проекта договора")
        self.assertContains(response, 'id="contract-dispatch-table"', html=False)
        self.assertContains(response, "Облако")
        self.assertContains(response, "Наименование файла DOCX")
        self.assertContains(response, "Наименование файла PDF")
        self.assertContains(response, "Договор 7001_Иванов ИИ.docx")
        self.assertContains(response, "https://cloud.example.com/s/contract-docx")
        self.assertContains(response, "Создать проект договора")
        self.assertContains(response, "Подписать проект договора")
        self.assertContains(response, "Отправить проект договора")
        self.assertContains(response, "Системная электронная почта")
        self.assertContains(response, "Подключенная электронная почта")
        self.assertContains(
            response,
            "На Nextcloud будут созданы папки договоров в структуре",
            html=False,
        )
        self.assertContains(response, 'id="contract-project-filter-toggle"', html=False)
        self.assertContains(response, 'id="signing-project-filter-toggle"', html=False)
        self.assertContains(response, 'class="form-check-input js-signing-filter"', html=False)
        self.assertContains(response, 'data-contract-status=', html=False)
        self.assertContains(response, 'data-request-url="%s"' % reverse("contract_request"), html=False)
        self.assertContains(
            response,
            'data-create-contract-url="%s"' % reverse("create_contract_project"),
            html=False,
        )
        self.assertContains(
            response,
            'data-sign-contract-url="%s"' % reverse("sign_contract_documents"),
            html=False,
        )
        self.assertContains(response, "Договорный проект")

        content = response.content.decode("utf-8")
        contract_drafting_table = content[
            content.index('id="contract-drafting-table"'):
            content.index('id="contract-dispatch-table"')
        ]
        self.assertNotIn("<th>Грейд</th>", contract_drafting_table)
        self.assertNotIn("col-grade", contract_drafting_table)
        contract_dispatch_table_start = content.index('id="contract-dispatch-table"')
        contract_dispatch_table = content[
            contract_dispatch_table_start:
            content.index("</table>", contract_dispatch_table_start)
        ]
        self.assertIn('data-contract-dispatch-ready="0"', contract_dispatch_table)
        self.assertNotIn("<th>Подписание</th>", contract_dispatch_table)
        self.assertNotIn(
            '<td class="cell-detail-val">Скрытая заметка отправки договора</td>',
            contract_dispatch_table,
        )
        self.assertLess(
            content.index('id="contracts-project-filter-toggle"'),
            content.index('id="contract-conclusion-section"'),
        )
        contracts_table = content[
            content.index('class="table table-sm align-middle contracts-table mb-0"'):
            content.index('id="contracts-edit-btn"')
        ]
        self.assertNotIn('<th class="text-nowrap">Ссылка</th>', contracts_table)
        self.assertLess(
            contracts_table.index('<th class="text-nowrap">Номер договора</th>'),
            contracts_table.index('>Цена</th>'),
        )
        self.assertLess(
            contracts_table.index('>Цена</th>'),
            contracts_table.index('<th class="text-nowrap">Дата договора</th>'),
        )
        self.assertLess(
            content.index("Составление проекта договора"),
            content.index("Отправка проекта договора"),
        )
        self.assertLess(
            content.index("Отправка проекта договора"),
            content.index("Подписание договора"),
        )
        contract_request_modal = content[
            content.index('id="contract-request-modal"'):
            content.index('class="modal-footer"', content.index('id="contract-request-modal"'))
        ]
        self.assertIn('value="system_email"', contract_request_modal)
        self.assertIn('value="connected_email"', contract_request_modal)
        self.assertNotIn("Дата отправки запроса", contract_request_modal)
        self.assertNotIn('id="contract-request-sent-at"', contract_request_modal)

    def test_contracts_partial_displays_multi_stage_contract_batch_as_single_row(self):
        product_b = Product.objects.create(
            short_name="TDD",
            name_en="Technical Due Diligence",
            name_ru="ТДД",
            display_name="Technical Due Diligence",
        )
        self.product.short_name = "RFR"
        self.product.display_name = "Due Diligence"
        self.product.save(update_fields=["short_name", "display_name"])
        self.project.name = "Тест 55"
        self.project.deadline = date(2026, 5, 10)
        self.project.save(update_fields=["name", "deadline"])
        ProjectRegistrationProduct.objects.create(registration=self.project, product=self.product, rank=1)
        second_project = ProjectRegistration.objects.create(
            number=self.project.number,
            type=product_b,
            name="Тест 55",
            year=2026,
            deadline=date(2026, 6, 15),
        )
        ProjectRegistrationProduct.objects.create(registration=second_project, product=product_b, rank=1)
        self.project.refresh_from_db()
        second_project.refresh_from_db()
        participation_batch_id = uuid.uuid4()
        self.performer.participation_batch_id = participation_batch_id
        self.performer.save(update_fields=["participation_batch_id"])
        Performer.objects.create(
            registration=second_project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            contract_batch_id=uuid.uuid4(),
            contract_number=self.performer.contract_number,
            contract_file=self.performer.contract_file,
            contract_project_disk_folder=self.performer.contract_project_disk_folder,
            agreed_amount=100,
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        correction_table_start = content.index('class="table table-sm align-middle contracts-table mb-0"')
        correction_table = content[
            correction_table_start:
            content.index("</table>", correction_table_start)
        ]
        self.assertEqual(correction_table.count("<tr data-project-id="), 1)
        self.assertIn('<th class="text-nowrap">Номер</th>', correction_table)
        self.assertIn('<th class="text-nowrap">Этап</th>', correction_table)
        self.assertIn("RFR-TDD", correction_table)
        self.assertIn("15.06.2026", correction_table)
        self.assertIn("70010RU", correction_table)
        self.assertNotIn(f">{self.project.short_uid}<", correction_table)
        self.assertNotIn(f">{second_project.short_uid}<", correction_table)

        drafting_table = content[
            content.index('id="contract-drafting-table"'):
            content.index('id="contract-dispatch-table"')
        ]
        self.assertEqual(drafting_table.count("<tr data-project-id="), 2)
        self.assertIn("RFR", drafting_table)
        self.assertIn("TDD", drafting_table)
        self.assertIn("70010RU", drafting_table)
        self.assertIn(f">{self.project.short_uid}<", drafting_table)
        self.assertIn(f">{second_project.short_uid}<", drafting_table)

        dispatch_table_start = content.index('id="contract-dispatch-table"')
        dispatch_table = content[
            dispatch_table_start:
            content.index("</table>", dispatch_table_start)
        ]
        self.assertEqual(dispatch_table.count("<tr data-project-id="), 1)
        self.assertIn("RFR-TDD", dispatch_table)
        self.assertIn("70010RU", dispatch_table)
        self.assertNotIn(f">{self.project.short_uid}<", dispatch_table)
        self.assertNotIn(f">{second_project.short_uid}<", dispatch_table)

        signing_table_start = content.index('class="table table-sm align-middle signing-table"')
        signing_table = content[
            signing_table_start:
            content.index("</table>", signing_table_start)
        ]
        self.assertEqual(signing_table.count("<tr data-project-id="), 1)
        self.assertIn("RFR-TDD", signing_table)
        self.assertIn("70010RU", signing_table)
        self.assertNotIn(f">{self.project.short_uid}<", signing_table)
        self.assertNotIn(f">{second_project.short_uid}<", signing_table)

        form_response = self.client.get(reverse("contracts_edit", args=[self.performer.pk]))
        self.assertEqual(form_response.status_code, 200)
        form_content = form_response.content.decode("utf-8")
        self.assertContains(form_response, '<label class="form-label">Номер</label>', html=False)
        self.assertContains(form_response, '<label class="form-label">Этап</label>', html=False)
        self.assertContains(form_response, '<label class="form-label">Продукт</label>', html=False)
        self.assertNotContains(form_response, '<label class="form-label">Проект</label>', html=False)
        self.assertIn('value="70010RU"', form_content)
        self.assertIn(f'value="{self.project.short_uid}"', form_content)
        self.assertIn(f'value="{second_project.short_uid}"', form_content)
        self.assertIn('value="RFR Due Diligence"', form_content)
        self.assertIn('value="TDD Technical Due Diligence"', form_content)
        self.assertIn('data-deadline="2026-06-15"', form_content)

    def test_participation_batch_contract_rows_are_separate_by_executor(self):
        participation_batch_id = uuid.uuid4()
        self.performer.participation_batch_id = participation_batch_id
        self.performer.agreed_amount = 100
        self.performer.save(update_fields=["participation_batch_id", "agreed_amount"])
        other_executor = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor="Петров Петр Петрович",
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            contract_batch_id=uuid.uuid4(),
            contract_number="OTHER-1",
            contract_date=date(2026, 5, 8),
            contract_file="other.docx",
            agreed_amount=200,
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        contracts_table = content[
            content.index('class="table table-sm align-middle contracts-table mb-0"'):
            content.index('id="contracts-edit-btn"')
        ]
        self.assertEqual(contracts_table.count("<tr data-project-id="), 2)
        self.assertIn("Иванов И.И.", contracts_table)
        self.assertIn("Петров П.П.", contracts_table)
        self.assertIn("100", contracts_table)
        self.assertIn("200", contracts_table)

        edit_response = self.client.post(
            reverse("contracts_edit", args=[self.performer.pk]),
            {
                "contract_number": "CUSTOM-42",
                "contract_date": "2026-05-07",
                "prepayment": "30",
                "contract_file": "custom.docx",
            },
        )

        self.assertEqual(edit_response.status_code, 200)
        other_executor.refresh_from_db()
        self.assertEqual(other_executor.contract_number, "OTHER-1")
        self.assertEqual(other_executor.contract_date, date(2026, 5, 8))
        self.assertIsNone(other_executor.prepayment)
        self.assertEqual(other_executor.contract_file, "other.docx")

    def test_contract_drafting_displays_multi_stage_participation_batch_as_detail_rows(self):
        product_b = Product.objects.create(
            short_name="TDD",
            name_en="Technical Due Diligence",
            name_ru="ТДД",
            display_name="Technical Due Diligence",
        )
        self.product.short_name = "RFR"
        self.product.display_name = "Due Diligence"
        self.product.save(update_fields=["short_name", "display_name"])
        self.project.name = "Тест 56"
        self.project.deadline = date(2026, 5, 10)
        self.project.save(update_fields=["name", "deadline"])
        ProjectRegistrationProduct.objects.create(registration=self.project, product=self.product, rank=1)
        second_project = ProjectRegistration.objects.create(
            number=self.project.number,
            type=product_b,
            name="Тест 56",
            year=2026,
            deadline=date(2026, 6, 15),
        )
        ProjectRegistrationProduct.objects.create(registration=second_project, product=product_b, rank=1)
        self.project.refresh_from_db()
        second_project.refresh_from_db()
        participation_batch_id = uuid.uuid4()
        self.performer.contract_batch_id = None
        self.performer.contract_file = ""
        self.performer.contract_project_link = ""
        self.performer.contract_project_disk_folder = ""
        self.performer.participation_batch_id = participation_batch_id
        self.performer.save(update_fields=[
            "contract_batch_id",
            "contract_file",
            "contract_project_link",
            "contract_project_disk_folder",
            "participation_batch_id",
        ])
        Performer.objects.create(
            registration=second_project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            agreed_amount=100,
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        drafting_table = content[
            content.index('id="contract-drafting-table"'):
            content.index('id="contract-dispatch-table"')
        ]
        self.assertEqual(drafting_table.count("<tr data-project-id="), 2)
        self.assertIn("RFR", drafting_table)
        self.assertIn("TDD", drafting_table)
        self.assertIn("70010RU", drafting_table)
        self.assertIn(f">{self.project.short_uid}<", drafting_table)
        self.assertIn(f">{second_project.short_uid}<", drafting_table)
        dispatch_table_start = content.index('id="contract-dispatch-table"')
        dispatch_table = content[
            dispatch_table_start:
            content.index("</table>", dispatch_table_start)
        ]
        self.assertEqual(dispatch_table.count("<tr data-project-id="), 1)
        self.assertIn("RFR-TDD", dispatch_table)

    def test_contract_conclusion_column_registry_excludes_grade(self):
        from core.column_registry import get_column_choices

        choices = get_column_choices("projects", "contract_conclusion")

        self.assertNotIn(("grade", "Грейд"), choices)

    def test_contract_conclusion_lists_unsent_addendum_batch_before_sent_contract(self):
        self.performer.contract_sent_at = timezone.now()
        self.performer.contract_signing_date = timezone.now()
        self.performer.contract_project_folder_link = "https://cloud.example.com/s/existing-contract-folder"
        self.performer.save(update_fields=[
            "contract_sent_at",
            "contract_signing_date",
            "contract_project_folder_link",
        ])
        addendum_batch_id = uuid.uuid4()
        addendum = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            contract_batch_id=addendum_batch_id,
            contract_is_addendum=True,
            contract_addendum_number=1,
            contract_number="IMCM/7001-ИИ/05-26 ДС1",
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        drafting_table = content[
            content.index('id="contract-drafting-table"'):
            content.index('id="contract-dispatch-table"')
        ]
        dispatch_table = content[
            content.index('id="contract-dispatch-table"'):
            content.index("</table>", content.index('id="contract-dispatch-table"'))
        ]
        self.assertLess(
            drafting_table.index(f'id="contract-drafting-sel-{addendum.pk}"'),
            drafting_table.index(f'id="contract-drafting-sel-{self.performer.pk}"'),
        )
        self.assertLess(
            dispatch_table.index(f'id="contract-dispatch-sel-{addendum.pk}"'),
            dispatch_table.index(f'id="contract-dispatch-sel-{self.performer.pk}"'),
        )
        addendum_checkbox = drafting_table[
            drafting_table.index(f'id="contract-drafting-sel-{addendum.pk}"'):
            drafting_table.index("aria-label", drafting_table.index(f'id="contract-drafting-sel-{addendum.pk}"'))
        ]
        self.assertIn('data-contract-sent="0"', addendum_checkbox)
        self.assertNotIn("disabled", addendum_checkbox)
        self.assertIn('data-contract-batch-id="%s"' % addendum_batch_id, drafting_table)
        self.assertIn("ДС1", drafting_table)
        self.assertIn("IMCM/7001-ИИ/05-26 ДС1", drafting_table)
        addendum_row = drafting_table[
            drafting_table.rindex("<tr", 0, drafting_table.index(f'id="contract-drafting-sel-{addendum.pk}"')):
            drafting_table.index("</tr>", drafting_table.index(f'id="contract-drafting-sel-{addendum.pk}"'))
        ]
        self.assertIn('href="https://cloud.example.com/s/existing-contract-folder"', addendum_row)

    def test_contract_dispatch_marks_ready_only_when_docx_and_pdf_exist(self):
        self.performer.contract_pdf_file = "Договор 7001_Иванов ИИ.pdf"
        self.performer.save(update_fields=["contract_pdf_file"])

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        dispatch_table = content[
            content.index('id="contract-dispatch-table"'):
            content.index("</table>", content.index('id="contract-dispatch-table"'))
        ]
        self.assertIn('data-contract-dispatch-ready="1"', dispatch_table)

    def test_contract_adjustment_prefills_addendum_file_name_with_kind_suffix(self):
        from contracts_app.services import prefill_contract_adjustment_fields

        self.performer.contract_project_created = True
        self.performer.save(update_fields=["contract_project_created"])
        addendum = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
        )

        prefill_contract_adjustment_fields([addendum.pk], confirmed_at=timezone.now())

        addendum.refresh_from_db()
        self.assertTrue(addendum.contract_is_addendum)
        self.assertEqual(addendum.contract_addendum_number, 1)
        self.assertEqual(
            addendum.contract_file,
            f"Договор {self.project.short_uid}_Иванов ИИ_ДС1.docx",
        )

    def test_contract_adjustment_prefills_multi_stage_file_name_with_number_display(self):
        from contracts_app.services import contract_project_number_display, prefill_contract_adjustment_fields

        product_b = Product.objects.create(
            short_name="TDD",
            name_en="Technical Due Diligence",
            name_ru="ТДД",
        )
        second_project = ProjectRegistration.objects.create(
            number=self.project.number,
            group_member=self.group_member,
            type=product_b,
            name=self.project.name,
            year=2026,
        )
        self.project.refresh_from_db()
        participation_batch_id = uuid.uuid4()
        self.performer.participation_batch_id = participation_batch_id
        self.performer.contract_batch_id = None
        self.performer.contract_file = ""
        self.performer.contract_project_disk_folder = ""
        self.performer.contract_project_created = False
        self.performer.save(update_fields=[
            "participation_batch_id",
            "contract_batch_id",
            "contract_file",
            "contract_project_disk_folder",
            "contract_project_created",
        ])
        Performer.objects.create(
            registration=second_project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
        )

        prefill_contract_adjustment_fields([self.performer.pk], confirmed_at=timezone.now())

        self.performer.refresh_from_db()
        self.assertEqual(
            self.performer.contract_file,
            f"Договор {contract_project_number_display(self.project)}_Иванов ИИ.docx",
        )

    def test_contract_dispatch_column_registry_excludes_signing_note(self):
        from core.column_registry import get_column_choices

        choices = get_column_choices("projects", "contract_dispatch")

        self.assertIn(("project", "Этап"), choices)
        self.assertNotIn(("signing", "Подписание"), choices)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_shows_only_current_expert_signing_rows(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        other_project = ProjectRegistration.objects.create(
            number=7002,
            group_member=self.group_member,
            type=self.product,
            name="Чужой договорный проект",
            year=2026,
        )
        other_user = get_user_model().objects.create_user(
            username="other-expert@example.com",
            email="other-expert@example.com",
            password="secret",
            is_staff=True,
        )
        other_employee = Employee.objects.create(user=other_user)
        Performer.objects.create(
            registration=other_project,
            employee=other_employee,
            executor="Петров Петр Петрович",
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            contract_batch_id=uuid.uuid4(),
            contract_file="Договор 7002_Петров ПП.docx",
            contract_project_disk_folder="/Corporate Root/2026/Project/09 Договоры/000 Петров ПП",
        )
        self.client.force_login(self.employee_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="signing-sign-contract-btn"', html=False)
        self.assertContains(response, 'data-expert-requires-contract-sent="1"', html=False)
        self.assertNotContains(response, 'data-expert-has-contract-sent="1"', html=False)
        self.assertContains(response, 'class="btn btn-primary btn-sm d-none align-items-center"', html=False)
        self.assertContains(response, 'data-contract-sent="0"', html=False)
        self.assertContains(response, "Подписать договор")
        self.assertNotContains(response, 'id="signing-edit-btn"', html=False)
        self.assertNotContains(response, 'id="signing-send-scan-btn"', html=False)
        self.assertContains(response, "Подписание договора")
        self.assertContains(response, 'id="signing-project-filter-toggle"', html=False)
        self.assertNotContains(response, "Корректировка условий договора")
        self.assertNotContains(response, "Составление проекта договора")
        self.assertNotContains(response, "Отправка проекта договора")
        self.assertNotContains(response, 'class="table table-sm align-middle contracts-table mb-0"', html=False)
        self.assertNotContains(response, 'id="contract-conclusion-section"', html=False)
        self.assertNotContains(response, 'id="contract-drafting-table"', html=False)
        self.assertNotContains(response, 'id="contract-dispatch-table"', html=False)
        self.assertContains(response, "Договорный проект")
        self.assertContains(response, "Иванов И.И.")
        self.assertNotContains(response, "Чужой договорный проект")
        self.assertNotContains(response, "Петров П.П.")
        content = response.content.decode("utf-8")
        signing_header = content[
            content.rindex("<div", 0, content.index("Подписание договора")):
            content.index("Подписание договора")
        ]
        self.assertIn("table-section-header", signing_header)
        self.assertNotIn("mt-4", signing_header)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_marks_expert_signing_row_after_contract_is_sent(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.performer.contract_sent_at = timezone.now()
        self.performer.contract_signing_note = "Отправлен проект договора"
        self.performer.save(update_fields=["contract_sent_at", "contract_signing_note"])
        self.client.force_login(self.employee_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-expert-requires-contract-sent="1"', html=False)
        self.assertContains(response, 'data-expert-has-contract-sent="1"', html=False)
        self.assertContains(response, 'class="btn btn-primary btn-sm d-flex align-items-center"', html=False)
        self.assertContains(response, 'data-contract-sent="1"', html=False)
        self.assertContains(response, "Отправлен проект договора")

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_adds_return_column_and_hides_signing_kind(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.performer.contract_sent_at = timezone.now()
        self.performer.contract_signing_note = "Отправлен проект договора"
        self.performer.save(update_fields=["contract_sent_at", "contract_signing_note"])
        self.client.force_login(self.employee_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        signing_table_start = content.index('class="table table-sm align-middle signing-table"')
        signing_table = content[
            signing_table_start:
            content.index("</table>", signing_table_start)
        ]
        self.assertNotIn('<th class="text-nowrap">Вид</th>', signing_table)
        self.assertIn('<th class="text-nowrap">Возврат</th>', signing_table)
        self.assertLess(
            signing_table.index('<th class="text-nowrap">Проект договора</th>'),
            signing_table.index('<th class="text-nowrap">Возврат</th>'),
        )
        self.assertLess(
            signing_table.index('<th class="text-nowrap">Возврат</th>'),
            signing_table.index('<th class="text-nowrap">Подписанный договор</th>'),
        )
        self.assertIn('data-contract-return-comment-trigger="1"', signing_table)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_shows_return_button_only_for_expert(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.performer.contract_sent_at = timezone.now()
        self.performer.contract_signing_note = "Отправлен проект договора"
        self.performer.save(update_fields=["contract_sent_at", "contract_signing_note"])
        self.client.force_login(self.employee_user)

        expert_response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(expert_response.status_code, 200)
        self.assertContains(expert_response, 'id="signing-return-contract-btn"', html=False)
        self.assertContains(expert_response, "Вернуть договор")

        self.client.force_login(self.user)
        admin_response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(admin_response.status_code, 200)
        self.assertNotContains(admin_response, 'id="signing-return-contract-btn"', html=False)

    def test_contract_return_comment_modal_uses_contract_title_without_asset_tabs(self):
        self.performer.contract_number = "IMCM/7001-ИИ/05-26"
        self.performer.contract_sent_at = timezone.now()
        self.performer.save(update_fields=["contract_number", "contract_sent_at"])
        self.client.force_login(self.employee_user)

        response = self.client.get(reverse("contracts_return_comment_modal", args=[self.performer.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Комментарии к проекту договора IMCM/7001-ИИ/05-26")
        self.assertNotContains(response, "Договорный проект")
        self.assertNotContains(response, "Все активы")
        self.assertNotContains(response, "data-comment-tab", html=False)
        self.assertContains(response, "Добавить комментарий")
        self.assertContains(response, "Вернуть договор")

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contract_return_comment_counts_lawyer_and_expert_roles(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="return-lawyer@example.com",
            email="return-lawyer@example.com",
            password="secret",
            is_staff=True,
        )
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        ContractReturnComment.objects.create(
            performer=self.performer,
            contract_batch_id=self.performer.contract_batch_id,
            text="Комментарий юриста",
            author=lawyer_user,
            author_role=ContractReturnComment.AuthorRole.LAWYER,
        )
        ContractReturnComment.objects.create(
            performer=self.performer,
            contract_batch_id=self.performer.contract_batch_id,
            text="Комментарий эксперта",
            author=self.employee_user,
            author_role=ContractReturnComment.AuthorRole.EXPERT,
        )
        self.client.force_login(self.employee_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        return_cell_start = content.index(f'id="contract-return-comment-{self.performer.pk}"')
        return_cell = content[return_cell_start:content.index("</div>", return_cell_start)]
        self.assertIn('chk-comment-counter--lawyer has-comments', return_cell)
        self.assertIn('>1</span>', return_cell)
        self.assertIn('chk-comment-counter--expert has-comments', return_cell)
        self.assertIn('contract-return-icon--expert', return_cell)

    def test_contract_return_endpoint_requires_expert_role(self):
        self.performer.contract_sent_at = timezone.now()
        self.performer.contract_signing_note = "Отправлен проект договора"
        self.performer.save(update_fields=["contract_sent_at", "contract_signing_note"])

        response = self.client.post(reverse("contracts_return_performer_contract", args=[self.performer.pk]))

        self.assertEqual(response.status_code, 403)
        self.performer.refresh_from_db()
        self.assertIsNotNone(self.performer.contract_sent_at)
        self.assertEqual(self.performer.contract_signing_note, "Отправлен проект договора")

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_expert_contract_return_moves_batch_back_to_project_drafting(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        sent_at = timezone.now()
        deadline_at = sent_at + timedelta(hours=96)
        self.performer.contract_sent_at = sent_at
        self.performer.contract_deadline_at = deadline_at
        self.performer.contract_conclusion_status = "В срок"
        self.performer.contract_signing_note = "Отправлен проект договора"
        self.performer.save(update_fields=[
            "contract_sent_at",
            "contract_deadline_at",
            "contract_conclusion_status",
            "contract_signing_note",
        ])
        self.client.force_login(self.employee_user)

        response = self.client.post(reverse("contracts_return_performer_contract", args=[self.performer.pk]))

        self.assertEqual(response.status_code, 200)
        self.performer.refresh_from_db()
        self.assertIsNone(self.performer.contract_sent_at)
        self.assertIsNone(self.performer.contract_deadline_at)
        self.assertEqual(self.performer.contract_conclusion_status, "")
        self.assertEqual(self.performer.contract_signing_note, "Разрабатывается проект договора")

        expert_response = self.client.get(reverse("contracts_partial"))
        self.assertEqual(expert_response.status_code, 200)
        expert_content = expert_response.content.decode("utf-8")
        signing_checkbox_start = expert_content.index(f'id="signing-sel-{self.performer.pk}"')
        signing_checkbox = expert_content[
            signing_checkbox_start:
            expert_content.index("aria-label", signing_checkbox_start)
        ]
        self.assertIn("disabled", signing_checkbox)
        self.assertIn("договор возвращён на составление проекта", signing_checkbox)

        self.client.force_login(self.user)
        admin_response = self.client.get(reverse("contracts_partial"))
        self.assertEqual(admin_response.status_code, 200)
        admin_content = admin_response.content.decode("utf-8")
        drafting_checkbox_start = admin_content.index(f'id="contract-drafting-sel-{self.performer.pk}"')
        drafting_checkbox = admin_content[
            drafting_checkbox_start:
            admin_content.index("aria-label", drafting_checkbox_start)
        ]
        self.assertIn('data-contract-sent="0"', drafting_checkbox)
        self.assertNotIn("disabled", drafting_checkbox)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contract_return_completes_expert_contract_notification(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        sent_at = timezone.now()
        self.performer.contract_sent_at = sent_at
        self.performer.contract_deadline_at = sent_at + timedelta(hours=96)
        self.performer.contract_signing_note = "Отправлен проект договора"
        self.performer.save(update_fields=[
            "contract_sent_at",
            "contract_deadline_at",
            "contract_signing_note",
        ])
        notification = Notification.objects.create(
            notification_type=Notification.NotificationType.PROJECT_CONTRACT_CONCLUSION,
            related_section=Notification.RelatedSection.CONTRACTS,
            recipient=self.employee_user,
            sender=self.user,
            project=self.project,
            title_text="Отправлен проект договора",
            content_text="Подписать договор",
            sent_at=sent_at,
            deadline_at=self.performer.contract_deadline_at,
            is_read=False,
            is_processed=False,
        )
        NotificationPerformerLink.objects.create(notification=notification, performer=self.performer)
        self.client.force_login(self.employee_user)

        counters_before = self.client.get(reverse("notifications_counters")).json()
        self.assertEqual(counters_before["sections"].get("contracts"), 1)
        response_before = self.client.get(reverse("contracts_partial"))
        self.assertContains(response_before, '<span class="badge rounded-pill text-bg-danger ms-1">1</span>', html=False)

        response = self.client.post(reverse("contracts_return_performer_contract", args=[self.performer.pk]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["completed_notifications"], 1)
        notification.refresh_from_db()
        self.assertTrue(notification.is_read)
        self.assertTrue(notification.is_processed)
        self.assertEqual(notification.action_by, self.employee_user)

        counters_after = self.client.get(reverse("notifications_counters")).json()
        self.assertEqual(counters_after["sections"].get("contracts", 0), 0)
        response_after = self.client.get(reverse("contracts_partial"))
        self.assertNotContains(response_after, '<span class="badge rounded-pill text-bg-danger ms-1">1</span>', html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_disables_expert_signing_checkbox_after_facsimile_signature(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.performer.contract_sent_at = timezone.now()
        self.performer.contract_signing_date = timezone.now()
        self.performer.contract_signed_pdf_file = "Договор 7001_Иванов ИИ_п.pdf"
        self.performer.contract_signing_note = "Договор подписан факсимиле"
        self.performer.save(update_fields=[
            "contract_sent_at",
            "contract_signing_date",
            "contract_signed_pdf_file",
            "contract_signing_note",
        ])
        self.client.force_login(self.employee_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        checkbox_start = content.index(f'id="signing-sel-{self.performer.pk}"')
        checkbox_html = content[checkbox_start:content.index("aria-label", checkbox_start)]
        self.assertIn("disabled", checkbox_html)
        self.assertIn("Строка недоступна: договор уже подписан факсимиле", checkbox_html)
        self.assertContains(response, "Договор подписан факсимиле")

    def test_contracts_partial_shows_all_signing_rows_for_admin(self):
        other_project = ProjectRegistration.objects.create(
            number=7002,
            group_member=self.group_member,
            type=self.product,
            name="Чужой договорный проект",
            year=2026,
        )
        Performer.objects.create(
            registration=other_project,
            executor="Петров Петр Петрович",
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            contract_batch_id=uuid.uuid4(),
            contract_file="Договор 7002_Петров ПП.docx",
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Договорный проект")
        self.assertContains(response, "Чужой договорный проект")
        self.assertContains(response, "Иванов И.И.")
        self.assertContains(response, "Петров П.П.")

    def test_contracts_partial_uses_non_empty_signing_note_for_batch_representative(self):
        self.performer.contract_signing_note = "Разрабатывается проект договора"
        self.performer.save(update_fields=["contract_signing_note"])
        Performer.objects.create(
            registration=self.project,
            executor="Иванов Иван Иванович",
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            contract_batch_id=self.performer.contract_batch_id,
            contract_file=self.performer.contract_file,
            position=0,
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        signing_table_start = content.index('class="table table-sm align-middle signing-table"')
        signing_table = content[
            signing_table_start:
            content.index("</table>", signing_table_start)
        ]
        self.assertIn("Разрабатывается проект договора", signing_table)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_hides_sign_contract_button_for_lawyer(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="lawyer-sign-hidden@example.com",
            email="lawyer-sign-hidden@example.com",
            password="secret",
            is_staff=True,
        )
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        self.client.force_login(lawyer_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'id="signing-sign-contract-btn"', html=False)
        self.assertNotContains(response, 'id="signing-edit-btn"', html=False)
        self.assertNotContains(response, 'id="signing-send-scan-btn"', html=False)
        content = response.content.decode("utf-8")
        checkbox_start = content.index(f'id="signing-sel-{self.performer.pk}"')
        checkbox_html = content[checkbox_start:content.index("aria-label", checkbox_start)]
        self.assertIn("disabled", checkbox_html)
        self.assertIn("signing-table-wrap", content)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_hides_signing_management_buttons_for_employee_lawyer_role(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="employee-lawyer-sign-hidden@example.com",
            email="employee-lawyer-sign-hidden@example.com",
            password="secret",
            is_staff=True,
        )
        Employee.objects.create(user=lawyer_user, role=LAWYER_GROUP)
        self.client.force_login(lawyer_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'id="signing-edit-btn"', html=False)
        self.assertNotContains(response, 'id="signing-send-scan-btn"', html=False)
        content = response.content.decode("utf-8")
        checkbox_start = content.index(f'id="signing-sel-{self.performer.pk}"')
        checkbox_html = content[checkbox_start:content.index("aria-label", checkbox_start)]
        self.assertIn("disabled", checkbox_html)

    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_hides_signing_management_buttons_for_superuser_lawyer_role(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="lawyer-superuser-sign-hidden@example.com",
            email="lawyer-superuser-sign-hidden@example.com",
            password="secret",
            is_staff=True,
            is_superuser=True,
        )
        Employee.objects.create(user=lawyer_user, role=LAWYER_GROUP)
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        self.client.force_login(lawyer_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'id="signing-sign-contract-btn"', html=False)
        self.assertNotContains(response, 'id="signing-edit-btn"', html=False)
        self.assertNotContains(response, 'id="signing-send-scan-btn"', html=False)
        content = response.content.decode("utf-8")
        checkbox_start = content.index(f'id="signing-sel-{self.performer.pk}"')
        checkbox_html = content[checkbox_start:content.index("aria-label", checkbox_start)]
        self.assertIn("disabled", checkbox_html)

    def test_contracts_partial_shows_signing_management_buttons_for_admin(self):
        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="signing-edit-btn"', html=False)
        self.assertContains(response, 'id="signing-send-scan-btn"', html=False)

    def test_contracts_partial_shows_signing_management_buttons_for_employee_admin_role(self):
        admin_user = get_user_model().objects.create_user(
            username="employee-admin@example.com",
            email="employee-admin@example.com",
            password="secret",
        )
        Employee.objects.create(user=admin_user, role=ADMIN_GROUP)
        self.client.force_login(admin_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="signing-edit-btn"', html=False)
        self.assertContains(response, 'id="signing-send-scan-btn"', html=False)
        self.assertContains(response, 'id="signing-master"', html=False)

    def test_contracts_partial_shows_signing_management_buttons_for_superuser(self):
        superuser = get_user_model().objects.create_user(
            username="contracts-superuser@example.com",
            email="contracts-superuser@example.com",
            password="secret",
            is_superuser=True,
        )
        self.client.force_login(superuser)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="signing-edit-btn"', html=False)
        self.assertContains(response, 'id="signing-send-scan-btn"', html=False)
        self.assertContains(response, 'id="signing-master"', html=False)

    def test_contract_signing_management_actions_reject_non_admin_staff_user(self):
        staff_user = get_user_model().objects.create_user(
            username="contracts-staff@example.com",
            email="contracts-staff@example.com",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(staff_user)

        edit_response = self.client.get(reverse("contracts_signing_edit", args=[self.performer.pk]))
        send_scan_response = self.client.post(
            reverse("signing_send_scan"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(edit_response.status_code, 302)
        self.assertEqual(send_scan_response.status_code, 302)

    def test_contract_signing_management_actions_reject_superuser_lawyer_role(self):
        lawyer_user = get_user_model().objects.create_user(
            username="contracts-superuser-lawyer@example.com",
            email="contracts-superuser-lawyer@example.com",
            password="secret",
            is_staff=True,
            is_superuser=True,
        )
        Employee.objects.create(user=lawyer_user, role=LAWYER_GROUP)
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        self.client.force_login(lawyer_user)

        edit_response = self.client.get(reverse("contracts_signing_edit", args=[self.performer.pk]))
        send_scan_response = self.client.post(
            reverse("signing_send_scan"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(edit_response.status_code, 302)
        self.assertEqual(send_scan_response.status_code, 302)

    def test_contracts_partial_renders_contract_pdf_file_like_proposal_pdf(self):
        signed_at = timezone.make_aware(datetime(2026, 5, 2, 13, 45))
        self.performer.contract_pdf_file = "Договор 7001_Иванов ИИ.pdf"
        self.performer.contract_pdf_link = "https://cloud.example.com/s/contract-pdf"
        self.performer.contract_signed_pdf_file = "Договор 7001_Иванов ИИ подписанный.pdf"
        self.performer.contract_signed_pdf_link = "https://cloud.example.com/s/signed-contract-pdf"
        self.performer.contract_signing_date = signed_at
        self.performer.contract_deadline_at = signed_at + timedelta(days=1)
        self.performer.contract_signing_note = "Отправлен проект договора"
        self.performer.save(update_fields=[
            "contract_pdf_file",
            "contract_pdf_link",
            "contract_signed_pdf_file",
            "contract_signed_pdf_link",
            "contract_signing_date",
            "contract_deadline_at",
            "contract_signing_note",
        ])

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<th class="text-nowrap">Проект договора</th>', html=False)
        self.assertContains(response, '<th class="text-nowrap">Дата подписания</th>', html=False)
        self.assertContains(response, '<th class="text-nowrap">Статус</th>', html=False)
        self.assertContains(response, '<th class="text-nowrap">Заключение договора</th>', html=False)
        self.assertContains(response, '<th class="text-nowrap">Подписанный договор</th>', html=False)
        self.assertContains(response, '<td class="text-nowrap">Отправлен проект договора</td>', html=False)
        self.assertContains(response, "https://cloud.example.com/s/contract-pdf")
        self.assertContains(response, "Договор 7001_Иванов ИИ.pdf")
        self.assertContains(response, "https://cloud.example.com/s/signed-contract-pdf")
        self.assertContains(response, "Договор 7001_Иванов ИИ подписанный.pdf")
        self.assertContains(response, "bi-file-pdf-fill")
        content = response.content.decode("utf-8")
        signing_section = content[content.index("Подписание договора"):]
        self.assertLess(
            signing_section.index('<th class="text-nowrap">Проект договора</th>'),
            signing_section.index('<th class="text-nowrap">Подписанный договор</th>'),
        )
        self.assertLess(
            signing_section.index('<th class="text-nowrap">Подписанный договор</th>'),
            signing_section.index('<th class="text-nowrap">Дата подписания</th>'),
        )
        self.assertLess(
            signing_section.index('<th class="text-nowrap">Дата подписания</th>'),
            signing_section.index('<th class="text-nowrap">Статус</th>'),
        )
        self.assertLess(
            signing_section.index('<th class="text-nowrap">Статус</th>'),
            signing_section.index('<th class="text-nowrap">Заключение договора</th>'),
        )
        self.assertNotIn('<th class="text-nowrap">Скан с подписью сотрудника</th>', signing_section)
        self.assertNotIn('<th class="text-nowrap">Облако</th>', signing_section)
        self.assertNotIn('<th class="text-nowrap">Дата загрузки</th>', signing_section)
        self.assertNotIn('<th class="text-nowrap">Дата отправки</th>', signing_section)
        self.assertNotIn('<th class="text-nowrap">Скан подписанного договора</th>', signing_section)
        self.assertIn("https://cloud.example.com/s/contract-pdf", signing_section)
        self.assertIn("Договор 7001_Иванов ИИ.pdf", signing_section)
        self.assertIn(timezone.localtime(signed_at).strftime("%d.%m.%Y %H:%M"), signing_section)
        self.assertIn("В срок", signing_section)
        self.assertIn("https://cloud.example.com/s/signed-contract-pdf", signing_section)
        self.assertIn("Договор 7001_Иванов ИИ подписанный.pdf", signing_section)

    def test_contract_request_sets_signing_note_for_signing_table(self):
        response = self.client.post(
            reverse("contract_request"),
            {
                "performer_ids[]": [self.performer.pk],
                "duration_hours": "24",
                "delivery_channels[]": ["system"],
            },
        )

        self.assertEqual(response.status_code, 200)
        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_signing_note, "Отправлен проект договора")

        partial_response = self.client.get(reverse("contracts_partial"))
        self.assertContains(
            partial_response,
            '<td class="text-nowrap">Отправлен проект договора</td>',
            html=False,
        )

    def test_contract_docx_source_is_available_to_onlyoffice_without_session(self):
        from projects_app.views import _build_contract_docx_source_token

        self.client.logout()
        token = _build_contract_docx_source_token(self.performer)
        doc = Document()
        doc.add_paragraph("contract body")
        buffer = BytesIO()
        doc.save(buffer)

        with patch(
            "projects_app.views.cloud_download_file",
            return_value=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                buffer.getvalue(),
            ),
        ):
            response = self.client.get(
                reverse("contract_onlyoffice_docx_source", args=[self.performer.pk]),
                {"token": token},
            )

        self.assertEqual(response.status_code, 200)
        response_doc = Document(BytesIO(response.content))
        self.assertEqual(response_doc.paragraphs[0].text, "contract body")
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_contract_actives_name_uses_russian_bullet_list_style(self):
        template_doc = Document()
        template_doc.styles.add_style("Маркированный список", WD_STYLE_TYPE.PARAGRAPH)
        template_doc.add_paragraph("[[actives_name]]")
        buffer = BytesIO()
        template_doc.save(buffer)

        generated_bytes = process_template(
            buffer.getvalue(),
            {},
            list_replacements={"[[actives_name]]": ["Актив 1", "Актив 2"]},
            default_language_code="ru-RU",
        )

        generated_doc = Document(BytesIO(generated_bytes))
        active_paragraphs = [
            paragraph
            for paragraph in generated_doc.paragraphs
            if paragraph.text in {"Актив 1", "Актив 2"}
        ]

        self.assertEqual(len(active_paragraphs), 2)
        for paragraph in active_paragraphs:
            self.assertEqual(paragraph.style.name, "Маркированный список")
            self.assertIn('w:lang w:val="ru-RU"', paragraph._element.xml)

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    def test_sign_contract_documents_generates_pdf_and_public_link(self):
        with (
            patch("projects_app.views.convert_docx_source_to_pdf", return_value=b"%PDF-1.4") as mocked_convert,
            patch("projects_app.views.cloud_upload_file", return_value=True) as mocked_upload,
            patch("projects_app.views.cloud_publish_resource", return_value="https://cloud.example.com/s/contract-pdf") as mocked_publish,
            patch("projects_app.views._resolve_contract_project_nextcloud_file_id", return_value="pdf-file-id"),
        ):
            response = self.client.post(
                reverse("sign_contract_documents"),
                {"performer_ids[]": [self.performer.pk]},
            )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["generated"], 1)
        self.assertEqual(data["updates"][0]["contract_pdf_file"], "Договор 7001_Иванов ИИ.pdf")

        source_url = mocked_convert.call_args.kwargs["source_url"]
        self.assertIn(
            reverse("contract_onlyoffice_docx_source", args=[self.performer.pk]),
            source_url,
        )
        expected_pdf_path = (
            f"{self.performer.contract_project_disk_folder}/"
            "Договор 7001_Иванов ИИ.pdf"
        )
        mocked_upload.assert_called_once_with(self.user, expected_pdf_path, b"%PDF-1.4")
        mocked_publish.assert_called_once_with(self.user, expected_pdf_path)

        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_pdf_file, "Договор 7001_Иванов ИИ.pdf")
        self.assertEqual(self.performer.contract_pdf_link, "https://cloud.example.com/s/contract-pdf")
        self.assertEqual(self.performer.contract_pdf_file_id, "pdf-file-id")

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    def test_sign_contract_documents_expands_selected_row_to_contract_batch(self):
        participation_batch_id = uuid.uuid4()
        self.performer.participation_batch_id = participation_batch_id
        self.performer.save(update_fields=["participation_batch_id"])
        sibling = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            contract_batch_id=uuid.uuid4(),
            contract_file=self.performer.contract_file,
            contract_project_disk_folder=self.performer.contract_project_disk_folder,
        )

        with (
            patch("projects_app.views.convert_docx_source_to_pdf", return_value=b"%PDF-1.4") as mocked_convert,
            patch("projects_app.views.cloud_upload_file", return_value=True),
            patch("projects_app.views.cloud_publish_resource", return_value="https://cloud.example.com/s/contract-pdf"),
            patch("projects_app.views._resolve_contract_project_nextcloud_file_id", return_value="pdf-file-id"),
        ):
            response = self.client.post(
                reverse("sign_contract_documents"),
                {"performer_ids[]": [self.performer.pk]},
            )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["generated"], 1)
        self.assertCountEqual(
            [item["id"] for item in data["updates"]],
            [self.performer.pk, sibling.pk],
        )
        mocked_convert.assert_called_once()
        self.performer.refresh_from_db()
        sibling.refresh_from_db()
        self.assertEqual(self.performer.contract_pdf_file, "Договор 7001_Иванов ИИ.pdf")
        self.assertEqual(sibling.contract_pdf_file, "Договор 7001_Иванов ИИ.pdf")
        self.assertEqual(self.performer.contract_pdf_link, "https://cloud.example.com/s/contract-pdf")
        self.assertEqual(sibling.contract_pdf_link, "https://cloud.example.com/s/contract-pdf")

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    def test_sign_performer_contract_documents_generates_signed_pdf_and_public_link(self):
        signed_at = timezone.now().replace(microsecond=0)
        self.performer.contract_deadline_at = signed_at + timedelta(days=1)
        self.performer.save(update_fields=["contract_deadline_at"])

        with (
            patch("projects_app.views.timezone.now", return_value=signed_at),
            patch("projects_app.views._load_performer_facsimile_bytes", return_value=b"facsimile"),
            patch("projects_app.views.convert_docx_source_to_pdf", return_value=b"%PDF-1.4") as mocked_convert,
            patch("projects_app.views.cloud_upload_file", return_value=True) as mocked_upload,
            patch("projects_app.views.cloud_publish_resource", return_value="https://cloud.example.com/s/signed-contract-pdf") as mocked_publish,
            patch("projects_app.views._resolve_contract_project_nextcloud_file_id", return_value="signed-pdf-file-id"),
        ):
            response = self.client.post(
                reverse("sign_performer_contract_documents"),
                {"performer_ids[]": [self.performer.pk]},
            )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["generated"], 1)
        self.assertEqual(
            data["updates"][0]["contract_signed_pdf_file"],
            "Договор 7001_Иванов ИИ_п.pdf",
        )

        source_url = mocked_convert.call_args.kwargs["source_url"]
        self.assertIn(
            reverse("contract_onlyoffice_docx_source", args=[self.performer.pk]),
            source_url,
        )
        expected_pdf_path = (
            f"{self.performer.contract_project_disk_folder}/"
            "Договор 7001_Иванов ИИ_п.pdf"
        )
        mocked_upload.assert_called_once_with(self.user, expected_pdf_path, b"%PDF-1.4")
        mocked_publish.assert_called_once_with(self.user, expected_pdf_path)

        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_signed_pdf_file, "Договор 7001_Иванов ИИ_п.pdf")
        self.assertEqual(self.performer.contract_signed_pdf_link, "https://cloud.example.com/s/signed-contract-pdf")
        self.assertEqual(self.performer.contract_signed_pdf_file_id, "signed-pdf-file-id")
        self.assertEqual(self.performer.contract_signing_date, signed_at)
        self.assertEqual(self.performer.contract_conclusion_status, "В срок")
        self.assertEqual(self.performer.contract_signing_note, "Договор подписан факсимиле")

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    def test_sign_performer_contract_documents_expands_selected_row_to_contract_batch(self):
        signed_at = timezone.now().replace(microsecond=0)
        self.performer.contract_deadline_at = signed_at + timedelta(days=1)
        participation_batch_id = uuid.uuid4()
        self.performer.participation_batch_id = participation_batch_id
        self.performer.save(update_fields=["contract_deadline_at", "participation_batch_id"])
        sibling = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            contract_batch_id=uuid.uuid4(),
            contract_file=self.performer.contract_file,
            contract_project_disk_folder=self.performer.contract_project_disk_folder,
            contract_deadline_at=signed_at + timedelta(days=1),
        )

        with (
            patch("projects_app.views.timezone.now", return_value=signed_at),
            patch("projects_app.views._load_performer_facsimile_bytes", return_value=b"facsimile"),
            patch("projects_app.views.convert_docx_source_to_pdf", return_value=b"%PDF-1.4") as mocked_convert,
            patch("projects_app.views.cloud_upload_file", return_value=True),
            patch("projects_app.views.cloud_publish_resource", return_value="https://cloud.example.com/s/signed-contract-pdf"),
            patch("projects_app.views._resolve_contract_project_nextcloud_file_id", return_value="signed-pdf-file-id"),
        ):
            response = self.client.post(
                reverse("sign_performer_contract_documents"),
                {"performer_ids[]": [self.performer.pk]},
            )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["generated"], 1)
        self.assertCountEqual(
            [item["id"] for item in data["updates"]],
            [self.performer.pk, sibling.pk],
        )
        mocked_convert.assert_called_once()
        self.performer.refresh_from_db()
        sibling.refresh_from_db()
        self.assertEqual(self.performer.contract_signed_pdf_file, "Договор 7001_Иванов ИИ_п.pdf")
        self.assertEqual(sibling.contract_signed_pdf_file, "Договор 7001_Иванов ИИ_п.pdf")
        self.assertEqual(self.performer.contract_signing_date, signed_at)
        self.assertEqual(sibling.contract_signing_date, signed_at)
        self.assertEqual(self.performer.contract_signing_note, "Договор подписан факсимиле")
        self.assertEqual(sibling.contract_signing_note, "Договор подписан факсимиле")

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    def test_sign_performer_contract_documents_returns_clear_error_without_facsimile(self):
        with (
            patch("projects_app.views._load_performer_facsimile_bytes", return_value=None),
            patch("projects_app.views.convert_docx_source_to_pdf") as mocked_convert,
        ):
            response = self.client.post(
                reverse("sign_performer_contract_documents"),
                {"performer_ids[]": [self.performer.pk]},
            )

        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertFalse(data["ok"])
        self.assertIn("факсимильная подпись исполнителя", data["error"])
        self.assertNotIn("Загрузите файл", data["error"])
        mocked_convert.assert_not_called()

    def test_sign_performer_contract_documents_rejects_regular_staff_user(self):
        staff_user = get_user_model().objects.create_user(
            username="regular-staff-sign@example.com",
            email="regular-staff-sign@example.com",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(staff_user)

        with patch("projects_app.views.convert_docx_source_to_pdf") as mocked_convert:
            response = self.client.post(
                reverse("sign_performer_contract_documents"),
                {"performer_ids[]": [self.performer.pk]},
            )

        self.assertEqual(response.status_code, 302)
        mocked_convert.assert_not_called()

    @override_settings(ONLYOFFICE_DOCUMENT_SERVER_URL="https://docs.example.com")
    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_sign_performer_contract_documents_completes_expert_contract_notification(
        self,
        _mocked_list_user_shares,
        _mocked_get_user_share,
    ):
        self.performer.contract_deadline_at = timezone.now() + timedelta(days=1)
        self.performer.save(update_fields=["contract_deadline_at"])
        notification = Notification.objects.create(
            notification_type=Notification.NotificationType.PROJECT_CONTRACT_CONCLUSION,
            related_section=Notification.RelatedSection.CONTRACTS,
            recipient=self.employee_user,
            sender=self.user,
            project=self.project,
            title_text="Отправлен проект договора",
            content_text="Подписать договор",
            sent_at=timezone.now(),
            deadline_at=self.performer.contract_deadline_at,
            is_read=False,
            is_processed=False,
        )
        NotificationPerformerLink.objects.create(notification=notification, performer=self.performer)
        self.client.force_login(self.employee_user)

        counters_before = self.client.get(reverse("notifications_counters")).json()
        self.assertEqual(counters_before["sections"].get("contracts"), 1)
        response_before = self.client.get(reverse("contracts_partial"))
        self.assertContains(response_before, '<span class="badge rounded-pill text-bg-danger ms-1">1</span>', html=False)

        with (
            patch("projects_app.views.convert_docx_source_to_pdf", return_value=b"%PDF-1.4"),
            patch("projects_app.views._load_performer_facsimile_bytes", return_value=b"facsimile"),
            patch("projects_app.views.cloud_upload_file", return_value=True),
            patch("projects_app.views.cloud_publish_resource", return_value="https://cloud.example.com/s/signed-contract-pdf"),
            patch("projects_app.views._resolve_contract_project_nextcloud_file_id", return_value="signed-pdf-file-id"),
        ):
            response = self.client.post(
                reverse("sign_performer_contract_documents"),
                {"performer_ids[]": [self.performer.pk]},
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["completed_notifications"], 1)
        notification.refresh_from_db()
        self.assertTrue(notification.is_read)
        self.assertTrue(notification.is_processed)
        self.assertEqual(notification.action_by, self.employee_user)

        counters_after = self.client.get(reverse("notifications_counters")).json()
        self.assertEqual(counters_after["sections"].get("contracts", 0), 0)
        response_after = self.client.get(reverse("contracts_partial"))
        self.assertNotContains(response_after, '<span class="badge rounded-pill text-bg-danger ms-1">1</span>', html=False)

    def test_contract_edit_updates_contract_date_for_batch_rows(self):
        self.project.deadline = date(2026, 5, 27)
        self.project.save(update_fields=["deadline"])
        self.performer.contract_date = date(2026, 5, 7)
        self.performer.save(update_fields=["contract_date"])
        sibling = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor="Иванов Иван Иванович",
            contract_batch_id=self.performer.contract_batch_id,
        )

        form_response = self.client.get(reverse("contracts_edit", args=[self.performer.pk]))
        form_content = form_response.content.decode("utf-8")

        self.assertEqual(form_response.status_code, 200)
        self.assertLess(
            form_content.index("Цена договора"),
            form_content.index("Дата договора"),
        )
        self.assertContains(form_response, 'name="prepayment"', html=False)
        self.assertContains(form_response, 'name="final_payment"', html=False)
        self.assertContains(form_response, 'id="contracts-term-days-field"', html=False)
        self.assertContains(form_response, 'data-deadline="2026-05-27"', html=False)
        self.assertContains(form_response, 'value="20 дн."', html=False)
        self.assertContains(form_response, 'data-contract-form-cancel-btn="1"', html=False)
        self.assertContains(form_response, 'data-contract-form-save-btn="1"', html=False)
        self.assertContains(form_response, "Сохранение...", html=False)
        self.assertContains(form_response, "spinner-border spinner-border-sm me-2", html=False)

        response = self.client.post(
            reverse("contracts_edit", args=[self.performer.pk]),
            {
                "contract_number": "CUSTOM-42",
                "contract_date": "2026-05-07",
                "prepayment": "30",
                "contract_file": "custom.docx",
            },
        )

        self.assertEqual(response.status_code, 200)
        sibling.refresh_from_db()
        self.assertEqual(sibling.contract_number, "CUSTOM-42")
        self.assertEqual(sibling.contract_date, date(2026, 5, 7))
        self.assertEqual(sibling.prepayment, 30)
        self.assertEqual(sibling.final_payment, 70)
        self.assertEqual(sibling.contract_file, "custom.docx")

    def test_contracts_partial_uses_direction_head_as_responsible(self):
        direction = OrgUnit.objects.create(
            company=self.group_member,
            level=2,
            department_name="Геология",
            unit_type="expertise",
        )
        head_user = get_user_model().objects.create_user(
            username="direction-head@example.com",
            password="secret",
            first_name="Петр",
            last_name="Петров",
            is_staff=True,
        )
        Employee.objects.create(
            user=head_user,
            patronymic="Петрович",
            department=direction,
            role=DEPARTMENT_HEAD_GROUP,
        )
        ExpertProfile.objects.create(
            employee=self.employee,
            expertise_direction=direction,
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Ответственный")
        self.assertContains(response, "Петров П.П.")

        form_response = self.client.get(reverse("contracts_edit", args=[self.performer.pk]))

        self.assertEqual(form_response.status_code, 200)
        self.assertContains(form_response, "Ответственный")
        self.assertContains(form_response, 'value="Петров Петр Петрович"', html=False)

    def test_contracts_partial_uses_project_manager_as_responsible_without_direction(self):
        self.project.project_manager = "Сидоров Сидор Сидорович"
        self.project.save(update_fields=["project_manager"])

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Ответственный")
        self.assertContains(response, "Сидоров С.С.")

        form_response = self.client.get(reverse("contracts_edit", args=[self.performer.pk]))

        self.assertEqual(form_response.status_code, 200)
        self.assertContains(form_response, 'value="Сидоров Сидор Сидорович"', html=False)

    def test_contracts_partial_uses_active_citizenship_country_as_default_group(self):
        kazakhstan = OKSMCountry.objects.create(
            number=398,
            code="398",
            short_name="Казахстан",
            alpha2="KZ",
            alpha3="KAZ",
        )
        GroupMember.objects.create(
            short_name="Казахстан",
            country_name="Казахстан",
            country_code="398",
            country_alpha2="KZ",
            position=2,
        )
        person = PersonRecord.objects.create(
            last_name="Иванов",
            first_name="Иван",
            middle_name="Иванович",
        )
        self.employee.person_record = person
        self.employee.save(update_fields=["person_record"])
        CitizenshipRecord.objects.create(
            person=person,
            country=kazakhstan,
            valid_to=None,
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Группа")
        content = response.content.decode("utf-8")
        contracts_table = content[
            content.index('class="table table-sm align-middle contracts-table mb-0"'):
            content.index('id="contracts-edit-btn"')
        ]
        self.assertLess(
            contracts_table.index(
                f'<span class="clf-id-droid-mono">{self.project.short_uid}</span>'
            ),
            contracts_table.index('<td class="text-nowrap">KZ</td>'),
        )

        form_response = self.client.get(reverse("contracts_edit", args=[self.performer.pk]))

        self.assertEqual(form_response.status_code, 200)
        self.assertContains(form_response, "contracts-group-select")
        self.assertContains(form_response, ">KZ Казахстан</option>", html=False)
        self.assertContains(form_response, "contracts-group-display")

    def test_contract_edit_saves_group_for_batch_rows(self):
        sibling = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor="Иванов Иван Иванович",
            contract_batch_id=self.performer.contract_batch_id,
        )

        response = self.client.post(
            reverse("contracts_edit", args=[self.performer.pk]),
            {
                "contract_group_member": str(self.group_member.pk),
                "contract_number": "CUSTOM-42",
                "contract_date": "2026-05-07",
                "contract_file": "custom.docx",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.performer.refresh_from_db()
        sibling.refresh_from_db()
        self.assertEqual(self.performer.contract_group_member_id, self.group_member.pk)
        self.assertEqual(sibling.contract_group_member_id, self.group_member.pk)

    def test_contracts_partial_uses_current_primary_cloud_label_in_disk_tooltip(self):
        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'title="Открыть папку на Nextcloud"', html=False)

    def test_contracts_partial_uses_public_folder_link_for_cloud_icon(self):
        self.performer.contract_project_folder_link = "https://cloud.example.com/s/contract-folder"
        self.performer.save(update_fields=["contract_project_folder_link"])

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'href="https://cloud.example.com/s/contract-folder"', html=False)
        content = response.content.decode("utf-8")
        section = content[
            content.index('id="contract-conclusion-section"'):
            content.index("Подписание договора")
        ]
        self.assertIn('href="https://cloud.example.com/s/contract-folder"', section)

    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_contracts_partial_uses_editor_docx_link_for_lawyer(
        self,
        mocked_list_user_shares,
        mocked_list_resources,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="lawyer-docx@example.com",
            email="lawyer-docx@example.com",
            password="secret",
            is_staff=True,
        )
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        NextcloudUserLink.objects.create(
            user=lawyer_user,
            nextcloud_user_id="nc-lawyer",
            nextcloud_username="nc-lawyer",
            nextcloud_email=lawyer_user.email,
        )
        mocked_list_user_shares.return_value = {
            self.performer.contract_project_disk_folder: NextcloudShare(
                share_id="56",
                path=self.performer.contract_project_disk_folder,
                share_with="nc-lawyer",
                permissions=15,
                target_path="/Shared/000 Иванов ИИ",
            )
        }
        mocked_list_resources.return_value = [
            {
                "name": self.performer.contract_file,
                "path": f"{self.performer.contract_project_disk_folder}/{self.performer.contract_file}",
                "type": "file",
                "file_id": "4474",
            }
        ]
        self.client.force_login(lawyer_user)

        response = self.client.get(reverse("contracts_partial"))

        expected_url = (
            "https://cloud.example.com/apps/files/files/4474?dir="
            + quote("/Shared/000 Иванов ИИ", safe="/")
            + "&amp;openfile=true"
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, expected_url, html=False)
        content = response.content.decode("utf-8")
        section = content[
            content.index('id="contract-conclusion-section"'):
            content.index("Подписание договора")
        ]
        self.assertIn(expected_url, section)
        self.assertNotIn('href="https://cloud.example.com/s/contract-docx"', section)

    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_contracts_partial_resolves_lawyer_docx_link_from_parent_share(
        self,
        mocked_list_user_shares,
        mocked_list_resources,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="lawyer-parent-share@example.com",
            email="lawyer-parent-share@example.com",
            password="secret",
            is_staff=True,
        )
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        NextcloudUserLink.objects.create(
            user=lawyer_user,
            nextcloud_user_id="nc-lawyer-parent",
            nextcloud_username="nc-lawyer-parent",
            nextcloud_email=lawyer_user.email,
        )
        parent_path = "/Corporate Root/2026/Project"
        mocked_list_user_shares.return_value = {
            parent_path: NextcloudShare(
                share_id="57",
                path=parent_path,
                share_with="nc-lawyer-parent",
                permissions=15,
                target_path="/Shared/Project",
            )
        }
        mocked_list_resources.return_value = [
            {
                "name": self.performer.contract_file,
                "path": f"{self.performer.contract_project_disk_folder}/{self.performer.contract_file}",
                "type": "file",
                "file_id": "4475",
            }
        ]
        self.client.force_login(lawyer_user)

        response = self.client.get(reverse("contracts_partial"))

        expected_target_dir = "/Shared/Project/09 Договоры/000 Иванов ИИ"
        expected_url = (
            "https://cloud.example.com/apps/files/files/4475?dir="
            + quote(expected_target_dir, safe="/")
            + "&amp;openfile=true"
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, expected_url, html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_contracts_partial_repairs_missing_lawyer_share_for_docx_link(
        self,
        mocked_list_user_shares,
        mocked_get_user_share,
        mocked_ensure_user_share,
        mocked_list_resources,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="lawyer-repair-share@example.com",
            email="lawyer-repair-share@example.com",
            password="secret",
            is_staff=True,
        )
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        NextcloudUserLink.objects.create(
            user=lawyer_user,
            nextcloud_user_id="nc-lawyer-repair",
            nextcloud_username="nc-lawyer-repair",
            nextcloud_email=lawyer_user.email,
        )
        mocked_list_user_shares.return_value = {}
        mocked_get_user_share.return_value = None
        mocked_ensure_user_share.return_value = NextcloudShare(
            share_id="58",
            path=self.performer.contract_project_disk_folder,
            share_with="nc-lawyer-repair",
            permissions=15,
            target_path="/Shared/000 Иванов ИИ",
        )
        mocked_list_resources.return_value = [
            {
                "name": self.performer.contract_file,
                "path": f"{self.performer.contract_project_disk_folder}/{self.performer.contract_file}",
                "type": "file",
                "file_id": "4476",
            }
        ]
        self.client.force_login(lawyer_user)

        response = self.client.get(reverse("contracts_partial"))

        expected_url = (
            "https://cloud.example.com/apps/files/files/4476?dir="
            + quote("/Shared/000 Иванов ИИ", safe="/")
            + "&amp;openfile=true"
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, expected_url, html=False)
        mocked_ensure_user_share.assert_called_once_with(
            "cloud-admin",
            self.performer.contract_project_disk_folder,
            "nc-lawyer-repair",
            permissions=15,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share", side_effect=NextcloudApiError("silent"))
    @patch("nextcloud_app.api.NextcloudApiClient.get_user_share", return_value=None)
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources")
    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", return_value={})
    def test_contracts_partial_falls_back_to_file_redirect_when_target_unknown(
        self,
        _mocked_list_user_shares,
        mocked_list_resources,
        _mocked_get_user_share,
        _mocked_ensure_user_share,
    ):
        lawyer_user = get_user_model().objects.create_user(
            username="lawyer-file-id@example.com",
            email="lawyer-file-id@example.com",
            password="secret",
            is_staff=True,
        )
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user.groups.add(lawyer_group)
        NextcloudUserLink.objects.create(
            user=lawyer_user,
            nextcloud_user_id="nc-lawyer-file-id",
            nextcloud_username="nc-lawyer-file-id",
            nextcloud_email=lawyer_user.email,
        )
        self.performer.contract_project_file_id = "4477"
        self.performer.contract_project_folder_file_id = "4478"
        self.performer.save(update_fields=["contract_project_file_id", "contract_project_folder_file_id"])
        self.client.force_login(lawyer_user)

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'href="https://cloud.example.com/f/4477"', html=False)
        self.assertContains(response, 'href="https://cloud.example.com/f/4478"', html=False)
        mocked_list_resources.assert_not_called()

    def test_contract_template_form_saves_multiple_groups_and_products(self):
        second_group = GroupMember.objects.create(
            short_name="Казахстан",
            country_name="Казахстан",
            country_code="398",
            country_alpha2="KZ",
            position=2,
        )
        country = OKSMCountry.objects.create(
            number=398,
            code="398",
            short_name="Казахстан",
            alpha2="KZ",
            alpha3="KAZ",
        )
        product = Product.objects.create(
            short_name="CT-TDD",
            name_en="TDD",
            name_ru="TDD",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        second_product = Product.objects.create(
            short_name="CT-OVR",
            name_en="OVR",
            name_ru="OVR",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Обзор",
        )
        form = ContractTemplateForm(
            data={
                "group_member_ids": [str(self.group_member.pk), str(second_group.pk)],
                "product_ids": [str(product.pk), str(second_product.pk)],
                "contract_type": "smz",
                "party": "individual",
                "country": str(country.pk),
                "sample_name": "",
                "version": "",
                "section_ids": ["__all__"],
            },
            files={
                "file": SimpleUploadedFile(
                    "template.docx",
                    b"docx",
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )

        self.assertTrue(form.is_valid(), form.errors)
        template = form.save()

        self.assertEqual(
            set(template.group_members.values_list("pk", flat=True)),
            {self.group_member.pk, second_group.pk},
        )
        self.assertEqual(template.group_member_id, self.group_member.pk)
        self.assertEqual(
            set(template.products.values_list("pk", flat=True)),
            {product.pk, second_product.pk},
        )
        self.assertEqual(template.product_id, product.pk)
        self.assertTrue(template.sample_name.startswith("RU-KZ Шаблон договора ФЗЛ СМЗ KAZ_CT-TDD-CT-OVR-Общий_v"))

    def test_contract_template_table_renders_all_group_for_unscoped_template(self):
        ContractTemplate.objects.create(
            group_member=None,
            product=self.product,
            contract_type="gph",
            party="individual",
            country_name="Россия",
            country_code="643",
            sample_name="Все Шаблон договора ФЗЛ ГПХ RUS_DD-Общий_v1",
            version="1",
            file=SimpleUploadedFile(
                "template.docx",
                b"docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )

        response = self.client.get(reverse("ct_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "<td class=\"text-nowrap\">Все</td>", html=False)

    def test_contract_template_table_renders_all_product_for_unscoped_template(self):
        ContractTemplate.objects.create(
            group_member=self.group_member,
            product=None,
            contract_type="gph",
            party="individual",
            country_name="Россия",
            country_code="643",
            sample_name="RU Шаблон договора ФЗЛ ГПХ RUS_Все-Общий_v1",
            version="1",
            file=SimpleUploadedFile(
                "template.docx",
                b"docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )

        response = self.client.get(reverse("ct_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "<td class=\"text-nowrap\">Все</td>", html=False)

    def test_contracts_development_partial_renders_client_contract_projects_table(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7101,
            sub_number=2,
            contract_number="IMC/7101-RU/05-26",
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Договорная строка",
            status="Разрабатывается проект договора",
            year=2026,
            evaluation_date=date(2026, 1, 1),
            service_term_months="1.0",
            preliminary_report_date=date(2026, 2, 15),
            final_report_term_weeks="2.0",
            final_report_date=date(2026, 3, 1),
            advance_percent="30",
            advance_term_days=5,
            preliminary_report_percent="20",
            preliminary_report_term_days=6,
            final_report_percent="50",
            final_report_term_days=14,
            registration_region="Москва",
            asset_owner="ООО Владелец",
        )

        response = self.client.get(reverse("contracts_development_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Проекты договоров с клиентами")
        self.assertContains(response, "Вид соглашения")
        self.assertContains(response, 'data-col="tkp-id"', html=False)
        self.assertContains(response, self.proposal.short_uid)
        self.assertContains(response, 'data-col="sub-number"', html=False)
        self.assertContains(response, '<td data-col="sub-number">2</td>', html=False)
        self.assertContains(response, 'data-col="contract-number"', html=False)
        self.assertContains(response, "Номер договора")
        self.assertContains(response, '<td data-col="contract-number">IMC/7101-RU/05-26</td>', html=False)
        self.assertContains(response, ">№</th>", html=False)
        self.assertContains(response, "Договор ID")
        self.assertContains(response, 'proposal-split-header-prefix">Заказчик: ', html=False)
        self.assertContains(response, 'proposal-split-header-prefix">Владелец: ', html=False)
        self.assertContains(response, 'id="contracts-drafts-colpicker-wrap"', html=False)
        self.assertContains(response, 'id="contracts-drafts-col-number"', html=False)
        self.assertContains(response, 'id="contracts-drafts-col-status"', html=False)
        self.assertContains(response, 'data-col="region"', html=False)
        self.assertContains(response, "Москва")
        self.assertContains(response, "ООО Владелец")
        self.assertContains(response, reverse("contracts_project_registration_create"))
        self.assertContains(response, reverse("contracts_project_registration_edit", args=[contract_project.pk]))
        self.assertContains(response, "data-queued-row-order", html=False)
        self.assertContains(response, reverse("contracts_project_registration_row_order"), html=False)
        self.assertContains(response, f'data-row-order-id="{contract_project.pk}"', html=False)
        self.assertContains(response, 'data-row-order-payload-field="ordered_contract_project_ids"', html=False)
        self.assertContains(response, 'data-row-order-status', html=False)
        content = response.content.decode()
        self.assertLess(content.index("contracts-drafts-table"), content.index("contracts-payment-schedule-table"))
        drafts_table_html = content[
            content.index("contracts-drafts-table") : content.index("contracts-payment-schedule-table")
        ]
        self.assertIn("contracts-payment-quick-edit me-1", drafts_table_html)
        self.assertContains(response, "Сроки и порядок платежей")
        self.assertContains(response, 'id="contracts-payment-colpicker-wrap"', html=False)
        self.assertContains(response, 'id="contracts-payment-col-evaluation-date"', html=False)
        self.assertContains(response, 'id="contracts-payment-col-advance-percent" data-default-hidden="true"', html=False)
        self.assertContains(response, "contracts-payment-quick-edit me-1", html=False)
        self.assertContains(response, 'data-col="project-id"', html=False)
        self.assertContains(response, "Этап 1")
        self.assertContains(response, "01.01.2026")
        self.assertContains(response, "15.01.2026")
        self.assertContains(response, "15.02.2026")
        self.assertContains(response, "1,0")
        self.assertContains(response, "2,0")
        self.assertContains(response, "30%")
        self.assertContains(response, "14")
        self.assertEqual(response.context["contract_payment_schedule_rows"][0]["start_date"], date(2026, 1, 15))
        self.assertContains(response, "Разрабатывается проект договора")
        self.assertNotContains(response, 'data-contract-project-status', html=False)
        self.assertNotContains(response, 'btn btn-link p-0 reg-status-btn', html=False)

    def test_contract_project_registration_defaults_to_draft_status(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7102,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Договор со статусом по умолчанию",
            year=2026,
        )

        self.assertEqual(contract_project.status, "Разрабатывается проект договора")

    def test_contract_project_registration_status_update_changes_status(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7103,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Договор для смены статуса",
            year=2026,
        )

        response = self.client.post(
            reverse("contracts_project_registration_status_update", args=[contract_project.pk]),
            {"status": "Договор подписан ЭЦП"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        contract_project.refresh_from_db()
        self.assertEqual(contract_project.status, "Договор подписан ЭЦП")

    def test_contract_project_registration_status_update_rejects_unknown_status(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7104,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Договор с ошибочным статусом",
            year=2026,
        )

        response = self.client.post(
            reverse("contracts_project_registration_status_update", args=[contract_project.pk]),
            {"status": "Неизвестно"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()["ok"])
        contract_project.refresh_from_db()
        self.assertEqual(contract_project.status, "Разрабатывается проект договора")

    def test_contracts_development_partial_groups_rows_by_number(self):
        second_proposal = ProposalRegistration.objects.create(
            number=7401,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Второй ТКП в группе",
            year=2026,
        )
        first = ContractProjectRegistration.objects.create(
            number=7401,
            sub_number=1,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            year=2026,
        )
        second = ContractProjectRegistration.objects.create(
            number=7401,
            sub_number=2,
            proposal_registration=second_proposal,
            group_member=self.group_member,
            type=self.product,
            name="Второй договор",
            year=2026,
        )
        third = ContractProjectRegistration.objects.create(
            number=7402,
            sub_number=1,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Третий договор",
            year=2026,
        )

        response = self.client.get(reverse("contracts_development_partial"))
        registrations = list(response.context["registrations"])
        payment_rows = list(response.context["contract_payment_schedule_rows"])

        self.assertEqual(response.status_code, 200)
        self.assertEqual([row.pk for row in registrations], [first.pk, second.pk, third.pk])
        self.assertTrue(registrations[0].is_first_for_number)
        self.assertTrue(registrations[0].has_next_for_number)
        self.assertTrue(registrations[0].has_next_for_different_contract_in_number_group)
        self.assertTrue(registrations[0].has_next_for_different_tkp_in_number_group)
        self.assertFalse(registrations[1].has_next_for_different_tkp_in_number_group)
        self.assertFalse(registrations[1].is_first_for_number)
        self.assertTrue(registrations[1].is_continuation)
        self.assertTrue(registrations[1].is_first_for_tkp)
        self.assertFalse(registrations[1].has_next_for_different_contract_in_number_group)
        self.assertTrue(registrations[2].is_first_for_number)
        self.assertContains(response, self.proposal.short_uid)
        self.assertContains(response, second_proposal.short_uid)

        self.assertEqual(len(payment_rows), 3)
        self.assertTrue(payment_rows[0]["is_first_for_number"])
        self.assertTrue(payment_rows[0]["has_next_for_number"])
        self.assertFalse(payment_rows[1]["is_first_for_number"])
        self.assertTrue(payment_rows[1]["is_number_continuation"])
        self.assertTrue(payment_rows[2]["is_first_for_number"])

        self.assertContains(response, "contracts-drafts-number-has-next", html=False)
        self.assertContains(response, "contracts-drafts-number-contract-has-next", html=False)
        self.assertContains(response, "contracts-drafts-tkp-separator-has-next", html=False)
        self.assertContains(response, "contracts-drafts-number-continuation", html=False)
        self.assertContains(response, "contracts-payment-number-has-next", html=False)
        self.assertContains(response, "contracts-payment-number-contract-has-next", html=False)
        self.assertNotContains(response, "contracts-payment-stage-has-next", html=False)
        self.assertContains(response, "contracts-payment-stage-continuation", html=False)
        self.assertContains(response, '<td class="col-project-number" data-col="number"></td>', html=False)
        self.assertContains(response, '<td data-col="number"></td>', html=False)

    def test_contracts_project_registration_row_order_persists_full_order(self):
        first = ContractProjectRegistration.objects.create(
            number=7501,
            position=1,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            year=2026,
        )
        second = ContractProjectRegistration.objects.create(
            number=7502,
            position=2,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Второй договор",
            year=2026,
        )
        third = ContractProjectRegistration.objects.create(
            number=7503,
            position=3,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Третий договор",
            year=2026,
        )

        response = self.client.post(
            reverse("contracts_project_registration_row_order"),
            data=json.dumps(
                {
                    "ordered_contract_project_ids": [second.pk, first.pk, third.pk],
                    "base_order_signature": f"{first.pk}:{second.pk}:{third.pk}",
                }
            ),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            list(ContractProjectRegistration.objects.order_by("position", "id").values_list("pk", flat=True)),
            [second.pk, first.pk, third.pk],
        )

    def test_contracts_project_registration_row_order_returns_conflict_for_stale_signature(self):
        first = ContractProjectRegistration.objects.create(
            number=7601,
            position=1,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            year=2026,
        )
        second = ContractProjectRegistration.objects.create(
            number=7602,
            position=2,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Второй договор",
            year=2026,
        )
        third = ContractProjectRegistration.objects.create(
            number=7603,
            position=3,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Третий договор",
            year=2026,
        )
        ContractProjectRegistration.objects.filter(pk=first.pk).update(position=2)
        ContractProjectRegistration.objects.filter(pk=second.pk).update(position=1)

        response = self.client.post(
            reverse("contracts_project_registration_row_order"),
            data=json.dumps(
                {
                    "ordered_contract_project_ids": [first.pk, second.pk, third.pk],
                    "base_order_signature": f"{first.pk}:{second.pk}:{third.pk}",
                }
            ),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 409)
        self.assertEqual(response.json()["current_contract_project_ids"], [second.pk, first.pk, third.pk])

    def test_contracts_development_partial_does_not_render_project_registry_rows(self):
        response = self.client.get(reverse("contracts_development_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Проекты договоров с клиентами")
        self.assertContains(response, "Пока нет данных.")
        self.assertNotContains(response, self.project.name)
        self.assertNotContains(response, self.proposal.short_uid)

    def test_contracts_project_registration_create_form_targets_contracts_drafts_pane(self):
        response = self.client.get(reverse("contracts_project_registration_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'hx-post="%s"' % reverse("contracts_project_registration_create"), html=False)
        self.assertContains(response, 'hx-target="#contracts-drafts-pane"', html=False)
        self.assertContains(response, 'id="contracts-sub-number-input"', html=False)
        self.assertContains(response, 'id="contracts-contract-number-input"', html=False)
        self.assertContains(response, 'name="contract_number"', html=False)
        self.assertContains(response, "Договор ID")
        self.assertContains(response, 'id="contracts-project-short-uid-input"', html=False)
        self.assertContains(response, 'id="contracts-project-group-order-map"', html=False)
        self.assertContains(response, 'id="contracts-project-group-alpha2-map"', html=False)
        self.assertContains(response, 'id="contracts-project-proposal-sub-number-map"', html=False)
        self.assertContains(response, "attachContractProjectShortUidPreview", html=False)
        self.assertContains(response, 'id="contracts-proposal-registration-select"', html=False)
        self.assertContains(response, 'id="contracts-proposal-registration-display"', html=False)
        self.assertContains(response, 'id="contracts-proposal-prefill-btn"', html=False)
        self.assertContains(response, "bi-arrow-clockwise", html=False)
        self.assertContains(
            response,
            reverse("contracts_project_registration_prefill_from_proposal", args=[0]),
            html=False,
        )
        self.assertContains(
            response,
            f'<option value="{self.proposal.pk}">{self.proposal.short_uid} DD {self.proposal.name}</option>',
            html=False,
        )
        self.assertContains(response, 'id="registration-products-container"', html=False)
        self.assertContains(response, 'id="registration-add-product"', html=False)
        self.assertContains(response, "Сроки")
        self.assertContains(response, "График оплаты")
        self.assertContains(response, "Общий для всех этапов", html=False)
        self.assertContains(response, 'name="payment_schedule_common"', html=False)
        self.assertContains(response, "proposal-terms-table")
        self.assertContains(response, "proposal-payment-schedule-editor")
        self.assertContains(response, "proposal-payment-schedule-editor-body")
        self.assertContains(response, 'name="evaluation_date"', html=False)
        self.assertContains(response, 'id="contract-stage-terms-tbody"', html=False)
        self.assertContains(response, "proposal-stage-terms-row")
        self.assertContains(response, 'name="advance_percent"', html=False)
        self.assertContains(response, 'name="final_report_percent"', html=False)
        self.assertContains(response, self.proposal.short_uid)
        self.assertNotContains(response, "Дедлайн")

    def test_contracts_project_registration_prefill_from_proposal_create_does_not_persist(self):
        self.proposal.name = "ТКП для prefill"
        self.proposal.customer = "ООО Заказчик prefill"
        self.proposal.sub_number = 3
        self.proposal.evaluation_date = date(2026, 1, 1)
        self.proposal.service_term_months = Decimal("1.5")
        self.proposal.preliminary_report_date = date(2026, 2, 15)
        self.proposal.final_report_term_weeks = Decimal("2.0")
        self.proposal.final_report_date = date(2026, 3, 1)
        self.proposal.advance_percent = Decimal("30")
        self.proposal.stage_payloads_json = [
            {
                "product_id": self.product.pk,
                "evaluation_date": "01.01.2026",
                "service_term_months": "1.5",
                "preliminary_report_date": "15.02.2026",
                "final_report_term_weeks": "2.0",
                "final_report_date": "01.03.2026",
                "next_stage_delay_days": "0",
                "payment_schedule_common": True,
                "advance_percent": "30",
                "advance_term_days": "5",
                "preliminary_report_percent": "20",
                "preliminary_report_term_days": "6",
                "final_report_percent": "50.00",
                "final_report_term_days": "14",
            }
        ]
        self.proposal.save()

        count_before = ContractProjectRegistration.objects.count()
        response = self.client.get(
            reverse("contracts_project_registration_prefill_from_proposal", args=[self.proposal.pk])
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(ContractProjectRegistration.objects.count(), count_before)
        self.assertContains(response, 'value="ТКП для prefill"', html=False)
        self.assertContains(response, 'value="ООО Заказчик prefill"', html=False)
        self.assertContains(response, f'value="{self.proposal.number}"', html=False)
        self.assertContains(response, f'option value="{self.proposal.pk}" selected', html=False)
        self.assertContains(response, 'value="01.01.2026"', html=False)
        self.assertContains(response, 'value="1.5"', html=False)
        self.assertContains(response, 'hx-post="%s"' % reverse("contracts_project_registration_create"), html=False)

    def test_contracts_project_registration_prefill_from_proposal_edit_preserves_contract_identity(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=8101,
            sub_number=1,
            contract_number="CUSTOM-NUM",
            group_member=self.group_member,
            type=self.product,
            name="Старое название",
            status="Отправлен проект договора",
            year=2026,
            customer="Старый заказчик",
        )
        second_proposal = ProposalRegistration.objects.create(
            number=8102,
            sub_number=2,
            group_member=self.group_member,
            type=self.product,
            name="Новое название из ТКП",
            year=2026,
            customer="Новый заказчик из ТКП",
            evaluation_date=date(2026, 6, 1),
            service_term_months=Decimal("3.0"),
            stage_payloads_json=[
                {
                    "product_id": self.product.pk,
                    "evaluation_date": "01.06.2026",
                    "service_term_months": "3.0",
                    "preliminary_report_date": "01.09.2026",
                    "final_report_term_weeks": "4.0",
                    "final_report_date": "01.12.2026",
                    "next_stage_delay_days": "0",
                    "payment_schedule_common": True,
                }
            ],
        )

        response = self.client.get(
            reverse("contracts_project_registration_prefill_from_proposal", args=[second_proposal.pk]),
            {"registration": contract_project.pk},
        )

        self.assertEqual(response.status_code, 200)
        contract_project.refresh_from_db()
        self.assertEqual(contract_project.name, "Старое название")
        self.assertEqual(contract_project.customer, "Старый заказчик")
        self.assertContains(response, 'value="8101"', html=False)
        self.assertContains(response, 'value="CUSTOM-NUM"', html=False)
        self.assertContains(response, f'value="{contract_project.short_uid}"', html=False)
        self.assertContains(response, 'value="Новое название из ТКП"', html=False)
        self.assertContains(response, 'value="Новый заказчик из ТКП"', html=False)
        self.assertContains(response, 'value="01.06.2026"', html=False)
        self.assertContains(
            response,
            'hx-post="%s"' % reverse("contracts_project_registration_edit", args=[contract_project.pk]),
            html=False,
        )

    def test_contracts_project_registration_prefill_from_proposal_returns_404_for_missing_proposal(self):
        response = self.client.get(
            reverse("contracts_project_registration_prefill_from_proposal", args=[999999])
        )

        self.assertEqual(response.status_code, 404)

    def test_contracts_project_registration_edit_form_shows_stage_row_per_product(self):
        second_product = Product.objects.create(
            short_name="QAQC",
            name_en="QAQC",
            name_ru="QAQC",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        contract_project = ContractProjectRegistration.objects.create(
            number=7304,
            group_member=self.group_member,
            type=self.product,
            name="Договор с двумя этапами",
            status="Разрабатывается проект договора",
            year=2026,
        )
        ContractProjectRegistrationProduct.objects.bulk_create(
            [
                ContractProjectRegistrationProduct(registration=contract_project, product=self.product, rank=1),
                ContractProjectRegistrationProduct(registration=contract_project, product=second_product, rank=2),
            ]
        )

        response = self.client.get(reverse("contracts_project_registration_edit", args=[contract_project.pk]))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertEqual(content.count("proposal-stage-terms-row"), 2)
        self.assertContains(response, 'value="Этап 1"', html=False)
        self.assertContains(response, 'value="Этап 2"', html=False)

    def test_contracts_project_registration_create_persists_stage_payloads_json(self):
        second_product = Product.objects.create(
            short_name="QAQC",
            name_en="QAQC",
            name_ru="QAQC",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        response = self.client.post(
            reverse("contracts_project_registration_create"),
            {
                "number": 7004,
                "sub_number": 0,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk), str(second_product.pk)],
                "name": "Договор с этапами",
                "status": "Разрабатывается проект договора",
                "year": 2026,
                "evaluation_date": ["01.01.2026", "01.07.2026"],
                "service_term_months": ["1.0", "2.0"],
                "preliminary_report_date": ["01.02.2026", "01.09.2026"],
                "final_report_term_weeks": ["2.0", "3.0"],
                "final_report_date": ["15.02.2026", "30.09.2026"],
                "next_stage_delay_days": ["0", "0"],
            },
        )

        self.assertEqual(response.status_code, 200)
        created = ContractProjectRegistration.objects.get(number=7004, name="Договор с этапами")
        self.assertEqual(len(created.stage_payloads_json), 2)
        self.assertEqual(created.stage_payloads_json[0]["product_id"], str(self.product.pk))
        self.assertEqual(created.stage_payloads_json[1]["product_id"], str(second_product.pk))
        self.assertEqual(created.stage_payloads_json[1]["service_term_months"], "2.0")

    def test_contracts_project_registration_create_persists_per_stage_payment_schedule(self):
        second_product = Product.objects.create(
            short_name="QAQC",
            name_en="QAQC",
            name_ru="QAQC",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        response = self.client.post(
            reverse("contracts_project_registration_create"),
            {
                "number": 7005,
                "sub_number": 0,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk), str(second_product.pk)],
                "name": "Договор с разным графиком оплаты",
                "status": "Разрабатывается проект договора",
                "year": 2026,
                "payment_schedule_common": "false",
                "evaluation_date": ["01.01.2026", "01.07.2026"],
                "service_term_months": ["1.0", "2.0"],
                "preliminary_report_date": ["01.02.2026", "01.09.2026"],
                "final_report_term_weeks": ["2.0", "3.0"],
                "final_report_date": ["15.02.2026", "30.09.2026"],
                "next_stage_delay_days": ["0", "0"],
                "advance_percent": ["30", "50"],
                "advance_term_days": ["5", "10"],
                "preliminary_report_percent": ["20", "30"],
                "preliminary_report_term_days": ["6", "7"],
                "final_report_term_days": ["14", "21"],
            },
        )

        self.assertEqual(response.status_code, 200)
        created = ContractProjectRegistration.objects.get(number=7005, name="Договор с разным графиком оплаты")
        self.assertFalse(created.stage_payloads_json[0]["payment_schedule_common"])
        self.assertEqual(created.stage_payloads_json[0]["advance_percent"], "30")
        self.assertEqual(created.stage_payloads_json[0]["final_report_percent"], "50.00")
        self.assertEqual(created.stage_payloads_json[1]["advance_percent"], "50")
        self.assertEqual(created.stage_payloads_json[1]["final_report_percent"], "20.00")
        self.assertEqual(created.advance_percent, Decimal("50"))
        self.assertEqual(created.final_report_term_days, 21)

    def test_contracts_project_registration_edit_form_shows_readonly_contract_id(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7303,
            group_member=self.group_member,
            type=self.product,
            name="Договор для ID",
            status="Разрабатывается проект договора",
            year=2026,
        )

        response = self.client.get(reverse("contracts_project_registration_edit", args=[contract_project.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Договор ID")
        self.assertContains(response, f'value="{contract_project.short_uid}"', html=False)
        self.assertContains(response, 'id="contracts-project-short-uid-input"', html=False)
        self.assertContains(response, "attachContractProjectShortUidPreview", html=False)

    def test_contract_project_short_uid_uses_proposal_and_contract_sub_numbers(self):
        proposal = ProposalRegistration.objects.create(
            number=7400,
            sub_number=4,
            group_member=self.group_member,
            type=self.product,
            name="ТКП с номером",
            year=2026,
        )
        contract_project = ContractProjectRegistration.objects.create(
            number=7401,
            sub_number=3,
            proposal_registration=proposal,
            group_member=self.group_member,
            type=self.product,
            name="Договор с ТКП",
            status="Разрабатывается проект договора",
            year=2026,
        )

        self.assertEqual(contract_project.short_uid, "7401430RU")

    def test_contract_project_short_uid_uses_zero_when_proposal_is_empty(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7402,
            sub_number=5,
            group_member=self.group_member,
            type=self.product,
            name="Договор без ТКП",
            status="Разрабатывается проект договора",
            year=2026,
        )

        self.assertEqual(contract_project.short_uid, "7402050RU")

    def test_contract_project_form_rejects_duplicate_sub_number_for_number_and_proposal(self):
        ContractProjectRegistration.objects.create(
            number=7403,
            sub_number=2,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            status="Разрабатывается проект договора",
            year=2026,
        )

        form = ContractProjectRegistrationForm(
            data={
                "number": 7403,
                "sub_number": 2,
                "proposal_registration": self.proposal.pk,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk)],
                "name": "Дубликат",
                "status": "Разрабатывается проект договора",
                "year": 2026,
            }
        )

        self.assertFalse(form.is_valid())
        self.assertIn("sub_number", form.errors)

    def test_contract_project_form_allows_same_sub_number_for_different_proposal(self):
        ContractProjectRegistration.objects.create(
            number=7404,
            sub_number=2,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            status="Разрабатывается проект договора",
            year=2026,
        )
        second_proposal = ProposalRegistration.objects.create(
            number=7404,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Другой ТКП",
            year=2026,
        )
        form = ContractProjectRegistrationForm(
            data={
                "number": 7404,
                "sub_number": 2,
                "proposal_registration": second_proposal.pk,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk)],
                "name": "Не дубликат",
                "status": "Разрабатывается проект договора",
                "year": 2026,
            }
        )

        self.assertTrue(form.is_valid(), form.errors)

    def test_contract_project_form_rejects_duplicate_generated_uid_for_same_proposal_sequence(self):
        ContractProjectRegistration.objects.create(
            number=7405,
            sub_number=2,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            status="Разрабатывается проект договора",
            year=2026,
        )
        second_proposal = ProposalRegistration.objects.create(
            number=7405,
            group_member=self.group_member,
            type=self.product,
            name="Другой ТКП с тем же №",
            year=2026,
        )
        form = ContractProjectRegistrationForm(
            data={
                "number": 7405,
                "sub_number": 2,
                "proposal_registration": second_proposal.pk,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk)],
                "name": "Дубликат по Договор ID",
                "status": "Разрабатывается проект договора",
                "year": 2026,
            }
        )

        self.assertFalse(form.is_valid())
        self.assertIn("sub_number", form.errors)

    def test_contracts_project_registration_create_accepts_multiple_products_on_one_row(self):
        second_product = Product.objects.create(
            short_name="QAQC",
            name_en="QAQC",
            name_ru="QAQC",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        second_proposal = ProposalRegistration.objects.create(
            number=7003,
            group_member=self.group_member,
            type=self.product,
            name="Второй ТКП",
            year=2026,
        )
        response = self.client.post(
            reverse("contracts_project_registration_create"),
            {
                "number": 7002,
                "sub_number": 3,
                "contract_number": "IMC/7002-RU/05-26",
                "proposal_registration": second_proposal.pk,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk), str(second_product.pk)],
                "name": "Новый договорный проект",
                "status": "Разрабатывается проект договора",
                "year": 2026,
            },
        )

        self.assertEqual(response.status_code, 200)
        created = ContractProjectRegistration.objects.get(number=7002, name="Новый договорный проект")
        self.assertEqual(created.sub_number, 3)
        self.assertEqual(created.contract_number, "IMC/7002-RU/05-26")
        self.assertEqual(created.proposal_registration_id, second_proposal.pk)
        self.assertEqual(created.type_short_display, "DD-QAQC")
        self.assertEqual(
            list(
                ContractProjectRegistrationProduct.objects.filter(registration=created)
                .order_by("rank")
                .values_list("product_id", flat=True)
            ),
            [self.product.pk, second_product.pk],
        )
        self.assertFalse(ProjectRegistration.objects.filter(number=7002, name="Новый договорный проект").exists())

    def test_contracts_project_registration_edit_updates_tkp_display_in_table(self):
        second_proposal = ProposalRegistration.objects.create(
            number=7601,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Новый ТКП после редактирования",
            year=2026,
        )
        ContractProjectRegistration.objects.create(
            number=7601,
            sub_number=1,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            year=2026,
        )
        contract_project = ContractProjectRegistration.objects.create(
            number=7601,
            sub_number=2,
            proposal_registration=self.proposal,
            group_member=self.group_member,
            type=self.product,
            name="Редактируемый договор",
            year=2026,
        )

        response = self.client.post(
            reverse("contracts_project_registration_edit", args=[contract_project.pk]),
            {
                "number": 7601,
                "sub_number": 2,
                "proposal_registration": second_proposal.pk,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk)],
                "name": "Редактируемый договор",
                "status": "Разрабатывается проект договора",
                "year": 2026,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            ContractProjectRegistration.objects.get(pk=contract_project.pk).proposal_registration_id,
            second_proposal.pk,
        )
        self.assertContains(response, second_proposal.short_uid)
        self.assertContains(response, self.proposal.short_uid)

    def test_contracts_development_partial_does_not_add_contract_separator_between_stages(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7501,
            group_member=self.group_member,
            type=self.product,
            name="Договор с этапами",
            year=2026,
        )
        second_product = Product.objects.create(
            short_name="QAQC",
            name_en="QAQC",
            name_ru="QAQC",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        ContractProjectRegistrationProduct.objects.bulk_create(
            [
                ContractProjectRegistrationProduct(registration=contract_project, product=self.product, rank=1),
                ContractProjectRegistrationProduct(registration=contract_project, product=second_product, rank=2),
            ]
        )

        response = self.client.get(reverse("contracts_development_partial"))
        payment_rows = list(response.context["contract_payment_schedule_rows"])

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(payment_rows), 2)
        self.assertTrue(payment_rows[0]["has_next_for_registration"])
        self.assertTrue(payment_rows[0]["has_next_for_number"])
        self.assertContains(response, "contracts-payment-stage-has-next", html=False)
        self.assertNotContains(response, "contracts-payment-number-contract-has-next", html=False)
        self.assertNotContains(response, "contracts-payment-tkp-separator-has-next", html=False)
        payment_table_html = response.content.decode()[
            response.content.decode().index("contracts-payment-schedule-table") :
        ]
        self.assertEqual(payment_table_html.count('aria-label="Строка сроков и порядка платежей"'), 2)

    def test_contracts_development_partial_renders_multiple_products_with_hyphen_in_type_column(self):
        second_product = Product.objects.create(
            short_name="QAQC",
            name_en="QAQC",
            name_ru="QAQC",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        contract_project = ContractProjectRegistration.objects.create(
            number=7201,
            group_member=self.group_member,
            type=self.product,
            name="Договорный мультипродукт",
            year=2026,
        )
        ContractProjectRegistrationProduct.objects.bulk_create(
            [
                ContractProjectRegistrationProduct(registration=contract_project, product=self.product, rank=1),
                ContractProjectRegistrationProduct(registration=contract_project, product=second_product, rank=2),
            ]
        )

        response = self.client.get(reverse("contracts_development_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, ">DD-QAQC<", html=False)

    def test_contracts_development_crud_does_not_mutate_project_registry(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=7301,
            group_member=self.group_member,
            type=self.product,
            name="Договорная строка до",
            status="Разрабатывается проект договора",
            year=2026,
            position=1,
        )
        ContractProjectRegistrationProduct.objects.create(
            registration=contract_project,
            product=self.product,
            rank=1,
        )
        second_contract_project = ContractProjectRegistration.objects.create(
            number=7302,
            group_member=self.group_member,
            type=self.product,
            name="Вторая договорная строка",
            status="Разрабатывается проект договора",
            year=2026,
            position=2,
        )
        project_name_before = self.project.name
        project_count_before = ProjectRegistration.objects.count()

        response = self.client.post(
            reverse("contracts_project_registration_edit", args=[contract_project.pk]),
            {
                "number": 7301,
                "sub_number": 4,
                "proposal_registration": self.proposal.pk,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [str(self.product.pk)],
                "name": "Договорная строка после",
                "status": "Отправлен проект договора",
                "year": 2027,
                "evaluation_date": "2026-01-01",
                "service_term_months": "1.5",
                "preliminary_report_date": "2026-02-15",
                "final_report_term_weeks": "2.0",
                "final_report_date": "2026-03-01",
                "advance_percent": "30",
                "advance_term_days": "5",
                "preliminary_report_percent": "20",
                "preliminary_report_term_days": "6",
                "final_report_term_days": "14",
            },
        )
        self.assertEqual(response.status_code, 200)
        contract_project.refresh_from_db()
        self.assertEqual(contract_project.evaluation_date, date(2026, 1, 1))
        self.assertEqual(str(contract_project.service_term_months), "1.5")
        self.assertEqual(contract_project.preliminary_report_date, date(2026, 2, 15))
        self.assertEqual(str(contract_project.final_report_term_weeks), "2.0")
        self.assertEqual(contract_project.final_report_date, date(2026, 3, 1))
        self.assertEqual(str(contract_project.advance_percent), "30.00")
        self.assertEqual(contract_project.advance_term_days, 5)
        self.assertEqual(str(contract_project.preliminary_report_percent), "20.00")
        self.assertEqual(contract_project.preliminary_report_term_days, 6)
        self.assertEqual(str(contract_project.final_report_percent), "50.00")
        self.assertEqual(contract_project.final_report_term_days, 14)
        self.project.refresh_from_db()
        self.assertEqual(self.project.name, project_name_before)
        self.assertEqual(ProjectRegistration.objects.count(), project_count_before)

        response = self.client.post(reverse("contracts_project_registration_move_down", args=[contract_project.pk]))
        self.assertEqual(response.status_code, 200)
        self.project.refresh_from_db()
        self.assertEqual(self.project.name, project_name_before)
        self.assertEqual(ProjectRegistration.objects.count(), project_count_before)

        response = self.client.post(reverse("contracts_project_registration_delete", args=[second_contract_project.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertTrue(ProjectRegistration.objects.filter(pk=self.project.pk).exists())
        self.assertFalse(ContractProjectRegistration.objects.filter(pk=second_contract_project.pk).exists())

    @patch("contracts_app.views._upload_scan_to_cloud_bytes", return_value="https://cloud.example.com/s/scan")
    def test_contract_scan_upload_returns_storage_label(self, _mock_upload):
        upload = SimpleUploadedFile("scan.pdf", b"pdf-data", content_type="application/pdf")
        with self.settings(MEDIA_ROOT=self.media_root):
            response = self.client.post(
                reverse("contract_scan_upload", args=[self.performer.pk]),
                {"contract_employee_scan": upload},
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["storage_label"], "Nextcloud")
        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_employee_scan.name, "")
        self.assertEqual(self.performer.contract_employee_scan_link, "https://cloud.example.com/s/scan")

    def test_contract_signing_modal_uses_cloud_link_for_current_file(self):
        self.performer.contract_scan_document = "Договор 7001_Иванов ИИ_1п.pdf"
        self.performer.contract_employee_scan_link = "https://cloud.example.com/s/current-scan"
        self.performer.save(update_fields=["contract_scan_document", "contract_employee_scan_link"])

        response = self.client.get(reverse("contracts_signing_edit", args=[self.performer.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "https://cloud.example.com/s/current-scan", html=False)
        self.assertContains(response, "Договор 7001_Иванов ИИ_1п.pdf", html=False)

    def test_contracts_partial_marks_signing_row_as_having_scan_from_cloud_fields(self):
        self.performer.contract_scan_document = "Договор 7001_Иванов ИИ_1п.pdf"
        self.performer.contract_employee_scan_link = "https://cloud.example.com/s/current-scan"
        self.performer.contract_employee_scan = ""
        self.performer.save(update_fields=["contract_scan_document", "contract_employee_scan_link", "contract_employee_scan"])

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-has-scan="1"', html=False)

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares")
    def test_contracts_partial_builds_nextcloud_disk_link_from_user_share_target(self, mocked_list_user_shares):
        admin_link = NextcloudUserLink.objects.create(
            user=self.user,
            nextcloud_user_id="nc-admin",
            nextcloud_username="nc-admin",
            nextcloud_email=self.user.email,
        )
        mocked_list_user_shares.return_value = {
            self.performer.contract_project_disk_folder: NextcloudShare(
                share_id="55",
                path=self.performer.contract_project_disk_folder,
                share_with=admin_link.nextcloud_user_id,
                permissions=1,
                target_path="/Shared/000 Иванов ИИ",
            )
        }

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "/apps/files/files?dir=/Shared/000%20%D0%98%D0%B2%D0%B0%D0%BD%D0%BE%D0%B2%20%D0%98%D0%98",
            html=False,
        )

    @patch("nextcloud_app.api.NextcloudApiClient.list_user_shares", side_effect=NextcloudApiError("temporary outage"))
    def test_contracts_partial_falls_back_to_generic_folder_url_when_share_resolution_fails(self, _mocked_list_user_shares):
        NextcloudUserLink.objects.create(
            user=self.user,
            nextcloud_user_id="nc-admin",
            nextcloud_username="nc-admin",
            nextcloud_email=self.user.email,
        )

        response = self.client.get(reverse("contracts_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "/apps/files/files?dir=/Corporate%20Root/2026/Project/09%20%D0%94%D0%BE%D0%B3%D0%BE%D0%B2%D0%BE%D1%80%D1%8B/000%20%D0%98%D0%B2%D0%B0%D0%BD%D0%BE%D0%B2%20%D0%98%D0%98",
            html=False,
        )

    @patch("contracts_app.views._upload_scan_to_cloud_bytes", return_value="")
    def test_contract_signing_edit_keeps_existing_local_scan_when_cloud_upload_fails(self, _mock_upload):
        with self.settings(MEDIA_ROOT=self.media_root):
            self.performer.contract_employee_scan.save("existing-scan.pdf", ContentFile(b"existing"), save=True)
            old_name = self.performer.contract_employee_scan.name
            old_path = self.performer.contract_employee_scan.path

            response = self.client.post(
                reverse("contracts_signing_edit", args=[self.performer.pk]),
                {
                    "contract_employee_scan": SimpleUploadedFile(
                        "new-scan.pdf",
                        b"new-data",
                        content_type="application/pdf",
                    ),
                },
            )

        self.assertEqual(response.status_code, 200)
        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_employee_scan.name, old_name)
        self.assertTrue(os.path.exists(old_path))
        self.assertIn("contract_employee_scan", response.context["form"].errors)

    def test_contract_signing_edit_noop_does_not_clear_sibling_local_scan_fields(self):
        sibling = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor="Иванов Иван Иванович",
            contract_batch_id=self.performer.contract_batch_id,
            contract_project_disk_folder=self.performer.contract_project_disk_folder,
        )
        with self.settings(MEDIA_ROOT=self.media_root):
            sibling.contract_employee_scan.save("sibling-employee.pdf", ContentFile(b"employee"), save=True)
            sibling.contract_signed_scan_file.save("sibling-signed.pdf", ContentFile(b"signed"), save=True)
            employee_name = sibling.contract_employee_scan.name
            signed_name = sibling.contract_signed_scan_file.name

            response = self.client.post(
                reverse("contracts_signing_edit", args=[self.performer.pk]),
                {},
            )

        self.assertEqual(response.status_code, 200)
        sibling.refresh_from_db()
        self.assertEqual(sibling.contract_employee_scan.name, employee_name)
        self.assertEqual(sibling.contract_signed_scan_file.name, signed_name)


class ContractVariableBindingDisplayTests(TestCase):
    def test_binding_display_uses_current_app_section_label(self):
        variable = ContractVariable(
            source_section="contracts",
            source_table="contract_details",
            source_column="full_name",
        )

        self.assertEqual(
            variable.binding_display,
            "Значения столбца «ФИО» "
            "из таблицы «Реквизиты физлиц-исполнителей» "
            "раздела «Договоры»",
        )

    def test_binding_display_handles_contract_details_inn(self):
        variable = ContractVariable(
            source_section="contracts",
            source_table="contract_details",
            source_column="inn",
        )

        self.assertEqual(
            variable.binding_display,
            "Значения столбца «ИНН» "
            "из таблицы «Реквизиты физлиц-исполнителей» "
            "раздела «Договоры»",
        )

    def test_performer_facsimile_variable_is_seeded_as_computed_field(self):
        variable = ContractVariable.objects.get(key="[[facsimile_prfrm]]")

        self.assertEqual(variable.description, "Подпись исполнителя")
        self.assertTrue(variable.is_computed)
        self.assertEqual(variable.binding_display, "Расчётное поле")

    def test_chapters_name_variable_description_mentions_stages(self):
        variable = ContractVariable.objects.get(key="[[chapters_name]]")

        self.assertEqual(
            variable.description,
            "Многоуровневый список: этапы, активы, разделы, подразделы",
        )

    def test_additional_contract_variables_are_seeded_as_computed_fields(self):
        expected = {
            "{{owner}}": "Владелец активов",
            "{{service_goal_genitive}}": "Цель оказания услуг в родительном падеже",
            "{{specialization}}": "Область специализации",
            "[[services]]": "Многоуровневый список: этапы, состав услуг",
        }

        variables = {
            variable.key: variable
            for variable in ContractVariable.objects.filter(key__in=expected)
        }

        self.assertEqual(set(variables), set(expected))
        for key, description in expected.items():
            self.assertEqual(variables[key].description, description)
            self.assertTrue(variables[key].is_computed)
            self.assertEqual(variables[key].binding_display, "Расчётное поле")


class ContractVariableResolverTests(TestCase):
    def test_contract_details_variables_resolve_from_expert_contract_details(self):
        user = get_user_model().objects.create_user(
            username="resolver-expert",
            password="secret",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=user, patronymic="Иванович")
        person = PersonRecord.objects.create(
            last_name="Иванов",
            first_name="Иван",
            middle_name="Иванович",
            full_name_genitive="Иванова Ивана Ивановича",
            gender="male",
            birth_date=date(1990, 5, 17),
            position=1,
        )
        employee.person_record = person
        employee.save(update_fields=["person_record"])
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            short_name_genitive="России",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        citizenship = CitizenshipRecord.objects.create(
            person=person,
            country=country,
            status="Гражданство",
            identifier="Паспорт",
            number="123456",
            position=1,
        )
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        ExpertContractDetails.objects.create(
            expert_profile=profile,
            citizenship_record=citizenship,
            inn="770123456789",
            passport_expiry_date=date(2030, 4, 27),
            bank_swift="SABRRUMM",
            bank_bik="044525225",
            settlement_account="40817810099910004312",
            corr_account="30101810400000000225",
            corr_bank_settlement_account="30101810945250000225",
            corr_bank_corr_account="30101810000000000000",
        )
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        project = ProjectRegistration.objects.create(
            number=7002,
            type=product,
            name="Договорный проект",
            customer='АО "Заказчик"',
            country=country,
            year=2026,
        )
        performer = Performer.objects.create(
            registration=project,
            employee=employee,
            executor="Иванов Иван Иванович",
        )

        variables = [
            ContractVariable(key="{{short_name}}", source_section="contacts", source_table="persons", source_column="short_name"),
            ContractVariable(key="{{name}}", source_section="projects", source_table="registration", source_column="customer"),
            ContractVariable(key="{{country}}", source_section="classifiers", source_table="oksm_countries", source_column="full_name"),
            ContractVariable(key="{{full_name_genitive}}", source_section="contracts", source_table="contract_details", source_column="full_name_genitive"),
            ContractVariable(key="{{citizenship_country}}", source_section="contracts", source_table="contract_details", source_column="citizenship_country"),
            ContractVariable(key="{{citizenship_identifier}}", source_section="contracts", source_table="contract_details", source_column="citizenship_identifier"),
            ContractVariable(key="{{passport_expiry_date}}", source_section="contracts", source_table="contract_details", source_column="passport_expiry"),
            ContractVariable(key="{{bank_swift}}", source_section="contracts", source_table="contract_details", source_column="swift"),
            ContractVariable(key="{{corr_account}}", source_section="contracts", source_table="contract_details", source_column="correspondent_account"),
            ContractVariable(key="{{corr_bank_settlement_account}}", source_section="contracts", source_table="contract_details", source_column="corr_bank_settlement"),
            ContractVariable(key="{{legacy_bank_swift}}", source_section="experts", source_table="contract_details", source_column="bank_swift"),
        ]

        replacements, lists = resolve_variables(performer, variables)

        self.assertEqual(lists, {})
        self.assertEqual(replacements["{{short_name}}"], "Иванов И.И.")
        self.assertEqual(replacements["{{name}}"], 'АО "Заказчик"')
        self.assertEqual(replacements["{{country}}"], "Российская Федерация")
        self.assertEqual(replacements["{{full_name_genitive}}"], "Иванова Ивана Ивановича")
        self.assertEqual(replacements["{{citizenship_country}}"], "Россия")
        self.assertEqual(replacements["{{citizenship_identifier}}"], "Паспорт")
        self.assertEqual(replacements["{{passport_expiry_date}}"], "27.04.2030")
        self.assertEqual(replacements["{{bank_swift}}"], "SABRRUMM")
        self.assertEqual(replacements["{{corr_account}}"], "30101810400000000225")
        self.assertEqual(replacements["{{corr_bank_settlement_account}}"], "30101810945250000225")
        self.assertEqual(replacements["{{legacy_bank_swift}}"], "SABRRUMM")

    def test_date_variables_use_saved_performer_contract_date(self):
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        project = ProjectRegistration.objects.create(
            number=7004,
            type=product,
            name="Проект с датой договора",
            year=2026,
        )
        performer = Performer.objects.create(
            registration=project,
            executor="Иванов Иван Иванович",
            contract_date=date(2026, 5, 7),
        )
        variables = [
            ContractVariable(key="{{year}}", is_computed=True),
            ContractVariable(key="{{day}}", is_computed=True),
            ContractVariable(key="{{month}}", is_computed=True),
        ]

        replacements, lists = resolve_variables(performer, variables)

        self.assertEqual(lists, {})
        self.assertEqual(replacements["{{year}}"], "2026")
        self.assertEqual(replacements["{{day}}"], "07")
        self.assertEqual(replacements["{{month}}"], "мая")

    def test_deadline_ru_uses_latest_gantt_deadline_for_performer_batch(self):
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        project = ProjectRegistration.objects.create(
            number=7005,
            type=product,
            name="Проект с дедлайнами графика",
            deadline=date(2030, 1, 1),
            year=2026,
        )
        performer_a = Performer.objects.create(
            registration=project,
            executor="Иванов Иван Иванович",
            asset_name="Карьер",
        )
        performer_b = Performer.objects.create(
            registration=project,
            executor="Иванов Иван Иванович",
            asset_name="Фабрика",
        )
        unrelated_performer = Performer.objects.create(
            registration=project,
            executor="Петров Петр Петрович",
            asset_name="Фабрика",
        )
        project.gantt_data = {
            "data": [
                {
                    "id": f"managed-performer-{performer_a.pk}",
                    "managed_source": "performer",
                    "performer_id": performer_a.pk,
                    "asset_name": "Карьер",
                    "deadline": "2026-05-12",
                },
                {
                    "id": f"managed-performer-{performer_b.pk}",
                    "managed_source": "performer",
                    "performer_id": performer_b.pk,
                    "asset_name": "Фабрика",
                    "deadline": "2026-05-20",
                },
                {
                    "id": f"managed-performer-{unrelated_performer.pk}",
                    "managed_source": "performer",
                    "performer_id": unrelated_performer.pk,
                    "asset_name": "Фабрика",
                    "deadline": "2026-06-30",
                },
            ],
            "links": [],
            "meta": {},
        }
        project.save(update_fields=["gantt_data"])
        project.refresh_from_db()
        performer_a.registration = project
        performer_b.registration = project
        variables = [ContractVariable(key="{{deadline_ru}}", is_computed=True)]

        replacements, lists = resolve_variables(
            performer_a,
            variables,
            all_performers=[performer_a, performer_b],
        )

        self.assertEqual(lists, {})
        self.assertEqual(replacements["{{deadline_ru}}"], "20 мая 2026 г.")

    def test_additional_scalar_variables_resolve_from_project_product_and_specialty(self):
        product = Product.objects.create(
            short_name="DD",
            display_name="Due Diligence",
            name_en="Due Diligence",
            name_ru="ДД",
        )
        ServiceGoalReport.objects.create(
            product=product,
            service_goal_genitive="Проведения due diligence",
            position=1,
        )
        project = ProjectRegistration.objects.create(
            number=7006,
            type=product,
            name="Проект с владельцем",
            asset_owner='АО "Владелец"',
            year=2026,
        )
        user = get_user_model().objects.create_user(
            username="specialist@example.com",
            password="secret",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=user, patronymic="Иванович")
        profile = ExpertProfile.objects.create(employee=employee, position=1)
        geology = ExpertSpecialty.objects.create(
            specialty="Геолог",
            specialization_area="Специалист по геологическому анализу",
            position=1,
        )
        mining = ExpertSpecialty.objects.create(
            specialty="Горный инженер",
            specialization_area="Специалист по горнотехнической экспертизе",
            position=2,
        )
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=geology, rank=1)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=mining, rank=2)
        geology_section = TypicalSection.objects.create(
            product=product,
            code="GEO",
            short_name="Geology",
            name_en="Geology",
            name_ru="Геология",
            position=1,
        )
        mining_section = TypicalSection.objects.create(
            product=product,
            code="MIN",
            short_name="Mining",
            name_en="Mining",
            name_ru="Горные работы",
            position=2,
        )
        TypicalSectionSpecialty.objects.create(section=geology_section, specialty=geology, rank=1)
        TypicalSectionSpecialty.objects.create(section=mining_section, specialty=mining, rank=1)
        geology_performer = Performer.objects.create(
            registration=project,
            employee=employee,
            executor="Иванов Иван Иванович",
            typical_section=geology_section,
            position=2,
        )
        mining_performer = Performer.objects.create(
            registration=project,
            employee=employee,
            executor="Иванов Иван Иванович",
            typical_section=mining_section,
            position=1,
        )
        variables = [
            ContractVariable(key="{{owner}}", is_computed=True),
            ContractVariable(key="{{service_goal_genitive}}", is_computed=True),
            ContractVariable(key="{{specialization}}", is_computed=True),
        ]

        replacements, lists = resolve_variables(
            geology_performer,
            variables,
            all_performers=[mining_performer, geology_performer],
        )

        self.assertEqual(lists, {})
        self.assertEqual(replacements["{{owner}}"], 'АО "Владелец"')
        self.assertEqual(replacements["{{service_goal_genitive}}"], "Проведения due diligence")
        self.assertEqual(replacements["{{specialization}}"], "горнотехнической экспертизе")

    def test_services_list_uses_typical_service_compositions_without_assets_or_sections(self):
        product = Product.objects.create(
            short_name="DD",
            display_name="Due Diligence",
            name_en="Due Diligence",
            name_ru="ДД",
        )
        project = ProjectRegistration.objects.create(
            number=7007,
            type=product,
            name="Проект с составом услуг",
            year=2026,
        )
        section_a = TypicalSection.objects.create(
            product=product,
            code="GEO",
            short_name="Geology",
            name_en="Geology",
            name_ru="Геология",
            position=1,
        )
        section_b = TypicalSection.objects.create(
            product=product,
            code="MIN",
            short_name="Mining",
            name_en="Mining",
            name_ru="Горные работы",
            position=2,
        )
        TypicalServiceComposition.objects.create(
            product=product,
            section=section_a,
            service_composition="Сбор данных\nАнализ геологии",
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=product,
            section=section_b,
            service_composition="",
            service_composition_editor_state={
                "html": "<p>Моделирование карьера</p>",
                "plain_text": "Моделирование карьера",
            },
            position=2,
        )
        performer_a = Performer.objects.create(
            registration=project,
            executor="Иванов Иван Иванович",
            asset_name="Карьер",
            typical_section=section_a,
            position=1,
        )
        performer_b = Performer.objects.create(
            registration=project,
            executor="Иванов Иван Иванович",
            asset_name="Фабрика",
            typical_section=section_b,
            position=2,
        )
        duplicate_section_performer = Performer.objects.create(
            registration=project,
            executor="Иванов Иван Иванович",
            asset_name="Другой актив",
            typical_section=section_a,
            position=3,
        )
        variables = [ContractVariable(key="[[services]]", is_computed=True)]

        replacements, lists = resolve_variables(
            performer_a,
            variables,
            all_performers=[performer_a, performer_b, duplicate_section_performer],
        )

        self.assertEqual(replacements, {})
        self.assertEqual(
            lists["[[services]]"],
            [
                (0, "Сбор данных"),
                (0, "Анализ геологии"),
                (0, "Моделирование карьера"),
            ],
        )

    def test_multistage_services_and_service_goal_use_stage_order(self):
        product_a = Product.objects.create(
            short_name="RFR",
            display_name="Red Flag Review",
            name_en="Red Flag Review",
            name_ru="РФР",
        )
        product_b = Product.objects.create(
            short_name="TDD",
            display_name="Technical Due Diligence",
            name_en="Technical Due Diligence",
            name_ru="ТДД",
        )
        ServiceGoalReport.objects.create(
            product=product_a,
            service_goal_genitive="Подготовки первого этапа",
            position=1,
        )
        ServiceGoalReport.objects.create(
            product=product_b,
            service_goal_genitive="Подготовки второго этапа",
            position=1,
        )
        project_a = ProjectRegistration.objects.create(
            number=7008,
            type=product_a,
            name="Многоэтапный проект",
            year=2026,
        )
        project_b = ProjectRegistration.objects.create(
            number=7008,
            type=product_b,
            name="Многоэтапный проект",
            year=2026,
        )
        project_a.refresh_from_db()
        project_b.refresh_from_db()
        section_a = TypicalSection.objects.create(
            product=product_a,
            code="RFR-01",
            short_name="Market",
            name_en="Market",
            name_ru="Рынок",
            position=1,
        )
        section_b = TypicalSection.objects.create(
            product=product_b,
            code="TDD-01",
            short_name="Tech",
            name_en="Technology",
            name_ru="Технология",
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=product_a,
            section=section_a,
            service_composition="Обзор рынка",
            position=1,
        )
        TypicalServiceComposition.objects.create(
            product=product_b,
            section=section_b,
            service_composition="Технический анализ",
            position=1,
        )
        performer_a = Performer.objects.create(
            registration=project_a,
            executor="Иванов Иван Иванович",
            asset_name="Карьер",
            typical_section=section_a,
        )
        performer_b = Performer.objects.create(
            registration=project_b,
            executor="Иванов Иван Иванович",
            asset_name="Фабрика",
            typical_section=section_b,
        )
        variables = [
            ContractVariable(key="{{service_goal_genitive}}", is_computed=True),
            ContractVariable(key="[[services]]", is_computed=True),
        ]

        replacements, lists = resolve_variables(
            performer_a,
            variables,
            all_performers=[performer_a, performer_b],
        )

        self.assertEqual(replacements["{{service_goal_genitive}}"], "Подготовки второго этапа")
        self.assertEqual(
            lists["[[services]]"],
            [
                (0, "Этап 1: RFR Red Flag Review"),
                (1, "Обзор рынка"),
                (0, "Этап 2: TDD Technical Due Diligence"),
                (1, "Технический анализ"),
            ],
        )

    def test_multi_stage_batch_chapters_name_groups_by_stage_and_contract_name_uses_last_stage(self):
        product_a = Product.objects.create(
            short_name="RFR",
            display_name="Red Flag Review",
            name_en="Red Flag Review",
            name_ru="РФР",
        )
        product_b = Product.objects.create(
            short_name="TDD",
            display_name="Technical Due Diligence",
            name_en="Technical Due Diligence",
            name_ru="ТДД",
        )
        project_a = ProjectRegistration.objects.create(
            number=7010,
            type=product_a,
            name="Multi-stage contract",
            year=2026,
        )
        project_b = ProjectRegistration.objects.create(
            number=7010,
            type=product_b,
            name="Multi-stage contract",
            year=2026,
        )
        project_a.refresh_from_db()
        project_b.refresh_from_db()
        section_a = TypicalSection.objects.create(
            product=product_a,
            code="RFR-01",
            short_name="Market",
            name_en="Market",
            name_ru="Рынок",
        )
        section_b = TypicalSection.objects.create(
            product=product_b,
            code="TDD-01",
            short_name="Tech",
            name_en="Technology",
            name_ru="Технология",
        )
        SectionStructure.objects.create(
            product=product_a,
            section=section_a,
            subsections="Обзор рынка",
            position=1,
        )
        SectionStructure.objects.create(
            product=product_b,
            section=section_b,
            subsections="Анализ технологии",
            position=1,
        )
        ContractSubject.objects.create(
            product=product_a,
            subject_text="предмет RFR",
            position=1,
        )
        ContractSubject.objects.create(
            product=product_b,
            subject_text="предмет TDD",
            position=2,
        )
        performer_a = Performer.objects.create(
            registration=project_a,
            executor="Иванов Иван Иванович",
            asset_name="Карьер",
            typical_section=section_a,
        )
        performer_b = Performer.objects.create(
            registration=project_b,
            executor="Иванов Иван Иванович",
            asset_name="Фабрика",
            typical_section=section_b,
        )
        variables = [
            ContractVariable(key="[[chapters_name]]", is_computed=True),
            ContractVariable(key="{{contract_name}}", is_computed=True),
        ]

        replacements, lists = resolve_variables(
            performer_a,
            variables,
            all_performers=[performer_a, performer_b],
        )

        self.assertEqual(replacements["{{contract_name}}"], "предмет TDD")
        self.assertEqual(
            lists["[[chapters_name]]"],
            [
                (0, "Этап 1: RFR Red Flag Review"),
                (1, "Карьер"),
                (2, "Рынок"),
                (3, "Обзор рынка"),
                (0, "Этап 2: TDD Technical Due Diligence"),
                (1, "Фабрика"),
                (2, "Технология"),
                (3, "Анализ технологии"),
            ],
        )

    def test_single_stage_chapters_name_and_contract_name_keep_legacy_behavior(self):
        product = Product.objects.create(
            short_name="DD",
            display_name="Due Diligence",
            name_en="Due Diligence",
            name_ru="ДД",
        )
        project = ProjectRegistration.objects.create(
            number=7011,
            type=product,
            name="Single-stage contract",
            year=2026,
        )
        section = TypicalSection.objects.create(
            product=product,
            code="DD-01",
            short_name="Market",
            name_en="Market",
            name_ru="Рынок",
        )
        SectionStructure.objects.create(
            product=product,
            section=section,
            subsections="Обзор рынка",
            position=1,
        )
        ContractSubject.objects.create(
            product=product,
            subject_text="предмет DD",
            position=1,
        )
        performer = Performer.objects.create(
            registration=project,
            executor="Иванов Иван Иванович",
            asset_name="Карьер",
            typical_section=section,
        )
        variables = [
            ContractVariable(key="[[chapters_name]]", is_computed=True),
            ContractVariable(key="{{contract_name}}", is_computed=True),
        ]

        replacements, lists = resolve_variables(performer, variables)

        self.assertEqual(replacements["{{contract_name}}"], "предмет DD")
        self.assertEqual(
            lists["[[chapters_name]]"],
            [
                (0, "Рынок"),
                (1, "Обзор рынка"),
            ],
        )

    def test_contacts_short_name_falls_back_to_performer_employee_person_record(self):
        user = get_user_model().objects.create_user(
            username="resolver-performer-person",
            password="secret",
        )
        person = PersonRecord.objects.create(
            last_name="Петров",
            first_name="Петр",
            middle_name="Петрович",
            position=1,
        )
        employee = Employee.objects.create(user=user, person_record=person)
        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        project = ProjectRegistration.objects.create(
            number=7003,
            type=product,
            name="Проект без профиля эксперта",
            year=2026,
        )
        performer = Performer.objects.create(
            registration=project,
            employee=employee,
            executor="Петров Петр Петрович",
        )
        variables = [
            ContractVariable(
                key="{{short_name}}",
                source_section="contacts",
                source_table="persons",
                source_column="short_name",
            ),
        ]

        replacements, lists = resolve_variables(performer, variables)

        self.assertEqual(lists, {})
        self.assertEqual(replacements["{{short_name}}"], "Петров П.П.")


class ContractsExecutionPartialTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="contracts-execution-admin",
            password="secret",
            is_staff=True,
        )
        admin_group, _ = Group.objects.get_or_create(name=ADMIN_GROUP)
        self.user.groups.add(admin_group)
        self.client.force_login(self.user)

        product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.project = ProjectRegistration.objects.create(
            number=8001,
            type=product,
            name="Проект исполнения договора",
            year=2026,
        )
        self.performer = Performer.objects.create(
            registration=self.project,
            employee=Employee.objects.create(user=self.user),
            executor="Иванов Иван Иванович",
            contract_batch_id=uuid.uuid4(),
            contract_number="EXEC-1",
            contract_signing_note="Договор подписан факсимиле",
        )

    def test_home_page_lists_performer_execution_after_conclusion(self):
        response = self.client.get(reverse("home"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        conclusion_index = content.index('data-contracts-section="in-progress"')
        execution_index = content.index('data-contracts-section="performer-execution"')
        self.assertLess(conclusion_index, execution_index)
        self.assertIn("'performer-execution':", content)
        self.assertIn("исполнение договора", content)
        self.assertContains(response, 'id="contracts-content-performer-execution"', html=False)

    def test_contracts_execution_partial_renders_payment_request_table(self):
        response = self.client.get(reverse("contracts_execution_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Заявки на оплату")
        content = response.content.decode("utf-8")
        self.assertIn('id="contracts-payment-request-section"', content)
        self.assertIn('class="payment-request-section payment-request-section--contracts"', content)
        self.assertIn(f'data-performer-id="{self.performer.pk}"', content)
        self.assertIn("payment-request-paid-toggle-icon", content)
        self.assertIn('data-payment-paid-toggle-url="/projects/performers/payment-paid-toggle/"', content)
        self.assertNotIn("Заключение договора", content)
        self.assertNotIn("payment-request-signing-status-cell", content)
        self.assertNotIn("payment-request-paid-toggle-header", content)
        section_start = content.index('id="contracts-payment-request-section"')
        section = content[section_start:section_start + 20000]
        self.assertEqual(section.count("payment-request-group-header"), 4)
        self.assertEqual(section.count("payment-request-paid-date-cell"), 2)
        self.assertIn('data-payment-paid-date="advance"', section)
        self.assertIn('data-payment-paid-date="final"', section)
        self.assertNotIn("Отправить заявку", section)
        self.assertNotIn('id="contracts-payment-request-send-btn"', section)
        self.assertNotIn("js-payment-request-actions", section)
        self.assertRegex(section, r'id="contracts-payment-request-master"[^>]*\sdisabled')
        self.assertRegex(section, rf'id="contracts-payment-request-sel-{self.performer.pk}"[^>]*\sdisabled')

