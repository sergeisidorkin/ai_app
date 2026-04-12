"""
Process a .docx template: replace ``{{variable}}`` placeholders with values
and ``[[variable]]`` placeholders with bulleted/numbered lists.

Handles the common Word issue where a single placeholder like ``{{inn}}``
is split across multiple XML runs (e.g. run1="{{", run2="inn", run3="}}").

Replacement text inherits the formatting of the run that contains the
opening ``{{`` of the placeholder, so bold/italic/font/size are preserved.
"""

from __future__ import annotations

import re
import uuid
from copy import deepcopy
from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


_PLACEHOLDER_RE = re.compile(r"\{\{[a-zA-Z][a-zA-Z0-9_]*\}\}")
_LIST_PLACEHOLDER_RE = re.compile(r"\[\[[a-zA-Z][a-zA-Z0-9_]*\]\]")
_CSS_COLOR_RE = re.compile(r"rgb\(\s*(\d{1,3})\s*,\s*(\d{1,3})\s*,\s*(\d{1,3})\s*\)", re.I)
_FONT_NAME_MAP = {
    "calibri": "Calibri",
    "cambria": "Cambria",
    "sans": "Arial",
    "serif": "Times New Roman",
    "monospace": "Courier New",
    "georgia": "Georgia",
    "times-new-roman": "Times New Roman",
}


def _parse_style_attr(value: str) -> dict[str, str]:
    result: dict[str, str] = {}
    for chunk in str(value or "").split(";"):
        if ":" not in chunk:
            continue
        key, item_value = chunk.split(":", 1)
        key = key.strip().lower()
        item_value = item_value.strip()
        if key:
            result[key] = item_value
    return result


def _normalize_css_color(value: str | None) -> str | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    if raw.startswith("#"):
        hex_value = raw[1:]
        if len(hex_value) == 3:
            hex_value = "".join(ch * 2 for ch in hex_value)
        if len(hex_value) == 6:
            return hex_value.upper()
        return None
    match = _CSS_COLOR_RE.fullmatch(raw)
    if match:
        rgb = [max(0, min(255, int(item))) for item in match.groups()]
        return "".join(f"{item:02X}" for item in rgb)
    named_colors = {
        "black": "000000",
        "white": "FFFFFF",
        "red": "FF0000",
        "green": "008000",
        "blue": "0000FF",
        "yellow": "FFFF00",
    }
    return named_colors.get(raw.lower())


def _merge_run_format(base: dict[str, object], **overrides) -> dict[str, object]:
    result = dict(base)
    result.update({key: value for key, value in overrides.items() if value is not None})
    return result


def _append_rich_run(runs: list[dict[str, object]], text: str, run_format: dict[str, object]) -> None:
    if text is None:
        return
    value = str(text)
    if not value:
        return
    if runs:
        style_keys = (set(runs[-1].keys()) | set(run_format.keys())) - {"text"}
    else:
        style_keys = set()
    if runs and all(runs[-1].get(key) == run_format.get(key) for key in style_keys):
        runs[-1]["text"] = str(runs[-1].get("text") or "") + value
        return
    payload = dict(run_format)
    payload["text"] = value
    runs.append(payload)


def _parse_quill_font(classes: set[str], styles: dict[str, str]) -> str | None:
    for class_name in classes:
        if class_name.startswith("ql-font-"):
            return _FONT_NAME_MAP.get(class_name.removeprefix("ql-font-"), class_name.removeprefix("ql-font-"))
    font_family = styles.get("font-family", "").split(",")[0].strip().strip("'\"")
    if not font_family:
        return None
    normalized = font_family.lower().replace(" ", "-")
    return _FONT_NAME_MAP.get(normalized, font_family)


def _collect_rich_runs(node, inherited: dict[str, object], runs: list[dict[str, object]]) -> None:
    tag = getattr(node, "tag", "")
    if not isinstance(tag, str):
        return
    classes = set(filter(None, str(node.get("class") or "").split()))
    styles = _parse_style_attr(node.get("style") or "")
    current = dict(inherited)
    if tag in {"strong", "b"}:
        current["bold"] = True
    if tag in {"em", "i"}:
        current["italic"] = True
    if tag == "u":
        current["underline"] = True
    if tag in {"s", "strike"}:
        current["strike"] = True
    font_name = _parse_quill_font(classes, styles)
    if font_name:
        current["font"] = font_name
    color = _normalize_css_color(styles.get("color"))
    if color:
        current["color"] = color
    background = _normalize_css_color(styles.get("background-color"))
    if background and background != "FFFFFF":
        current["background"] = background
    if node.text:
        _append_rich_run(runs, node.text, current)
    for child in node:
        child_tag = getattr(child, "tag", "")
        if child_tag == "br":
            _append_rich_run(runs, "\n", current)
        else:
            _collect_rich_runs(child, current, runs)
        if child.tail:
            _append_rich_run(runs, child.tail, current)


def _paragraph_alignment(node) -> str | None:
    styles = _parse_style_attr(node.get("style") or "")
    classes = set(filter(None, str(node.get("class") or "").split()))
    align = styles.get("text-align", "").strip().lower()
    if align in {"left", "center", "right", "justify"}:
        return align
    for class_name in classes:
        if class_name == "ql-align-center":
            return "center"
        if class_name == "ql-align-right":
            return "right"
        if class_name == "ql-align-justify":
            return "justify"
    return None


def _list_level(node) -> int:
    classes = set(filter(None, str(node.get("class") or "").split()))
    for class_name in classes:
        if class_name.startswith("ql-indent-"):
            suffix = class_name.removeprefix("ql-indent-")
            if suffix.isdigit():
                return int(suffix)
    return 0


def _build_rich_paragraph(node, *, list_type: str | None = None, list_level: int = 0) -> dict[str, object]:
    runs: list[dict[str, object]] = []
    _collect_rich_runs(node, {}, runs)
    return {
        "runs": runs,
        "alignment": _paragraph_alignment(node),
        "list_type": list_type or "",
        "list_level": list_level,
    }


def _rich_paragraphs_from_html(html: str) -> list[dict[str, object]]:
    from lxml import html as lxml_html

    wrapper = lxml_html.fragment_fromstring(str(html or ""), create_parent="div")
    paragraphs: list[dict[str, object]] = []
    if wrapper.text and wrapper.text.strip():
        paragraphs.append({"runs": [{"text": wrapper.text}], "alignment": None, "list_type": "", "list_level": 0})
    for child in wrapper:
        tag = getattr(child, "tag", "")
        if tag in {"p", "div"}:
            paragraphs.append(_build_rich_paragraph(child))
        elif tag in {"ul", "ol"}:
            for li in child.findall("./li"):
                item_list_type = str(li.get("data-list") or "").strip().lower()
                if item_list_type not in {"ordered", "bullet"}:
                    item_list_type = "ordered" if tag == "ol" else "bullet"
                paragraphs.append(
                    _build_rich_paragraph(
                        li,
                        list_type=item_list_type,
                        list_level=_list_level(li),
                    )
                )
        else:
            paragraphs.append(_build_rich_paragraph(child))
        if child.tail and child.tail.strip():
            paragraphs.append({"runs": [{"text": child.tail}], "alignment": None, "list_type": "", "list_level": 0})
    return paragraphs


def _rich_paragraphs_from_item(item) -> list[dict[str, object]] | None:
    if isinstance(item, dict):
        html = str(item.get("html") or "").strip()
        if html:
            return _rich_paragraphs_from_html(html)
        runs = item.get("runs")
        if isinstance(runs, list):
            return [item]
        return None
    return None


def _set_on_off_property(r_pr, tag_name: str, enabled: bool | None) -> None:
    if enabled is None:
        return
    existing = r_pr.find(qn(f"w:{tag_name}"))
    if enabled:
        if existing is None:
            from docx.oxml import OxmlElement

            existing = OxmlElement(f"w:{tag_name}")
            r_pr.append(existing)
        existing.set(qn("w:val"), "single" if tag_name == "u" else "1")
    elif existing is not None:
        r_pr.remove(existing)


def _set_color_property(r_pr, color_value: str | None) -> None:
    existing = r_pr.find(qn("w:color"))
    if color_value:
        if existing is None:
            from docx.oxml import OxmlElement

            existing = OxmlElement("w:color")
            r_pr.append(existing)
        existing.set(qn("w:val"), color_value)
    elif existing is not None:
        r_pr.remove(existing)


def _set_shading_property(r_pr, color_value: str | None) -> None:
    existing = r_pr.find(qn("w:shd"))
    if color_value:
        if existing is None:
            from docx.oxml import OxmlElement

            existing = OxmlElement("w:shd")
            r_pr.append(existing)
        existing.set(qn("w:val"), "clear")
        existing.set(qn("w:color"), "auto")
        existing.set(qn("w:fill"), color_value)
    elif existing is not None:
        r_pr.remove(existing)


def _set_font_property(r_pr, font_name: str | None) -> None:
    existing = r_pr.find(qn("w:rFonts"))
    if font_name:
        if existing is None:
            from docx.oxml import OxmlElement

            existing = OxmlElement("w:rFonts")
            r_pr.append(existing)
        for attr_name in ("ascii", "hAnsi", "cs"):
            existing.set(qn(f"w:{attr_name}"), font_name)
    elif existing is not None:
        r_pr.remove(existing)


def _set_language_property(r_pr, language_code: str | None) -> None:
    if not language_code:
        return
    existing = r_pr.find(qn("w:lang"))
    if existing is None:
        from docx.oxml import OxmlElement

        existing = OxmlElement("w:lang")
        r_pr.append(existing)
    for attr_name in ("val", "bidi", "eastAsia"):
        existing.set(qn(f"w:{attr_name}"), language_code)


def _append_text_segments(run_element, text: str) -> None:
    from docx.oxml import OxmlElement

    parts = str(text or "").split("\n")
    for index, part in enumerate(parts):
        if index:
            run_element.append(OxmlElement("w:br"))
        text_element = OxmlElement("w:t")
        text_element.set(qn("xml:space"), "preserve")
        text_element.text = part
        run_element.append(text_element)


def _append_formatted_run(
    paragraph_element,
    run_data: dict[str, object],
    source_rPr,
    *,
    language_code: str | None = None,
) -> None:
    from docx.oxml import OxmlElement

    run_element = OxmlElement("w:r")
    r_pr = deepcopy(source_rPr) if source_rPr is not None else OxmlElement("w:rPr")
    _set_on_off_property(r_pr, "b", run_data.get("bold"))
    _set_on_off_property(r_pr, "i", run_data.get("italic"))
    _set_on_off_property(r_pr, "u", run_data.get("underline"))
    _set_on_off_property(r_pr, "strike", run_data.get("strike"))
    _set_color_property(r_pr, str(run_data.get("color") or "").strip() or None)
    _set_shading_property(r_pr, str(run_data.get("background") or "").strip() or None)
    _set_font_property(r_pr, str(run_data.get("font") or "").strip() or None)
    _set_language_property(r_pr, language_code)
    if len(r_pr):
        run_element.append(r_pr)
    _append_text_segments(run_element, str(run_data.get("text") or ""))
    paragraph_element.append(run_element)


def _apply_paragraph_alignment(p_pr, alignment: str | None) -> None:
    if alignment not in {"left", "center", "right", "justify"}:
        return
    from docx.oxml import OxmlElement

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both" if alignment == "justify" else alignment)
    p_pr.append(jc)


def _insert_rich_paragraphs(
    anchor,
    parent,
    rich_items,
    *,
    bullet_num_id: str,
    multilevel_num_id: str,
    bullet_style_id: str,
    source_rPr,
    render_plain: bool = False,
    language_code: str | None = None,
):
    from docx.oxml import OxmlElement

    current_anchor = anchor
    for item in rich_items:
        new_p = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        list_type = "" if render_plain else str(item.get("list_type") or "").strip()
        list_level = 0 if render_plain else int(item.get("list_level") or 0)
        if list_type == "bullet":
            if bullet_style_id:
                p_style = OxmlElement("w:pStyle")
                p_style.set(qn("w:val"), bullet_style_id)
                p_pr.append(p_style)
            elif bullet_num_id:
                num_pr = OxmlElement("w:numPr")
                ilvl = OxmlElement("w:ilvl")
                ilvl.set(qn("w:val"), str(max(0, list_level)))
                num_id = OxmlElement("w:numId")
                num_id.set(qn("w:val"), bullet_num_id)
                num_pr.append(ilvl)
                num_pr.append(num_id)
                p_pr.append(num_pr)
        elif list_type == "ordered" and multilevel_num_id:
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), str(max(0, min(list_level, 2))))
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), multilevel_num_id)
            num_pr.append(ilvl)
            num_pr.append(num_id)
            p_pr.append(num_pr)
        _apply_paragraph_alignment(p_pr, item.get("alignment"))
        if len(p_pr):
            new_p.append(p_pr)
        runs = item.get("runs") if isinstance(item.get("runs"), list) else []
        if runs:
            for run_data in runs:
                if not isinstance(run_data, dict):
                    _append_formatted_run(new_p, {"text": str(run_data)}, source_rPr, language_code=language_code)
                    continue
                _append_formatted_run(new_p, run_data, source_rPr, language_code=language_code)
        else:
            _append_formatted_run(new_p, {"text": ""}, source_rPr, language_code=language_code)
        current_anchor.addnext(new_p)
        current_anchor = new_p
    parent.remove(anchor)


def _replace_in_paragraph(
    paragraph,
    replacements: dict[str, str],
    *,
    language_code: str | None = None,
) -> None:
    """Replace placeholders in a single paragraph, preserving run formatting.

    The algorithm:
    1. Build a mapping from character offset → run index.
    2. Find each placeholder span in the joined text.
    3. For each match, identify which runs it spans.
    4. Put the replacement value into the run where ``{{`` starts
       (preserving that run's formatting), then trim/clear the other
       runs that were consumed by the placeholder.
    """
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(r.text for r in runs)
    if not _PLACEHOLDER_RE.search(full_text):
        return

    run_texts = [r.text for r in runs]
    run_starts: list[int] = []
    offset = 0
    for t in run_texts:
        run_starts.append(offset)
        offset += len(t)

    def _run_at(char_pos: int) -> int:
        for i in range(len(run_starts) - 1, -1, -1):
            if char_pos >= run_starts[i]:
                return i
        return 0

    matches = list(_PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    for m in reversed(matches):
        key = m.group()
        value = replacements.get(key)
        if value is None:
            continue

        start, end = m.start(), m.end()
        first_run_idx = _run_at(start)
        last_run_idx = _run_at(end - 1)

        if first_run_idx == last_run_idx:
            local_start = start - run_starts[first_run_idx]
            local_end = end - run_starts[first_run_idx]
            t = run_texts[first_run_idx]
            run_texts[first_run_idx] = t[:local_start] + value + t[local_end:]
        else:
            ft = run_texts[first_run_idx]
            local_start = start - run_starts[first_run_idx]
            run_texts[first_run_idx] = ft[:local_start] + value

            lt = run_texts[last_run_idx]
            local_end = end - run_starts[last_run_idx]
            run_texts[last_run_idx] = lt[local_end:]

            for mid in range(first_run_idx + 1, last_run_idx):
                run_texts[mid] = ""

        target_run = runs[first_run_idx]
        r_pr = target_run._element.find(qn("w:rPr"))
        if r_pr is None and language_code:
            from docx.oxml import OxmlElement

            r_pr = OxmlElement("w:rPr")
            target_run._element.insert(0, r_pr)
        if r_pr is not None:
            _set_language_property(r_pr, language_code)

    for i, run in enumerate(runs):
        run.text = run_texts[i]


def _replace_list_in_paragraph(
    paragraph,
    list_replacements: dict,
    bullet_num_id: str,
    multilevel_num_id: str,
    bullet_style_id: str,
    plain_list_keys: set[str] | None = None,
    default_language_code: str | None = None,
) -> bool:
    """If the paragraph contains a ``[[list_var]]`` placeholder, replace the
    entire paragraph with a list of items.  Returns True if replaced.

    Items may be:
    - ``list[str]`` — flat bullet list or plain paragraphs
    - ``list[tuple[int, str]]`` — multi-level numbered list

    Formatting (font, size, bold, italic) is copied from the run containing
    the opening ``[[`` of the placeholder.
    """
    if not list_replacements:
        return False
    runs = paragraph.runs
    if not runs:
        return False
    full_text = "".join(r.text for r in runs)
    m = _LIST_PLACEHOLDER_RE.search(full_text)
    if not m:
        return False
    key = m.group()
    items = list_replacements.get(key)
    if items is None:
        return False
    language_code = default_language_code

    rich_items = []
    is_rich = False
    for item in items:
        rich_paragraphs = _rich_paragraphs_from_item(item)
        if rich_paragraphs is None:
            if isinstance(item, dict):
                continue
            rich_items.append(
                {
                    "runs": [{"text": str(item)}],
                    "alignment": None,
                    "list_type": "",
                    "list_level": 0,
                }
            )
        else:
            is_rich = True
            rich_items.extend(rich_paragraphs)

    is_multilevel = not is_rich and items and isinstance(items[0], (tuple, list))
    is_plain_list = (is_rich or not is_multilevel) and key in (plain_list_keys or set())

    if is_multilevel and not multilevel_num_id:
        return False
    if not is_rich and not is_multilevel and not is_plain_list and not bullet_num_id:
        return False

    run_texts = [r.text for r in runs]
    run_starts: list[int] = []
    offset = 0
    for t in run_texts:
        run_starts.append(offset)
        offset += len(t)

    def _run_at(char_pos: int) -> int:
        for i in range(len(run_starts) - 1, -1, -1):
            if char_pos >= run_starts[i]:
                return i
        return 0

    source_run = runs[_run_at(m.start())]
    source_rPr = source_run._element.find(qn("w:rPr"))

    parent = paragraph._element.getparent()
    anchor = paragraph._element

    if is_rich:
        _insert_rich_paragraphs(
            anchor,
            parent,
            rich_items,
            bullet_num_id=bullet_num_id,
            multilevel_num_id=multilevel_num_id,
            bullet_style_id=bullet_style_id,
            source_rPr=source_rPr,
            render_plain=is_plain_list,
            language_code=language_code,
        )
        return True

    from docx.oxml import OxmlElement

    for item in items:
        if is_multilevel:
            level, item_text = item
            num_id_val = multilevel_num_id
        else:
            level = 0
            item_text = item
            num_id_val = bullet_num_id

        new_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")

        if is_plain_list:
            pass
        elif not is_multilevel and bullet_style_id:
            pStyle = OxmlElement("w:pStyle")
            pStyle.set(qn("w:val"), bullet_style_id)
            pPr.append(pStyle)
        else:
            numPr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), str(level))
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), num_id_val)
            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)
        new_p.append(pPr)

        new_r = OxmlElement("w:r")
        if source_rPr is not None:
            r_pr = deepcopy(source_rPr)
        else:
            r_pr = OxmlElement("w:rPr")
        _set_language_property(r_pr, language_code)
        if len(r_pr):
            new_r.append(r_pr)
        new_t = OxmlElement("w:t")
        new_t.set(qn("xml:space"), "preserve")
        new_t.text = str(item_text)
        new_r.append(new_t)
        new_p.append(new_r)

        anchor.addnext(new_p)
        anchor = new_p

    parent.remove(paragraph._element)
    return True


def _normalize_table_cell_spec(value) -> dict[str, object]:
    if isinstance(value, dict):
        return {
            "text": str(value.get("text") or ""),
            "colspan": max(1, int(value.get("colspan") or 1)),
            "rowspan": max(1, int(value.get("rowspan") or 1)),
            "bold": bool(value.get("bold")),
            "align": str(value.get("align") or "").strip().lower() or "left",
            "header": bool(value.get("header")),
            "vertical_align": str(value.get("vertical_align") or "").strip().lower() or "top",
            "margins_cm": value.get("margins_cm") or {},
            "no_wrap": bool(value.get("no_wrap")),
        }
    return {
        "text": str(value or ""),
        "colspan": 1,
        "rowspan": 1,
        "bold": False,
        "align": "left",
        "header": False,
        "vertical_align": "top",
        "margins_cm": {},
        "no_wrap": False,
    }


def _table_alignment(value: str) -> WD_ALIGN_PARAGRAPH:
    mapping = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(str(value or "").strip().lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _measure_table_columns(rows: list[list[dict[str, object]]]) -> int:
    occupancy: list[list[bool]] = []
    max_cols = 0
    for row_idx, row in enumerate(rows):
        while len(occupancy) <= row_idx:
            occupancy.append([])
        col_idx = 0
        for raw_cell in row:
            cell = _normalize_table_cell_spec(raw_cell)
            while col_idx < len(occupancy[row_idx]) and occupancy[row_idx][col_idx]:
                col_idx += 1
            colspan = int(cell["colspan"])
            rowspan = int(cell["rowspan"])
            for rr in range(row_idx, row_idx + rowspan):
                while len(occupancy) <= rr:
                    occupancy.append([])
                while len(occupancy[rr]) < col_idx + colspan:
                    occupancy[rr].append(False)
                for cc in range(col_idx, col_idx + colspan):
                    occupancy[rr][cc] = True
            max_cols = max(max_cols, col_idx + colspan)
            col_idx += colspan
    return max_cols


def _apply_cell_text(cell, spec: dict[str, object], font_size_pt: int | float | None, language_code: str | None) -> None:
    from docx.oxml import OxmlElement

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = _table_alignment(str(spec.get("align") or "left"))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph._element.insert(0, p_pr)
    paragraph_r_pr = p_pr.find(qn("w:rPr"))
    if paragraph_r_pr is None:
        paragraph_r_pr = OxmlElement("w:rPr")
        p_pr.append(paragraph_r_pr)
    if font_size_pt:
        half_points = str(int(round(float(font_size_pt) * 2)))
        for tag_name in ("w:sz", "w:szCs"):
            existing = paragraph_r_pr.find(qn(tag_name))
            if existing is None:
                existing = OxmlElement(tag_name)
                paragraph_r_pr.append(existing)
            existing.set(qn("w:val"), half_points)
    if language_code:
        _set_language_property(paragraph_r_pr, language_code)
    text_parts = str(spec.get("text") or "").split("\n")
    runs = list(paragraph.runs)
    if runs:
        run = runs[0]
        run.text = text_parts[0] if text_parts else ""
        for extra_run in runs[1:]:
            extra_run.text = ""
    else:
        run = paragraph.add_run(text_parts[0] if text_parts else "")
    for text_part in text_parts[1:]:
        run.add_break()
        run.add_text(text_part)
    for paragraph_run in paragraph.runs:
        paragraph_run.bold = bool(spec.get("bold"))
        if font_size_pt:
            paragraph_run.font.size = Pt(font_size_pt)
        if language_code:
            r_pr = paragraph_run._element.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                paragraph_run._element.insert(0, r_pr)
            _set_language_property(r_pr, language_code)
    vertical_align = str(spec.get("vertical_align") or "top").strip().lower()
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if vertical_align == "center"
        else WD_CELL_VERTICAL_ALIGNMENT.TOP
    )


def _set_cell_margins(cell, *, top_cm=0, right_cm=0.1, bottom_cm=0, left_cm=0.1) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value_cm in (
        ("top", top_cm),
        ("right", right_cm),
        ("bottom", bottom_cm),
        ("left", left_cm),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(int(round(Cm(value_cm).twips))))
        node.set(qn("w:type"), "dxa")


def _set_cell_no_wrap(cell, enabled: bool) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:noWrap"))
    if enabled:
        if existing is None:
            existing = OxmlElement("w:noWrap")
            tc_pr.append(existing)
        existing.set(qn("w:val"), "1")
    elif existing is not None:
        tc_pr.remove(existing)


def _set_table_autofit(table) -> None:
    try:
        table.autofit = True
    except Exception:
        pass
    try:
        table.allow_autofit = True
    except Exception:
        pass

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        from docx.oxml import OxmlElement

        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        from docx.oxml import OxmlElement

        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "auto")
    tbl_w.set(qn("w:w"), "0")


def _set_table_fixed_pct_widths(table, column_widths_pct: list[float]) -> None:
    from docx.oxml import OxmlElement

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return

    try:
        table.autofit = False
    except Exception:
        pass
    try:
        table.allow_autofit = False
    except Exception:
        pass

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")


def _set_cell_width_pct(cell, width_pct: float) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "pct")
    tc_w.set(qn("w:w"), str(int(round(float(width_pct) * 50))))


def _insert_table_after_paragraph(paragraph, table_spec: dict, language_code: str | None = None) -> bool:
    rows = table_spec.get("rows") if isinstance(table_spec, dict) else None
    if not isinstance(rows, list) or not rows:
        return False
    column_count = _measure_table_columns(rows)
    if column_count <= 0:
        return False

    container = getattr(paragraph, "_parent", None)
    add_table = getattr(container, "add_table", None)
    if add_table is None:
        return False

    try:
        table = add_table(rows=len(rows), cols=column_count)
    except TypeError:
        block_width = getattr(getattr(paragraph.part, "document", None), "_block_width", None)
        if block_width is None:
            return False
        table = add_table(len(rows), column_count, block_width)
    try:
        if table_spec.get("style"):
            table.style = str(table_spec.get("style"))
    except Exception:
        pass
    column_widths_pct = list(table_spec.get("column_widths_pct") or []) if isinstance(table_spec, dict) else []
    if column_widths_pct and len(column_widths_pct) == column_count:
        _set_table_fixed_pct_widths(table, column_widths_pct)
    else:
        _set_table_autofit(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    font_size_pt = table_spec.get("font_size_pt") if isinstance(table_spec, dict) else None
    occupied = [[False] * column_count for _ in range(len(rows))]
    for row_idx, row in enumerate(rows):
        col_idx = 0
        for raw_cell in row:
            spec = _normalize_table_cell_spec(raw_cell)
            while col_idx < column_count and occupied[row_idx][col_idx]:
                col_idx += 1
            colspan = min(int(spec["colspan"]), column_count - col_idx)
            rowspan = min(int(spec["rowspan"]), len(rows) - row_idx)
            start_cell = table.cell(row_idx, col_idx)
            for rr in range(row_idx, row_idx + rowspan):
                for cc in range(col_idx, col_idx + colspan):
                    occupied[rr][cc] = True
            if rowspan > 1 or colspan > 1:
                start_cell = start_cell.merge(table.cell(row_idx + rowspan - 1, col_idx + colspan - 1))
            margins = spec.get("margins_cm") if isinstance(spec, dict) else {}
            _set_cell_margins(
                start_cell,
                top_cm=float((margins or {}).get("top", 0)),
                right_cm=float((margins or {}).get("right", 0.1)),
                bottom_cm=float((margins or {}).get("bottom", 0)),
                left_cm=float((margins or {}).get("left", 0.1)),
            )
            if column_widths_pct and len(column_widths_pct) == column_count:
                _set_cell_width_pct(start_cell, sum(column_widths_pct[col_idx:col_idx + colspan]))
            _set_cell_no_wrap(start_cell, bool(spec.get("no_wrap")))
            _apply_cell_text(start_cell, spec, font_size_pt, language_code)
            col_idx += colspan

    anchor = paragraph._element
    anchor.addnext(table._tbl)
    anchor.getparent().remove(anchor)
    return True


def _replace_table_in_paragraph(
    paragraph,
    table_replacements: dict | None,
    *,
    default_language_code: str | None = None,
) -> bool:
    if not table_replacements:
        return False
    full_text = "".join(run.text for run in paragraph.runs)
    key = full_text.strip()
    if key not in table_replacements:
        return False
    return _insert_table_after_paragraph(
        paragraph,
        table_replacements.get(key) or {},
        language_code=default_language_code,
    )


def _process_table_placeholders(
    paragraphs,
    table_replacements: dict | None = None,
    *,
    default_language_code: str | None = None,
) -> None:
    for para in list(paragraphs):
        _replace_table_in_paragraph(
            para,
            table_replacements,
            default_language_code=default_language_code,
        )


# ---------------------------------------------------------------------------
#  Numbering helpers — always allocate NEW numIds / abstractNumIds so that
#  existing template numbering definitions are never altered.
# ---------------------------------------------------------------------------

def _unique_nsid() -> str:
    """Return an 8-char uppercase hex suitable for ``w:nsid``."""
    return uuid.uuid4().hex[:8].upper()


def _max_num_id(numbering_elm) -> int:
    ids = [
        int(n.get(qn("w:numId")))
        for n in numbering_elm.findall(qn("w:num"))
        if n.get(qn("w:numId"), "").isdigit()
    ]
    return max(ids, default=0)


def _max_abstract_num_id(numbering_elm) -> int:
    ids = [
        int(a.get(qn("w:abstractNumId")))
        for a in numbering_elm.findall(qn("w:abstractNum"))
        if a.get(qn("w:abstractNumId"), "").isdigit()
    ]
    return max(ids, default=0)


def _insert_abstract_num(numbering_elm, abstract_num):
    """Insert ``w:abstractNum`` before the first ``w:num`` for valid OOXML."""
    first_num = numbering_elm.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_elm.append(abstract_num)


def _ensure_bullet_numbering(doc) -> str:
    """Create a new bullet numbering definition and return its ``numId``."""
    from docx.oxml import OxmlElement

    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    new_abs_id = str(_max_abstract_num_id(numbering_elm) + 1)
    new_num_id = str(_max_num_id(numbering_elm) + 1)

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), new_abs_id)

    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), _unique_nsid())
    abstract_num.append(nsid)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), _unique_nsid())
    abstract_num.append(tmpl)
    ml_type = OxmlElement("w:multiLevelType")
    ml_type.set(qn("w:val"), "singleLevel")
    abstract_num.append(ml_type)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)
    lvl.append(pPr)
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Symbol")
    rFonts.set(qn("w:hAnsi"), "Symbol")
    rFonts.set(qn("w:hint"), "default")
    rPr.append(rFonts)
    lvl.append(rPr)
    abstract_num.append(lvl)

    _insert_abstract_num(numbering_elm, abstract_num)

    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), new_num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), new_abs_id)
    num_el.append(abstract_ref)
    numbering_elm.append(num_el)

    return new_num_id


def _ensure_multilevel_numbering(doc) -> str:
    """Create a new 3-level decimal numbering definition and return its ``numId``.

    Levels: ``1.``, ``1.1``, ``1.1.1`` with increasing indentation.
    """
    from docx.oxml import OxmlElement

    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    new_abs_id = str(_max_abstract_num_id(numbering_elm) + 1)
    new_num_id = str(_max_num_id(numbering_elm) + 1)

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), new_abs_id)

    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), _unique_nsid())
    abstract_num.append(nsid)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), _unique_nsid())
    abstract_num.append(tmpl)
    ml_type = OxmlElement("w:multiLevelType")
    ml_type.set(qn("w:val"), "multilevel")
    abstract_num.append(ml_type)

    level_defs = [
        {"ilvl": "0", "fmt": "decimal", "text": "%1.", "indent": "360", "hanging": "360"},
        {"ilvl": "1", "fmt": "decimal", "text": "%1.%2", "indent": "720", "hanging": "360"},
        {"ilvl": "2", "fmt": "decimal", "text": "%1.%2.%3", "indent": "1080", "hanging": "360"},
    ]
    for ld in level_defs:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), ld["ilvl"])
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), ld["fmt"])
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), ld["text"])
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), ld["indent"])
        ind.set(qn("w:hanging"), ld["hanging"])
        pPr.append(ind)
        lvl.append(pPr)
        abstract_num.append(lvl)

    _insert_abstract_num(numbering_elm, abstract_num)

    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), new_num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), new_abs_id)
    num_el.append(abstract_ref)
    numbering_elm.append(num_el)

    return new_num_id


# ---------------------------------------------------------------------------
#  Style helpers
# ---------------------------------------------------------------------------

def _find_bullet_style_id(doc) -> str:
    """Find the built-in 'List Bullet' style ID in the document.

    Searches by canonical name (works for both English and Russian Word).
    If absent, creates the style so paragraphs display as
    "Маркированный список" / "List Bullet".
    """
    for style in doc.styles:
        if style.name == "List Bullet":
            return style.style_id
    try:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        return style.style_id
    except Exception:
        return ""


# ---------------------------------------------------------------------------

def _process_paragraphs(
    paragraphs,
    replacements: dict[str, str],
    *,
    language_code: str | None = None,
) -> None:
    for para in paragraphs:
        _replace_in_paragraph(para, replacements, language_code=language_code)


def _process_list_paragraphs(
    paragraphs,
    list_replacements: dict,
    bullet_num_id: str,
    multilevel_num_id: str,
    bullet_style_id: str,
    plain_list_keys: set[str] | None = None,
    default_language_code: str | None = None,
) -> None:
    for para in list(paragraphs):
        _replace_list_in_paragraph(
            para, list_replacements,
            bullet_num_id, multilevel_num_id, bullet_style_id,
            plain_list_keys, default_language_code,
        )


def _process_tables(
    tables,
    replacements: dict[str, str],
    table_replacements: dict | None = None,
    list_replacements: dict | None = None,
    bullet_num_id: str = "",
    multilevel_num_id: str = "",
    bullet_style_id: str = "",
    plain_list_keys: set[str] | None = None,
    default_language_code: str | None = None,
) -> None:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                if table_replacements:
                    _process_table_placeholders(
                        cell.paragraphs,
                        table_replacements,
                        default_language_code=default_language_code,
                    )
                if list_replacements:
                    _process_list_paragraphs(
                        cell.paragraphs, list_replacements,
                        bullet_num_id, multilevel_num_id, bullet_style_id,
                        plain_list_keys, default_language_code,
                    )
                _process_paragraphs(cell.paragraphs, replacements, language_code=default_language_code)
                _process_tables(
                    cell.tables, replacements, table_replacements, list_replacements,
                    bullet_num_id, multilevel_num_id, bullet_style_id,
                    plain_list_keys, default_language_code,
                )


def _iter_table_paragraphs(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                yield from _iter_table_paragraphs(cell.tables)


def _iter_document_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    yield from _iter_table_paragraphs(doc.tables)
    for section in doc.sections:
        for header_footer in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if header_footer and header_footer.is_linked_to_previous:
                continue
            if header_footer:
                for paragraph in header_footer.paragraphs:
                    yield paragraph
                yield from _iter_table_paragraphs(header_footer.tables)


def _replace_literal_in_paragraph(paragraph, literal: str, replacement: str = "") -> int:
    runs = paragraph.runs
    if not runs:
        return 0

    full_text = "".join(run.text for run in runs)
    matches = list(re.finditer(re.escape(str(literal or "")), full_text))
    if not matches:
        return 0

    run_texts = [run.text for run in runs]
    run_starts: list[int] = []
    offset = 0
    for text in run_texts:
        run_starts.append(offset)
        offset += len(text)

    def _run_at(char_pos: int) -> int:
        for index in range(len(run_starts) - 1, -1, -1):
            if char_pos >= run_starts[index]:
                return index
        return 0

    for match in reversed(matches):
        start, end = match.start(), match.end()
        first_run_idx = _run_at(start)
        last_run_idx = _run_at(end - 1)

        if first_run_idx == last_run_idx:
            local_start = start - run_starts[first_run_idx]
            local_end = end - run_starts[first_run_idx]
            current = run_texts[first_run_idx]
            run_texts[first_run_idx] = current[:local_start] + replacement + current[local_end:]
            continue

        first_text = run_texts[first_run_idx]
        local_start = start - run_starts[first_run_idx]
        run_texts[first_run_idx] = first_text[:local_start] + replacement

        last_text = run_texts[last_run_idx]
        local_end = end - run_starts[last_run_idx]
        run_texts[last_run_idx] = last_text[local_end:]

        for middle_idx in range(first_run_idx + 1, last_run_idx):
            run_texts[middle_idx] = ""

    for index, run in enumerate(runs):
        run.text = run_texts[index]
    return len(matches)


def _build_anchor_from_inline(
    inline,
    *,
    x_offset_emu: int = 0,
    y_offset_emu: int = 0,
    x_relative_from: str = "column",
    x_align: str | None = None,
):
    from docx.oxml import OxmlElement

    anchor = OxmlElement("wp:anchor")
    anchor.set("behindDoc", "1")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", x_relative_from)
    if x_align:
        align_h = OxmlElement("wp:align")
        align_h.text = str(x_align)
        position_h.append(align_h)
    else:
        pos_offset_h = OxmlElement("wp:posOffset")
        pos_offset_h.text = str(int(x_offset_emu))
        position_h.append(pos_offset_h)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    pos_offset_v = OxmlElement("wp:posOffset")
    pos_offset_v.text = str(int(y_offset_emu))
    position_v.append(pos_offset_v)
    anchor.append(position_v)

    extent = inline.find(qn("wp:extent"))
    if extent is not None:
        anchor.append(deepcopy(extent))
    effect_extent = inline.find(qn("wp:effectExtent"))
    if effect_extent is not None:
        anchor.append(deepcopy(effect_extent))
    else:
        effect_extent = OxmlElement("wp:effectExtent")
        effect_extent.set("l", "0")
        effect_extent.set("t", "0")
        effect_extent.set("r", "0")
        effect_extent.set("b", "0")
        anchor.append(effect_extent)

    anchor.append(OxmlElement("wp:wrapNone"))

    doc_pr = inline.find(qn("wp:docPr"))
    if doc_pr is not None:
        anchor.append(deepcopy(doc_pr))
    else:
        generated_doc_pr = OxmlElement("wp:docPr")
        generated_doc_pr.set("id", str(uuid.uuid4().int % 100000))
        generated_doc_pr.set("name", "FloatingImage")
        anchor.append(generated_doc_pr)

    c_nv_graphic_frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))
    if c_nv_graphic_frame_pr is not None:
        anchor.append(deepcopy(c_nv_graphic_frame_pr))
    else:
        anchor.append(OxmlElement("wp:cNvGraphicFramePr"))

    graphic = inline.find(qn("a:graphic"))
    if graphic is not None:
        anchor.append(deepcopy(graphic))
    return anchor


def _append_floating_image_run(
    paragraph,
    image_bytes: bytes,
    *,
    width_cm: float | None = None,
    x_offset_cm: float = 0,
    y_offset_cm: float = 0,
    x_relative_from: str = "page",
    x_align: str | None = "center",
) -> None:
    from docx.image.exceptions import UnexpectedEndOfFileError, UnrecognizedImageError

    run = paragraph.add_run()
    try:
        if width_cm is None:
            run.add_picture(BytesIO(image_bytes))
        else:
            run.add_picture(BytesIO(image_bytes), width=Cm(width_cm))
    except (UnrecognizedImageError, UnexpectedEndOfFileError) as exc:
        raise RuntimeError(
            "Факсимиле должно быть изображением в поддерживаемом формате Word (например PNG или JPEG)."
        ) from exc

    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline")) if drawing is not None else None
    if drawing is None or inline is None:
        raise RuntimeError("Не удалось вставить факсимиле в DOCX-документ.")

    anchor = _build_anchor_from_inline(
        inline,
        x_offset_emu=int(Cm(x_offset_cm).emu),
        y_offset_emu=int(Cm(y_offset_cm).emu),
        x_relative_from=x_relative_from,
        x_align=x_align,
    )
    drawing.remove(inline)
    drawing.append(anchor)


def insert_floating_image_at_placeholder(
    file_bytes: bytes,
    image_bytes: bytes,
    *,
    placeholder: str = "[[facsimile]]",
    width_cm: float | None = None,
    x_offset_cm: float = 0,
    y_offset_cm: float = 0,
    x_relative_from: str = "page",
    x_align: str | None = "center",
) -> bytes:
    if not file_bytes or not image_bytes or not str(placeholder or "").strip():
        return file_bytes

    doc = Document(BytesIO(file_bytes))
    inserted = False
    for paragraph in _iter_document_paragraphs(doc):
        occurrences = _replace_literal_in_paragraph(paragraph, placeholder, "")
        if not occurrences:
            continue
        for _ in range(occurrences):
            _append_floating_image_run(
                paragraph,
                image_bytes,
                width_cm=width_cm,
                x_offset_cm=x_offset_cm,
                y_offset_cm=y_offset_cm,
                x_relative_from=x_relative_from,
                x_align=x_align,
            )
        inserted = True

    if not inserted:
        return file_bytes

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def process_template(
    file_bytes: bytes,
    replacements: dict[str, str],
    table_replacements: dict | None = None,
    list_replacements: dict | None = None,
    plain_list_keys: set[str] | None = None,
    default_language_code: str | None = None,
) -> bytes:
    """Return modified .docx bytes with all placeholders substituted.

    *file_bytes*: raw bytes of the source .docx template.
    *replacements*: mapping ``{"{{key}}": "value", ...}``.
    *list_replacements*: mapping ``{"[[key]]": ["item1", ...], ...}``.
    *plain_list_keys*: list placeholders that should be inserted as plain
    paragraphs without bullet formatting.
    *default_language_code*: language tag to apply to inserted/replaced text
    (for example ``ru-RU``).
    """
    if not replacements and not list_replacements and not table_replacements:
        return file_bytes

    doc = Document(BytesIO(file_bytes))

    bullet_num_id = ""
    multilevel_num_id = ""
    bullet_style_id = ""

    if table_replacements:
        _process_table_placeholders(
            doc.paragraphs,
            table_replacements,
            default_language_code=default_language_code,
        )

    if list_replacements:
        try:
            bullet_num_id = _ensure_bullet_numbering(doc)
        except Exception:
            pass
        try:
            multilevel_num_id = _ensure_multilevel_numbering(doc)
        except Exception:
            pass
        bullet_style_id = _find_bullet_style_id(doc)
        _process_list_paragraphs(
            doc.paragraphs, list_replacements,
            bullet_num_id, multilevel_num_id, bullet_style_id,
            plain_list_keys, default_language_code,
        )

    _process_paragraphs(doc.paragraphs, replacements, language_code=default_language_code)
    _process_tables(
        doc.tables, replacements, table_replacements, list_replacements,
        bullet_num_id, multilevel_num_id, bullet_style_id, plain_list_keys, default_language_code,
    )

    for section in doc.sections:
        for header_footer in (section.header, section.footer,
                              section.first_page_header, section.first_page_footer,
                              section.even_page_header, section.even_page_footer):
            if header_footer and header_footer.is_linked_to_previous:
                continue
            if header_footer:
                if table_replacements:
                    _process_table_placeholders(
                        header_footer.paragraphs,
                        table_replacements,
                        default_language_code=default_language_code,
                    )
                if list_replacements:
                    _process_list_paragraphs(
                        header_footer.paragraphs, list_replacements,
                        bullet_num_id, multilevel_num_id, bullet_style_id,
                        plain_list_keys, default_language_code,
                    )
                _process_paragraphs(
                    header_footer.paragraphs,
                    replacements,
                    language_code=default_language_code,
                )
                _process_tables(
                    header_footer.tables, replacements, table_replacements, list_replacements,
                    bullet_num_id, multilevel_num_id, bullet_style_id,
                    plain_list_keys, default_language_code,
                )

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
