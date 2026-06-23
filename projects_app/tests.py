import base64
import copy
import json
import uuid
from decimal import Decimal
from importlib import import_module
from io import BytesIO
from datetime import date, timedelta

from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.contrib.auth import get_user_model
from django.contrib.auth.models import Group
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from checklists_app.models import ChecklistItem, ChecklistStatus, SourceDataItemFolder, SourceDataSectionFolder, SourceDataWorkspace
from classifiers_app.models import BusinessEntityIdentifierRecord, BusinessEntityRecord, LegalEntityRecord, OKSMCountry
from contacts_app.models import CitizenshipRecord, PersonRecord
from core.models import CloudStorageSettings
from contracts_app.models import ContractProjectRegistration, ContractTemplate
from nextcloud_app.models import NextcloudUserLink
from notifications_app.models import Notification, NotificationPerformerLink
from policy_app.models import (
    DIRECTOR_GROUP,
    DIRECTION_DIRECTOR_GROUP,
    EXPERT_GROUP,
    LAWYER_GROUP,
    PROJECTS_HEAD_GROUP,
    Product,
    TypicalServiceTerm,
    TypicalSectionSpecialty,
)
from experts_app.models import ExpertContractDetails, ExpertProfile, ExpertProfileSpecialty, ExpertSpecialty
from projects_app.models import (
    LegalEntity,
    PaymentRequestPerformer,
    Performer,
    ProjectRegistration,
    ProjectRegistrationProduct,
    RegistrationWorkspaceFolder,
    SourceDataTargetFolder,
    WorkVolume,
)
from projects_app.forms import ContractConditionsForm, LegalEntityForm, PerformerForm, ProjectRegistrationForm, WorkVolumeForm
from group_app.models import GroupMember, OrgUnit
from users_app.forms import FREELANCER_LABEL
from users_app.models import Employee
from unittest.mock import call, patch


TEST_CONTRACT_IMAGE_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMBAAZ/qh8AAAAASUVORK5CYII="
)


class ParticipationBatchViewTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="staff@example.com",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.recipient_user = get_user_model().objects.create_user(
            username="expert-batch@example.com",
            email="expert-batch@example.com",
            password="secret",
            is_staff=True,
        )
        self.employee = Employee.objects.create(user=self.recipient_user)
        self.project_a = ProjectRegistration.objects.create(number=8234, name="Продукт A", year=2026)
        self.project_b = ProjectRegistration.objects.create(number=8234, name="Продукт B", year=2026)
        self.other_number_project = ProjectRegistration.objects.create(number=8235, name="Другой номер", year=2026)
        self.project_a.refresh_from_db()
        self.project_b.refresh_from_db()
        self.performer_a = Performer.objects.create(
            registration=self.project_a,
            employee=self.employee,
            executor="Иванов Иван Иванович",
            asset_name="Актив A",
        )
        self.performer_b = Performer.objects.create(
            registration=self.project_b,
            employee=self.employee,
            executor="Иванов Иван Иванович",
            asset_name="Актив B",
        )
        self.performer_other_number = Performer.objects.create(
            registration=self.other_number_project,
            employee=self.employee,
            executor="Иванов Иван Иванович",
        )

    def test_participation_batch_merge_and_split(self):
        response = self.client.post(
            reverse("participation_batch_merge"),
            {"performer_ids": [str(self.performer_a.pk), str(self.performer_b.pk)]},
        )

        self.assertEqual(response.status_code, 200)
        self.performer_a.refresh_from_db()
        self.performer_b.refresh_from_db()
        self.assertIsNotNone(self.performer_a.participation_batch_id)
        self.assertEqual(self.performer_a.participation_batch_id, self.performer_b.participation_batch_id)

        response = self.client.post(
            reverse("participation_batch_split"),
            {"performer_ids": [str(self.performer_a.pk)]},
        )

        self.assertEqual(response.status_code, 200)
        self.performer_a.refresh_from_db()
        self.performer_b.refresh_from_db()
        self.assertIsNone(self.performer_a.participation_batch_id)
        self.assertIsNone(self.performer_b.participation_batch_id)

    def test_participation_batch_merge_rejects_different_numbers(self):
        response = self.client.post(
            reverse("participation_batch_merge"),
            {"performer_ids": [str(self.performer_a.pk), str(self.performer_other_number.pk)]},
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("одного четырехзначного номера", response.json()["error"])

    def test_participation_request_expands_saved_batch(self):
        batch_id = uuid.uuid4()
        Performer.objects.filter(pk__in=[self.performer_a.pk, self.performer_b.pk]).update(
            participation_batch_id=batch_id,
        )
        request_sent_at = timezone.localtime().replace(second=0, microsecond=0)

        with patch("projects_app.views.create_participation_notifications") as mocked_notifications:
            mocked_notifications.return_value = {
                "delivery_channels": ("system",),
                "email_delivery": {"requested": False, "attempted": 0, "sent": 0, "failed": 0, "errors": []},
            }
            response = self.client.post(
                reverse("participation_request"),
                {
                    "performer_ids": [str(self.performer_a.pk)],
                    "duration_hours": "4",
                    "request_sent_at": request_sent_at.isoformat(timespec="minutes"),
                    "delivery_channels": ["system"],
                },
            )

        self.assertEqual(response.status_code, 200)
        sent_ids = {performer.pk for performer in mocked_notifications.call_args.kwargs["performers"]}
        self.assertEqual(sent_ids, {self.performer_a.pk, self.performer_b.pk})
        self.performer_a.refresh_from_db()
        self.performer_b.refresh_from_db()
        self.assertIsNotNone(self.performer_a.participation_request_sent_at)
        self.assertEqual(self.performer_a.participation_request_sent_at, self.performer_b.participation_request_sent_at)

    def test_direction_recipient_can_rebatch_active_direction_notification_rows(self):
        sent_at = timezone.now()
        Performer.objects.filter(pk__in=[self.performer_a.pk, self.performer_b.pk]).update(
            participation_request_sent_at=sent_at,
            participation_deadline_at=sent_at + timedelta(hours=4),
        )
        notification = Notification.objects.create(
            notification_type=Notification.NotificationType.PROJECT_PARTICIPATION_CONFIRMATION,
            recipient=self.recipient_user,
            sender=self.user,
            project=self.project_a,
            title_text="Подтверждение по направлению",
            payload={"letter_template_type": "direction_confirmation"},
        )
        NotificationPerformerLink.objects.create(notification=notification, performer=self.performer_a, position=1)
        NotificationPerformerLink.objects.create(notification=notification, performer=self.performer_b, position=2)

        response = self.client.post(
            reverse("participation_batch_merge"),
            {"performer_ids": [str(self.performer_a.pk), str(self.performer_b.pk)]},
        )
        self.assertEqual(response.status_code, 400)

        self.client.force_login(self.recipient_user)
        response = self.client.post(
            reverse("participation_batch_merge"),
            {"performer_ids": [str(self.performer_a.pk), str(self.performer_b.pk)]},
        )

        self.assertEqual(response.status_code, 200)
        self.performer_a.refresh_from_db()
        self.performer_b.refresh_from_db()
        self.assertIsNotNone(self.performer_a.participation_batch_id)
        self.assertEqual(self.performer_a.participation_batch_id, self.performer_b.participation_batch_id)

        response = self.client.post(
            reverse("participation_batch_split"),
            {"performer_ids": [str(self.performer_a.pk)]},
        )

        self.assertEqual(response.status_code, 200)
        self.performer_a.refresh_from_db()
        self.performer_b.refresh_from_db()
        self.assertIsNone(self.performer_a.participation_batch_id)
        self.assertIsNone(self.performer_b.participation_batch_id)


class SourceDataTargetFolderViewTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="staff-user",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)

    def test_load_excludes_unresolved_template_paths(self):
        RegistrationWorkspaceFolder.objects.bulk_create(
            [
                RegistrationWorkspaceFolder(user=self.user, level=1, name="05 Исходные данные", position=0),
                RegistrationWorkspaceFolder(user=self.user, level=2, name="{project_label}", position=1),
                RegistrationWorkspaceFolder(user=self.user, level=2, name="01 Запросы", position=2),
            ]
        )

        response = self.client.get(reverse("source_data_target_folder_load"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["options"], ["05 Исходные данные", "05 Исходные данные/01 Запросы"])

    def test_save_rejects_unresolved_template_paths(self):
        response = self.client.post(
            reverse("source_data_target_folder_save"),
            data=json.dumps({"folder_name": "05 Исходные данные/{project_label}"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()["ok"])
        self.assertFalse(SourceDataTargetFolder.objects.filter(user=self.user).exists())

    def test_load_keeps_second_level_options_when_nextcloud_selected(self):
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.save()
        RegistrationWorkspaceFolder.objects.bulk_create(
            [
                RegistrationWorkspaceFolder(user=self.user, level=1, name="05 Исходные данные", position=0),
                RegistrationWorkspaceFolder(user=self.user, level=2, name="01 Запросы", position=1),
            ]
        )

        response = self.client.get(reverse("source_data_target_folder_load"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["options"], ["05 Исходные данные", "05 Исходные данные/01 Запросы"])


class ProjectRegistrationFormTests(TestCase):
    def setUp(self):
        self.group_member = GroupMember.objects.create(
            short_name="IMC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=0,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.second_product = Product.objects.create(
            short_name="QAQC_FORM",
            name_en="QAQC Form",
            name_ru="QAQC Form",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Контроль качества",
            position=2,
        )

    def test_type_ids_required_and_deadline_optional(self):
        form = ProjectRegistrationForm(
            data={
                "number": 4444,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "name": "Проект Альфа",
                "status": "Не начат",
                "deadline": "",
                "year": 2026,
            }
        )

        self.assertFalse(form.is_valid())
        self.assertIn("type_ids", form.errors)
        self.assertNotIn("deadline", form.errors)

    def test_number_accepts_zero_and_rejects_negative_values(self):
        valid_form = ProjectRegistrationForm(
            data={
                "number": 0,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [self.product.pk],
                "name": "Проект Ноль",
                "status": "Не начат",
                "deadline": "2026-01-10",
                "year": 2026,
            }
        )
        invalid_form = ProjectRegistrationForm(
            data={
                "number": -1,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [self.product.pk],
                "name": "Проект Минус",
                "status": "Не начат",
                "deadline": "2026-01-10",
                "year": 2026,
            }
        )

        self.assertTrue(valid_form.is_valid())
        self.assertFalse(invalid_form.is_valid())
        self.assertIn("number", invalid_form.errors)

    def test_number_uses_numeric_widget_and_cleans_to_int(self):
        form = ProjectRegistrationForm(
            data={
                "number": 10,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [self.product.pk],
                "name": "Проект Десять",
                "status": "Не начат",
                "deadline": "2026-01-10",
                "year": 2026,
            }
        )

        self.assertEqual(form.fields["number"].widget.input_type, "number")
        self.assertEqual(form.fields["number"].widget.attrs["id"], "registration-number-input")
        self.assertTrue(form.is_valid())
        self.assertEqual(form.cleaned_data["number"], 10)

    def test_existing_instance_keeps_numeric_number_value(self):
        registration = ProjectRegistration.objects.create(
            number=1,
            group_member=self.group_member,
            type=self.product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Проект Один",
            status="Не начат",
        )

        form = ProjectRegistrationForm(instance=registration)

        self.assertEqual(form["number"].value(), 1)

    def test_contract_schedule_labels_use_report_terms(self):
        registration_form = ProjectRegistrationForm()
        contract_form = ContractConditionsForm()

        for form in (registration_form, contract_form):
            self.assertEqual(
                form.fields["stage1_weeks"].label,
                "Срок подготовки Предварительного отчёта, мес.",
            )
            self.assertTrue(form.fields["stage1_weeks"].widget.attrs["readonly"])
            self.assertIn("readonly-field", form.fields["stage1_weeks"].widget.attrs["class"])
            self.assertEqual(
                form.fields["stage1_end"].label,
                "Дата Предварительного отчёта",
            )
            self.assertEqual(form.fields["stage1_date"].label, "Дата Предварительного отчёта")
            self.assertNotIn("readonly-field", form.fields["stage1_date"].widget.attrs["class"])
            self.assertEqual(
                form.fields["stage2_weeks"].label,
                "Срок подготовки Итогового отчёта, нед.",
            )
            self.assertTrue(form.fields["stage2_weeks"].widget.attrs["readonly"])
            self.assertIn("readonly-field", form.fields["stage2_weeks"].widget.attrs["class"])
            self.assertEqual(form.fields["stage2_end"].label, "Дата Итогового отчёта")
            self.assertNotIn("readonly-field", form.fields["stage2_end"].widget.attrs["class"])

    def test_contract_schedule_calculates_preliminary_report_in_months(self):
        registration = ProjectRegistration.objects.create(
            number=2,
            group_member=self.group_member,
            type=self.product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Проект со сроками",
            status="Не начат",
            contract_start=date(2026, 1, 31),
            input_data=10,
            stage1_weeks="1.0",
            stage2_weeks="2.0",
            stage3_weeks="4.0",
        )

        self.assertEqual(registration.stage1_end, date(2026, 2, 28))
        self.assertEqual(registration.stage2_end, date(2026, 3, 14))
        self.assertEqual(registration.completion_calc, date(2026, 3, 14))
        self.assertEqual(str(registration.term_weeks), "6.3")

    def test_stage1_migration_converts_legacy_weeks_to_months(self):
        migration = import_module("projects_app.migrations.0069_migrate_stage1_weeks_to_months")

        stage1_months = migration._weeks_to_contract_months(Decimal("6.0"))

        self.assertEqual(stage1_months, Decimal("1.4"))
        self.assertEqual(migration._term_weeks(stage1_months, Decimal("2.0")), Decimal("8.0"))

    def test_project_short_uid_uses_zero_padded_number(self):
        registration = ProjectRegistration.objects.create(
            number=1,
            group_member=self.group_member,
            type=self.product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Проект ID",
            status="Не начат",
        )

        self.assertEqual(registration.short_uid, "00010000RU")
        self.assertTrue(registration.display_identifier.startswith("0001 "))

    def test_project_short_uid_stage_digit_uses_product_row_sequence(self):
        first = ProjectRegistration.objects.create(
            number=55,
            group_member=self.group_member,
            type=self.product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Первый продукт",
            status="Не начат",
            position=1,
        )
        second = ProjectRegistration.objects.create(
            number=55,
            group_member=self.group_member,
            type=self.second_product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Второй продукт",
            status="Не начат",
            position=2,
        )

        first.refresh_from_db()
        second.refresh_from_db()

        self.assertEqual(first.short_uid, "00550010RU")
        self.assertEqual(second.short_uid, "00550020RU")
        self.assertEqual(first.agreement_sequence, 1)
        self.assertEqual(second.agreement_sequence, 2)

    def test_project_short_uid_stage_digit_restarts_for_each_contract_project(self):
        first_contract = ContractProjectRegistration.objects.create(
            number=7401,
            sub_number=1,
            group_member=self.group_member,
            name="Договор 01",
        )
        second_contract = ContractProjectRegistration.objects.create(
            number=7402,
            sub_number=2,
            group_member=self.group_member,
            name="Договор 02",
        )
        first = ProjectRegistration.objects.create(
            number=55,
            group_member=self.group_member,
            contract_project_registration=first_contract,
            type=self.product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Первый продукт",
            status="Не начат",
            position=1,
        )
        second = ProjectRegistration.objects.create(
            number=55,
            group_member=self.group_member,
            contract_project_registration=second_contract,
            type=self.second_product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Второй продукт",
            status="Не начат",
            position=2,
        )
        third = ProjectRegistration.objects.create(
            number=55,
            group_member=self.group_member,
            contract_project_registration=first_contract,
            type=self.second_product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Третий продукт",
            status="Не начат",
            position=3,
        )

        first.refresh_from_db()
        second.refresh_from_db()
        third.refresh_from_db()

        self.assertEqual(first.short_uid, "00550110RU")
        self.assertEqual(second.short_uid, "00550210RU")
        self.assertEqual(third.short_uid, "00550120RU")
        self.assertEqual(first.agreement_sequence, 1)
        self.assertEqual(second.agreement_sequence, 1)
        self.assertEqual(third.agreement_sequence, 2)

        third.contract_project_registration = second_contract
        third.save(update_fields=["contract_project_registration"])
        first.refresh_from_db()
        second.refresh_from_db()
        third.refresh_from_db()

        self.assertEqual(first.short_uid, "00550110RU")
        self.assertEqual(second.short_uid, "00550210RU")
        self.assertEqual(third.short_uid, "00550220RU")
        self.assertEqual(first.agreement_sequence, 1)
        self.assertEqual(second.agreement_sequence, 1)
        self.assertEqual(third.agreement_sequence, 2)

    def test_form_allows_duplicate_project_identity_values_and_keeps_project_id_unique(self):
        first = ProjectRegistration.objects.create(
            number=4444,
            group_member=self.group_member,
            type=self.product,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            name="Проект Дубль 1",
            status="Не начат",
        )

        form = ProjectRegistrationForm(
            data={
                "number": 4444,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [self.product.pk],
                "name": "Проект Дубль 2",
                "status": "Не начат",
                "deadline": "2026-01-10",
                "year": 2026,
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        duplicate = form.save()
        duplicate.refresh_from_db()

        self.assertEqual(first.number, duplicate.number)
        self.assertEqual(first.group_member_id, duplicate.group_member_id)
        self.assertEqual(first.agreement_type, duplicate.agreement_type)
        self.assertEqual(first.agreement_number, duplicate.agreement_number)
        self.assertNotEqual(first.pk, duplicate.pk)
        self.assertNotEqual(first.short_uid, duplicate.short_uid)

    def test_cleaned_type_ids_keep_ranked_order_with_duplicates(self):
        form = ProjectRegistrationForm()
        bound_form = ProjectRegistrationForm(
            data={
                "number": 11,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [self.second_product.pk, self.product.pk, self.second_product.pk],
                "name": "Проект Ранги",
                "status": "Не начат",
                "deadline": "2026-01-10",
                "year": 2026,
            }
        )

        self.assertNotIn("type", form.fields)
        self.assertTrue(bound_form.is_valid())
        self.assertEqual(
            bound_form.cleaned_type_ids,
            [self.second_product.pk, self.product.pk, self.second_product.pk],
        )

    def test_edit_form_rejects_multiple_products(self):
        bound_form = ProjectRegistrationForm(
            data={
                "number": 11,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [self.product.pk, self.second_product.pk],
                "name": "Проект Редактирование",
                "status": "Не начат",
                "deadline": "2026-01-10",
                "year": 2026,
            },
            allow_multiple_products=False,
        )

        self.assertFalse(bound_form.is_valid())
        self.assertIn("type_ids", bound_form.errors)

    def test_projects_date_fields_use_native_date_inputs(self):
        forms_and_fields = [
            (ProjectRegistrationForm(), ["deadline", "evaluation_date", "registration_date"]),
            (ContractConditionsForm(), ["contract_start", "contract_end"]),
            (WorkVolumeForm(), ["registration_date"]),
            (LegalEntityForm(), ["registration_date"]),
        ]

        for form, field_names in forms_and_fields:
            for field_name in field_names:
                widget = form.fields[field_name].widget
                with self.subTest(form=form.__class__.__name__, field=field_name):
                    self.assertEqual(widget.input_type, "date")
                    self.assertEqual(widget.format, "%Y-%m-%d")
                    self.assertNotIn("js-date", widget.attrs.get("class", ""))

    def test_project_manager_choices_include_direction_directors(self):
        project_manager_user = get_user_model().objects.create_user(
            username="project-manager-choice",
            first_name="Иван",
            last_name="Проектов",
            is_staff=True,
        )
        direction_director_user = get_user_model().objects.create_user(
            username="direction-director-choice",
            first_name="Дарья",
            last_name="Директорова",
            is_staff=True,
        )
        expert_user = get_user_model().objects.create_user(
            username="expert-choice",
            first_name="Егор",
            last_name="Экспертов",
            is_staff=True,
        )
        Employee.objects.create(
            user=project_manager_user,
            patronymic="Иванович",
            role=PROJECTS_HEAD_GROUP,
        )
        Employee.objects.create(
            user=direction_director_user,
            patronymic="Дмитриевна",
            role=DIRECTION_DIRECTOR_GROUP,
        )
        Employee.objects.create(
            user=expert_user,
            patronymic="Егорович",
            role=EXPERT_GROUP,
        )

        registration_choices = [label for _value, label in ProjectRegistrationForm().fields["project_manager"].choices]
        work_choices = [label for _value, label in WorkVolumeForm().fields["manager"].choices]

        for choices in (registration_choices, work_choices):
            self.assertTrue(any(label.startswith("Проектов Иван Иванович") for label in choices))
            self.assertTrue(any(label.startswith("Директорова Дарья Дмитриевна") for label in choices))
            self.assertFalse(any(label.startswith("Экспертов Егор Егорович") for label in choices))

    def test_project_manager_form_saves_selected_prs_id(self):
        person = PersonRecord.objects.create(
            last_name="Иванов",
            first_name="Иван",
            middle_name="Иванович",
            position=1,
        )
        user = get_user_model().objects.create_user(
            username="project-manager-prs",
            first_name="Иван",
            last_name="Иванов",
            is_staff=True,
        )
        Employee.objects.create(
            user=user,
            person_record=person,
            patronymic="Иванович",
            role=PROJECTS_HEAD_GROUP,
        )
        group_member = GroupMember.objects.create(
            short_name="IMCM",
            country_name="Россия",
            country_alpha2="RU",
        )
        product = Product.objects.create(short_name="PRS-DD", name_en="Due Diligence PRS", name_ru="ДД PRS")

        choices = dict(ProjectRegistrationForm().fields["project_manager"].choices)
        self.assertEqual(choices[person.formatted_id], "Иванов Иван Иванович")
        legacy_value = f"Иванов Иван Иванович ({person.formatted_id})"
        legacy_form = ProjectRegistrationForm(
            instance=ProjectRegistration(project_manager=legacy_value)
        )
        self.assertEqual(
            dict(legacy_form.fields["project_manager"].choices)[legacy_value],
            "Иванов Иван Иванович",
        )

        form = ProjectRegistrationForm(
            data={
                "number": "6201",
                "group_member": str(group_member.pk),
                "agreement_type": ProjectRegistration.AgreementType.MAIN,
                "name": "Проект с руководителем",
                "status": "Не начат",
                "deadline": "2026-05-01",
                "year": "2026",
                "country": "",
                "customer": "",
                "identifier": "",
                "registration_number": "",
                "registration_date": "",
                "project_manager": person.formatted_id,
                "type_id": str(product.pk),
            }
        )

        self.assertTrue(form.is_valid(), form.errors)
        obj = form.save(commit=False)
        self.assertEqual(obj.project_manager, "Иванов Иван Иванович")
        self.assertEqual(obj.project_manager_prs_id, person.formatted_id)

    def test_project_manager_form_selects_saved_prs_id_on_edit(self):
        person = PersonRecord.objects.create(
            last_name="Петров",
            first_name="Петр",
            middle_name="Петрович",
            position=2,
        )
        user = get_user_model().objects.create_user(
            username="project-manager-edit-prs",
            first_name="Петр",
            last_name="Петров",
            is_staff=True,
        )
        Employee.objects.create(
            user=user,
            person_record=person,
            patronymic="Петрович",
            role=PROJECTS_HEAD_GROUP,
        )
        registration = ProjectRegistration(
            project_manager="Петров Петр Петрович",
            project_manager_prs_id=person.formatted_id,
        )

        form = ProjectRegistrationForm(instance=registration)

        self.assertEqual(form["project_manager"].value(), person.formatted_id)


class PerformerFormExecutorFilteringTests(TestCase):
    def setUp(self):
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
        )
        self.project = ProjectRegistration.objects.create(
            number=6101,
            type=self.product,
            name="Проект исполнителей",
            year=2026,
        )
        self.section = self.product.sections.create(
            code="TAX",
            short_name="Tax",
            short_name_ru="Налоги",
            name_en="Tax",
            name_ru="Налоги",
            accounting_type="Раздел",
        )
        self.matching_specialty = ExpertSpecialty.objects.create(specialty="Налоги")
        self.other_specialty = ExpertSpecialty.objects.create(specialty="Геология")
        TypicalSectionSpecialty.objects.create(
            section=self.section,
            specialty=self.matching_specialty,
            rank=1,
        )
        self.matching_name = "Иванов Иван Иванович"
        self.other_name = "Петров Петр Петрович"
        self.matching_employee = self._create_employee(
            username="tax.executor",
            first_name="Иван",
            last_name="Иванов",
            patronymic="Иванович",
            specialty=self.matching_specialty,
        )
        self.other_employee = self._create_employee(
            username="geo.executor",
            first_name="Петр",
            last_name="Петров",
            patronymic="Петрович",
            specialty=self.other_specialty,
        )

    def _create_employee(self, *, username, first_name, last_name, patronymic, specialty):
        user = get_user_model().objects.create_user(
            username=username,
            password="secret",
            first_name=first_name,
            last_name=last_name,
            is_staff=True,
        )
        employee = Employee.objects.create(user=user, patronymic=patronymic)
        profile = ExpertProfile.objects.create(employee=employee)
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=specialty, rank=1)
        return employee

    def test_executor_choices_are_filtered_by_typical_section_specialties(self):
        form = PerformerForm(data={
            "registration": self.project.pk,
            "typical_section": self.section.pk,
        })

        choices = [value for value, _label in form.fields["executor"].choices]

        self.assertIn(self.matching_name, choices)
        self.assertNotIn(self.other_name, choices)

    def test_executor_choices_are_unfiltered_until_typical_section_is_selected(self):
        form = PerformerForm()

        choices = [value for value, _label in form.fields["executor"].choices]

        self.assertIn(self.matching_name, choices)
        self.assertIn(self.other_name, choices)


class ProjectRegistrationFormViewTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="projects-form-user",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.group_member = GroupMember.objects.create(
            short_name="IMC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        self.extra_product = Product.objects.create(
            short_name="CTRL",
            name_en="Control",
            name_ru="Контроль",
            consulting_type="Горный",
            service_category="Контроль",
            service_subtype="Контроль качества",
            position=2,
        )

    def test_registration_form_renders_linked_type_block(self):
        response = self.client.get(reverse("registration_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Добавить проект")
        self.assertContains(response, 'id="registration-type-meta"', html=False)
        self.assertContains(response, 'name="type_consulting"', html=False)
        self.assertContains(response, 'name="type_service_category"', html=False)
        self.assertContains(response, 'name="type_service_subtype"', html=False)
        self.assertContains(response, 'name="type_id"', html=False)
        self.assertContains(response, "Вид консалтинга")
        self.assertContains(response, "Тип услуги")
        self.assertContains(response, "Подтип услуги")
        self.assertContains(response, "Продукт")
        self.assertContains(response, "Контроль")
        self.assertContains(response, "Контроль качества")
        content = response.content.decode("utf-8")
        schedule_field_order = [
            '<label class="form-label">Год</label>',
            '<label class="form-label">Дата оценки</label>',
            '<label class="form-label">Дедлайн</label>',
            '<label class="form-label">Руководитель проекта</label>',
            '<label class="form-label">Статус</label>',
        ]
        schedule_field_positions = [content.index(item) for item in schedule_field_order]
        self.assertEqual(schedule_field_positions, sorted(schedule_field_positions))
        self.assertContains(response, '<label class="form-label mb-2">Сроки</label>', html=False)
        self.assertContains(response, '<th class="registration-contract-terms-stage-col">Этап</th>', html=False)
        self.assertContains(response, 'value="Итого"', html=False)
        self.assertContains(response, "js-registration-report-terms-lock", html=False)
        self.assertContains(response, "registration-stage-delay-add", html=False)
        self.assertContains(response, "registration-stage-next-delay-days", html=False)
        self.assertNotContains(response, "Дата Предварительного отчёта, оконч. расчет")

    def test_registration_form_renders_contract_id_field(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=8101,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Договор для выбора",
            year=2026,
        )

        response = self.client.get(reverse("registration_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Договор ID")
        self.assertContains(response, 'name="contract_project_registration"', html=False)
        self.assertContains(response, contract_project.short_uid)

    def test_projects_registry_shows_contract_id_column(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=8102,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Договор в реестре",
            year=2026,
        )
        registration = ProjectRegistration.objects.create(
            number=8103,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект с договором",
            status="Не начат",
            year=2026,
            contract_project_registration=contract_project,
        )

        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-col="contract-id"', html=False)
        self.assertContains(response, contract_project.short_uid)
        self.assertContains(response, registration.short_uid)
        self.assertEqual(registration.short_uid, "81030100RU")

    def test_registration_form_product_meta_includes_typical_terms(self):
        TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months="1.5",
            final_report_weeks=2,
            position=1,
        )

        response = self.client.get(reverse("registration_form_create"))

        self.assertEqual(response.status_code, 200)
        meta = json.loads(response.context["registration_type_meta_json"])
        product_meta = next(item for item in meta["products"] if item["id"] == self.product.pk)
        self.assertEqual(
            product_meta["typical_service_terms"],
            {
                "source_data_weeks": "0.0",
                "source_data_term_unit": "weeks",
                "preliminary_report_months": "1.5",
                "preliminary_report_term_unit": "months",
                "final_report_weeks": "2.0",
                "final_report_term_unit": "weeks",
            },
        )

    def test_registration_edit_form_preserves_custom_stage_terms_in_html(self):
        TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months="1.0",
            final_report_weeks=2,
            position=1,
        )
        registration = ProjectRegistration.objects.create(
            number=6212,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект с индивидуальными сроками",
            status="Не начат",
            year=2026,
            deadline=date(2026, 1, 10),
            stage1_weeks="5.0",
            stage2_weeks="10.0",
        )
        ProjectRegistrationProduct.objects.create(
            registration=registration,
            product=self.product,
            rank=1,
        )

        response = self.client.get(reverse("registration_form_edit", args=[registration.pk]))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        self.assertIn(f'data-selected-product-id="{self.product.pk}"', content)
        self.assertIn('name="stage1_weeks" value="5.0"', content)
        self.assertIn('name="stage2_weeks" value="10.0"', content)
        self.assertIn("productRow.dataset.selectedProductId", content)

    def test_registration_edit_rejects_multiple_products(self):
        registration = ProjectRegistration.objects.create(
            number=6210,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект для редактирования",
            status="Не начат",
            year=2026,
            deadline=date(2026, 1, 10),
        )
        ProjectRegistrationProduct.objects.create(
            registration=registration,
            product=self.product,
            rank=1,
        )

        response = self.client.post(
            reverse("registration_form_edit", args=[registration.pk]),
            {
                "number": 6210,
                "group_member": self.group_member.pk,
                "agreement_type": "MAIN",
                "type_id": [self.product.pk, self.extra_product.pk],
                "name": "Проект для редактирования",
                "status": "Не начат",
                "deadline": "2026-01-10",
                "year": 2026,
                "project_manager": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Для строки проекта можно выбрать только один продукт.")
        self.assertEqual(
            list(registration.product_links.order_by("rank").values_list("product_id", flat=True)),
            [self.product.pk],
        )

    def test_projects_registry_groups_repeated_numbers_visually(self):
        first = ProjectRegistration.objects.create(
            number=6211,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Первый продукт",
            status="Не начат",
            year=2026,
            position=1,
        )
        second = ProjectRegistration.objects.create(
            number=6211,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.extra_product,
            name="Второй продукт",
            status="Не начат",
            year=2026,
            position=2,
        )

        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertContains(response, 'class="registration-number-has-next"', html=False)
        self.assertContains(response, 'class="registration-number-continuation"', html=False)
        self.assertContains(response, 'data-reg-group-number="6211"', html=False)
        self.assertContains(response, 'data-reg-group-cell="number"', html=False)
        self.assertContains(response, 'querySelector(\'[data-reg-group-cell="agreement-type"]\')', html=False)
        self.assertContains(response, "__refreshRegistrationVisibleGroups", html=False)
        self.assertContains(response, first.short_uid)
        self.assertContains(response, second.short_uid)

    def test_projects_registry_keeps_grouped_labels_hidden_for_same_contract_id(self):
        contract_project = ContractProjectRegistration.objects.create(
            number=6212,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Общий договор",
            year=2026,
        )
        ProjectRegistration.objects.create(
            number=6212,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Первый продукт",
            status="Не начат",
            year=2026,
            position=1,
            contract_project_registration=contract_project,
        )
        ProjectRegistration.objects.create(
            number=6212,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.extra_product,
            name="Второй продукт",
            status="Не начат",
            year=2026,
            position=2,
            contract_project_registration=contract_project,
        )

        response = self.client.get(reverse("projects_partial"))
        content = response.content.decode("utf-8")

        self.assertEqual(response.status_code, 200)
        self.assertIn('class="registration-number-has-next"', content)
        self.assertNotIn('class="registration-number-has-next registration-number-contract-has-next"', content)
        self.assertEqual(
            content.count(
                f'<td class="contract-id-cell" data-col="contract-id" data-reg-group-cell="contract-id"><span class="clf-id-droid-mono">{contract_project.short_uid}</span></td>'
            ),
            1,
        )
        self.assertEqual(
            content.count('<td data-col="agreement-type" data-reg-group-cell="agreement-type">Основной договор</td>'),
            1,
        )
        self.assertEqual(
            content.count('<td data-col="group" data-reg-group-cell="group">RU</td>'),
            1,
        )

    def test_projects_registry_shows_grouped_labels_for_different_contract_ids(self):
        first_contract = ContractProjectRegistration.objects.create(
            number=6213,
            sub_number=1,
            group_member=self.group_member,
            type=self.product,
            name="Первый договор",
            year=2026,
        )
        second_contract = ContractProjectRegistration.objects.create(
            number=6213,
            sub_number=2,
            group_member=self.group_member,
            type=self.extra_product,
            name="Второй договор",
            year=2026,
        )
        ProjectRegistration.objects.create(
            number=6213,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Первый продукт",
            status="Не начат",
            year=2026,
            position=1,
            contract_project_registration=first_contract,
        )
        ProjectRegistration.objects.create(
            number=6213,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.extra_product,
            name="Второй продукт",
            status="Не начат",
            year=2026,
            position=2,
            contract_project_registration=second_contract,
        )

        response = self.client.get(reverse("projects_partial"))
        content = response.content.decode("utf-8")

        self.assertEqual(response.status_code, 200)
        self.assertIn('class="registration-number-has-next registration-number-contract-has-next"', content)
        self.assertContains(response, f'data-reg-group-contract-id="{first_contract.pk}"', html=False)
        self.assertContains(response, f'data-reg-group-contract-id="{second_contract.pk}"', html=False)
        self.assertEqual(
            content.count(
                f'<td class="contract-id-cell" data-col="contract-id" data-reg-group-cell="contract-id"><span class="clf-id-droid-mono">{first_contract.short_uid}</span></td>'
            ),
            1,
        )
        self.assertEqual(
            content.count(
                f'<td class="contract-id-cell" data-col="contract-id" data-reg-group-cell="contract-id"><span class="clf-id-droid-mono">{second_contract.short_uid}</span></td>'
            ),
            1,
        )
        self.assertEqual(
            content.count('<td data-col="agreement-type" data-reg-group-cell="agreement-type">Основной договор</td>'),
            2,
        )
        self.assertEqual(
            content.count('<td data-col="group" data-reg-group-cell="group">RU</td>'),
            2,
        )

    def test_projects_registry_customer_details_are_hidden_by_default_in_colpicker(self):
        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        default_hidden_columns = [
            "country",
            "identifier",
            "registration-number",
            "region",
            "date",
            "asset-owner-country",
            "asset-owner-identifier",
            "asset-owner-reg-number",
            "asset-owner-region",
            "asset-owner-date",
        ]
        for column in default_hidden_columns:
            self.assertContains(
                response,
                f'id="registration-col-{column}" data-default-hidden="true"',
                html=False,
            )
        self.assertContains(
            response,
            'id="registration-col-customer" checked',
            html=False,
        )
        self.assertContains(
            response,
            'id="registration-col-asset-owner" checked',
            html=False,
        )

    def test_registration_launch_moves_not_started_project_to_in_progress(self):
        registration = ProjectRegistration.objects.create(
            number=6201,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект для запуска",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(reverse("registration_launch", args=[registration.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"ok": True, "status": "В работе"})
        registration.refresh_from_db()
        self.assertEqual(registration.status, "В работе")

    def test_registration_launch_rejects_already_started_project(self):
        registration = ProjectRegistration.objects.create(
            number=6202,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Уже запущенный проект",
            status="В работе",
            year=2026,
        )

        response = self.client.post(reverse("registration_launch", args=[registration.pk]))

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()["ok"])
        registration.refresh_from_db()
        self.assertEqual(registration.status, "В работе")

    def test_registration_status_update_changes_project_status(self):
        registration = ProjectRegistration.objects.create(
            number=6203,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект со сменой статуса",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(
            reverse("registration_status_update", args=[registration.pk]),
            {"status": "На проверке"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"ok": True, "status": "На проверке"})
        registration.refresh_from_db()
        self.assertEqual(registration.status, "На проверке")

    def test_registration_status_update_copies_gantt_to_russian_production_calendar(self):
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )
        source_gantt_data = {
            "data": [
                {
                    "id": "task-1",
                    "text": "Работа",
                    "start_date": "2026-04-30",
                    "end_date": "2026-05-04",
                    "deadline": "2026-05-04",
                    "constraint_type": "fnlt",
                    "constraint_date": "2026-05-04",
                    "duration": 2,
                    "progress": 0,
                    "type": "task",
                }
            ],
            "links": [],
            "meta": {
                "base_date": "2026-04-30",
                "project_start": "2026-04-30",
                "project_end": "2026-05-04",
                "calendar_kind": "abstract",
                "executor_display": "resource_name",
                "version": 1,
            },
        }
        term = TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months="0",
            final_report_weeks=0,
            position=1,
            gantt_data=copy.deepcopy(source_gantt_data),
        )
        registration = ProjectRegistration.objects.create(
            number=6213,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект с типовым графиком",
            status="Не начат",
            year=2026,
        )

        with patch("projects_app.views.date") as mocked_date:
            mocked_date.today.return_value = date(2026, 4, 30)
            response = self.client.post(
                reverse("registration_status_update", args=[registration.pk]),
                {"status": "В работе"},
            )

        self.assertEqual(response.status_code, 200)
        registration.refresh_from_db()
        task = registration.gantt_data["data"][0]
        meta = registration.gantt_data["meta"]
        self.assertEqual(meta["calendar_kind"], "production")
        self.assertEqual(meta["calendar_country_id"], country.pk)
        self.assertEqual(meta["executor_display"], "executor")
        self.assertEqual(task["start_date"], "2026-04-30")
        self.assertEqual(task["end_date"], "2026-05-05")
        self.assertEqual(task["deadline"], "2026-05-05")
        self.assertEqual(task["constraint_date"], "2026-05-05")
        self.assertEqual(task["duration"], 2)
        term.refresh_from_db()
        self.assertEqual(term.gantt_data, source_gantt_data)

    def _create_project_gantt_asset_sync_fixture(self, *, deadline=None):
        section_a = self.product.sections.create(
            code="SEC-A",
            short_name="Section A",
            short_name_ru="Раздел A",
            name_en="Section A",
            name_ru="Раздел A",
            accounting_type="Раздел",
            position=1,
        )
        section_b = self.product.sections.create(
            code="SEC-B",
            short_name="Section B",
            short_name_ru="Раздел B",
            name_en="Section B",
            name_ru="Раздел B",
            accounting_type="Раздел",
            position=2,
        )
        specialty_a = ExpertSpecialty.objects.create(specialty="Специальность A")
        specialty_b = ExpertSpecialty.objects.create(specialty="Специальность B")
        TypicalSectionSpecialty.objects.create(section=section_a, specialty=specialty_a, rank=1)
        TypicalSectionSpecialty.objects.create(section=section_b, specialty=specialty_b, rank=1)
        TypicalServiceTerm.objects.create(
            product=self.product,
            preliminary_report_months="0",
            final_report_weeks=0,
            position=1,
            gantt_data={
                "data": [
                    {
                        "id": "source-data",
                        "text": "Исходные данные",
                        "start_date": "2025-12-20",
                        "end_date": "2026-01-01",
                        "duration": 8,
                        "progress": 0,
                        "type": "project",
                        "system_key": "source_data",
                        "$open": True,
                    },
                    {
                        "id": "source-data-asset-template",
                        "text": "Актив",
                        "start_date": "2025-12-20",
                        "end_date": "2026-01-01",
                        "duration": 8,
                        "progress": 0,
                        "type": "project",
                        "parent": "source-data",
                        "system_key": "source_data_asset",
                        "$open": True,
                    },
                    {
                        "id": "source-data-sec-a-template",
                        "text": "Раздел A",
                        "service_section_name": "Раздел A",
                        "start_date": "2025-12-22",
                        "end_date": "2025-12-26",
                        "duration": 4,
                        "progress": 0,
                        "type": "service_section",
                        "parent": "source-data-asset-template",
                    },
                    {
                        "id": "source-data-sec-b-template",
                        "text": "Раздел B",
                        "service_section_name": "Раздел B",
                        "start_date": "2025-12-26",
                        "end_date": "2025-12-30",
                        "duration": 4,
                        "progress": 0,
                        "type": "service_section",
                        "parent": "source-data-asset-template",
                    },
                    {
                        "id": "source-data-note-template",
                        "text": "Вспомогательная задача исходных данных",
                        "start_date": "2025-12-30",
                        "end_date": "2026-01-01",
                        "duration": 2,
                        "progress": 0,
                        "type": "task",
                        "parent": "source-data-asset-template",
                    },
                    {
                        "id": "preliminary",
                        "text": "Предварительный отчёт",
                        "start_date": "2026-01-01",
                        "end_date": "2026-01-20",
                        "duration": 13,
                        "progress": 0,
                        "type": "project",
                        "system_key": "preliminary_report",
                        "$open": True,
                    },
                    {
                        "id": "before-asset",
                        "text": "До актива",
                        "start_date": "2026-01-01",
                        "end_date": "2026-01-02",
                        "duration": 1,
                        "progress": 0,
                        "type": "task",
                        "parent": "preliminary",
                    },
                    {
                        "id": "asset-template",
                        "text": "Актив",
                        "start_date": "2026-01-01",
                        "end_date": "2026-01-20",
                        "duration": 13,
                        "progress": 0,
                        "type": "project",
                        "parent": "preliminary",
                        "system_key": "preliminary_report_asset",
                        "$open": True,
                    },
                    {
                        "id": "sec-a-template",
                        "text": "Раздел A",
                        "service_section_name": "Раздел A",
                        "start_date": "2026-01-02",
                        "end_date": "2026-01-06",
                        "duration": 2,
                        "progress": 0,
                        "type": "service_section",
                        "parent": "asset-template",
                    },
                    {
                        "id": "sec-b-template",
                        "text": "Раздел B",
                        "service_section_name": "Раздел B",
                        "start_date": "2026-01-06",
                        "end_date": "2026-01-08",
                        "duration": 2,
                        "progress": 0,
                        "type": "service_section",
                        "parent": "asset-template",
                    },
                    {
                        "id": "preliminary-note-template",
                        "text": "Вспомогательная задача предварительного отчета",
                        "start_date": "2026-01-08",
                        "end_date": "2026-01-10",
                        "duration": 2,
                        "progress": 0,
                        "type": "task",
                        "parent": "asset-template",
                    },
                    {
                        "id": "after-asset",
                        "text": "После актива",
                        "start_date": "2026-01-08",
                        "end_date": "2026-01-09",
                        "duration": 1,
                        "progress": 0,
                        "type": "task",
                        "parent": "preliminary",
                    },
                    {
                        "id": "final",
                        "text": "Итоговый отчёт",
                        "start_date": "2026-01-20",
                        "end_date": "2026-01-20",
                        "duration": 0,
                        "progress": 0,
                        "type": "milestone",
                        "system_key": "final_report",
                    },
                ],
                "links": [
                    {
                        "id": "sec-a-to-final",
                        "source": "sec-a-template",
                        "target": "final",
                        "type": "0",
                        "lag": 3,
                        "lag_mode": "auto",
                    },
                    {
                        "id": "source-data-sec-a-to-preliminary",
                        "source": "source-data-sec-a-template",
                        "target": "preliminary",
                        "type": "0",
                        "lag": 2,
                        "lag_mode": "fixed",
                    },
                    {
                        "id": "source-data-sec-b-to-note",
                        "source": "source-data-sec-b-template",
                        "target": "source-data-note-template",
                        "type": "0",
                        "lag": 0,
                        "lag_mode": "fixed",
                    },
                    {
                        "id": "preliminary-sec-b-to-note",
                        "source": "sec-b-template",
                        "target": "preliminary-note-template",
                        "type": "0",
                        "lag": 0,
                        "lag_mode": "fixed",
                    },
                ],
                "meta": {
                    "base_date": "2026-01-01",
                    "project_start": "2026-01-01",
                    "project_end": "2026-01-20",
                    "calendar_kind": "abstract",
                    "executor_display": "resource_name",
                    "version": 1,
                },
            },
        )
        registration = ProjectRegistration.objects.create(
            number=6214,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект с активами в графике",
            status="Не начат",
            year=2026,
            deadline=deadline,
        )
        ProjectRegistrationProduct.objects.create(
            registration=registration,
            product=self.product,
            rank=1,
        )
        asset_a = WorkVolume.objects.create(
            project=registration,
            name="Карьер",
            asset_name="Карьер",
            manager="Менеджер A",
        )
        asset_b = WorkVolume.objects.create(
            project=registration,
            name="Фабрика",
            asset_name="Фабрика",
            manager="Менеджер B",
        )
        Performer.objects.filter(work_item=asset_a).exclude(typical_section=section_a).delete()
        Performer.objects.filter(work_item=asset_b).exclude(typical_section=section_b).delete()
        performer_a = Performer.objects.get(work_item=asset_a, typical_section=section_a)
        performer_a.executor = "Иванов Иван Иванович"
        performer_a.save()
        performer_b = Performer.objects.get(work_item=asset_b, typical_section=section_b)
        performer_b.executor = "Петров Петр Петрович"
        performer_b.save()

        response = self.client.post(
            reverse("registration_status_update", args=[registration.pk]),
            {"status": "В работе"},
        )
        self.assertEqual(response.status_code, 200)
        registration.refresh_from_db()
        return registration, asset_a, asset_b, performer_a, performer_b

    def test_registration_launch_creates_managed_asset_and_performer_gantt_tasks(self):
        registration, asset_a, asset_b, performer_a, performer_b = self._create_project_gantt_asset_sync_fixture()

        tasks = registration.gantt_data["data"]
        task_ids = {str(task.get("id")) for task in tasks}
        broken_links = [
            link for link in registration.gantt_data["links"]
            if str(link.get("source")) not in task_ids
            or str(link.get("target")) not in task_ids
        ]
        asset_tasks = [
            task for task in tasks
            if task.get("managed_source") == "work_volume"
            and task.get("managed_scope") == "preliminary_report"
        ]
        performer_tasks = [task for task in tasks if task.get("managed_source") == "performer"]

        self.assertEqual(broken_links, [])
        self.assertCountEqual([task["text"] for task in asset_tasks], ["Карьер", "Фабрика"])
        self.assertCountEqual(
            [task["work_volume_id"] for task in asset_tasks],
            [asset_a.pk, asset_b.pk],
        )
        self.assertCountEqual(
            [(task["work_volume_id"], task["performer_id"], task["service_section_name"], task["executor"]) for task in performer_tasks],
            [
                (asset_a.pk, performer_a.pk, "Раздел A", "Иванов И.И."),
                (asset_b.pk, performer_b.pk, "Раздел B", "Петров П.П."),
            ],
        )
        for task in performer_tasks:
            parent = next(item for item in asset_tasks if item["id"] == task["parent"])
            self.assertEqual(task["asset_name"], parent["text"])

    def test_registration_launch_preserves_asset_template_position_among_preliminary_children(self):
        registration, _asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()

        preliminary_children = [
            task["text"]
            for task in registration.gantt_data["data"]
            if task.get("parent") == "preliminary"
        ]

        self.assertEqual(preliminary_children, ["До актива", "Карьер", "Фабрика", "После актива"])

    def test_registration_launch_preserves_service_section_template_links(self):
        registration, _asset_a, _asset_b, performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()

        managed_section_id = f"managed-performer-{performer_a.pk}"
        link = next(
            link for link in registration.gantt_data["links"]
            if str(link.get("source")) == managed_section_id and str(link.get("target")) == "final"
        )

        self.assertEqual(link["type"], "0")
        self.assertEqual(link["lag"], 3)
        self.assertEqual(link["lag_mode"], "auto")

    def test_registration_gantt_sync_creates_source_data_sections_from_checklists(self):
        registration, asset_a, asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        section_a = self.product.sections.get(code="SEC-A")
        section_b = self.product.sections.get(code="SEC-B")

        ChecklistItem.objects.create(
            project=registration,
            section=section_a,
            code="A",
            number=1,
            short_name="Паспорт",
            name="Паспорт по разделу A",
            position=1,
        )
        ChecklistItem.objects.create(
            project=registration,
            section=section_b,
            code="B",
            number=1,
            short_name="Отчет",
            name="Отчет по разделу B",
            position=2,
        )
        registration.refresh_from_db()

        tasks = registration.gantt_data["data"]
        source_assets = [
            task for task in tasks
            if task.get("managed_source") == "work_volume"
            and task.get("managed_scope") == "source_data"
        ]
        source_sections = [
            task for task in tasks
            if task.get("managed_source") == "checklist_section"
        ]

        self.assertCountEqual([task["text"] for task in source_assets], ["Карьер", "Фабрика"])
        self.assertCountEqual(
            [
                (
                    task["work_volume_id"],
                    task["typical_section_id"],
                    task["service_section_name"],
                    task["executor"],
                    task["specialty"],
                )
                for task in source_sections
            ],
            [
                (asset_a.pk, section_a.pk, "Раздел A", "", ""),
                (asset_a.pk, section_b.pk, "Раздел B", "", ""),
                (asset_b.pk, section_a.pk, "Раздел A", "", ""),
                (asset_b.pk, section_b.pk, "Раздел B", "", ""),
            ],
        )

    def test_project_gantt_sync_reflects_checklist_section_changes(self):
        registration, asset_a, asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        section_a = self.product.sections.get(code="SEC-A")
        item = ChecklistItem.objects.create(
            project=registration,
            section=section_a,
            code="A",
            number=1,
            short_name="Паспорт",
            name="Паспорт по разделу A",
            position=1,
        )
        registration.refresh_from_db()
        self.assertCountEqual(
            [
                task["work_volume_id"]
                for task in registration.gantt_data["data"]
                if task.get("managed_source") == "checklist_section"
                and task.get("typical_section_id") == section_a.pk
            ],
            [asset_a.pk, asset_b.pk],
        )

        item.delete()
        registration.refresh_from_db()
        self.assertFalse(
            any(
                task.get("managed_source") == "checklist_section"
                and task.get("typical_section_id") == section_a.pk
                for task in registration.gantt_data["data"]
            )
        )

    def test_source_data_section_progress_reflects_imcm_provided_status_share(self):
        registration, asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        section_a = self.product.sections.get(code="SEC-A")
        legal_entity = asset_a.legal_entities.get()
        item_a = ChecklistItem.objects.create(
            project=registration,
            section=section_a,
            code="A",
            number=1,
            short_name="Паспорт",
            name="Паспорт по разделу A",
            position=1,
        )
        item_b = ChecklistItem.objects.create(
            project=registration,
            section=section_a,
            code="A",
            number=2,
            short_name="Отчет",
            name="Отчет по разделу A",
            position=2,
        )

        ChecklistStatus.objects.create(
            checklist_item=item_a,
            legal_entity=legal_entity,
            status=ChecklistStatus.Status.PROVIDED,
        )
        ChecklistStatus.objects.create(
            checklist_item=item_b,
            legal_entity=legal_entity,
            status=ChecklistStatus.Status.PARTIAL,
        )
        registration.refresh_from_db()
        managed_section = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "checklist_section"
            and task.get("work_volume_id") == asset_a.pk
            and task.get("typical_section_id") == section_a.pk
        )
        self.assertEqual(managed_section["progress"], 0.5)

        second_status = ChecklistStatus.objects.get(checklist_item=item_b, legal_entity=legal_entity)
        second_status.status = ChecklistStatus.Status.PROVIDED
        second_status.save()
        registration.refresh_from_db()
        managed_section = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "checklist_section"
            and task.get("work_volume_id") == asset_a.pk
            and task.get("typical_section_id") == section_a.pk
        )
        self.assertEqual(managed_section["progress"], 1)

    def test_project_gantt_save_rejects_manual_source_data_section_progress_change(self):
        registration, asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        section_a = self.product.sections.get(code="SEC-A")
        ChecklistItem.objects.create(
            project=registration,
            section=section_a,
            code="A",
            number=1,
            short_name="Паспорт",
            name="Паспорт по разделу A",
            position=1,
        )
        registration.refresh_from_db()
        submitted_payload = copy.deepcopy(registration.gantt_data)
        managed_section = next(
            task for task in submitted_payload["data"]
            if task.get("managed_source") == "checklist_section"
            and task.get("work_volume_id") == asset_a.pk
            and task.get("typical_section_id") == section_a.pk
        )
        managed_section["progress"] = 0.75

        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("Управляемые поля", response.json()["error"])

    def test_registration_launch_preserves_source_data_section_template_links(self):
        registration, _asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        section_a = self.product.sections.get(code="SEC-A")
        ChecklistItem.objects.create(
            project=registration,
            section=section_a,
            code="A",
            number=1,
            short_name="Паспорт",
            name="Паспорт по разделу A",
            position=1,
        )
        registration.refresh_from_db()

        managed_section = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "checklist_section"
            and task.get("typical_section_id") == section_a.pk
        )
        link = next(
            link for link in registration.gantt_data["links"]
            if str(link.get("source")) == managed_section["id"]
            and str(link.get("target")) == "preliminary"
        )

        self.assertEqual(link["type"], "0")
        self.assertEqual(link["lag"], 2)
        self.assertEqual(link["lag_mode"], "fixed")

    def test_registration_launch_sets_preliminary_submission_deadline_constraint(self):
        registration, _asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture(
            deadline=date(2026, 6, 15),
        )

        submission = next(
            task for task in registration.gantt_data["data"]
            if task.get("system_key") == "preliminary_report_submission"
        )

        self.assertEqual(submission["text"], "Отправка Предварительного отчёта")
        self.assertEqual(submission["type"], "milestone")
        self.assertEqual(submission["constraint_type"], "mfo")
        self.assertEqual(submission["constraint_date"], "2026-06-15")

    def test_registration_deadline_update_syncs_preliminary_submission_constraint(self):
        registration, _asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()

        response = self.client.post(
            reverse("registration_deadline_update", args=[registration.pk]),
            {"deadline": "2026-06-15"},
        )

        self.assertEqual(response.status_code, 200)
        registration.refresh_from_db()
        submission = next(
            task for task in registration.gantt_data["data"]
            if task.get("system_key") == "preliminary_report_submission"
        )
        self.assertEqual(submission["constraint_type"], "mfo")
        self.assertEqual(submission["constraint_date"], "2026-06-15")

        response = self.client.post(
            reverse("registration_deadline_update", args=[registration.pk]),
            {"deadline": ""},
        )

        self.assertEqual(response.status_code, 200)
        registration.refresh_from_db()
        submission = next(
            task for task in registration.gantt_data["data"]
            if task.get("system_key") == "preliminary_report_submission"
        )
        self.assertNotIn("constraint_type", submission)
        self.assertNotIn("constraint_date", submission)

    def test_project_gantt_sync_reflects_performer_and_work_volume_changes(self):
        registration, asset_a, asset_b, performer_a, performer_b = self._create_project_gantt_asset_sync_fixture()

        performer_a.executor = "Сидоров Сидор Сидорович"
        performer_a.save()
        registration.refresh_from_db()
        updated_task = next(
            task for task in registration.gantt_data["data"]
            if task.get("performer_id") == performer_a.pk
        )
        self.assertEqual(updated_task["executor"], "Сидоров С.С.")

        performer_b_id = performer_b.pk
        performer_b.delete()
        registration.refresh_from_db()
        self.assertFalse(
            any(task.get("performer_id") == performer_b_id for task in registration.gantt_data["data"])
        )

        asset_a_id = asset_a.pk
        asset_a.delete()
        registration.refresh_from_db()
        self.assertFalse(
            any(task.get("work_volume_id") == asset_a_id for task in registration.gantt_data["data"])
        )
        self.assertTrue(
            any(task.get("work_volume_id") == asset_b.pk for task in registration.gantt_data["data"])
        )

        asset_b_id = asset_b.pk
        asset_b.delete()
        registration.refresh_from_db()
        self.assertFalse(
            any(task.get("work_volume_id") == asset_b_id for task in registration.gantt_data["data"])
        )
        self.assertFalse(
            any(task.get("managed_source") in {"work_volume", "performer"} for task in registration.gantt_data["data"])
        )
        self.assertTrue(
            any(task.get("system_key") == "preliminary_report_asset" for task in registration.gantt_data["data"])
        )

    def test_source_data_section_assignment_does_not_track_preliminary_performer(self):
        registration, asset_a, _asset_b, performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        section_a = self.product.sections.get(code="SEC-A")
        ChecklistItem.objects.create(
            project=registration,
            section=section_a,
            code="A",
            number=1,
            short_name="Паспорт",
            name="Паспорт по разделу A",
            position=1,
        )
        registration.refresh_from_db()

        source_section = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "checklist_section"
            and task.get("work_volume_id") == asset_a.pk
            and task.get("typical_section_id") == section_a.pk
        )
        self.assertEqual(source_section["executor"], "")
        self.assertEqual(source_section["specialty"], "")

        performer_a.executor = "Сидоров Сидор Сидорович"
        performer_a.save()
        registration.refresh_from_db()

        preliminary_section = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "performer"
            and task.get("performer_id") == performer_a.pk
        )
        source_section = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "checklist_section"
            and task.get("work_volume_id") == asset_a.pk
            and task.get("typical_section_id") == section_a.pk
        )
        self.assertEqual(preliminary_section["executor"], "Сидоров С.С.")
        self.assertEqual(source_section["executor"], "")
        self.assertEqual(source_section["specialty"], "")

    def test_managed_gantt_tasks_cannot_be_deleted_through_schedule_api(self):
        registration, asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        asset_task = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "work_volume"
            and task.get("managed_scope") == "preliminary_report"
            and task.get("work_volume_id") == asset_a.pk
        )

        from projects_app.services import gantt_tasks

        self.assertFalse(gantt_tasks.delete_task(registration, asset_task["id"]))
        registration.refresh_from_db()
        self.assertTrue(
            any(task.get("id") == asset_task["id"] for task in registration.gantt_data["data"])
        )

        submitted_payload = copy.deepcopy(registration.gantt_data)
        submitted_payload["data"] = [
            task for task in submitted_payload["data"]
            if task.get("id") != asset_task["id"]
        ]
        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("Управляемые задачи", response.json()["error"])

    def test_project_gantt_save_allows_managed_performer_executor_text(self):
        registration, _asset_a, _asset_b, performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        submitted_payload = copy.deepcopy(registration.gantt_data)
        unmanaged_task = next(task for task in submitted_payload["data"] if task.get("id") == "final")
        unmanaged_task["progress"] = 25

        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200, response.content.decode("utf-8"))
        registration.refresh_from_db()
        managed_task = next(
            task for task in registration.gantt_data["data"]
            if task.get("performer_id") == performer_a.pk
        )
        self.assertEqual(managed_task["executor"], "Иванов И.И.")
        final_task = next(task for task in registration.gantt_data["data"] if task.get("id") == "final")
        self.assertEqual(final_task["progress"], 25)

    def test_project_gantt_save_accepts_legacy_full_name_managed_executor(self):
        registration, _asset_a, _asset_b, performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        submitted_payload = copy.deepcopy(registration.gantt_data)
        managed_task = next(
            task for task in submitted_payload["data"]
            if task.get("performer_id") == performer_a.pk
        )
        managed_task["executor"] = "Иванов Иван Иванович"

        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200, response.content.decode("utf-8"))
        registration.refresh_from_db()
        saved_task = next(
            task for task in registration.gantt_data["data"]
            if task.get("performer_id") == performer_a.pk
        )
        self.assertEqual(saved_task["executor"], "Иванов И.И.")

    def test_project_gantt_save_accepts_managed_executor_profile_value(self):
        registration, _asset_a, _asset_b, performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        employee_user = get_user_model().objects.create_user(
            username="managed-executor-profile",
            first_name="Иван",
            last_name="Иванов",
        )
        employee = Employee.objects.create(user=employee_user, patronymic="Иванович")
        profile = ExpertProfile.objects.create(employee=employee)
        specialty = ExpertSpecialty.objects.get(specialty="Специальность A")
        ExpertProfileSpecialty.objects.create(profile=profile, specialty=specialty, rank=1)
        submitted_payload = copy.deepcopy(registration.gantt_data)
        managed_task = next(
            task for task in submitted_payload["data"]
            if task.get("performer_id") == performer_a.pk
        )
        managed_task["executor"] = f"expert-profile:{profile.pk}"

        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200, response.content.decode("utf-8"))
        registration.refresh_from_db()
        saved_task = next(
            task for task in registration.gantt_data["data"]
            if task.get("performer_id") == performer_a.pk
        )
        self.assertEqual(saved_task["executor"], "Иванов И.И.")

    def test_project_gantt_save_accepts_browser_flattened_managed_asset_type(self):
        registration, asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        submitted_payload = copy.deepcopy(registration.gantt_data)
        source_asset = next(
            task for task in submitted_payload["data"]
            if task.get("managed_source") == "work_volume"
            and task.get("managed_scope") == "source_data"
            and task.get("work_volume_id") == asset_a.pk
        )
        source_asset["type"] = "task"
        managed_task = next(
            task for task in submitted_payload["data"]
            if task.get("managed_source") == "performer"
        )
        managed_task["deadline"] = "2026-09-02"

        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200, response.content.decode("utf-8"))
        registration.refresh_from_db()
        saved_source_asset = next(
            task for task in registration.gantt_data["data"]
            if task.get("id") == source_asset["id"]
        )
        self.assertEqual(saved_source_asset["type"], "project")
        saved_managed_task = next(
            task for task in registration.gantt_data["data"]
            if task.get("id") == managed_task["id"]
        )
        self.assertEqual(saved_managed_task["deadline"], "2026-09-02")

    def test_project_gantt_save_allows_existing_template_executor_assignments(self):
        registration, _asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        submitted_payload = copy.deepcopy(registration.gantt_data)
        submitted_payload["data"].append({
            "id": "template-assigned-task",
            "text": "Шаблонная задача с исполнителем",
            "start_date": "2026-01-10",
            "end_date": "2026-01-12",
            "type": "task",
            "specialty": "Специальность A",
            "executor": "Пичугин В.А.",
            "resource_id": "resource-template",
        })
        submitted_payload.setdefault("meta", {})["resources"] = [
            {
                "id": "resource-template",
                "specialty": "Специальность A",
                "executor": "Пичугин В.А.",
                "task_ids": ["template-assigned-task", "missing-task"],
            },
            {
                "id": "resource-template",
                "specialty": "Специальность B",
                "executor": "Дворников А.В.",
                "task_ids": [],
            },
        ]

        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200, response.content.decode("utf-8"))
        registration.refresh_from_db()
        saved_task = next(task for task in registration.gantt_data["data"] if task["id"] == "template-assigned-task")
        self.assertEqual(saved_task["executor"], "Пичугин В.А.")
        resources = registration.gantt_data["meta"]["resources"]
        self.assertEqual([resource["id"] for resource in resources], ["resource-template", "resource-template-2"])

    def test_project_gantt_save_drops_browser_resource_for_managed_task_without_specialty(self):
        registration, asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        section_without_specialties = self.product.sections.create(
            code="SEC-C",
            short_name="Section C",
            short_name_ru="Раздел C",
            name_en="Section C",
            name_ru="Раздел C",
            accounting_type="Раздел",
            position=3,
        )
        performer = Performer.objects.create(
            work_item=asset_a,
            registration=registration,
            asset_name=asset_a.asset_name,
            typical_section=section_without_specialties,
            executor="Соколова Мария Александровна",
        )
        asset_task = next(
            task for task in registration.gantt_data["data"]
            if task.get("managed_source") == "work_volume"
            and task.get("managed_scope") == "preliminary_report"
            and task.get("work_volume_id") == asset_a.pk
        )
        task_id = f"managed-performer-{performer.pk}"
        browser_resource_id = "resource-browser-managed-without-specialty"
        registration.gantt_data["data"].append({
            "id": task_id,
            "text": "Раздел C",
            "service_section_name": "Раздел C",
            "start_date": "2026-01-10",
            "end_date": "2026-01-12",
            "duration": 2,
            "progress": 0,
            "type": "service_section",
            "parent": asset_task["id"],
            "managed_source": "performer",
            "performer_id": performer.pk,
            "work_volume_id": asset_a.pk,
            "typical_section_id": section_without_specialties.pk,
            "asset_name": asset_a.asset_name,
            "specialty": "",
            "executor": "Соколова М.А.",
            "resource_id": browser_resource_id,
            "resource_name": "Соколова М.А.",
        })
        registration.gantt_data.setdefault("meta", {}).setdefault("resources", []).append({
            "id": browser_resource_id,
            "specialty": "",
            "executor": "Соколова М.А.",
            "resource_name": "Соколова М.А.",
            "task_ids": [task_id],
        })
        registration.save(update_fields=["gantt_data"])
        submitted_payload = copy.deepcopy(registration.gantt_data)

        response = self.client.post(
            reverse("project_schedule_gantt", args=[registration.pk]),
            data=json.dumps(submitted_payload),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200, response.content.decode("utf-8"))
        registration.refresh_from_db()
        saved_task = next(task for task in registration.gantt_data["data"] if task["id"] == task_id)
        self.assertEqual(saved_task["specialty"], "")
        self.assertNotIn("resource_id", saved_task)
        self.assertFalse(
            any(resource["id"] == browser_resource_id for resource in registration.gantt_data["meta"]["resources"])
        )

    def test_project_gantt_get_includes_managed_performer_executor_options(self):
        registration, _asset_a, _asset_b, _performer_a, _performer_b = self._create_project_gantt_asset_sync_fixture()
        conflict_user = get_user_model().objects.create_user(
            username="managed.executor.conflict",
            first_name="Иван",
            last_name="Иванов",
            is_staff=True,
        )
        conflict_employee = Employee.objects.create(user=conflict_user, patronymic="Иванович")
        conflict_profile = ExpertProfile.objects.create(employee=conflict_employee)
        conflict_specialty = ExpertSpecialty.objects.create(specialty="Другая специальность")
        ExpertProfileSpecialty.objects.create(
            profile=conflict_profile,
            specialty=conflict_specialty,
            rank=1,
        )

        response = self.client.get(reverse("project_schedule_gantt", args=[registration.pk]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        executor_options = payload["executor_options"]
        self.assertTrue(
            any(
                item["value"] == "Иванов И.И."
                and item["label"] == "Иванов И.И."
                and "Специальность A" in item["specialties"]
                for item in executor_options
            )
        )
        self.assertTrue(
            any(item["value"] == "Петров П.П." and item["label"] == "Петров П.П." for item in executor_options)
        )

    def test_registration_status_update_rejects_unknown_status(self):
        registration = ProjectRegistration.objects.create(
            number=6204,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект с ошибочным статусом",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(
            reverse("registration_status_update", args=[registration.pk]),
            {"status": "Неизвестно"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()["ok"])
        registration.refresh_from_db()
        self.assertEqual(registration.status, "Не начат")

    def test_registration_manager_update_changes_project_manager(self):
        manager_user = get_user_model().objects.create_user(
            username="project.manager",
            password="secret",
            first_name="Иван",
            last_name="Иванов",
        )
        manager_employee = Employee.objects.create(
            user=manager_user,
            patronymic="Иванович",
            role=PROJECTS_HEAD_GROUP,
        )
        registration = ProjectRegistration.objects.create(
            number=6205,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект со сменой руководителя",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(
            reverse("registration_manager_update", args=[registration.pk]),
            {"project_manager": "Иванов Иван Иванович"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json(),
            {
                "ok": True,
                "manager": "Иванов Иван Иванович",
                "managerValue": manager_employee.formatted_prs_id,
                "managerLabel": "Иванов И.И.",
            },
        )
        registration.refresh_from_db()
        self.assertEqual(registration.project_manager, "Иванов Иван Иванович")
        self.assertEqual(registration.project_manager_prs_id, manager_employee.formatted_prs_id)

    def test_registration_manager_update_allows_clearing_project_manager(self):
        registration = ProjectRegistration.objects.create(
            number=6206,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект без руководителя",
            status="Не начат",
            year=2026,
            project_manager="Иванов Иван Иванович",
        )

        response = self.client.post(
            reverse("registration_manager_update", args=[registration.pk]),
            {"project_manager": ""},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["managerLabel"], "")
        registration.refresh_from_db()
        self.assertEqual(registration.project_manager, "")
        self.assertEqual(registration.project_manager_prs_id, "")

    def test_registration_deadline_update_changes_project_deadline(self):
        registration = ProjectRegistration.objects.create(
            number=6207,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект со сменой дедлайна",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(
            reverse("registration_deadline_update", args=[registration.pk]),
            {"deadline": "2026-06-15"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json(),
            {"ok": True, "deadline": "2026-06-15", "deadlineLabel": "15.06.2026"},
        )
        registration.refresh_from_db()
        self.assertEqual(registration.deadline, date(2026, 6, 15))

    def test_registration_deadline_update_allows_clearing_project_deadline(self):
        registration = ProjectRegistration.objects.create(
            number=6208,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект без дедлайна",
            status="Не начат",
            year=2026,
            deadline=date(2026, 6, 15),
        )

        response = self.client.post(
            reverse("registration_deadline_update", args=[registration.pk]),
            {"deadline": ""},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"ok": True, "deadline": "", "deadlineLabel": ""})
        registration.refresh_from_db()
        self.assertIsNone(registration.deadline)

    def test_registration_deadline_update_rejects_invalid_date(self):
        registration = ProjectRegistration.objects.create(
            number=6209,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект с ошибочным дедлайном",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(
            reverse("registration_deadline_update", args=[registration.pk]),
            {"deadline": "15.06.2026"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()["ok"])
        registration.refresh_from_db()
        self.assertIsNone(registration.deadline)

    def test_registration_evaluation_date_update_changes_project_evaluation_date(self):
        registration = ProjectRegistration.objects.create(
            number=6210,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект со сменой даты оценки",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(
            reverse("registration_evaluation_date_update", args=[registration.pk]),
            {"evaluation_date": "2026-06-15"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json(),
            {"ok": True, "evaluationDate": "2026-06-15", "evaluationDateLabel": "15.06.2026"},
        )
        registration.refresh_from_db()
        self.assertEqual(registration.evaluation_date, date(2026, 6, 15))

    def test_registration_evaluation_date_update_allows_clearing_project_evaluation_date(self):
        registration = ProjectRegistration.objects.create(
            number=6211,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект без даты оценки",
            status="Не начат",
            year=2026,
            evaluation_date=date(2026, 6, 15),
        )

        response = self.client.post(
            reverse("registration_evaluation_date_update", args=[registration.pk]),
            {"evaluation_date": ""},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"ok": True, "evaluationDate": "", "evaluationDateLabel": ""})
        registration.refresh_from_db()
        self.assertIsNone(registration.evaluation_date)

    def test_registration_evaluation_date_update_rejects_invalid_date(self):
        registration = ProjectRegistration.objects.create(
            number=6212,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            type=self.product,
            name="Проект с ошибочной датой оценки",
            status="Не начат",
            year=2026,
        )

        response = self.client.post(
            reverse("registration_evaluation_date_update", args=[registration.pk]),
            {"evaluation_date": "15.06.2026"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(response.json()["ok"])
        registration.refresh_from_db()
        self.assertIsNone(registration.evaluation_date)


class ProjectRegistrationRegistrySyncTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="projects-registry-user",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.group_member = GroupMember.objects.create(
            short_name="IMC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=1,
        )
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
            position=1,
        )
        self.second_product = Product.objects.create(
            short_name="QAQC_SYNC",
            name_en="QAQC Sync",
            name_ru="QAQC Sync",
            consulting_type="Горный",
            service_category="Контроль",
            service_subtype="Контроль качества",
            position=2,
        )
        self.country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Российская Федерация",
            alpha2="RU",
            alpha3="RUS",
            position=1,
        )

    def _post_registration(self, **extra):
        payload = {
            "number": 4444,
            "group_member": self.group_member.pk,
            "agreement_type": "MAIN",
            "type_id": [self.product.pk],
            "name": "Проект Альфа",
            "status": "Не начат",
            "deadline": "2026-01-10",
            "year": 2026,
            "customer": 'ООО "Синхронизация"',
            "country": self.country.pk,
            "identifier": "ОГРН",
            "registration_number": "1234567890",
            "registration_date": "2025-02-10",
            "project_manager": "",
            "customer_autocomplete_identifier_record_id": "",
            "customer_autocomplete_selected_from_autocomplete": "0",
        }
        payload.update(extra)
        return self.client.post(reverse("registration_form_create"), payload)

    def test_manual_registration_save_creates_new_registry_chain(self):
        existing_entity = BusinessEntityRecord.objects.create(name='ООО "Синхронизация"', position=1)
        existing_identifier = BusinessEntityIdentifierRecord.objects.create(
            business_entity=existing_entity,
            identifier_type="ОГРН",
            registration_country=self.country,
            registration_region="Москва",
            registration_date=timezone.datetime(2025, 2, 10).date(),
            number="1234567890",
            valid_from=timezone.datetime(2025, 2, 10).date(),
            position=1,
        )
        LegalEntityRecord.objects.create(
            attribute=LegalEntityRecord.ATTRIBUTE_NAME,
            identifier_record=existing_identifier,
            short_name='ООО "Синхронизация"',
            registration_country=self.country,
            identifier="ОГРН",
            registration_number="1234567890",
            registration_date=timezone.datetime(2025, 2, 10).date(),
            position=1,
        )

        response = self._post_registration()

        self.assertEqual(response.status_code, 200)
        self.assertEqual(BusinessEntityRecord.objects.count(), 2)
        self.assertEqual(
            LegalEntityRecord.objects.filter(
                attribute=LegalEntityRecord.ATTRIBUTE_NAME,
                short_name='ООО "Синхронизация"',
            ).count(),
            2,
        )
        newest_entity = BusinessEntityRecord.objects.order_by("-id").first()
        project = ProjectRegistration.objects.get(number=4444)
        self.assertEqual(newest_entity.source, "[Проекты / Заказчик]")
        self.assertEqual(newest_entity.record_author, "projects-registry-user")
        self.assertEqual(project.type_short_display, "DD")
        self.assertEqual(
            list(project.product_links.order_by("rank").values_list("product_id", flat=True)),
            [self.product.pk],
        )

    def test_manual_registration_save_allows_duplicate_identity_values(self):
        first = ProjectRegistration.objects.create(
            number=4444,
            group_member=self.group_member,
            agreement_type=ProjectRegistration.AgreementType.MAIN,
            agreement_number="IMCM/4444",
            type=self.product,
            name="Проект Альфа",
            status="Не начат",
            year=2026,
            position=1,
        )

        response = self._post_registration()

        self.assertEqual(response.status_code, 200)
        duplicates = list(
            ProjectRegistration.objects
            .filter(
                number=4444,
                group_member=self.group_member,
                agreement_type=ProjectRegistration.AgreementType.MAIN,
                agreement_number="IMCM/4444",
                type=self.product,
            )
            .order_by("id")
        )
        self.assertEqual(len(duplicates), 2)
        self.assertNotEqual(duplicates[0].pk, duplicates[1].pk)
        self.assertNotEqual(duplicates[0].short_uid, duplicates[1].short_uid)
        first.refresh_from_db()
        self.assertEqual(first.short_uid, duplicates[0].short_uid)
        self.assertEqual(duplicates[0].short_uid, "44440010RU")
        self.assertEqual(duplicates[1].short_uid, "44440020RU")
        self.assertEqual(
            list(duplicates[1].product_links.order_by("rank").values_list("product_id", flat=True)),
            [self.product.pk],
        )

    def test_selected_autocomplete_registration_save_does_not_create_new_chain(self):
        existing_entity = BusinessEntityRecord.objects.create(name='ООО "Связать"', position=1)
        existing_identifier = BusinessEntityIdentifierRecord.objects.create(
            business_entity=existing_entity,
            identifier_type="ОГРН",
            registration_country=self.country,
            registration_region="Москва",
            registration_date=timezone.datetime(2025, 2, 10).date(),
            number="1234567890",
            valid_from=timezone.datetime(2025, 2, 10).date(),
            position=1,
        )
        LegalEntityRecord.objects.create(
            attribute=LegalEntityRecord.ATTRIBUTE_NAME,
            identifier_record=existing_identifier,
            short_name='ООО "Связать"',
            registration_country=self.country,
            identifier="ОГРН",
            registration_number="1234567890",
            registration_date=timezone.datetime(2025, 2, 10).date(),
            position=1,
        )

        response = self._post_registration(
            customer='ООО "Связать"',
            customer_autocomplete_identifier_record_id=str(existing_identifier.pk),
            customer_autocomplete_selected_from_autocomplete="1",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(BusinessEntityRecord.objects.count(), 1)

    def test_registration_create_and_work_deps_use_ranked_joined_products(self):
        response = self._post_registration(type_id=[self.product.pk, self.second_product.pk])

        self.assertEqual(response.status_code, 200)
        projects = list(ProjectRegistration.objects.filter(number=4444).order_by("position", "id"))
        self.assertEqual(len(projects), 2)
        self.assertEqual([project.type_short_display for project in projects], ["DD", self.second_product.short_name])
        self.assertEqual([project.short_uid for project in projects], ["44440010RU", "44440020RU"])
        self.assertEqual(
            list(projects[0].product_links.order_by("rank").values_list("product_id", flat=True)),
            [self.product.pk],
        )
        self.assertEqual(
            list(projects[1].product_links.order_by("rank").values_list("product_id", flat=True)),
            [self.second_product.pk],
        )
        self.assertEqual(BusinessEntityRecord.objects.count(), 1)

        deps_response = self.client.get(reverse("work_deps"), {"project": projects[1].pk})

        self.assertEqual(deps_response.status_code, 200)
        self.assertEqual(deps_response.json()["type_short"], self.second_product.short_name)

    def test_registration_create_preserves_duplicate_product_stages_as_rows(self):
        response = self._post_registration(
            type_id=[self.product.pk, self.product.pk],
            stage1_weeks=["1.0", "2.0"],
            stage2_weeks=["3.0", "4.0"],
        )

        self.assertEqual(response.status_code, 200)
        projects = list(ProjectRegistration.objects.filter(number=4444).order_by("position", "id"))
        self.assertEqual(len(projects), 2)
        self.assertEqual([project.type_short_display for project in projects], ["DD", "DD"])
        self.assertEqual([project.short_uid for project in projects], ["44440010RU", "44440020RU"])
        self.assertEqual([str(project.stage1_weeks) for project in projects], ["1.0", "2.0"])
        self.assertEqual([str(project.stage2_weeks) for project in projects], ["3.0", "4.0"])
        self.assertEqual(
            [
                list(project.product_links.order_by("rank").values_list("product_id", flat=True))
                for project in projects
            ],
            [[self.product.pk], [self.product.pk]],
        )

    def test_registration_create_applies_schedule_terms_by_product_row(self):
        response = self._post_registration(
            type_id=[self.product.pk, self.second_product.pk],
            contract_start="2026-01-01",
            stage1_weeks=["1.0", "2.0"],
            stage2_weeks=["3.0", "4.0"],
        )

        self.assertEqual(response.status_code, 200)
        projects = list(ProjectRegistration.objects.filter(number=4444).order_by("position", "id"))
        self.assertEqual([str(project.stage1_weeks) for project in projects], ["1.0", "2.0"])
        self.assertEqual([str(project.stage2_weeks) for project in projects], ["3.0", "4.0"])


class WorkVolumePerformerCreationTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="performers-view-user",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.second_product = Product.objects.create(
            short_name="QAQC_WORK",
            name_en="QAQC Work",
            name_ru="QAQC Work",
            consulting_type="Горный",
            service_category="Контроль",
            service_subtype="Контроль качества",
            position=2,
        )
        self.project_manager_employee = self._create_employee(
            username="project.manager",
            first_name="Иван",
            last_name="Иванов",
            patronymic="Иванович",
        )
        self.first_manager_employee = self._create_employee(
            username="asset.manager.one",
            first_name="Петр",
            last_name="Петров",
            patronymic="Петрович",
        )
        self.second_manager_employee = self._create_employee(
            username="asset.manager.two",
            first_name="Сидор",
            last_name="Сидоров",
            patronymic="Сидорович",
        )
        self.project_manager_name = "Иванов Иван Иванович"
        self.first_manager_name = "Петров Петр Петрович"
        self.second_manager_name = "Сидоров Сидор Сидорович"
        self.project = ProjectRegistration.objects.create(
            number=6001,
            type=self.product,
            name="Проект активов",
            year=2026,
            project_manager=self.project_manager_name,
        )
        ProjectRegistrationProduct.objects.create(
            registration=self.project,
            product=self.product,
            rank=1,
        )
        ProjectRegistrationProduct.objects.create(
            registration=self.project,
            product=self.second_product,
            rank=2,
        )
        self.product.sections.create(
            code="PRD",
            short_name="Project Director",
            short_name_ru="Руководитель проекта",
            name_en="Project Director",
            name_ru="Руководитель проекта",
            accounting_type="Раздел",
            position=1,
        )
        self.product.sections.create(
            code="PRJ",
            short_name="Project Manager",
            short_name_ru="Менеджер проекта",
            name_en="Project Manager",
            name_ru="Менеджер проекта",
            accounting_type="Раздел",
            position=2,
        )
        self.product.sections.create(
            code="CRD",
            short_name="Coordinator",
            short_name_ru="Координатор",
            name_en="Coordinator",
            name_ru="Координатор",
            accounting_type="Раздел",
            position=3,
        )
        self.second_product.sections.create(
            code="QA",
            short_name="QA Section",
            short_name_ru="QA раздел",
            name_en="QA Section",
            name_ru="QA раздел",
            accounting_type="Раздел",
            position=1,
        )

    def _create_employee(self, *, username, first_name, last_name, patronymic):
        user = get_user_model().objects.create_user(
            username=username,
            password="secret",
            first_name=first_name,
            last_name=last_name,
            is_staff=True,
        )
        return Employee.objects.create(user=user, patronymic=patronymic)

    def _create_payment_request_notification(self, *, performers, request_number, recipient=None):
        notification = Notification.objects.create(
            notification_type=Notification.NotificationType.PROJECT_PAYMENT_REQUEST,
            related_section=Notification.RelatedSection.PROJECTS,
            recipient=recipient or self.user,
            sender=self.user,
            project=self.project,
            title_text=f"Заявка на оплату №{request_number}",
            payload={"number_of_request": str(request_number)},
            is_read=False,
            is_processed=False,
        )
        NotificationPerformerLink.objects.bulk_create(
            [
                NotificationPerformerLink(
                    notification=notification,
                    performer=performer,
                    position=index,
                )
                for index, performer in enumerate(performers, start=1)
            ]
        )
        return notification

    def test_performer_move_actions_are_scoped_to_project(self):
        other_project = ProjectRegistration.objects.create(
            number=6002,
            type=self.product,
            name="Другой проект",
            year=2026,
        )
        first = Performer.objects.create(
            registration=self.project,
            executor="Первый исполнитель",
            position=1,
        )
        second = Performer.objects.create(
            registration=self.project,
            executor="Второй исполнитель",
            position=2,
        )
        other = Performer.objects.create(
            registration=other_project,
            executor="Исполнитель другого проекта",
            position=1,
        )

        response = self.client.post(reverse("performer_move_up", args=[second.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            list(
                Performer.objects
                .filter(registration=self.project)
                .order_by("position", "id")
                .values_list("pk", flat=True)
            ),
            [second.pk, first.pk],
        )
        other.refresh_from_db()
        self.assertEqual(other.position, 1)

        response = self.client.post(reverse("performer_move_down", args=[second.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            list(
                Performer.objects
                .filter(registration=self.project)
                .order_by("position", "id")
                .values_list("pk", flat=True)
            ),
            [first.pk, second.pk],
        )

    def test_performer_row_order_persists_order_within_each_project(self):
        other_project = ProjectRegistration.objects.create(
            number=6002,
            type=self.product,
            name="Другой проект",
            year=2026,
        )
        first = Performer.objects.create(
            registration=self.project,
            executor="Первый исполнитель",
            position=1,
        )
        second = Performer.objects.create(
            registration=self.project,
            executor="Второй исполнитель",
            position=2,
        )
        other = Performer.objects.create(
            registration=other_project,
            executor="Исполнитель другого проекта",
            position=1,
        )

        response = self.client.post(
            reverse("performer_row_order"),
            data=json.dumps({"ordered_performer_ids": [other.pk, second.pk, first.pk]}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["ok"], True)
        self.assertEqual(
            list(
                Performer.objects
                .filter(registration=self.project)
                .order_by("position", "id")
                .values_list("pk", flat=True)
            ),
            [second.pk, first.pk],
        )
        other.refresh_from_db()
        self.assertEqual(other.position, 1)

    def test_performers_partial_enables_queued_row_order(self):
        performer = Performer.objects.create(
            registration=self.project,
            executor="Исполнитель",
            position=1,
        )

        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "data-queued-row-order", html=False)
        self.assertContains(response, reverse("performer_row_order"), html=False)
        self.assertContains(response, 'data-row-order-payload-field="ordered_performer_ids"', html=False)
        self.assertContains(response, f'data-row-order-id="{performer.pk}"', html=False)

    def test_work_volume_form_uses_project_manager_label(self):
        form = WorkVolumeForm()

        self.assertEqual(form.fields["manager"].label, "Менеджер проекта")

    def test_equal_project_manager_and_manager_skips_prj_and_assigns_prd_to_project_manager(self):
        work_item = WorkVolume.objects.create(
            project=self.project,
            name="Актив 1",
            asset_name="Актив 1",
            manager=self.project_manager_name,
        )

        self.assertFalse(
            Performer.objects.filter(work_item=work_item, typical_section__code="PRJ").exists()
        )
        prd = Performer.objects.get(work_item=work_item, typical_section__code="PRD")
        self.assertEqual(prd.executor, self.project_manager_name)
        self.assertEqual(prd.employee, self.project_manager_employee)

    def test_different_manager_creates_prj_and_does_not_duplicate_prd_or_crd(self):
        first_work_item = WorkVolume.objects.create(
            project=self.project,
            name="Актив 1",
            asset_name="Актив 1",
            manager=self.first_manager_name,
        )

        self.assertTrue(
            Performer.objects.filter(work_item=first_work_item, typical_section__code="PRD").exists()
        )
        self.assertTrue(
            Performer.objects.filter(work_item=first_work_item, typical_section__code="PRJ").exists()
        )

        second_work_item = WorkVolume.objects.create(
            project=self.project,
            name="Актив 2",
            asset_name="Актив 2",
            manager=self.second_manager_name,
        )

        self.assertEqual(
            Performer.objects.filter(registration=self.project, typical_section__code="PRD").count(),
            1,
        )
        self.assertEqual(
            Performer.objects.filter(registration=self.project, typical_section__code="CRD").count(),
            1,
        )
        self.assertEqual(
            Performer.objects.filter(registration=self.project, typical_section__code="PRJ").count(),
            2,
        )

        first_prj = Performer.objects.get(work_item=first_work_item, typical_section__code="PRJ")
        second_prj = Performer.objects.get(work_item=second_work_item, typical_section__code="PRJ")
        prd = Performer.objects.get(registration=self.project, typical_section__code="PRD")

        self.assertEqual(first_prj.executor, self.first_manager_name)
        self.assertEqual(first_prj.employee, self.first_manager_employee)
        self.assertEqual(second_prj.executor, self.second_manager_name)
        self.assertEqual(second_prj.employee, self.second_manager_employee)
        self.assertEqual(prd.executor, self.project_manager_name)
        self.assertEqual(prd.employee, self.project_manager_employee)

    def test_multi_product_work_item_creates_sections_grouped_by_ranked_products(self):
        work_item = WorkVolume.objects.create(
            project=self.project,
            name="Актив 3",
            asset_name="Актив 3",
            manager=self.first_manager_name,
        )

        codes = list(
            Performer.objects
            .filter(work_item=work_item)
            .order_by("position", "id")
            .values_list("typical_section__code", flat=True)
        )

        self.assertEqual(codes, ["PRD", "PRJ", "CRD", "QA"])

    def test_system_dsc_does_not_create_performer_rows_for_new_asset(self):
        self.product.sections.create(
            code="DSC",
            short_name="Description",
            short_name_ru="Описание",
            name_en="Product description",
            name_ru="Описание продукта",
            accounting_type="Раздел",
            is_system=True,
            position=0,
        )

        work_item = WorkVolume.objects.create(
            project=self.project,
            name="Актив DSC",
            asset_name="Актив DSC",
            manager=self.first_manager_name,
        )

        self.assertFalse(
            Performer.objects.filter(work_item=work_item, typical_section__code="DSC").exists()
        )

    def test_performer_modal_sections_map_excludes_system_dsc(self):
        self.product.sections.create(
            code="DSC",
            short_name="Description",
            short_name_ru="Описание",
            name_en="Product description",
            name_ru="Описание продукта",
            accounting_type="Раздел",
            is_system=True,
            position=0,
        )

        response = self.client.get(reverse("performer_form_create"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, '"code": "DSC"')
        self.assertNotContains(response, "Описание продукта")

    def test_performers_partial_shows_section_product_in_main_table(self):
        WorkVolume.objects.create(
            project=self.project,
            name="Актив 3",
            asset_name="Актив 3",
            manager=self.first_manager_name,
        )

        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "<th class=\"text-nowrap\">Продукт</th>", html=False)
        self.assertContains(response, self.second_product.short_name)
        content = response.content.decode("utf-8")
        main_table = content[
            content.index('id="performers-main-section"'):
            content.index('id="participation-confirmation-section"')
        ]
        self.assertNotIn("<th>Аванс</th>", main_table)
        self.assertNotIn("<th>Окон. платёж</th>", main_table)
        self.assertNotIn("<th>№ договора</th>", main_table)

        performer = Performer.objects.filter(registration=self.project).first()
        form_response = self.client.get(reverse("performer_form_edit", args=[performer.pk]))

        self.assertEqual(form_response.status_code, 200)
        form_content = form_response.content.decode("utf-8")
        self.assertNotIn('name="prepayment"', form_content)
        self.assertNotIn('name="final_payment"', form_content)
        self.assertNotIn('name="contract_number"', form_content)

    def test_performers_partial_splits_project_subsections(self):
        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="projects-content-team" class="projects-section-content d-none"', html=False)
        self.assertContains(response, 'id="projects-content-info-request" class="projects-section-content d-none"', html=False)
        self.assertContains(response, 'id="projects-content-performer-payments" class="projects-section-content d-none"', html=False)
        self.assertContains(response, 'id="performers-main-section"', html=False)
        self.assertContains(response, 'id="participation-confirmation-section"', html=False)
        self.assertContains(response, 'id="info-request-approval-section"', html=False)
        self.assertContains(response, 'id="payment-request-section"', html=False)

        content = response.content.decode("utf-8")
        self.assertLess(
            content.index('id="projects-content-team"'),
            content.index('id="performers-main-section"'),
        )
        self.assertLess(
            content.index('id="participation-confirmation-section"'),
            content.index('id="projects-content-info-request"'),
        )
        self.assertLess(
            content.index('id="projects-content-info-request"'),
            content.index('id="info-request-approval-section"'),
        )
        self.assertLess(
            content.index('id="info-request-approval-section"'),
            content.index('id="projects-content-performer-payments"'),
        )
        self.assertLess(
            content.index('id="projects-content-performer-payments"'),
            content.index('id="payment-request-section"'),
        )

    def test_home_page_lists_performer_payments_after_info_request(self):
        response = self.client.get(reverse("home"))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        info_request_index = content.index('data-projects-section="info-request"')
        payments_index = content.index('data-projects-section="performer-payments"')
        self.assertLess(info_request_index, payments_index)
        self.assertIn("'performer-payments': 'Платежи исполнителям'", content)

    def test_performers_partial_renders_payment_requests_from_signing_rows(self):
        from classifiers_app.models import OKVCurrency

        rub = OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
        )
        contracted = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-42",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("30"),
            final_payment=Decimal("70"),
            agreed_amount=Decimal("1234567.89"),
            currency=rub,
        )
        draft_only = Performer.objects.create(
            registration=self.project,
            employee=self.first_manager_employee,
            executor=self.first_manager_name,
        )

        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Заявки на оплату")
        content = response.content.decode("utf-8")
        section = content[
            content.index('id="payment-request-section"'):
            content.index("<script>", content.index('id="payment-request-section"'))
        ]
        for header in (
            "Номер",
            "Тип",
            "Название",
            "Исполнитель",
            "Номер договора",
            "Заключение договора",
            "Цена",
            "Аванс",
            "Ок. пл.",
            "Заявка на аванс",
            "Заявка на ок. пл.",
            "Оплата",
            "Отправитель",
        ):
            self.assertIn(header, section)
        self.assertNotIn(">Дата<", section)
        self.assertNotIn("Ок. платеж", section)
        self.assertIn("payment-request-signing-status-cell", section)
        self.assertIn("Договор заключен", section)
        self.assertEqual(section.count("payment-request-group-header"), 2)
        self.assertEqual(section.count("payment-request-toggle-cell"), 2)
        self.assertEqual(section.count("payment-request-toggle-header"), 2)
        self.assertEqual(section.count("payment-request-paid-toggle-cell"), 2)
        self.assertEqual(section.count("payment-request-paid-toggle-header"), 2)
        self.assertNotIn("payment-request-paid-date-cell", section)
        self.assertNotIn("payment-request-stage-inner", section)
        self.assertEqual(section.count(">№<"), 2)
        self.assertIn('data-payment-date="advance"', section)
        self.assertIn('data-payment-date="final"', section)
        self.assertNotIn('data-payment-paid-date="advance"', section)
        self.assertNotIn('data-payment-paid-date="final"', section)
        self.assertIn('data-payment-number="advance"', section)
        self.assertIn('data-payment-number="final"', section)
        self.assertIn('data-payment-sender', section)
        self.assertIn('id="payment-request-send-btn"', section)
        self.assertIn('id="payment-request-settings-btn"', section)
        self.assertIn('id="payment-request-modal"', section)
        self.assertIn('bi-toggle-off', section)
        self.assertIn('data-payment-toggle="advance"', section)
        self.assertIn('data-payment-toggle="final"', section)
        self.assertIn('data-payment-paid-toggle="advance"', section)
        self.assertIn('data-payment-paid-toggle="final"', section)
        self.assertIn("payment-request-paid-toggle-icon", section)
        self.assertIn('id="payment-request-project-filter-toggle"', section)
        self.assertIn(f'id="payment-request-filter-{self.project.pk}"', section)
        self.assertIn(f'id="payment-request-sel-{contracted.pk}"', section)
        self.assertNotIn(f'id="payment-request-sel-{draft_only.pk}"', section)
        self.assertIn("AGR-42", section)
        self.assertIn("30%", section)
        self.assertIn("70%", section)
        self.assertIn("1\u00a0234\u00a0567,89 RUB", section)
        self.assertIn('data-request-url="/projects/performers/payment-request/"', section)
        self.assertNotIn('data-payment-paid-toggle-url="/projects/performers/payment-paid-toggle/"', section)
        self.assertNotIn("payment-request-paid-toggle-icon is-interactive", section)
        self.assertIn('data-advance-paid="0"', section)
        self.assertIn('data-final-paid="0"', section)

    def test_payment_request_admin_filters_to_contract_rows(self):
        from django.contrib import admin as django_admin

        contract_batch_id = uuid.uuid4()
        contracted = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=contract_batch_id,
            contract_number="AGR-ADMIN",
            contract_signing_note="Договор заключен",
            agreed_amount=Decimal("1000.00"),
        )
        contracted_sibling = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=contract_batch_id,
            contract_number="AGR-ADMIN",
            contract_signing_note="Договор заключен",
            agreed_amount=Decimal("2000.00"),
        )
        draft_only = Performer.objects.create(
            registration=self.project,
            employee=self.first_manager_employee,
            executor=self.first_manager_name,
        )

        admin_instance = django_admin.site._registry[PaymentRequestPerformer]
        queryset = admin_instance.get_queryset(None)
        queryset_ids = set(queryset.values_list("pk", flat=True))

        self.assertIn(contracted.pk, queryset_ids)
        self.assertIn(contracted_sibling.pk, queryset_ids)
        self.assertNotIn(draft_only.pk, queryset_ids)
        self.assertEqual(admin_instance.contract_total_price(contracted), Decimal("3000.00"))

    def test_payment_paid_toggle_updates_performer_flags(self):
        performer = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_signing_note="Договор заключен",
        )

        enable_response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(performer.pk),
                "kind": "advance",
                "paid": "1",
            },
        )
        self.assertEqual(enable_response.status_code, 200)
        enable_payload = enable_response.json()
        self.assertTrue(enable_payload["ok"])
        self.assertTrue(enable_payload["paid_at"])
        performer.refresh_from_db()
        self.assertTrue(performer.advance_payment_paid)
        self.assertIsNotNone(performer.advance_payment_paid_at)
        self.assertFalse(performer.final_payment_paid)
        self.assertIsNone(performer.final_payment_paid_at)

        disable_response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(performer.pk),
                "kind": "advance",
                "paid": "0",
            },
        )
        self.assertEqual(disable_response.status_code, 200)
        disable_payload = disable_response.json()
        self.assertFalse(disable_payload["paid"])
        self.assertEqual(disable_payload["paid_at"], "")
        performer.refresh_from_db()
        self.assertFalse(performer.advance_payment_paid)
        self.assertIsNone(performer.advance_payment_paid_at)

        final_response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(performer.pk),
                "kind": "final",
                "paid": "1",
            },
        )
        self.assertEqual(final_response.status_code, 200)
        final_payload = final_response.json()
        self.assertTrue(final_payload["paid_at"])
        performer.refresh_from_db()
        self.assertTrue(performer.final_payment_paid)
        self.assertIsNotNone(performer.final_payment_paid_at)

    def test_payment_paid_toggle_allows_zero_prepayment_advance(self):
        performer = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            prepayment=Decimal("0"),
        )
        response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(performer.pk),
                "kind": "advance",
                "paid": "1",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        performer.refresh_from_db()
        self.assertTrue(performer.advance_payment_paid)

    def test_payment_paid_toggle_processes_notification_after_all_request_positions_paid(self):
        request_number = 501
        sent_at = timezone.now()
        advance_performer = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            advance_payment_request_sent_at=sent_at,
            advance_payment_request_number=request_number,
        )
        final_performer = Performer.objects.create(
            registration=self.project,
            employee=self.first_manager_employee,
            executor=self.first_manager_name,
            contract_batch_id=uuid.uuid4(),
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            advance_payment_request_sent_at=sent_at - timedelta(days=1),
            advance_payment_request_number=400,
            advance_payment_paid=True,
            final_payment_request_sent_at=sent_at,
            final_payment_request_number=request_number,
        )
        notification = self._create_payment_request_notification(
            performers=[advance_performer, final_performer],
            request_number=request_number,
        )

        response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(advance_performer.pk),
                "kind": "advance",
                "paid": "1",
            },
        )
        self.assertEqual(response.status_code, 200)
        notification.refresh_from_db()
        self.assertFalse(notification.is_processed)

        response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(final_performer.pk),
                "kind": "final",
                "paid": "1",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["notifications_changed"], 1)
        notification.refresh_from_db()
        self.assertTrue(notification.is_processed)
        self.assertTrue(notification.is_read)
        self.assertEqual(Notification.objects.for_user(self.user).pending_attention().count(), 0)

    def test_payment_paid_toggle_reopens_processed_notification_when_payment_disabled(self):
        request_number = 502
        sent_at = timezone.now()
        performer = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            advance_payment_request_sent_at=sent_at,
            advance_payment_request_number=request_number,
            advance_payment_paid=True,
        )
        notification = self._create_payment_request_notification(
            performers=[performer],
            request_number=request_number,
        )
        notification.is_read = True
        notification.is_processed = True
        notification.action_at = sent_at
        notification.action_by = self.user
        notification.save(update_fields=["is_read", "is_processed", "action_at", "action_by"])

        response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(performer.pk),
                "kind": "advance",
                "paid": "0",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["notifications_changed"], 1)
        performer.refresh_from_db()
        notification.refresh_from_db()
        self.assertFalse(performer.advance_payment_paid)
        self.assertFalse(notification.is_processed)
        self.assertTrue(notification.is_read)
        self.assertIsNone(notification.action_at)
        self.assertEqual(Notification.objects.for_user(self.user).pending_attention().count(), 1)

    def test_payment_paid_toggle_updates_contract_batch_and_processes_notification(self):
        request_number = 503
        sent_at = timezone.now()
        contract_batch_id = uuid.uuid4()
        representative = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=contract_batch_id,
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            advance_payment_request_sent_at=sent_at,
            advance_payment_request_number=request_number,
        )
        sibling = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=contract_batch_id,
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            advance_payment_request_sent_at=sent_at,
            advance_payment_request_number=request_number,
        )
        notification = self._create_payment_request_notification(
            performers=[representative, sibling],
            request_number=request_number,
        )

        response = self.client.post(
            reverse("payment_paid_toggle"),
            {
                "performer_id": str(representative.pk),
                "kind": "advance",
                "paid": "1",
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["updated"], 2)
        self.assertEqual(set(payload["row_ids"]), {representative.pk, sibling.pk})
        representative.refresh_from_db()
        sibling.refresh_from_db()
        notification.refresh_from_db()
        self.assertTrue(representative.advance_payment_paid)
        self.assertTrue(sibling.advance_payment_paid)
        self.assertTrue(notification.is_processed)

    def test_payment_request_sender_display_formats_initials_without_space(self):
        from projects_app.services.payment_request import payment_request_sender_display

        user = get_user_model().objects.create_user(
            username="payment-request-sender@example.com",
            first_name="Иван",
            last_name="Петров",
        )
        Employee.objects.create(user=user, patronymic="Петрович")
        self.assertEqual(payment_request_sender_display(user), "Петров И.П.")

    def test_payment_request_sends_notification(self):
        from classifiers_app.models import OKVCurrency
        from django.contrib.auth.models import Group
        from notifications_app.models import Notification
        from policy_app.models import LAWYER_GROUP

        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user = get_user_model().objects.create_user(
            username="payment-request-lawyer@example.com",
            email="payment-request-lawyer@example.com",
            password="secret",
            first_name="Анна",
            last_name="Юрина",
            is_staff=True,
        )
        Employee.objects.create(user=lawyer_user, patronymic="Юрьевна")
        lawyer_user.groups.add(lawyer_group)

        rub = OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
        )
        contracted = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-42",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("30"),
            final_payment=Decimal("70"),
            agreed_amount=Decimal("1000.00"),
            currency=rub,
        )
        sent_at = timezone.localtime().replace(second=0, microsecond=0)
        sent_at_value = sent_at.strftime("%Y-%m-%dT%H:%M")

        response = self.client.post(
            reverse("payment_request"),
            {
                "performer_ids[]": [str(contracted.pk)],
                "delivery_channels[]": ["system"],
                "request_sent_at": sent_at_value,
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])
        request_number = data["request_number"]
        self.assertIsInstance(request_number, int)
        contracted.refresh_from_db()
        self.assertIsNotNone(contracted.advance_payment_request_sent_at)
        self.assertEqual(contracted.advance_payment_request_number, request_number)
        self.assertEqual(contracted.advance_payment_request_sender, data["sender_display"])
        self.assertIsNone(contracted.final_payment_request_sent_at)
        self.assertEqual(
            Notification.objects.filter(
                notification_type=Notification.NotificationType.PROJECT_PAYMENT_REQUEST,
            ).count(),
            1,
        )
        notification = Notification.objects.get(
            notification_type=Notification.NotificationType.PROJECT_PAYMENT_REQUEST,
        )
        self.assertEqual(notification.recipient_id, lawyer_user.pk)
        self.assertIn(f"№{request_number}", notification.title_text)

        response_final = self.client.post(
            reverse("payment_request"),
            {
                "performer_ids[]": [str(contracted.pk)],
                "delivery_channels[]": ["system"],
                "request_sent_at": sent_at_value,
            },
        )
        self.assertEqual(response_final.status_code, 200)
        self.assertTrue(response_final.json()["ok"])
        contracted.refresh_from_db()
        self.assertIsNotNone(contracted.advance_payment_request_sent_at)
        self.assertIsNotNone(contracted.final_payment_request_sent_at)

    def test_payment_request_forbids_expert_direct_post(self):
        expert_user = get_user_model().objects.create_user(
            username="payment-request-expert@example.com",
            password="secret",
            is_staff=True,
        )
        Employee.objects.create(user=expert_user, role=EXPERT_GROUP)
        contracted = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-EXPERT",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("30"),
            final_payment=Decimal("70"),
            agreed_amount=Decimal("1000.00"),
        )
        sent_at_value = timezone.localtime().replace(second=0, microsecond=0).strftime("%Y-%m-%dT%H:%M")
        self.client.force_login(expert_user)

        response = self.client.post(
            reverse("payment_request"),
            {
                "performer_ids[]": [str(contracted.pk)],
                "delivery_channels[]": ["system"],
                "request_sent_at": sent_at_value,
            },
        )

        self.assertEqual(response.status_code, 403)
        contracted.refresh_from_db()
        self.assertIsNone(contracted.advance_payment_request_sent_at)
        self.assertIsNone(contracted.advance_payment_request_number)
        self.assertEqual(contracted.advance_payment_request_sender, "")

    def test_payment_request_participation_batch_expands_only_selected_executor(self):
        from classifiers_app.models import OKVCurrency
        from django.contrib.auth.models import Group
        from policy_app.models import LAWYER_GROUP

        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        lawyer_user = get_user_model().objects.create_user(
            username="payment-request-lawyer-executor-split@example.com",
            email="payment-request-lawyer-executor-split@example.com",
            password="secret",
            first_name="Анна",
            last_name="Юрина",
            is_staff=True,
        )
        Employee.objects.create(user=lawyer_user, patronymic="Юрьевна")
        lawyer_user.groups.add(lawyer_group)
        rub = OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
        )
        participation_batch_id = uuid.uuid4()
        selected = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-SEL",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            agreed_amount=Decimal("1000.00"),
            currency=rub,
        )
        same_executor = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-SAME",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            agreed_amount=Decimal("2000.00"),
            currency=rub,
        )
        other_executor = Performer.objects.create(
            registration=self.project,
            employee=self.first_manager_employee,
            executor=self.first_manager_name,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=participation_batch_id,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-OTHER",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            agreed_amount=Decimal("3000.00"),
            currency=rub,
        )
        sent_at_value = timezone.localtime().replace(second=0, microsecond=0).strftime("%Y-%m-%dT%H:%M")

        response = self.client.post(
            reverse("payment_request"),
            {
                "performer_ids[]": [str(selected.pk)],
                "delivery_channels[]": ["system"],
                "request_sent_at": sent_at_value,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["updated"], 2)
        selected.refresh_from_db()
        same_executor.refresh_from_db()
        other_executor.refresh_from_db()
        self.assertIsNotNone(selected.advance_payment_request_sent_at)
        self.assertIsNotNone(same_executor.advance_payment_request_sent_at)
        self.assertIsNone(other_executor.advance_payment_request_sent_at)
        self.assertIsNone(other_executor.advance_payment_request_number)
        notification = Notification.objects.get(
            notification_type=Notification.NotificationType.PROJECT_PAYMENT_REQUEST,
        )
        payment_request_text = notification.payload["payment_request"]
        self.assertIn("Цена договора: 3 000,00 RUB", payment_request_text)
        self.assertIn("Оплатить: 1 500,00 RUB (50% аванс)", payment_request_text)
        self.assertNotIn("Цена договора: 1 000,00 RUB", payment_request_text)

    def test_payment_request_mixed_stages_in_one_batch(self):
        from classifiers_app.models import OKVCurrency
        from django.contrib.auth.models import Group
        from notifications_app.models import Notification
        from policy_app.models import LAWYER_GROUP

        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        if not get_user_model().objects.filter(groups__name=LAWYER_GROUP, is_staff=True).exists():
            lawyer_user = get_user_model().objects.create_user(
                username="payment-request-lawyer-mixed@example.com",
                email="payment-request-lawyer-mixed@example.com",
                password="secret",
                first_name="Анна",
                last_name="Юрина",
                is_staff=True,
            )
            Employee.objects.create(user=lawyer_user, patronymic="Юрьевна")
            lawyer_user.groups.add(lawyer_group)

        rub = OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
        )
        advance_row = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-ADV",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            agreed_amount=Decimal("100000.00"),
            currency=rub,
        )
        final_row = Performer.objects.create(
            registration=self.project,
            employee=self.first_manager_employee,
            executor=self.first_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-FIN",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("50"),
            final_payment=Decimal("50"),
            agreed_amount=Decimal("50000.00"),
            currency=rub,
            advance_payment_request_sent_at=timezone.now(),
            advance_payment_request_number=99,
        )
        sent_at_value = timezone.localtime().replace(second=0, microsecond=0).strftime("%Y-%m-%dT%H:%M")

        response = self.client.post(
            reverse("payment_request"),
            {
                "performer_ids[]": [str(advance_row.pk), str(final_row.pk)],
                "delivery_channels[]": ["system"],
                "request_sent_at": sent_at_value,
            },
        )

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data["ok"])
        self.assertEqual(len(data["row_updates"]), 2)
        request_number = data["request_number"]
        advance_row.refresh_from_db()
        final_row.refresh_from_db()
        self.assertEqual(advance_row.advance_payment_request_number, request_number)
        self.assertEqual(final_row.final_payment_request_number, request_number)
        self.assertEqual(
            Notification.objects.filter(
                notification_type=Notification.NotificationType.PROJECT_PAYMENT_REQUEST,
            ).count(),
            1,
        )

    def test_payment_request_zero_prepayment_sends_final_only(self):
        from classifiers_app.models import OKVCurrency
        from django.contrib.auth.models import Group
        from policy_app.models import LAWYER_GROUP

        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        if not get_user_model().objects.filter(groups__name=LAWYER_GROUP, is_staff=True).exists():
            lawyer_user = get_user_model().objects.create_user(
                username="payment-request-lawyer-zero@example.com",
                email="payment-request-lawyer-zero@example.com",
                password="secret",
                first_name="Анна",
                last_name="Юрина",
                is_staff=True,
            )
            Employee.objects.create(user=lawyer_user, patronymic="Юрьевна")
            lawyer_user.groups.add(lawyer_group)

        rub = OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
        )
        contracted = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-0",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("0"),
            final_payment=Decimal("100"),
            agreed_amount=Decimal("1000.00"),
            currency=rub,
        )
        sent_at_value = timezone.localtime().replace(second=0, microsecond=0).strftime("%Y-%m-%dT%H:%M")

        response = self.client.post(
            reverse("payment_request"),
            {
                "performer_ids[]": [str(contracted.pk)],
                "delivery_channels[]": ["system"],
                "request_sent_at": sent_at_value,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        contracted.refresh_from_db()
        self.assertIsNone(contracted.advance_payment_request_sent_at)
        self.assertIsNotNone(contracted.final_payment_request_sent_at)

    def test_performers_partial_renders_skipped_advance_toggle_for_zero_prepayment(self):
        from classifiers_app.models import OKVCurrency

        rub = OKVCurrency.objects.create(
            code_numeric="643",
            code_alpha="RUB",
            name="Российский рубль",
        )
        contracted = Performer.objects.create(
            registration=self.project,
            employee=self.project_manager_employee,
            executor=self.project_manager_name,
            contract_batch_id=uuid.uuid4(),
            contract_number="AGR-0",
            contract_signing_note="Договор заключен",
            prepayment=Decimal("0"),
            final_payment=Decimal("100"),
            agreed_amount=Decimal("1000.00"),
            currency=rub,
        )

        response = self.client.get(reverse("performers_partial"))
        content = response.content.decode("utf-8")
        section = content[
            content.index('id="payment-request-section"'):
            content.index("<script>", content.index('id="payment-request-section"'))
        ]
        row_start = section.index(f'data-performer-id="{contracted.pk}"')
        row_end = section.index("</tr>", row_start)
        row_html = section[row_start:row_end]
        self.assertIn('data-prepayment="0"', row_html)
        self.assertEqual(row_html.count('is-skipped'), 2)
        self.assertIn('data-payment-toggle="advance"', row_html)
        self.assertIn('data-payment-paid-toggle="advance"', row_html)
        self.assertIn('bi-toggle-on', row_html)

    def test_direction_director_uses_project_manager_executor_locking(self):
        Employee.objects.create(user=self.user, role=DIRECTION_DIRECTOR_GROUP)
        company = GroupMember.objects.create(
            short_name="IMC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
            position=10,
        )
        direction = OrgUnit.objects.create(
            company=company,
            level=2,
            department_name="Налоговое направление",
            short_name="TAX",
            unit_type="expertise",
        )
        section = self.product.sections.create(
            code="TAX",
            short_name="Tax",
            short_name_ru="Налоги",
            name_en="Tax",
            name_ru="Налоги",
            accounting_type="Раздел",
            expertise_direction=direction,
            position=4,
        )
        Performer.objects.create(
            registration=self.project,
            asset_name="Актив с направлением",
            executor="Эксперт",
            typical_section=section,
        )

        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "performer-locked-icon")


class ProjectProductLinkSyncTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="project-link-staff",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.second_product = Product.objects.create(
            short_name="QAQC_LINK",
            name_en="QAQC Link",
            name_ru="QAQC Link",
            consulting_type="Горный",
            service_category="Контроль",
            service_subtype="Контроль качества",
            position=2,
        )
        self.project = ProjectRegistration.objects.create(
            number=6101,
            type=self.product,
            name="Связанный проект",
            year=2026,
        )
        ProjectRegistrationProduct.objects.create(
            registration=self.project,
            product=self.product,
            rank=1,
        )
        ProjectRegistrationProduct.objects.create(
            registration=self.project,
            product=self.second_product,
            rank=2,
        )
        self.work_item = WorkVolume.objects.create(
            project=self.project,
            asset_name="Актив 1",
            manager="Менеджер",
            position=1,
        )
        self.legal_entity = LegalEntity.objects.get(work_item=self.work_item)

    def test_product_short_name_change_updates_joined_work_and_legal_rows(self):
        self.product.short_name = "CONS"
        self.product.save()

        self.work_item.refresh_from_db()
        self.legal_entity.refresh_from_db()

        self.project.refresh_from_db()

        expected_type = f"CONS-{self.second_product.short_name}"
        self.assertEqual(self.project.type_short_display, expected_type)
        self.assertEqual(self.work_item.type, expected_type)
        self.assertEqual(self.legal_entity.work_type, expected_type)

        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, expected_type)
        self.assertNotContains(response, '<td class="text-nowrap">DD</td>', html=False)

    def test_projects_partial_splits_scope_tables_to_separate_subsection(self):
        project = ProjectRegistration.objects.create(
            number=6102,
            type=self.product,
            name="Проект в работе",
            status="В работе",
            year=2026,
        )
        project.gantt_data = {
            "data": [
                {
                    "id": "task-1",
                    "parent": 0,
                    "text": "Подготовка данных",
                    "start_date": "2026-05-14",
                    "end_date": "2026-05-21",
                    "specialty": "Оценка",
                    "executor": "Иванов И.И.",
                    "deadline": "2026-05-25",
                    "constraint_type": "fnlt",
                    "constraint_date": "2026-05-25",
                    "duration": 5,
                    "progress": 50,
                }
            ],
            "links": [],
            "meta": {},
        }
        project.save(update_fields=["gantt_data"])
        new_project = ProjectRegistration.objects.create(
            number=6104,
            type=self.product,
            name="Новый проект без графика",
            status="Не начат",
            year=2026,
        )

        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="projects-content-launch" class="projects-section-content"', html=False)
        self.assertContains(response, 'id="projects-content-scope" class="projects-section-content d-none"', html=False)
        self.assertContains(response, "Реестр проектов")
        self.assertContains(response, "Иконки статуса")
        self.assertContains(response, 'data-registration-launch', html=False)
        self.assertContains(response, 'data-registration-status', html=False)
        self.assertContains(response, 'data-registration-manager', html=False)
        self.assertContains(response, 'data-registration-deadline', html=False)
        self.assertContains(response, 'data-registration-date', html=False)
        self.assertContains(response, "Дата оценки")
        self.assertContains(response, 'id="registration-status-editor"', html=False)
        self.assertContains(response, 'id="registration-manager-editor"', html=False)
        self.assertContains(response, 'id="registration-deadline-editor"', html=False)
        self.assertContains(response, 'bi bi-play-circle', html=False)
        self.assertContains(response, 'bi bi-circle', html=False)
        self.assertContains(response, 'reg-launch-status--work', html=False)
        self.assertContains(response, "Сроки проекта по договору")
        self.assertContains(response, "Срок предв. отчёта, мес.")
        self.assertContains(response, "Дата предв. отчёта")
        self.assertContains(response, "Срок итог. отчёта, нед.")
        self.assertContains(response, "Дата итог. отчёта")
        self.assertNotContains(response, "Этап 3, нед.")
        self.assertContains(response, "График проекта")
        self.assertContains(response, "Подготовка данных")
        self.assertContains(response, "14.05.2026")
        self.assertContains(response, "50%")
        self.assertContains(response, 'id="project-schedule-filter-dropdown"', html=False)
        self.assertContains(response, 'id="project-schedule-filter-none"', html=False)
        self.assertContains(response, f'id="project-schedule-filter-{new_project.pk}"', html=False)
        self.assertContains(response, "Не выбран")
        self.assertContains(response, 'type="radio"', html=False)
        self.assertContains(response, 'name="project-schedule-project-radio"', html=False)
        self.assertContains(response, 'proposal-payment-view-option', html=False)
        self.assertNotContains(response, 'id="project-schedule-filter-all"', html=False)
        self.assertContains(response, 'name="project-schedule-select"', html=False)
        self.assertContains(response, 'id="project-schedule-actions"', html=False)
        self.assertContains(response, 'data-target-name="project-schedule-select"', html=False)
        self.assertContains(response, "Объем услуг: активы")
        self.assertContains(response, "Объем услуг: юрлица")

        content = response.content.decode("utf-8")
        self.assertLess(
            content.index('id="projects-content-launch"'),
            content.index("Реестр проектов"),
        )
        self.assertLess(
            content.index("Сроки проекта по договору"),
            content.index('id="projects-content-scope"'),
        )
        self.assertLess(
            content.index('id="projects-content-scope"'),
            content.index("График проекта"),
        )
        self.assertLess(
            content.index("График проекта"),
            content.index("Объем услуг: активы"),
        )
        self.assertLess(
            content.index("Объем услуг: активы"),
            content.index("Объем услуг: юрлица"),
        )
        registration_header_order = [
            '<th class="nowrap" data-col="name">Название</th>',
            '<th data-col="year">Год</th>',
            '<th data-col="evaluation-date">Дата оценки</th>',
            '<th data-col="deadline">Дедлайн</th>',
            '<th data-col="manager">Руководитель проекта</th>',
            '<th class="reg-launch-cell" data-col="launch"><span class="visually-hidden">Запуск</span></th>',
            '<th data-col="status">Статус</th>',
        ]
        registration_header_positions = [content.index(item) for item in registration_header_order]
        self.assertEqual(registration_header_positions, sorted(registration_header_positions))

    def test_project_schedule_crud_and_move_actions(self):
        project = ProjectRegistration.objects.create(
            number=6103,
            type=self.product,
            name="Проект с графиком",
            status="В работе",
            year=2026,
        )

        create_payload = {
            "task": "Старт",
            "start_date": "2026-05-14",
            "end_date": "2026-05-15",
            "specialty": "",
            "executor": "",
            "deadline": "2026-05-16",
            "constraint_type": "",
            "constraint_date": "",
            "duration": "2",
            "duration_star": "2",
            "predecessors": "",
            "progress": "10",
        }
        response = self.client.post(
            reverse("project_schedule_form_create", args=[project.pk]), create_payload
        )
        self.assertEqual(response.status_code, 200)
        project.refresh_from_db()
        tasks = project.gantt_data.get("data", [])
        self.assertEqual(len(tasks), 1)
        first_task = tasks[0]
        self.assertEqual(first_task["text"], "Старт")
        self.assertEqual(int(first_task.get("progress") or 0), 10)
        first_task_id = str(first_task["id"])

        # Append a second task directly via the same service the view uses, then
        # move the first one down.
        from projects_app.services import gantt_tasks as _gantt_tasks
        second_task_id = _gantt_tasks.add_task(project, {"text": "Финиш"})

        response = self.client.post(
            reverse("project_schedule_move_down", args=[project.pk, first_task_id])
        )
        self.assertEqual(response.status_code, 200)
        project.refresh_from_db()
        ordered_ids = [str(t["id"]) for t in project.gantt_data["data"]]
        self.assertEqual(ordered_ids, [second_task_id, first_task_id])

        edit_payload = {**create_payload, "task": "Обновлённый старт", "progress": "75"}
        response = self.client.post(
            reverse("project_schedule_form_edit", args=[project.pk, first_task_id]),
            edit_payload,
        )
        self.assertEqual(response.status_code, 200)
        project.refresh_from_db()
        edited = next(
            t for t in project.gantt_data["data"] if str(t["id"]) == first_task_id
        )
        self.assertEqual(edited["text"], "Обновлённый старт")
        self.assertEqual(int(edited.get("progress") or 0), 75)

        response = self.client.post(
            reverse("project_schedule_delete", args=[project.pk, first_task_id])
        )
        self.assertEqual(response.status_code, 200)
        project.refresh_from_db()
        remaining_ids = [str(t["id"]) for t in project.gantt_data["data"]]
        self.assertNotIn(first_task_id, remaining_ids)

    def test_project_schedule_edit_accepts_posted_legacy_executor_choice(self):
        project = ProjectRegistration.objects.create(
            number=6104,
            type=self.product,
            name="Проект с устаревшим исполнителем",
            status="В работе",
            year=2026,
        )
        project.gantt_data = {
            "data": [
                {
                    "id": "legacy-task",
                    "text": "Старая задача",
                    "type": "task",
                    "start_date": "2026-05-14",
                    "end_date": "2026-05-15",
                    "specialty": "Архивная специальность",
                    "executor": "Архивный Исполнитель",
                    "progress": 10,
                }
            ],
            "links": [],
            "meta": {},
        }
        project.save(update_fields=["gantt_data"])

        response = self.client.post(
            reverse("project_schedule_form_edit", args=[project.pk, "legacy-task"]),
            {
                "type": "task",
                "service_section_name": "",
                "task": "Обновлённая задача",
                "start_date": "2026-05-14",
                "end_date": "2026-05-15",
                "specialty": "Архивная специальность",
                "executor": "Архивный Исполнитель",
                "deadline": "",
                "constraint_type": "",
                "constraint_date": "",
                "duration": "2",
                "duration_star": "2",
                "predecessors": "",
                "progress": "75",
            },
        )

        self.assertEqual(response.status_code, 200)
        project.refresh_from_db()
        edited = next(t for t in project.gantt_data["data"] if t["id"] == "legacy-task")
        self.assertEqual(edited["text"], "Обновлённая задача")
        self.assertEqual(edited["executor"], "Архивный Исполнитель")
        self.assertEqual(int(edited["progress"]), 75)


class ExpertProjectVisibilityTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="projects-expert",
            password="secret",
            is_staff=True,
            first_name="Иван",
            last_name="Эксперт",
        )
        self.employee = Employee.objects.create(
            user=self.user,
            patronymic="Иванович",
            role=EXPERT_GROUP,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.confirmed_project = ProjectRegistration.objects.create(
            number=7001,
            type=self.product,
            name="Подтвержденный проект",
            year=2026,
        )
        self.declined_project = ProjectRegistration.objects.create(
            number=7002,
            type=self.product,
            name="Отклоненный проект",
            year=2026,
        )
        self.requested_project = ProjectRegistration.objects.create(
            number=7003,
            type=self.product,
            name="Проект с запросом участия",
            year=2026,
        )
        WorkVolume.objects.create(
            project=self.confirmed_project,
            name="Подтвержденный актив",
            asset_name="Подтвержденный актив",
            manager="Менеджер",
        )
        WorkVolume.objects.create(
            project=self.declined_project,
            name="Отклоненный актив",
            asset_name="Отклоненный актив",
            manager="Менеджер",
        )
        WorkVolume.objects.create(
            project=self.requested_project,
            name="Актив с запросом участия",
            asset_name="Актив с запросом участия",
            manager="Менеджер",
        )
        Performer.objects.create(
            registration=self.confirmed_project,
            asset_name="Подтвержденный актив",
            executor=Performer.employee_full_name(self.employee),
            employee=self.employee,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
        )
        Performer.objects.create(
            registration=self.declined_project,
            asset_name="Отклоненный актив",
            executor=Performer.employee_full_name(self.employee),
            employee=self.employee,
            participation_response=Performer.ParticipationResponse.DECLINED,
        )
        Performer.objects.create(
            registration=self.requested_project,
            asset_name="Актив с запросом участия",
            executor=Performer.employee_full_name(self.employee),
            employee=self.employee,
            participation_request_sent_at=timezone.now(),
        )

    def test_projects_partial_for_expert_shows_confirmed_and_requested_participation_projects(self):
        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, self.confirmed_project.name)
        self.assertContains(response, self.confirmed_project.short_uid)
        self.assertContains(response, self.requested_project.name)
        self.assertContains(response, self.requested_project.short_uid)
        self.assertNotContains(response, self.declined_project.name)
        self.assertNotContains(response, self.declined_project.short_uid)
        self.assertNotContains(response, "Отклоненный актив")

    def test_performers_partial_for_expert_shows_confirmed_and_requested_participation_projects(self):
        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, self.confirmed_project.name)
        self.assertContains(response, self.confirmed_project.short_uid)
        self.assertContains(response, self.requested_project.name)
        self.assertContains(response, self.requested_project.short_uid)
        self.assertNotContains(response, self.declined_project.name)
        self.assertNotContains(response, self.declined_project.short_uid)
        self.assertNotContains(response, "Отклоненный актив")

    def test_projects_partial_for_expert_is_readonly(self):
        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "products-actions-row")
        content = response.content.decode("utf-8")
        for checkbox_id in (
            "registrations-master",
            f"reg-sel-{self.confirmed_project.pk}",
            "contract-conditions-master",
            f"contract-sel-{self.confirmed_project.pk}",
            "work-master",
            "legal-entities-master",
        ):
            checkbox_start = content.index(f'id="{checkbox_id}"')
            checkbox_html = content[checkbox_start:content.index("aria-label", checkbox_start)]
            self.assertIn("disabled", checkbox_html)
        work_item = WorkVolume.objects.get(project=self.confirmed_project)
        checkbox_start = content.index(f'id="work-sel-{work_item.pk}"')
        checkbox_html = content[checkbox_start:content.index("aria-label", checkbox_start)]
        self.assertIn("disabled", checkbox_html)

    def test_performers_partial_for_expert_is_readonly(self):
        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "products-actions-row")
        self.assertContains(response, "performer-locked-icon")
        self.assertNotContains(response, "performer-quick-edit")
        self.assertContains(response, 'id="participation-confirmation-section" class="mt-5" data-expert-readonly="1"', html=False)
        self.assertNotContains(response, "Скоррект. затраты")
        self.assertNotContains(response, "Расч. затраты")
        self.assertNotContains(response, "Согласовано")
        self.assertNotContains(response, 'id="perf-total-actual"', html=False)
        content = response.content.decode("utf-8")
        performer = Performer.objects.get(registration=self.confirmed_project)
        for checkbox_id in (
            "performers-master",
            f"perf-sel-{performer.pk}",
            "participation-master",
            f"participation-sel-{performer.pk}",
            "info-request-master",
            f"info-request-sel-{performer.pk}",
        ):
            checkbox_start = content.index(f'id="{checkbox_id}"')
            checkbox_html = content[checkbox_start:content.index("aria-label", checkbox_start)]
            self.assertIn("disabled", checkbox_html)

    def test_group_expert_with_empty_employee_role_is_still_filtered(self):
        self.employee.role = ""
        self.employee.save(update_fields=["role"])
        expert_group, _ = Group.objects.get_or_create(name=EXPERT_GROUP)
        self.user.groups.add(expert_group)

        projects_response = self.client.get(reverse("projects_partial"))
        performers_response = self.client.get(reverse("performers_partial"))

        self.assertEqual(projects_response.status_code, 200)
        self.assertContains(projects_response, self.confirmed_project.name)
        self.assertNotContains(projects_response, self.declined_project.name)
        self.assertNotContains(projects_response, "Отклоненный актив")
        self.assertEqual(performers_response.status_code, 200)
        self.assertContains(performers_response, self.confirmed_project.name)
        self.assertNotContains(performers_response, self.declined_project.name)
        self.assertNotContains(performers_response, "Отклоненный актив")


class CloudStorageProjectRoutingTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="project-staff",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.project = ProjectRegistration.objects.create(
            number=5001,
            type=self.product,
            name="Проект маршрутизации",
            year=2026,
        )

    def test_create_workspace_returns_controlled_error_when_nextcloud_selected(self):
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.save()

        response = self.client.post(
            reverse("create_workspace"),
            {"project_id": self.project.pk},
        )

        self.assertEqual(response.status_code, 400)
        payload = response.json()
        self.assertFalse(payload["ok"])
        self.assertIn("Nextcloud", payload["error"])

    def test_projects_partial_uses_current_primary_cloud_label_in_workspace_modal(self):
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.save()

        response = self.client.get(reverse("projects_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="reg-create-workspace-storage-label">Nextcloud', html=False)

    def test_workspace_folders_list_returns_current_primary_cloud_label(self):
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.save()

        response = self.client.get(reverse("workspace_folders_list"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["storage_label"], "Nextcloud")

    def test_performers_partial_no_longer_renders_contract_conclusion_table(self):
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.save()

        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "Составление проекта договора")
        self.assertNotContains(response, "Отправка проекта договора")
        self.assertNotContains(response, "create-contract-progress-modal")

    def test_performers_partial_uses_current_primary_cloud_label_in_source_data_modal(self):
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.save()

        response = self.client.get(reverse("performers_partial"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "На Nextcloud в целевой директории будет создана структура папок для выбранного проекта в соответствии с согласованным информационным запросом.",
            html=False,
        )


@override_settings(
    NEXTCLOUD_PROVISIONING_BASE_URL="https://cloud.example.com",
    NEXTCLOUD_PROVISIONING_USERNAME="cloud-admin",
    NEXTCLOUD_PROVISIONING_TOKEN="token",
    NEXTCLOUD_OIDC_PROVIDER_ID=1,
)
class NextcloudSourceDataWorkspaceFlowTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="source-data-staff",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.nextcloud_root_path = "/Corporate Root"
        settings_obj.save()
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.project = ProjectRegistration.objects.create(
            number=5003,
            type=self.product,
            name="Исходные данные",
            year=2026,
        )
        self.section = self.product.sections.create(
            code="FIN",
            short_name="Finance",
            short_name_ru="Финансы",
            name_en="Finance",
            name_ru="Финансы",
            accounting_type="Раздел",
            position=1,
        )
        self.item = ChecklistItem.objects.create(
            project=self.project,
            section=self.section,
            code="REQ",
            number=1,
            short_name="ОСВ",
            name="Оборотно-сальдовая ведомость",
            position=1,
        )
        Performer.objects.create(
            registration=self.project,
            asset_name="ООО Ромашка",
            typical_section=self.section,
            info_approval_status=Performer.InfoApprovalStatus.APPROVED,
        )
        SourceDataTargetFolder.objects.create(user=self.user, folder_name="05 Исходные данные/01 Запросы")

    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder")
    def test_create_source_data_workspace_uses_nextcloud_and_stores_public_links(
        self,
        mocked_ensure_folder,
        mocked_public_share,
    ):
        expected_project_folder = f"{self.project.short_uid} DD Исходные данные"
        base_path = f"/Corporate Root/03 Проекты/2026/{expected_project_folder}/05 Исходные данные/01 Запросы"
        section_path = f"{base_path}/01 FIN Финансы"
        item_path = f"{section_path}/REQ-01 ОСВ"
        mocked_ensure_folder.side_effect = [section_path, item_path]
        mocked_public_share.return_value = "https://cloud.example.com/s/item-folder"

        response = self.client.post(
            reverse("create_source_data_workspace"),
            {"project_id": self.project.pk},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertTrue(chunks[-1]["ok"])
        self.assertIn("Nextcloud", chunks[-1]["message"])

        mocked_ensure_folder.assert_has_calls(
            [
                call("cloud-admin", section_path),
                call("cloud-admin", item_path),
            ]
        )
        mocked_public_share.assert_has_calls(
            [
                call("cloud-admin", item_path, _quick=True),
            ]
        )

        section_folder = SourceDataSectionFolder.objects.get(project=self.project, section=self.section, asset_name="ООО Ромашка")
        item_folder = SourceDataItemFolder.objects.get(
            project=self.project,
            checklist_item=self.item,
            asset_name="ООО Ромашка",
        )
        workspace = SourceDataWorkspace.objects.get(project=self.project)
        self.assertEqual(section_folder.disk_path, section_path)
        self.assertEqual(section_folder.public_url, "")
        self.assertEqual(item_folder.disk_path, item_path)
        self.assertEqual(item_folder.public_url, "https://cloud.example.com/s/item-folder")
        self.assertEqual(workspace.disk_path, base_path)
        self.assertEqual(workspace.created_by, self.user)


@override_settings(
    NEXTCLOUD_PROVISIONING_BASE_URL="https://cloud.example.com",
    NEXTCLOUD_PROVISIONING_USERNAME="cloud-admin",
    NEXTCLOUD_PROVISIONING_TOKEN="token",
    NEXTCLOUD_OIDC_PROVIDER_ID=1,
)
class NextcloudContractProjectFlowTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="contracts-staff",
            password="secret",
            is_staff=True,
        )
        self.client.force_login(self.user)
        settings_obj = CloudStorageSettings.get_solo()
        settings_obj.primary_storage = CloudStorageSettings.PrimaryStorage.NEXTCLOUD
        settings_obj.nextcloud_root_path = "/Corporate Root"
        settings_obj.save()
        self.product = Product.objects.create(
            short_name="DD",
            name_en="Due Diligence",
            name_ru="ДД",
            consulting_type="Горный",
            service_category="Аудит",
            service_subtype="Аудит соответствия стандартам",
        )
        self.project = ProjectRegistration.objects.create(
            number=5002,
            type=self.product,
            name="Контрактный проект",
            year=2026,
        )
        self.recipient_user = get_user_model().objects.create_user(
            username="performer@example.com",
            email="performer@example.com",
            password="secret",
            first_name="Иван",
            last_name="Иванов",
            is_staff=True,
        )
        self.employee = Employee.objects.create(
            user=self.recipient_user,
            patronymic="Иванович",
            employment="Фрилансер",
        )
        self.lawyer_user = get_user_model().objects.create_user(
            username="lawyer@example.com",
            email="lawyer@example.com",
            password="secret",
            is_staff=True,
        )
        lawyer_group, _ = Group.objects.get_or_create(name=LAWYER_GROUP)
        self.lawyer_user.groups.add(lawyer_group)
        self.performer = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor="Иванов Иван Иванович",
        )
        self.template = ContractTemplate.objects.create(
            product=self.product,
            contract_type="gph",
            party="individual",
            country_name="",
            sample_name="Базовый шаблон",
            version="1",
            file=SimpleUploadedFile(
                "contract.docx",
                b"fake-docx-content",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )
        self.executor_link = NextcloudUserLink.objects.create(
            user=self.recipient_user,
            nextcloud_user_id=f"ncstaff-{self.recipient_user.pk}",
            nextcloud_username=f"ncstaff-{self.recipient_user.pk}",
            nextcloud_email=self.recipient_user.email,
        )
        self.lawyer_link = NextcloudUserLink.objects.create(
            user=self.lawyer_user,
            nextcloud_user_id=f"ncstaff-{self.lawyer_user.pk}",
            nextcloud_username=f"ncstaff-{self.lawyer_user.pk}",
            nextcloud_email=self.lawyer_user.email,
        )

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/public-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/50020RU DD Контрактный проект/5002010RU DD Контрактный проект/02 Исполнители/5002010RU 001 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_uses_nextcloud_and_stores_public_link(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        contract_project = ContractProjectRegistration.objects.create(
            number=self.project.number,
            sub_number=1,
            type=self.product,
            name=self.project.name,
            year=2026,
        )
        self.project.contract_project_registration = contract_project
        self.project.save(update_fields=["contract_project_registration"])
        self.project.refresh_from_db()
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )
        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)

        self.performer.refresh_from_db()
        expected_month = timezone.localtime(timezone.now()).strftime("%m-%y")
        expected_contract_number = f"IMCM/5002/1-ИИ/{expected_month}"
        expected_project_folder = "50020RU DD Контрактный проект"
        expected_contract_folder = "5002010RU DD Контрактный проект"
        expected_base_path = f"/Corporate Root/02 Договоры/2026/{expected_project_folder}/{expected_contract_folder}/02 Исполнители"
        expected_folder_path = f"{expected_base_path}/5002010RU 001 Иванов ИИ"
        expected_docx_name = "Договор 5002010RU_Иванов ИИ.docx"
        expected_upload_path = f"{expected_folder_path}/{expected_docx_name}"

        mocked_list_resources.assert_any_call("cloud-admin", expected_base_path, limit=1000)
        mocked_ensure_folder.assert_has_calls(
            [
                call("cloud-admin", "/Corporate Root"),
                call("cloud-admin", "/Corporate Root/02 Договоры"),
                call("cloud-admin", "/Corporate Root/02 Договоры/2026"),
                call("cloud-admin", f"/Corporate Root/02 Договоры/2026/{expected_project_folder}"),
                call("cloud-admin", f"/Corporate Root/02 Договоры/2026/{expected_project_folder}/{expected_contract_folder}"),
                call("cloud-admin", expected_base_path),
                call("cloud-admin", expected_folder_path),
            ]
        )
        mocked_upload_file.assert_called_once()
        self.assertEqual(mocked_upload_file.call_args.args[0], "cloud-admin")
        self.assertEqual(mocked_upload_file.call_args.args[1], expected_upload_path)
        mocked_public_share.assert_has_calls(
            [
                call("cloud-admin", expected_folder_path),
                call("cloud-admin", expected_upload_path),
            ]
        )
        mocked_ensure_user_share.assert_has_calls(
            [
                call("cloud-admin", expected_folder_path, self.executor_link.nextcloud_user_id, permissions=1),
                call("cloud-admin", expected_folder_path, self.lawyer_link.nextcloud_user_id, permissions=15),
            ],
            any_order=True,
        )
        self.assertTrue(self.performer.contract_project_created)
        self.assertEqual(self.performer.contract_project_link, "https://cloud.example.com/s/public-doc")
        self.assertEqual(self.performer.contract_project_folder_link, "https://cloud.example.com/s/public-doc")
        self.assertEqual(self.performer.contract_project_disk_folder, expected_folder_path)
        self.assertEqual(self.performer.contract_file, expected_docx_name)
        self.assertEqual(self.performer.contract_number, expected_contract_number)
        self.assertEqual(mocked_resolve_variables.call_args.args[0].contract_number, expected_contract_number)
        self.assertIsNotNone(self.performer.contract_project_created_at)
        self.assertIsNotNone(self.performer.contract_date)

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/public-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/50020RU DD Контрактный проект/5002000RU DD Контрактный проект/02 Исполнители/5002000RU 001 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_uses_saved_contract_number_and_date(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        pending_batch_id = uuid.uuid4()
        self.performer.contract_batch_id = pending_batch_id
        self.performer.contract_number = "CUSTOM-42"
        self.performer.contract_date = date(2026, 4, 15)
        self.performer.save(update_fields=["contract_batch_id", "contract_number", "contract_date"])
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        mocked_resolve_variables.assert_called()
        resolved_performer = mocked_resolve_variables.call_args.args[0]
        self.assertEqual(resolved_performer.contract_number, "CUSTOM-42")
        self.assertEqual(resolved_performer.contract_date, date(2026, 4, 15))
        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_batch_id, pending_batch_id)
        self.assertFalse(self.performer.contract_is_addendum)
        self.assertEqual(self.performer.contract_number, "CUSTOM-42")
        self.assertEqual(self.performer.contract_date, date(2026, 4, 15))

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/public-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/50020RU DD-RFR Контрактный проект/5002010RU DD-RFR Контрактный проект/02 Исполнители/5002010RU 001 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_expands_participation_batch(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        batch_id = uuid.uuid4()
        self.employee.employment = FREELANCER_LABEL
        self.employee.save(update_fields=["employment"])
        product_b = Product.objects.create(
            short_name="RFR",
            name_en="Review",
            name_ru="Ревью",
        )
        contract_project = ContractProjectRegistration.objects.create(
            number=self.project.number,
            sub_number=1,
            type=self.product,
            name=self.project.name,
            year=2026,
        )
        self.project.contract_project_registration = contract_project
        self.project.save(update_fields=["contract_project_registration"])
        second_project = ProjectRegistration.objects.create(
            number=self.project.number,
            type=product_b,
            contract_project_registration=contract_project,
            name="Контрактный проект",
            year=2026,
        )
        second_performer = Performer.objects.create(
            registration=second_project,
            employee=self.employee,
            executor=self.performer.executor,
            participation_response=Performer.ParticipationResponse.CONFIRMED,
            participation_batch_id=batch_id,
        )
        self.performer.participation_response = Performer.ParticipationResponse.CONFIRMED
        self.performer.participation_batch_id = batch_id
        self.performer.save(update_fields=["participation_response", "participation_batch_id"])
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )
        expected_docx_name = "Договор 5002010RU_Иванов ИИ.docx"
        expected_folder_name = "5002010RU 001 Иванов ИИ"
        docx_folder_calls = 0

        def list_resources_side_effect(_owner_user_id, path, *, limit=100):
            nonlocal docx_folder_calls
            if path.endswith(f"/{expected_folder_name}"):
                docx_folder_calls += 1
                if docx_folder_calls == 1:
                    return []
                return [
                    {
                        "name": expected_docx_name,
                        "path": f"{path}/{expected_docx_name}",
                        "type": "file",
                        "file_id": "docx-file-id",
                    }
                ]
            return []

        mocked_list_resources.side_effect = list_resources_side_effect

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        self.performer.refresh_from_db()
        second_performer.refresh_from_db()
        self.project.refresh_from_db()
        self.assertIsNotNone(self.performer.contract_batch_id)
        self.assertEqual(self.performer.contract_batch_id, second_performer.contract_batch_id)
        self.assertEqual(self.performer.contract_file, expected_docx_name)
        self.assertEqual(second_performer.contract_file, expected_docx_name)
        self.assertEqual(self.performer.contract_project_file_id, "docx-file-id")
        self.assertEqual(second_performer.contract_project_file_id, "docx-file-id")
        self.assertEqual(docx_folder_calls, 2)
        upload_path = mocked_upload_file.call_args.args[1]
        self.assertIn("/50020RU DD-RFR Контрактный проект/5002010RU DD-RFR Контрактный проект/02 Исполнители/", upload_path)
        self.assertIn(f"/{expected_folder_name}/", upload_path)
        self.assertIn(expected_docx_name, upload_path)
        self.assertTrue(self.performer.contract_project_created)
        self.assertTrue(second_performer.contract_project_created)

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/reused-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/50020RU DD Контрактный проект/5002010RU DD Контрактный проект/02 Исполнители/5002010RU 001 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_reuses_existing_executor_folder_for_regeneration(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        contract_project = ContractProjectRegistration.objects.create(
            number=self.project.number,
            sub_number=1,
            type=self.product,
            name=self.project.name,
            year=2026,
        )
        self.project.contract_project_registration = contract_project
        self.project.save(update_fields=["contract_project_registration"])
        existing_batch_id = uuid.uuid4()
        existing_folder_path = "/Corporate Root/02 Договоры/2026/50020RU DD Контрактный проект/5002010RU DD Контрактный проект/02 Исполнители/5002010RU 001 Иванов ИИ"
        self.performer.contract_batch_id = existing_batch_id
        self.performer.contract_project_created = True
        self.performer.contract_project_disk_folder = existing_folder_path
        self.performer.contract_project_folder_link = "https://cloud.example.com/s/existing-folder"
        self.performer.contract_project_folder_file_id = "folder-file-id"
        self.performer.contract_pdf_file = "old-contract.pdf"
        self.performer.contract_pdf_link = "https://cloud.example.com/s/old-pdf"
        self.performer.contract_pdf_file_id = "old-pdf-id"
        self.performer.contract_signed_pdf_file = "old-signed-contract.pdf"
        self.performer.contract_signed_pdf_link = "https://cloud.example.com/s/old-signed-pdf"
        self.performer.contract_signed_pdf_file_id = "old-signed-pdf-id"
        self.performer.save(update_fields=[
            "contract_batch_id",
            "contract_project_created",
            "contract_project_disk_folder",
            "contract_project_folder_link",
            "contract_project_folder_file_id",
            "contract_pdf_file",
            "contract_pdf_link",
            "contract_pdf_file_id",
            "contract_signed_pdf_file",
            "contract_signed_pdf_link",
            "contract_signed_pdf_file_id",
        ])
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        mocked_upload_file.assert_called_once()
        self.assertEqual(
            mocked_upload_file.call_args.args[1],
            f"{existing_folder_path}/Договор 5002010RU_Иванов ИИ.docx",
        )
        self.assertFalse(
            any("5002010RU 002 Иванов ИИ" in str(call_args) for call_args in mocked_ensure_folder.call_args_list),
            mocked_ensure_folder.call_args_list,
        )
        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_batch_id, existing_batch_id)
        self.assertEqual(self.performer.contract_project_disk_folder, existing_folder_path)
        self.assertEqual(self.performer.contract_pdf_file, "")
        self.assertEqual(self.performer.contract_pdf_link, "")
        self.assertEqual(self.performer.contract_pdf_file_id, "")
        self.assertEqual(self.performer.contract_signed_pdf_file, "")
        self.assertEqual(self.performer.contract_signed_pdf_link, "")
        self.assertEqual(self.performer.contract_signed_pdf_file_id, "")

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/new-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/50020RU DD Контрактный проект/5002010RU DD Контрактный проект/02 Исполнители/5002010RU 001 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[{"name": "5002010RU 001 Иванов ИИ"}])
    def test_create_contract_project_reuses_sent_batch_folder_when_creating_addendum(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        contract_project = ContractProjectRegistration.objects.create(
            number=self.project.number,
            sub_number=1,
            type=self.product,
            name=self.project.name,
            year=2026,
        )
        self.project.contract_project_registration = contract_project
        self.project.save(update_fields=["contract_project_registration"])
        old_batch_id = uuid.uuid4()
        pending_batch_id = uuid.uuid4()
        expected_project_folder = "50020RU DD Контрактный проект"
        expected_contract_folder = "5002010RU DD Контрактный проект"
        old_folder_path = f"/Corporate Root/02 Договоры/2026/{expected_project_folder}/{expected_contract_folder}/02 Исполнители/5002010RU 001 Иванов ИИ"
        self.performer.contract_batch_id = old_batch_id
        self.performer.contract_project_created = True
        self.performer.contract_project_disk_folder = old_folder_path
        self.performer.contract_project_folder_link = "https://cloud.example.com/s/existing-folder"
        self.performer.contract_file = "Договор 5002_Иванов ИИ.docx"
        self.performer.contract_sent_at = timezone.now()
        self.performer.save(update_fields=[
            "contract_batch_id",
            "contract_project_created",
            "contract_project_disk_folder",
            "contract_project_folder_link",
            "contract_file",
            "contract_sent_at",
        ])
        addendum_performer = Performer.objects.create(
            registration=self.project,
            employee=self.employee,
            executor="Иванов Иван Иванович",
            contract_batch_id=pending_batch_id,
        )
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk, addendum_performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        mocked_upload_file.assert_called_once()
        self.assertEqual(
            mocked_upload_file.call_args.args[1],
            f"{old_folder_path}/Договор 5002010RU_Иванов ИИ_ДС1.docx",
        )
        self.performer.refresh_from_db()
        addendum_performer.refresh_from_db()
        self.assertEqual(self.performer.contract_batch_id, old_batch_id)
        self.assertEqual(self.performer.contract_project_disk_folder, old_folder_path)
        self.assertEqual(addendum_performer.contract_batch_id, pending_batch_id)
        self.assertTrue(addendum_performer.contract_is_addendum)
        self.assertEqual(addendum_performer.contract_addendum_number, 1)
        self.assertEqual(addendum_performer.contract_project_disk_folder, old_folder_path)
        self.assertEqual(addendum_performer.contract_project_folder_link, "https://cloud.example.com/s/existing-folder")
        self.assertEqual(addendum_performer.contract_file, "Договор 5002010RU_Иванов ИИ_ДС1.docx")

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/image-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/02 Исполнители/000 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_preserves_seal_and_director_facsimile_placeholders(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        group_member = GroupMember.objects.create(
            short_name="IMCM",
            full_name="IMCM LLC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
        )
        group_member.seal_file.save("seal.png", ContentFile(TEST_CONTRACT_IMAGE_PNG_BYTES), save=True)
        self.project.group_member = group_member
        self.project.save(update_fields=["group_member"])

        director_user = get_user_model().objects.create_user(
            username="director@example.com",
            email="director@example.com",
            password="secret",
            first_name="Дина",
            last_name="Директорова",
            is_staff=True,
        )
        director_department = OrgUnit.objects.create(
            company=group_member,
            department_name="Руководство",
            level=1,
        )
        director_employee = Employee.objects.create(
            user=director_user,
            patronymic="Дмитриевна",
            role=DIRECTOR_GROUP,
            department=director_department,
        )
        director_person = PersonRecord.objects.create(
            last_name="Директорова",
            first_name="Дина",
            middle_name="Дмитриевна",
            position=1,
        )
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            alpha2="RU",
            alpha3="RUS",
        )
        citizenship = CitizenshipRecord.objects.create(
            person=director_person,
            country=country,
            position=1,
        )
        director_profile = ExpertProfile.objects.create(employee=director_employee, position=1)
        director_details = ExpertContractDetails.objects.create(
            expert_profile=director_profile,
            citizenship_record=citizenship,
        )
        director_details.facsimile_file.save(
            "facsimile.png",
            ContentFile(TEST_CONTRACT_IMAGE_PNG_BYTES),
            save=True,
        )

        template_doc = Document()
        template_doc.add_paragraph("Печать [[seal]]")
        template_doc.add_paragraph("Подпись [[facsimile_imcm]]")
        template_buffer = BytesIO()
        template_doc.save(template_buffer)
        self.template.file.save("contract-images.docx", ContentFile(template_buffer.getvalue()), save=True)

        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        mocked_upload_file.assert_called_once()
        generated_doc = Document(BytesIO(mocked_upload_file.call_args.args[2]))
        self.assertEqual(generated_doc.paragraphs[0].text, "Печать [[seal]]")
        self.assertEqual(generated_doc.paragraphs[1].text, "Подпись [[facsimile_imcm]]")
        combined_xml = "\n".join(paragraph._p.xml for paragraph in generated_doc.paragraphs)
        self.assertNotIn("<wp:anchor", combined_xml)
        self.assertIn("[[seal]]", combined_xml)
        self.assertIn("[[facsimile_imcm]]", combined_xml)

    def test_contract_docx_source_inserts_images_and_clears_highlighting_for_pdf(self):
        from projects_app.views import _build_contract_docx_source_token

        group_member = GroupMember.objects.create(
            short_name="IMCM",
            full_name="IMCM LLC",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
        )
        group_member.seal_file.save("seal.png", ContentFile(TEST_CONTRACT_IMAGE_PNG_BYTES), save=True)
        self.project.group_member = group_member
        self.project.save(update_fields=["group_member"])

        director_user = get_user_model().objects.create_user(
            username="pdf-director@example.com",
            email="pdf-director@example.com",
            password="secret",
            first_name="Дина",
            last_name="Директорова",
            is_staff=True,
        )
        director_department = OrgUnit.objects.create(
            company=group_member,
            department_name="Руководство",
            level=1,
        )
        director_employee = Employee.objects.create(
            user=director_user,
            patronymic="Дмитриевна",
            role=DIRECTOR_GROUP,
            department=director_department,
        )
        director_person = PersonRecord.objects.create(
            last_name="Директорова",
            first_name="Дина",
            middle_name="Дмитриевна",
            position=1,
        )
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            alpha2="RU",
            alpha3="RUS",
        )
        citizenship = CitizenshipRecord.objects.create(
            person=director_person,
            country=country,
            position=1,
        )
        director_profile = ExpertProfile.objects.create(employee=director_employee, position=1)
        director_details = ExpertContractDetails.objects.create(
            expert_profile=director_profile,
            citizenship_record=citizenship,
        )
        director_details.facsimile_file.save(
            "facsimile.png",
            ContentFile(TEST_CONTRACT_IMAGE_PNG_BYTES),
            save=True,
        )

        source_doc = Document()
        seal_run = source_doc.add_paragraph().add_run("Печать [[seal]]")
        seal_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        facsimile_run = source_doc.add_paragraph().add_run("Подпись [[facsimile_imcm]]")
        facsimile_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        performer_facsimile_run = source_doc.add_paragraph().add_run("Исполнитель [[facsimile_prfrm]]")
        performer_facsimile_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        source_buffer = BytesIO()
        source_doc.save(source_buffer)

        self.performer.contract_file = "Договор 5002_Иванов ИИ.docx"
        self.performer.contract_project_disk_folder = (
            "/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/"
            "02 Исполнители/000 Иванов ИИ"
        )
        self.performer.save(update_fields=["contract_file", "contract_project_disk_folder"])
        token = _build_contract_docx_source_token(self.performer)
        self.client.logout()

        with patch(
            "projects_app.views.cloud_download_file",
            return_value=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                source_buffer.getvalue(),
            ),
        ):
            response = self.client.get(
                reverse("contract_onlyoffice_docx_source", args=[self.performer.pk]),
                {"token": token},
            )

        self.assertEqual(response.status_code, 200)
        signed_doc = Document(BytesIO(response.content))
        self.assertEqual(signed_doc.paragraphs[0].text, "Печать ")
        self.assertEqual(signed_doc.paragraphs[1].text, "Подпись ")
        self.assertEqual(signed_doc.paragraphs[2].text, "Исполнитель ")
        combined_xml = "\n".join(paragraph._p.xml for paragraph in signed_doc.paragraphs)
        self.assertEqual(combined_xml.count("<wp:anchor"), 2)
        self.assertEqual(combined_xml.count('behindDoc="1"'), 2)
        self.assertIn('relativeHeight="0"', signed_doc.paragraphs[0]._p.xml)
        self.assertIn('wp:positionH relativeFrom="column"', signed_doc.paragraphs[0]._p.xml)
        self.assertIn("<wp:align>center</wp:align>", signed_doc.paragraphs[0]._p.xml)
        self.assertIn('relativeHeight="1"', signed_doc.paragraphs[1]._p.xml)
        self.assertNotIn("[[seal]]", combined_xml)
        self.assertNotIn("[[facsimile_imcm]]", combined_xml)
        self.assertNotIn("[[facsimile_prfrm]]", combined_xml)
        self.assertNotIn("w:highlight", combined_xml)

    def test_contract_docx_source_inserts_performer_facsimile_for_signed_pdf(self):
        from projects_app.views import _build_contract_docx_source_token

        person = PersonRecord.objects.create(
            last_name="Иванов",
            first_name="Иван",
            middle_name="Иванович",
            position=1,
        )
        country = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            alpha2="RU",
            alpha3="RUS",
        )
        citizenship = CitizenshipRecord.objects.create(
            person=person,
            country=country,
            position=1,
        )
        expert_profile = ExpertProfile.objects.create(employee=self.employee, position=1)
        contract_details = ExpertContractDetails.objects.create(
            expert_profile=expert_profile,
            citizenship_record=citizenship,
        )
        contract_details.facsimile_file.save(
            "performer-facsimile.png",
            ContentFile(TEST_CONTRACT_IMAGE_PNG_BYTES),
            save=True,
        )

        source_doc = Document()
        source_doc.add_paragraph("Исполнитель [[facsimile_prfrm]]")
        source_buffer = BytesIO()
        source_doc.save(source_buffer)

        self.performer.contract_file = "Договор 5002_Иванов ИИ.docx"
        self.performer.contract_project_disk_folder = (
            "/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/"
            "02 Исполнители/000 Иванов ИИ"
        )
        self.performer.save(update_fields=["contract_file", "contract_project_disk_folder"])
        token = _build_contract_docx_source_token(self.performer, include_performer_facsimile=True)
        self.client.logout()

        with patch(
            "projects_app.views.cloud_download_file",
            return_value=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                source_buffer.getvalue(),
            ),
        ):
            response = self.client.get(
                reverse("contract_onlyoffice_docx_source", args=[self.performer.pk]),
                {"token": token},
            )

        self.assertEqual(response.status_code, 200)
        signed_doc = Document(BytesIO(response.content))
        self.assertEqual(signed_doc.paragraphs[0].text, "Исполнитель ")
        combined_xml = "\n".join(paragraph._p.xml for paragraph in signed_doc.paragraphs)
        self.assertEqual(combined_xml.count("<wp:anchor"), 1)
        self.assertIn('behindDoc="1"', signed_doc.paragraphs[0]._p.xml)
        self.assertIn('relativeHeight="1"', signed_doc.paragraphs[0]._p.xml)
        self.assertIn('wp:positionH relativeFrom="column"', signed_doc.paragraphs[0]._p.xml)
        self.assertIn("<wp:align>center</wp:align>", signed_doc.paragraphs[0]._p.xml)
        self.assertNotIn("[[facsimile_prfrm]]", combined_xml)

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/profile-country")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/02 Исполнители/000 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_falls_back_to_profile_country_without_contract_details(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        ru = OKSMCountry.objects.create(
            number=643,
            code="643",
            short_name="Россия",
            full_name="Россия",
            alpha2="RU",
            alpha3="RUS",
        )
        profile = ExpertProfile.objects.create(employee=self.employee, position=1)
        ExpertProfile.objects.filter(pk=profile.pk).update(country=ru)
        self.template.country_name = "Россия"
        self.template.country_code = "643"
        self.template.save(update_fields=["country_name", "country_code"])
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        self.assertTrue(mocked_upload_file.called, chunks)
        mocked_upload_file.assert_called_once()
        self.assertEqual(mocked_upload_file.call_args.args[2], b"fake-docx-content")

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/smz-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/02 Исполнители/000 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_matches_smz_template_by_contract_details_country(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        kz_member = GroupMember.objects.create(
            short_name="KZ",
            country_name="Казахстан",
            country_code="398",
            country_alpha2="KZ",
        )
        self.project.group_member = kz_member
        self.project.save()
        person = PersonRecord.objects.create(
            last_name="Смирнова",
            first_name="Элина",
            middle_name="Юрьевна",
            position=1,
        )
        self.employee.person_record = person
        self.employee.save(update_fields=["person_record"])
        ru = OKSMCountry.objects.create(number=643, code="643", short_name="Россия", alpha2="RU", alpha3="RUS")
        kz = OKSMCountry.objects.create(number=398, code="398", short_name="Казахстан", alpha2="KZ", alpha3="KAZ")
        kz_citizenship = CitizenshipRecord.objects.create(person=person, country=kz, position=1)
        ru_citizenship = CitizenshipRecord.objects.create(person=person, country=ru, position=2)
        profile = ExpertProfile.objects.create(employee=self.employee, position=1)
        ExpertContractDetails.objects.create(
            expert_profile=profile,
            citizenship_record=kz_citizenship,
        )
        ru_details = ExpertContractDetails.objects.create(
            expert_profile=profile,
            citizenship_record=ru_citizenship,
            self_employed=date(2026, 1, 1),
        )
        smz_template = ContractTemplate.objects.create(
            group_member=kz_member,
            product=self.product,
            contract_type="smz",
            party="individual",
            country_name="Россия",
            country_code="643",
            sample_name="KZ Шаблон договора ФЗЛ СМЗ RUS_TDD-Общий_v1",
            version="1",
            file=SimpleUploadedFile(
                "smz-contract.docx",
                b"fake-smz-docx-content",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )
        smz_template.group_members.set([kz_member])
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        self.assertNotIn("warnings", chunks[-1])
        mocked_resolve_variables.assert_called()
        self.assertEqual(mocked_resolve_variables.call_args.kwargs["contract_details"], ru_details)
        self.assertEqual(mocked_upload_file.call_args.args[2], b"fake-smz-docx-content")
        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_project_link, "https://cloud.example.com/s/smz-doc")

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/smz-all-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/02 Исполнители/000 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_matches_smz_template_with_all_group_and_all_product(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        person = PersonRecord.objects.create(
            last_name="Смирнова",
            first_name="Элина",
            middle_name="Юрьевна",
            position=1,
        )
        self.employee.person_record = person
        self.employee.save(update_fields=["person_record"])
        ru = OKSMCountry.objects.create(number=643, code="643", short_name="Россия", alpha2="RU", alpha3="RUS")
        ru_citizenship = CitizenshipRecord.objects.create(person=person, country=ru, position=1)
        profile = ExpertProfile.objects.create(employee=self.employee, position=1)
        ru_details = ExpertContractDetails.objects.create(
            expert_profile=profile,
            citizenship_record=ru_citizenship,
            self_employed=date(2026, 1, 1),
        )
        ContractTemplate.objects.create(
            group_member=None,
            product=None,
            contract_type="smz",
            party="individual",
            country_name="Россия",
            country_code="643",
            sample_name="Все Шаблон договора ФЗЛ СМЗ RUS_Все-Общий_v1",
            version="1",
            file=SimpleUploadedFile(
                "smz-all-contract.docx",
                b"fake-smz-all-docx-content",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True)
        self.assertNotIn("warnings", chunks[-1])
        self.assertEqual(mocked_resolve_variables.call_args.kwargs["contract_details"], ru_details)
        self.assertEqual(mocked_upload_file.call_args.args[2], b"fake-smz-all-docx-content")

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/smz-edited-all-doc")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/02 Исполнители/000 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_matches_template_edited_to_all_group_and_all_product(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        from contracts_app.forms import ContractTemplateForm

        person = PersonRecord.objects.create(
            last_name="Смирнова",
            first_name="Элина",
            middle_name="Юрьевна",
            position=1,
        )
        self.employee.person_record = person
        self.employee.save(update_fields=["person_record"])
        ru = OKSMCountry.objects.create(number=643, code="643", short_name="Россия", alpha2="RU", alpha3="RUS")
        ru_citizenship = CitizenshipRecord.objects.create(person=person, country=ru, position=1)
        profile = ExpertProfile.objects.create(employee=self.employee, position=1)
        ru_details = ExpertContractDetails.objects.create(
            expert_profile=profile,
            citizenship_record=ru_citizenship,
            self_employed=date(2026, 1, 1),
        )
        smz_template = ContractTemplate.objects.create(
            group_member=self.project.group_member,
            product=self.product,
            contract_type="smz",
            party="individual",
            country_name="Россия",
            country_code="643",
            sample_name="RU Шаблон договора ФЗЛ СМЗ RUS_DD-Общий_v1",
            version="1",
            file=SimpleUploadedFile(
                "smz-specific-contract.docx",
                b"fake-edited-smz-docx-content",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )
        smz_template.products.set([self.product])
        form = ContractTemplateForm(
            data={
                "group_member_ids": ["__all__"],
                "product_ids": ["__all__"],
                "contract_type": "smz",
                "party": "individual",
                "country": str(ru.pk),
                "sample_name": smz_template.sample_name,
                "version": smz_template.version,
                "section_ids": ["__all__"],
            },
            instance=smz_template,
        )
        self.assertTrue(form.is_valid(), form.errors)
        form.save()
        smz_template.refresh_from_db()
        self.assertIsNone(smz_template.group_member_id)
        self.assertFalse(smz_template.group_members.exists())
        self.assertIsNone(smz_template.product_id)
        self.assertFalse(smz_template.products.exists())
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True, chunks)
        self.assertNotIn("warnings", chunks[-1])
        self.assertEqual(mocked_resolve_variables.call_args.kwargs["contract_details"], ru_details)
        self.assertEqual(mocked_upload_file.call_args.args[2], b"fake-edited-smz-docx-content")

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/group-specific")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/02 Исполнители/000 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_prefers_specific_group_over_all_group_template(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        project_group = GroupMember.objects.create(
            short_name="RU",
            country_name="Россия",
            country_code="643",
            country_alpha2="RU",
        )
        self.project.group_member = project_group
        self.project.save(update_fields=["group_member"])
        group_specific_template = ContractTemplate.objects.create(
            group_member=project_group,
            product=self.product,
            contract_type="gph",
            party="individual",
            country_name="",
            sample_name="RU Шаблон договора ФЗЛ ГПХ RUS_DD-Общий_v1",
            version="1",
            file=SimpleUploadedFile(
                "specific-contract.docx",
                b"group-specific-docx-content",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )
        group_specific_template.group_members.set([project_group])
        ContractTemplate.objects.create(
            group_member=None,
            product=self.product,
            contract_type="gph",
            party="individual",
            country_name="",
            sample_name="Все Шаблон договора ФЗЛ ГПХ RUS_DD-Общий_v99",
            version="99",
            file=SimpleUploadedFile(
                "all-contract.docx",
                b"all-group-docx-content",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True)
        self.assertEqual(mocked_upload_file.call_args.args[2], b"group-specific-docx-content")

    @patch("nextcloud_app.provisioning.ensure_nextcloud_account")
    @patch("contracts_app.variable_resolver.resolve_variables", return_value=({}, {}))
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_user_share")
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_public_link_share", return_value="https://cloud.example.com/s/product-specific")
    @patch("nextcloud_app.api.NextcloudApiClient.upload_file", return_value=True)
    @patch("nextcloud_app.api.NextcloudApiClient.ensure_folder", return_value="/Corporate Root/02 Договоры/2026/5002 DD Контрактный проект/02 Исполнители/000 Иванов ИИ")
    @patch("nextcloud_app.api.NextcloudApiClient.list_resources", return_value=[])
    def test_create_contract_project_prefers_specific_product_over_all_product_template(
        self,
        mocked_list_resources,
        mocked_ensure_folder,
        mocked_upload_file,
        mocked_public_share,
        mocked_ensure_user_share,
        mocked_resolve_variables,
        mocked_ensure_nextcloud_account,
    ):
        ContractTemplate.objects.create(
            group_member=None,
            product=None,
            contract_type="gph",
            party="individual",
            country_name="",
            sample_name="Все Шаблон договора ФЗЛ ГПХ RUS_Все-Общий_v99",
            version="99",
            file=SimpleUploadedFile(
                "all-product-contract.docx",
                b"all-product-docx-content",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            is_all_sections=True,
        )
        self.template.products.set([self.product])
        mocked_ensure_nextcloud_account.side_effect = lambda user, client=None: (
            self.executor_link if user.pk == self.recipient_user.pk else self.lawyer_link
        )

        response = self.client.post(
            reverse("create_contract_project"),
            {"performer_ids[]": [self.performer.pk]},
        )

        self.assertEqual(response.status_code, 200)
        chunks = [json.loads(chunk.decode("utf-8").strip()) for chunk in response.streaming_content]
        self.assertEqual(chunks[-1]["ok"], True)
        self.assertEqual(mocked_upload_file.call_args.args[2], b"fake-docx-content")

    def test_contract_request_notification_uses_existing_nextcloud_link(self):
        self.performer.contract_project_link = "https://cloud.example.com/s/public-doc"
        self.performer.contract_pdf_link = "https://cloud.example.com/s/public-pdf"
        self.performer.save(update_fields=["contract_project_link", "contract_pdf_link"])
        request_sent_at = timezone.localtime().replace(second=0, microsecond=0)

        response = self.client.post(
            reverse("contract_request"),
            {
                "performer_ids[]": [self.performer.pk],
                "duration_hours": "24",
                "request_sent_at": request_sent_at.isoformat(),
                "delivery_channels[]": ["system"],
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["delivery_channels"], ["system"])
        self.assertFalse(payload["email_delivery"]["requested"])
        self.performer.refresh_from_db()
        self.assertEqual(self.performer.contract_sent_at, request_sent_at)
        notification = Notification.objects.get(notification_type=Notification.NotificationType.PROJECT_CONTRACT_CONCLUSION)
        self.assertEqual(notification.payload["document_docx_link"], "https://cloud.example.com/s/public-doc")
        self.assertEqual(notification.payload["document_pdf_link"], "https://cloud.example.com/s/public-pdf")
        self.assertEqual(notification.payload["document_link"], "https://cloud.example.com/s/public-doc")
        self.assertIn("https://cloud.example.com/s/public-doc", notification.content_text)
        self.assertIn("https://cloud.example.com/s/public-pdf", notification.content_text)

    def test_contract_request_can_send_only_system_email(self):
        self.performer.contract_project_link = "https://cloud.example.com/s/public-doc"
        self.performer.save(update_fields=["contract_project_link"])
        request_sent_at = timezone.localtime().replace(second=0, microsecond=0)

        with patch("notifications_app.services.send_notification_email") as mocked_send:
            with self.captureOnCommitCallbacks(execute=True):
                response = self.client.post(
                    reverse("contract_request"),
                    {
                        "performer_ids[]": [self.performer.pk],
                        "duration_hours": "24",
                        "request_sent_at": request_sent_at.isoformat(),
                        "delivery_channels[]": ["system_email"],
                    },
                )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["delivery_channels"], ["system_email"])
        self.assertTrue(payload["email_delivery"]["requested"])
        self.assertEqual(payload["email_delivery"]["attempted"], 1)
        self.assertFalse(
            Notification.objects.filter(
                notification_type=Notification.NotificationType.PROJECT_CONTRACT_CONCLUSION,
            ).exists()
        )
        mocked_send.assert_called_once()
        call_kwargs = mocked_send.call_args.kwargs
        self.assertEqual(call_kwargs["recipient"], self.recipient_user)
        self.assertIn("https://cloud.example.com/s/public-doc", call_kwargs["content"])
