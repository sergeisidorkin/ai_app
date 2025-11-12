# macroops_app/docx_builder.py
from __future__ import annotations
from io import BytesIO
from base64 import b64encode
from typing import List, Dict, Any
import os
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm
from docx.enum.style import WD_STYLE_TYPE
from collections import defaultdict

try:
    from logs_app import utils as plog
except Exception:
    plog = None

TEMPLATE_WITH_FOOTNOTES = os.path.join(
    os.path.dirname(__file__), "assets", "blank_with_footnotes.docx"
)

SRC_MARKER = "&SRC&"  # единый маркер-делимитер

# Пытаемся импортировать адаптер сносок; если файла нет — будет мягкий фолбэк
try:
    from .word_footnotes import add_footnote_to_paragraph  # адаптер
except Exception:
    add_footnote_to_paragraph = None  # type: ignore


# ─────────────────────────────────────────────────────────────────────────────
# Централизованный логгер с пробросом trace_id/request_id/project_code6/user

def _get_plog():
    global plog
    if plog is None:
        try:
            from logs_app import utils as _plog
            plog = _plog
        except Exception:
            plog = None
    return plog

def _log(level: str, event: str, *, message: str = "",
         log_ctx: Dict[str, Any] | None = None,
         data: Dict[str, Any] | None = None):
    p = _get_plog()
    if not p:
        return
    ctx = log_ctx or {}

    # Только разрешённые именованные поля в сам вызов:
    trace_id = ctx.get("trace_id")
    email    = ctx.get("email")

    # Всё остальное (request_id, project_code6 и пр.) — в data
    payload = dict(data or {})
    if ctx.get("request_id") is not None:
        payload.setdefault("request_id", ctx["request_id"])
    if ctx.get("project_code6") is not None:
        payload.setdefault("project_code6", ctx["project_code6"])

    fn = getattr(p, level, p.info)
    try:
        fn(
            None,
            phase="docx",
            event=event,
            message=message,
            trace_id=trace_id,
            email=email,
            data=payload,
        )
    except TypeError:
        # ultra-safe fallback: минимум параметров
        try:
            p.info(None, phase="docx", event=event,
                   message=f"{message} [fallback]",
                   data=payload)
        except Exception:
            pass
    except Exception:
        pass
# ─────────────────────────────────────────────────────────────────────────────

def _clear_document_body(doc):
    body = doc._element.body
    for el in list(body):
        if el.tag in (qn('w:p'), qn('w:tbl')):
            body.remove(el)
    doc.add_paragraph()

def _load_document(log_ctx: Dict[str, Any] | None = None):
    from docx import Document
    try:
        if os.path.exists(TEMPLATE_WITH_FOOTNOTES):
            doc = Document(TEMPLATE_WITH_FOOTNOTES)
            _log("info", "template.used", log_ctx=log_ctx, data={"path": TEMPLATE_WITH_FOOTNOTES})
            _clear_document_body(doc)
            return doc
    except Exception as e:
        _log("warn", "template.failed", message=str(e), log_ctx=log_ctx)
    return Document()

def _force_paragraph_pstyle(paragraph, style_id: str):
    """Принудительно ставит w:pStyle для абзаца по styleId (без опоры на doc.styles[])."""
    pPr = paragraph._p.get_or_add_pPr()
    # убрать предыдущий pStyle, если был
    for node in list(pPr):
        if node.tag == qn('w:pStyle'):
            pPr.remove(node)
            break
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)

def _find_paragraph_style_by_id(doc, style_id: str):
    try:
        for s in doc.styles:
            try:
                if s.type == WD_STYLE_TYPE.PARAGRAPH and getattr(s, "style_id", None) == style_id:
                    return s
            except Exception:
                continue
    except Exception:
        pass
    return None

def _find_paragraph_style_by_name(doc, names: list[str]):
    try:
        targets = set(n.strip() for n in names if n)
        for s in doc.styles:
            try:
                if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name in targets:
                    return s
            except Exception:
                continue
    except Exception:
        pass
    return None

def _get_builtin_caption_style(doc):
    """
    Возвращает ОБЪЕКТ встроенного стиля подписи с styleId='Caption'.
    Ничего не создаёт и не переименовывает.
    """
    # 1) Прямой доступ по styleId (python-docx умеет по display name и по id)
    try:
        return doc.styles['Caption']
    except Exception:
        pass

    # 2) По локализованным display-name (на всякий случай)
    for nm in ("Название объекта", "Подпись"):
        try:
            s = doc.styles[nm]
            # Проверим, что это именно Caption под капотом (по возможности)
            try:
                if getattr(s, "style_id", None) == "Caption":
                    return s
            except Exception:
                return s
        except Exception:
            continue

    # 3) Последний шанс — найти в XML styleId='Caption'
    try:
        for s in doc.styles:
            try:
                if s.type == WD_STYLE_TYPE.PARAGRAPH and getattr(s, "style_id", None) == "Caption":
                    return s
            except Exception:
                continue
    except Exception:
        pass

    # 4) Нет встроенного Caption в коллекции — не создаём ничего, вернём None
    return None

def _apply_caption_style(paragraph, doc):
    _ensure_caption_style_in_styles_part(doc)
    _force_paragraph_pstyle(paragraph, "Caption")

def _ensure_caption_style(doc):
    """
    Гарантируем наличие параграфного стиля 'Название объекта'.
    Если его нет — создаём через официальное API (add_style), чтобы python-docx «видел» стиль.
    """
    try:
        styles = doc.styles
        # уже есть?
        for s in styles:
            try:
                if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name == "Название объекта":
                    return
            except Exception:
                continue
        # создаём
        s = styles.add_style("Название объекта", WD_STYLE_TYPE.PARAGRAPH)
        try:
            s.base_style = styles["Normal"]
        except Exception:
            pass
    except Exception:
        pass

def _ensure_caption_style_in_styles_part(doc):
    """
    Гарантируем наличие paragraph-стиля с styleId='Caption' и display-name 'Caption'.
    Не создаём/не держим в сниппете стиль с именем «Название объекта», чтобы Word не плодил «…1».
    """
    styles = doc.styles

    # 1) Уже есть стиль с ID 'Caption'?
    for s in styles:
        try:
            if s.type == WD_STYLE_TYPE.PARAGRAPH and getattr(s, "style_id", None) == "Caption":
                # Если у него имя не 'Caption' — поправим на 'Caption'
                try:
                    name_el = s._element.find(qn("w:name"))
                    if name_el is not None:
                        name_el.set(qn("w:val"), "Caption")
                except Exception:
                    pass
                return s
        except Exception:
            continue

    # 2) Есть стиль по имени 'Caption' — выравниваем ID
    try:
        s = styles["Caption"]
        if s.type == WD_STYLE_TYPE.PARAGRAPH:
            try:
                s._element.set(qn("w:styleId"), "Caption")
                name_el = s._element.find(qn("w:name"))
                if name_el is not None:
                    name_el.set(qn("w:val"), "Caption")
                if s._element.find(qn("w:qFormat")) is None:
                    s._element.append(OxmlElement("w:qFormat"))
            except Exception:
                pass
            return s
    except Exception:
        pass

    # 3) Создаём новый 'Caption' (имя и id совпадают)
    s = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    try:
        s.base_style = styles["Normal"]
    except Exception:
        pass
    try:
        s._element.set(qn("w:styleId"), "Caption")
        name_el = s._element.find(qn("w:name"))
        if name_el is not None:
            name_el.set(qn("w:val"), "Caption")
        if s._element.find(qn("w:qFormat")) is None:
            s._element.append(OxmlElement("w:qFormat"))
    except Exception:
        pass
    return s

def _ensure_builtin_caption_style_stub(doc):
    """
    Гарантирует наличие в styles.xml минимального paragraph-стиля с ID 'Caption'.
    Создаём стиль с display-name 'Caption' (англ.), qFormat. Ничего не переименовываем.
    Это нужно, чтобы w:pStyle w:val="Caption" НЕ падал в «Обычный» при merge.
    """
    try:
        styles_part = doc.part._styles  # StylesPart
        root = styles_part.element
        W = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Уже есть стиль с таким ID?
        if root.xpath("//w:style[@w:type='paragraph' and @w:styleId='Caption']", namespaces=W):
            return

        # Создаём минимальный <w:style w:type="paragraph" w:styleId="Caption">
        st = OxmlElement('w:style')
        st.set(qn('w:type'), 'paragraph')
        st.set(qn('w:styleId'), 'Caption')

        nm = OxmlElement('w:name'); nm.set(qn('w:val'), 'Caption'); st.append(nm)
        qf = OxmlElement('w:qFormat'); st.append(qf)

        # (не задаём жирных свойств — примем оформление из принимающего документа)
        root.append(st)
    except Exception:
        pass

def _resolve_caption_style_name(doc) -> str:
    """Возвращает лучшее имя стиля подписи (в порядке приоритета)."""
    try:
        for s in doc.styles:
            try:
                if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name == "Название объекта":
                    return "Название объекта"
            except Exception:
                continue
        for fallback in ("Caption", "Подпись"):
            for s in doc.styles:
                try:
                    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name == fallback:
                        return fallback
                except Exception:
                    continue
    except Exception:
        pass
    return "Caption"

def _apply_paragraph_style_strict(paragraph, doc, preferred_name: str | None):
    """
    Ставит стиль параграфу надёжно:
    1) пробуем объект стиля с именем preferred_name;
    2) иначе — 'Название объекта';
    3) иначе — 'Caption'/'Подпись';
    4) в конце — попытка присвоить строкой (на случай кастомных шаблонов).
    """
    names_try = []
    if preferred_name:
        names_try.append(preferred_name)
    names_try.extend(["Название объекта", "Caption", "Подпись"])

    try:
        for nm in names_try:
            for s in doc.styles:
                try:
                    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name == nm:
                        paragraph.style = s  # ставим объектом — это самый надёжный способ
                        return
                except Exception:
                    continue
        # последняя попытка — строкой
        if preferred_name:
            try:
                paragraph.style = preferred_name
                return
            except Exception:
                pass
        # и ещё раз строкой с вычисленным фолбэком
        try:
            paragraph.style = _resolve_caption_style_name(doc)
        except Exception:
            pass
    except Exception:
        pass

def _append_field(paragraph, instr_text: str, displayed_text: str = ""):
    """Вставляет поле { <instr_text> } с показанным значением displayed_text."""
    def w(tag): return qn(f"w:{tag}")

    r_begin = OxmlElement('w:r'); fld_begin = OxmlElement('w:fldChar'); fld_begin.set(w('fldCharType'), 'begin'); r_begin.append(fld_begin)
    paragraph._p.append(r_begin)

    r_instr = OxmlElement('w:r'); instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = instr_text
    r_instr.append(instr)
    paragraph._p.append(r_instr)

    r_sep = OxmlElement('w:r'); fld_sep = OxmlElement('w:fldChar'); fld_sep.set(w('fldCharType'), 'separate'); r_sep.append(fld_sep)
    paragraph._p.append(r_sep)

    r_val = OxmlElement('w:r'); t_val = OxmlElement('w:t'); t_val.text = displayed_text or "1"; r_val.append(t_val)
    paragraph._p.append(r_val)

    r_end = OxmlElement('w:r'); fld_end = OxmlElement('w:fldChar'); fld_end.set(w('fldCharType'), 'end'); r_end.append(fld_end)
    paragraph._p.append(r_end)

def _append_caption_number(paragraph, *, seq_label: str, seq_display: int = 1):
    """
    Вставляет номер вида: { STYLEREF 1 \s } - { SEQ <seq_label> \* ARABIC \s 1 }.
    seq_label: 'Табл.' | 'Рис.' и т.п.
    """
    # № главы (из Heading 1) → { STYLEREF 1 \s }
    _append_field(paragraph, " STYLEREF 1 \\s ", "1")
    # литерал «-»
    paragraph.add_run("-")
    # локальный порядок с \s 1 → сброс внутри главы 1
    _append_field(paragraph, f" SEQ {seq_label} \\* ARABIC \\s 1 ", str(int(seq_display)))

def _append_seq_field(paragraph, label: str, chapter_level: int | None, seq_value: int):
    """Вставляет в paragraph поле { SEQ <label> [\s N] } с отображённым номером seq_value."""
    def w(tag): return qn(f"w:{tag}")

    # { SEQ Label \* ARABIC [\s N] }
    instr = f" SEQ {label} \\* ARABIC"
    if isinstance(chapter_level, int) and chapter_level > 0:
        instr += f" \\s {int(chapter_level)}"

    r_begin = OxmlElement('w:r'); fld_begin = OxmlElement('w:fldChar'); fld_begin.set(w('fldCharType'), 'begin'); r_begin.append(fld_begin)
    r_instr = OxmlElement('w:r'); instr_text = OxmlElement('w:instrText'); instr_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve'); instr_text.text = instr; r_instr.append(instr_text)
    r_sep   = OxmlElement('w:r'); fld_sep   = OxmlElement('w:fldChar');  fld_sep.set(w('fldCharType'), 'separate'); r_sep.append(fld_sep)
    r_val   = OxmlElement('w:r'); t_val     = OxmlElement('w:t');         t_val.text = str(seq_value); r_val.append(t_val)
    r_end   = OxmlElement('w:r'); fld_end   = OxmlElement('w:fldChar');   fld_end.set(w('fldCharType'), 'end'); r_end.append(fld_end)

    paragraph._p.append(r_begin)
    paragraph._p.append(r_instr)
    paragraph._p.append(r_sep)
    paragraph._p.append(r_val)
    paragraph._p.append(r_end)

def _add_caption_paragraph(
    doc,
    *,
    label: str,                   # "Табл." / "Рис."
    seq_label: str | None = None, # метка для SEQ
    text: str,
    place_above_of=None
):
    p = doc.add_paragraph()

    # стиль подписи
    _apply_caption_style(p, doc)

    # "Табл. " + номер + " " + текст
    p.add_run(f"{label} ")
    key = (seq_label or label)
    if not hasattr(doc, "_seq_counters"):
        from collections import defaultdict
        doc._seq_counters = defaultdict(int)  # type: ignore[attr-defined]
    doc._seq_counters[key] += 1               # type: ignore[attr-defined]
    _append_caption_number(p, seq_label=key, seq_display=doc._seq_counters[key])  # type: ignore[attr-defined]
    p.add_run(" ")
    if text:
        p.add_run(text)

    if place_above_of is not None:
        try:
            place_above_of._element.addprevious(p._p)
        except Exception:
            pass
    return p

def _inject_footnote_style_primer(doc):
    """
    Вставляет в body два скрытых абзаца, чтобы Word гарантированно импортировал
    paragraph-стиль 'FootnoteText' и character-стиль 'FootnoteReference'
    при insertFileFromBase64. Эти абзацы состоят из скрытого нулевой ширины символа.
    """

    # 1) Абзац со стилем FootnoteText (Текст сноски) — скрытый
    p1 = doc.add_paragraph()
    for name in ("FootnoteText", "Footnote Text", "Текст сноски"):
        try:
            p1.style = name
            break
        except Exception:
            continue
    r1 = p1.add_run("\u200B")  # zero-width space
    try:
        r1.font.hidden = True
    except Exception:
        pass
    # на всякий случай дубль скрытия через oxml
    try:
        rPr = r1._r.get_or_add_rPr()
        vanish = OxmlElement("w:vanish")
        rPr.append(vanish)
    except Exception:
        pass

    # 2) Абзац с run, у которого char-style FootnoteReference (Знак сноски) — скрытый
    p2 = doc.add_paragraph()
    r2 = p2.add_run("\u200B")
    try:
        r2.font.hidden = True
    except Exception:
        pass

    # пробуем назначить стиль по имени (локаль-френдли)…
    applied = False
    for name in ("Footnote Reference", "Знак сноски"):
        try:
            r2.style = name
            applied = True
            break
        except Exception:
            continue
    # …если не получилось — задаём styleId напрямую
    if not applied:
        try:
            rPr = r2._r.get_or_add_rPr()
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), "FootnoteReference")
            rPr.append(rStyle)
        except Exception:
            pass

def build_docx_from_paragraphs(paragraphs: list[dict], *, log_ctx: Dict[str, Any] | None = None) -> str:
    try:
        doc = _load_document(log_ctx=log_ctx)
    except Exception as e:
        raise RuntimeError("python-docx is required to build .docx") from e

    for p in (paragraphs or []):
        text = (p.get("text") or "").strip()
        if text == "":
            doc.add_paragraph("")
            continue

        para = doc.add_paragraph()
        try:
            style = (p.get("styleBuiltIn") or p.get("style") or "").strip() or "Normal"
            para.style = style
        except Exception:
            try:
                para.style = "Normal"
            except Exception:
                pass
        para.add_run(text)

    if not doc.paragraphs:
        doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio)
    raw = bio.getvalue()
    _log("debug", "zip.signature", message="docx zip signature", log_ctx=log_ctx, data={"sig": raw[:2].hex()})
    return b64encode(raw).decode("ascii")

def build_docx_from_ops(ops: List[Dict[str, Any]], *, log_ctx: Dict[str, Any] | None = None) -> str:
    try:
        doc = _load_document(log_ctx=log_ctx)
        _ensure_caption_style_in_styles_part(doc)
    except Exception as e:
        raise RuntimeError("python-docx is required to build .docx") from e

    saw_footnote = False

    def _set_style(paragraph, style_name: str | None):
        if not style_name:
            return
        try:
            paragraph.style = style_name
        except Exception:
            pass

    def _list_style_from_listType(list_type: str | None, fallback: str = "List Bullet") -> str:
        lt = (list_type or "").strip()
        if lt in ("ListBullet", "List Bullet"):
            return "List Bullet"
        if lt in ("ListNumber", "List Number", "Numbered"):
            return "List Number"
        return fallback

    current_list_style: str | None = None
    table = None
    table_cols = 0
    cur_row_idx = -1
    cur_col_idx = 0
    row_bold = False
    table_style = "Table Grid"
    last_paragraph = None

    def _add_paragraph(text: str, style_name: str | None = None):
        nonlocal last_paragraph
        p = doc.add_paragraph()
        if style_name:
            _set_style(p, style_name)
        if text:
            p.add_run(text)
        last_paragraph = p
        return p

    for op in (ops or []):
        if not isinstance(op, dict):
            continue
        kind = op.get("op")

        if kind == "paragraph.insert":
            text = (op.get("text") or "")
            style_name = (op.get("styleBuiltIn") or op.get("style") or "Normal")
            last_paragraph = _add_paragraph(text, style_name)

        elif kind == "list.start":
            list_type = (op.get("listType") or op.get("list_type"))
            style_name = (op.get("styleBuiltIn") or op.get("style") or None)
            current_list_style = style_name or _list_style_from_listType(list_type)

        elif kind == "list.item":
            text = (op.get("text") or "")
            style_name = (op.get("styleBuiltIn") or op.get("style") or current_list_style or "List Bullet")
            last_paragraph = _add_paragraph(text, style_name)

        elif kind == "list.end":
            current_list_style = None

        elif kind == "image.insert":
            import base64 as _b64
            raw = _b64.b64decode((op.get("base64") or "").encode("ascii"))
            stream = BytesIO(raw)
            pic = doc.add_picture(stream)
            # размеры
            try:
                wmm = op.get("widthMm"); hmm = op.get("heightMm")
                if wmm: pic.width = Mm(float(wmm))
                if hmm: pic.height = Mm(float(hmm))
            except Exception:
                pass
            last_paragraph = doc.paragraphs[-1] if doc.paragraphs else None
            last_picture = pic  # можете сохранить, если захотите «above»
            continue

        elif kind == "caption.add":
            visible_label = (op.get("label") or ("Таблица" if (op.get("target") == "table") else "Рисунок"))
            seq_label = (
                    op.get("seqLabel")
                    or ("Табл." if (op.get("target") == "table" or "таб" in visible_label.lower()) else "Рис.")
            )
            text = (op.get("text") or "")
            placement = (op.get("placement") or "below").lower()
            style_name = (op.get("style") or None)

            place_above_of = None
            if placement == "above":
                if table is not None:
                    place_above_of = table._tbl
                elif last_paragraph is not None:
                    place_above_of = last_paragraph

            pcap = _add_caption_paragraph(
                doc,
                label=visible_label,
                seq_label=seq_label,
                text=text,
                place_above_of=place_above_of
            )
            last_paragraph = pcap
            continue

        elif kind == "table.start":
            table = None
            cur_row_idx = -1
            cur_col_idx = 0
            row_bold = False

            table_cols = int(op.get("cols") or op.get("columns") or 0) or 2
            table_style = (op.get("tableStyle") or op.get("table_style") or "Table Grid")

            table = doc.add_table(rows=1, cols=table_cols)
            try:
                table.style = table_style
            except Exception:
                pass

            cur_row_idx = -1
            cur_col_idx = 0
            row_bold = False

        elif kind == "table.row":
            if table is None:
                table_cols = 2
                table_style = "Table Grid"
                table = doc.add_table(rows=1, cols=table_cols)
                try:
                    table.style = table_style
                except Exception:
                    pass
                cur_row_idx = -1

            cur_row_idx += 1
            _row = table.rows[0] if cur_row_idx == 0 else table.add_row()
            cur_col_idx = 0
            row_bold = bool(op.get("header") or op.get("isHeader") or op.get("is_header"))

        elif kind == "table.cell":
            if table is None or cur_row_idx < 0:
                continue
            text = (op.get("text") or "")
            style_name = (op.get("styleBuiltIn") or op.get("style") or None)

            if 0 <= cur_col_idx < table_cols:
                cell = table.rows[cur_row_idx].cells[cur_col_idx]
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                if style_name:
                    _set_style(p, style_name)
                run = p.add_run(text)
                if row_bold:
                    run.bold = True
                last_paragraph = p
            cur_col_idx += 1

        elif kind == "table.end":
            table = None
            table_cols = 0
            cur_row_idx = -1
            cur_col_idx = 0
            row_bold = False

        elif kind == "footnote.add":
            # ВМЕСТО реальной сноски — вставляем в последний параграф: " &SRC&{TEXT}&SRC&"
            text = (op.get("text") or "")
            if last_paragraph is None:
                last_paragraph = _add_paragraph("", "Normal")
            # отделяем пробелом от слова, чтобы не слипалось
            last_paragraph.add_run(" " + SRC_MARKER + text + SRC_MARKER)

        else:
            continue

    if not doc.paragraphs and (not doc.tables):
        doc.add_paragraph("")

    if saw_footnote:
        try:
            _inject_footnote_style_primer(doc)
            _log("debug", "footnote.style.primer", log_ctx=log_ctx, data={"ok": True})
        except Exception as e:
            _log("warn", "footnote.style.primer", message=str(e), log_ctx=log_ctx, data={"ok": False})

    bio = BytesIO()
    doc.save(bio)
    raw = bio.getvalue()
    _log("debug", "zip.signature", message="docx zip signature", log_ctx=log_ctx, data={"sig": raw[:2].hex()})
    return b64encode(raw).decode("ascii")