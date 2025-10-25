import msal, json, requests
from io import BytesIO
from django.conf import settings
from docx import Document
from .models import OneDriveAccount
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

GRAPH_API = "https://graph.microsoft.com/v1.0"

def _load_cache(od_acc):
    cache = msal.SerializableTokenCache()
    if od_acc and od_acc.token_cache:
        cache.deserialize(od_acc.token_cache)
    return cache

def _save_cache(od_acc, cache):
    if cache.has_state_changed:
        od_acc.token_cache = cache.serialize()
        od_acc.save(update_fields=["token_cache"])

def _app(cache=None):
    return msal.ConfidentialClientApplication(
        client_id=settings.MS_CLIENT_ID,
        client_credential=settings.MS_CLIENT_SECRET,
        authority=settings.MSAL_AUTHORITY,
        token_cache=cache,
    )

# def get_auth_url(state="connect"):
#     app = _app()
#     return app.get_authorization_request_url(
#         scopes=settings.MS_SCOPES,
#         redirect_uri=settings.MS_REDIRECT_URI,
#         state=state,
#         prompt="select_account",
#     )
#
# def exchange_code(user, code: str):
#     od, _ = OneDriveAccount.objects.get_or_create(user=user)
#     cache = _load_cache(od)
#     app = _app(cache)
#     res = app.acquire_token_by_authorization_code(
#         code=code, scopes=settings.MS_SCOPES, redirect_uri=settings.MS_REDIRECT_URI
#     )
#     if "access_token" not in res:
#         raise RuntimeError(res.get("error_description", "Token error"))
#     _save_cache(od, cache)

def get_auth_url(state="connect", redirect_uri: str | None = None):
    app = _app()
    ru = redirect_uri or settings.MS_REDIRECT_URI
    return app.get_authorization_request_url(
        scopes=settings.MS_SCOPES,
        redirect_uri=ru,
        state=state,
        prompt="select_account",
    )

def exchange_code(user, code: str, redirect_uri: str | None = None):
    od, _ = OneDriveAccount.objects.get_or_create(user=user)
    cache = _load_cache(od)
    app = _app(cache)
    ru = redirect_uri or settings.MS_REDIRECT_URI
    res = app.acquire_token_by_authorization_code(
        code=code, scopes=settings.MS_SCOPES, redirect_uri=ru
    )
    if "access_token" not in res:
        raise RuntimeError(res.get("error_description", "Token error"))
    _save_cache(od, cache)


def _access_token(user):
    od = OneDriveAccount.objects.filter(user=user).first()
    if not od:
        return None
    cache = _load_cache(od)
    app = _app(cache)
    accs = app.get_accounts()
    res = app.acquire_token_silent(scopes=settings.MS_SCOPES, account=accs[0]) if accs else None
    if not res:
        return None
    _save_cache(od, cache)
    return res["access_token"]

def graph_get_json(user, url):
    token = _access_token(user)
    if not token: return None
    r = requests.get(url, headers={"Authorization": f"Bearer {token}"})
    r.raise_for_status()
    return r.json()

def list_children(user, item_id=None):
    url = f"{GRAPH_API}/me/drive/root/children" if not item_id else f"{GRAPH_API}/me/drive/items/{item_id}/children"
    return graph_get_json(user, url)

def download_item_bytes(user, item_id):
    token = _access_token(user)
    if not token: raise RuntimeError("No token")
    r = requests.get(f"{GRAPH_API}/me/drive/items/{item_id}/content",
                     headers={"Authorization": f"Bearer {token}"})
    r.raise_for_status()
    return r.content  # bytes

def upload_item_bytes(user, item_id, data: bytes):
    token = _access_token(user)
    if not token: raise RuntimeError("No token")
    r = requests.put(f"{GRAPH_API}/me/drive/items/{item_id}/content",
                     headers={"Authorization": f"Bearer {token}"}, data=data)
    r.raise_for_status()
    return r.json()

def append_text_to_docx(user, item_id, text: str):
    """Скачать .docx, добавить текст (не стирая существующий), загрузить обратно."""
    src = download_item_bytes(user, item_id)
    doc = Document(BytesIO(src))  # открываем существующий документ
    for line in (text or "").splitlines() or [""]:
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return upload_item_bytes(user, item_id, buf.read())

def _insert_paragraph_after(paragraph, text="", copy_style=True):
    """
    Вставляет новый абзац сразу после переданного paragraph.
    Возвращает созданный Paragraph.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if copy_style:
        try:
            new_para.style = paragraph.style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para

def upload_new_file_in_folder(user, folder_id: str, name: str, data: bytes):
    """
    Создать НОВЫЙ файл в указанной папке OneDrive и залить содержимое.
    Для небольших файлов (<4 МБ) достаточно простого PUT:
    PUT /me/drive/items/{folder_id}:/{name}:/content
    """
    token = _access_token(user)
    if not token:
        raise RuntimeError("No token")

    # безопасно экранировать имя (Graph сам обработает запрещённые символы)
    url = f"{GRAPH_API}/me/drive/items/{folder_id}:/{name}:/content"
    r = requests.put(url, headers={"Authorization": f"Bearer {token}"}, data=data)
    r.raise_for_status()
    return r.json()

def append_text_after_marker(user, item_id: str, marker: str, text: str, insert_after_all=False):
    """
    Ищет `marker` в тексте абзацев (включая абзацы в таблицах).
    Вставляет новый(ые) абзац(ы) сразу после абзаца с маркером.
    По умолчанию — только первое совпадение (insert_after_all=False).
    """
    src = download_item_bytes(user, item_id)
    doc = Document(BytesIO(src))

    found = False

    def process_paragraphs(paragraphs):
        nonlocal found
        for p in paragraphs:
            if marker and (marker in (p.text or "")) and (insert_after_all or not found):
                # вставляем построчно: каждая строка промпта — отдельный абзац
                prev = p
                lines = (text or "").splitlines() or [""]
                for line in lines:
                    prev = _insert_paragraph_after(prev, line, copy_style=True)
                found = True

    # 1) Обычные абзацы
    process_paragraphs(doc.paragraphs)

    # 2) Абзацы внутри таблиц (простая рекурсия по ячейкам)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    if not found:
        raise RuntimeError(f"Маркер '{marker}' не найден в документе")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return upload_item_bytes(user, item_id, buf.read())

def get_access_token_for_user(user):
    """Публичный враппер, чтобы _auth_headers не падал."""
    return _access_token(user)

def list_permissions(user, drive_id, item_id):
    """GET /drives/{drive-id}/items/{item-id}/permissions"""
    token = get_access_token_for_user(user)
    if not token:
        raise RuntimeError("No token")
    url = f"{GRAPH_API}/drives/{drive_id}/items/{item_id}/permissions"
    r = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=30)
    r.raise_for_status()
    return r.json()

def search_in_folder(user, folder_id, query):
    """GET /me/drive/items/{folder_id}/search(q='...')"""
    token = _access_token(user)
    if not token:
        raise RuntimeError("No token")
    url = f"{GRAPH_API}/me/drive/items/{folder_id}/search(q='{query}')"
    r = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=30)
    r.raise_for_status()
    return r.json()

def _auth_headers(user):
    # у тебя уже должен быть хелпер получения токена Microsoft Graph.
    # Используй тот же, что в list_children (достаёт access_token для user).
    token = get_access_token_for_user(user)  # твоя существующая функция
    return {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

def create_link(user, drive_id, item_id, link_type="edit", scope="organization"):
    """
    POST /drives/{drive-id}/items/{item-id}/createLink
    body: {"type": "edit"|"view", "scope": "anonymous"|"organization"|"users"}
    """
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/createLink"
    payload = {"type": link_type, "scope": scope}
    r = requests.post(url, headers=_auth_headers(user), json=payload, timeout=30)
    r.raise_for_status()
    return r.json()  # {"link": {"webUrl": "...", "type": "...", ...}, ...}

def invite_recipients(user, drive_id, item_id, emails, role="write", send_invitation=False, message=None):
    """
    POST /drives/{drive-id}/items/{item-id}/invite
    Даст прямые права конкретным людям.
    """
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/invite"
    recipients = [{"email": e} for e in emails]
    body = {
        "recipients": recipients,
        "requireSignIn": True,
        "sendInvitation": bool(send_invitation),
        "message": message or "",
        "roles": ["write" if role == "write" else "read"],
    }
    r = requests.post(url, headers=_auth_headers(user), json=body, timeout=30)
    try:
        r.raise_for_status()
    except requests.HTTPError as e:
        # полезная диагностика в консоли
        try:
            print("Graph invite error:", r.status_code, r.text[:2000])
        except Exception:
            pass
        raise
    return r.json()
